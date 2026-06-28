import json
import csv
import html
import io
import math
import random
import re
import zipfile
import unicodecsv
import time
import datetime
import sys
from collections import defaultdict

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment

from django.shortcuts import render
from django.forms.formsets import formset_factory
from django.http import HttpResponse, HttpResponseRedirect, JsonResponse

from django.contrib.auth import logout
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .models import *
from .forms import *
from .model_utils import *
from .utils import *
from .packet_parser import parse_packet_data
from .duplicate_checker import find_duplicates, find_internal_issues, find_topic_repeats, find_answer_matches, CRITICAL, WARNING, INFO
from django.utils.safestring import mark_safe
from django_comments.models import Comment
from django.db.models import Q
from django.db import connection
from django.utils import timezone
from django.core.cache import cache


def fulltext_filter(queryset, query):
    """Full-text match against the maintained search_question_content /
    search_question_answers fields. Uses Postgres full-text search in
    production and falls back to icontains on SQLite (local dev)."""
    if not query:
        return queryset.none()
    if connection.vendor == 'postgresql':
        from django.contrib.postgres.search import SearchVector, SearchQuery
        vector = SearchVector('search_question_content', 'search_question_answers')
        return queryset.annotate(qsearch=vector).filter(qsearch=SearchQuery(query, search_type='websearch'))
    return queryset.filter(
        Q(search_question_content__icontains=query) | Q(search_question_answers__icontains=query))

from django.contrib.contenttypes.models import ContentType


def main (request):
    # Logged-out visitors get the public splash page; members get their
    # question-set dashboard.
    if not request.user.is_authenticated:
        return render(request, 'splash.html', {})
    return question_sets(request)

@login_required
def about (request):
    return render(request, 'about.html', {'user': request.user.writer})

@login_required
def help_page (request):
    return render(request, 'help.html', {'user': request.user.writer})

@login_required
def sidebar (request):
    writer = request.user.writer
    # the tournaments for which this user is a writer
    writer_sets = writer.question_set_writer.all()
    # all the tournaments owned by this user
    owned_sets = QuestionSet.objects.filter(owner=writer)
    # the tournaments for which this user is an editor
    editor_sets = writer.question_set_editor.all()

    all_sets = editor_sets
    print('All sets object:')
    print(all_sets)
        
    return render(request, 'sidebar.html', {'question_sets': all_sets, 'user': writer})

@login_required
def question_sets (request):
    writer = request.user.writer

    # all the tournaments owned by this user
    owned_sets = QuestionSet.objects.filter(owner=writer)
    # the tournaments for which this user is an editor
    editor_sets = writer.question_set_editor.all()
    
    all_sets = owned_sets | editor_sets | writer.question_set_writer.all()
    all_sets = all_sets.order_by('date')
    
    # Sets that are upcoming or finished within the last month
    upcoming_sets = {}

    # Sets that are in the past
    completed_sets = {}

    # Keep recently-finished sets in "upcoming" for a month after the
    # tournament date before moving them to "completed".
    from datetime import timedelta
    cutoff = datetime.now().date() - timedelta(days=30)

    for qset in (all_sets):
        if (qset.date >= cutoff):
            upcoming_sets[qset.id] = qset
        else:
            completed_sets[qset.id] = qset
            
    upcoming_sets = upcoming_sets
    completed_sets = completed_sets
            
    upcoming_sets = upcoming_sets.values()
    completed_sets = completed_sets.values()
    
    upcoming_sets = sorted(upcoming_sets, key=lambda qset: qset.date)
    completed_sets = sorted(completed_sets, key=lambda qset: qset.date)

    all_sets  = [{'header': 'Upcoming question sets', 'qsets': upcoming_sets, 'id': 'qsets-write'},
                 {'header': 'Completed question sets', 'qsets': completed_sets, 'id': 'qsets-complete'}]

    # Public sets the user isn't already part of — they can request to join.
    my_set_ids = {qset.id for qset in all_sets[0]['qsets']} | {qset.id for qset in all_sets[1]['qsets']}
    public_sets = [qs for qs in QuestionSet.objects.filter(public=True).order_by('-date')
                   if qs.id not in my_set_ids]

    return render(request, 'question_sets.html',
                  {'question_set_list': all_sets, 'public_sets': public_sets, 'user': writer})

@login_required
def request_to_join(request):
    """A logged-in user asks to join a public question set; emails the owner(s)."""
    user = request.user.writer
    if request.method != 'POST':
        return HttpResponse(json.dumps({'success': False, 'message': 'Invalid request'}))
    try:
        qset = QuestionSet.objects.get(id=int(request.POST['qset_id']))
    except (KeyError, ValueError, QuestionSet.DoesNotExist):
        return HttpResponse(json.dumps({'success': False, 'message': 'Set not found.'}))
    if not qset.public:
        return HttpResponse(json.dumps({'success': False, 'message': 'This set is not public.'}))
    if qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all():
        return HttpResponse(json.dumps({'success': False, 'message': 'You are already part of this set.'}))

    note = (request.POST.get('message') or '').strip()[:1000]
    requester_name = user.get_real_name().strip() or user.user.username
    requester_email = user.user.email
    recipients = [o.user.email for o in qset.all_owners() if o.user and o.user.email]
    if not recipients:
        return HttpResponse(json.dumps({'success': False,
            'message': 'The set owner has no email on file, so the request could not be sent.'}))

    from django.core.mail import EmailMessage
    from django.conf import settings
    approve_url = '{0}/approve_join/{1}/{2}/'.format(
        settings.BASE_URL.rstrip('/'), qset.id, user.id)
    subject = 'QEMS3: {0} requests to join "{1}"'.format(requester_name, qset.name)
    body = ('{0} (@{1}{2}) has requested to join your question set "{3}" on QEMS3.\n\n'
            '{4}\n\n'
            'Approve this request (add them as a writer or editor):\n{5}\n\n'
            'You can also open the set on QEMS3 and use "Add Writer" or "Add Editor".').format(
        requester_name, user.user.username,
        ', ' + requester_email if requester_email else '', qset.name,
        ('Their message: ' + note) if note else '(No message included.)',
        approve_url)
    try:
        EmailMessage(subject, body, settings.DEFAULT_FROM_EMAIL, recipients,
                     reply_to=[requester_email] if requester_email else None).send(fail_silently=False)
    except Exception as ex:
        return HttpResponse(json.dumps({'success': False,
            'message': 'Could not send the request email ({0}).'.format(ex)}))
    return HttpResponse(json.dumps({'success': True,
        'message': 'Your request has been emailed to the set owner.'}))


@login_required
def approve_join(request, qset_id, writer_id):
    """Owner-facing approval of a join request, linked directly from the request
    email. Shows a small confirmation page; on POST adds the requester to the
    set as a writer (default) or editor and emails them that they were added."""
    user = request.user.writer
    try:
        qset = QuestionSet.objects.get(id=int(qset_id))
        requester = Writer.objects.get(id=int(writer_id))
    except (ValueError, QuestionSet.DoesNotExist, Writer.DoesNotExist):
        return render(request, 'failure.html',
                      {'message': 'That set or user no longer exists.',
                       'message_class': 'alert-box alert'})

    if not qset.is_owner(user):
        return render(request, 'failure.html',
                      {'message': 'Only an owner of this set can approve join requests.',
                       'message_class': 'alert-box alert'})

    already = (qset.is_owner(requester) or requester in qset.editor.all()
               or requester in qset.writer.all())

    if request.method == 'POST' and not already:
        role = 'editor' if request.POST.get('role') == 'editor' else 'writer'
        if role == 'editor':
            qset.editor.add(requester)
            GroupRoleGrant.objects.filter(
                question_set=qset, writer=requester, role='editor').delete()
            if qset.writer.filter(id=requester.id).exists():
                qset.writer.remove(requester)
        else:
            qset.writer.add(requester)
            GroupRoleGrant.objects.filter(
                question_set=qset, writer=requester, role='writer').delete()
        qset.save()
        _notify_added_to_set(requester, qset, role, user)
        cache.clear()
        return render(request, 'approve_join.html',
                      {'qset': qset, 'requester': requester, 'added_role': role, 'user': user})

    return render(request, 'approve_join.html',
                  {'qset': qset, 'requester': requester, 'already': already, 'user': user})


def _writer_email(writer):
    return writer.user.email if (writer and writer.user and writer.user.email) else ''


def _actor_name(writer):
    if writer is None:
        return 'An organizer'
    return writer.get_real_name().strip() or writer.user.username


def _notify_added_to_set(writer, qset, role, by_writer):
    """Email a writer that they were added to a set (editor/writer/co-owner)."""
    email = _writer_email(writer)
    if not email:
        return
    from .signals import _send_mail_async
    from django.conf import settings as dj_settings
    role_label = {'editor': 'an editor', 'writer': 'a writer', 'co-owner': 'a co-owner'}.get(role, role)
    subject = 'QEMS3: you were added to "{0}"'.format(qset.name)
    body = ('{0} added you as {1} on the QEMS3 question set "{2}".\n\n'
            '{3}/edit_question_set/{4}/').format(
        _actor_name(by_writer), role_label, qset.name, dj_settings.BASE_URL, qset.id)
    _send_mail_async(subject, body, [email])


def _notify_added_to_group(writer, group, by_writer):
    """Email a writer that they were added to a role group."""
    email = _writer_email(writer)
    if not email:
        return
    from .signals import _send_mail_async
    from django.conf import settings as dj_settings
    subject = 'QEMS3: you were added to the role group "{0}"'.format(group.name)
    body = ('{0} added you to the QEMS3 role group "{1}". Members of a role group '
            'automatically get its role on every question set the group is attached to.\n\n'
            '{2}/role_groups/').format(_actor_name(by_writer), group.name, dj_settings.BASE_URL)
    _send_mail_async(subject, body, [email])


def _notify_group_join_request(group, requester):
    """Email the group owner that someone has requested to join. Returns False if
    the owner has no email on file."""
    email = _writer_email(group.created_by)
    if not email:
        return False
    from .signals import _send_mail_async
    from django.conf import settings as dj_settings
    rname = _actor_name(requester)
    remail = requester.user.email or ''
    subject = 'QEMS3: {0} requests to join role group "{1}"'.format(rname, group.name)
    body = ('{0} (@{1}{2}) has requested to join your QEMS3 role group "{3}".\n\n'
            'To add them, open Role Groups and use "Add member":\n{4}/role_groups/').format(
        rname, requester.user.username, ', ' + remail if remail else '',
        group.name, dj_settings.BASE_URL)
    _send_mail_async(subject, body, [email])
    return True


@login_required
def import_set(request):
    """Admin-only: create a new question set from an uploaded TSV/CSV in the
    export format, including comments, and index it for search."""
    user = request.user.writer

    if not request.user.is_superuser:
        messages.error(request, 'Only the admin account may import sets.')
        return HttpResponseRedirect('/failure.html/')

    message = ''
    message_class = ''
    summary = None

    if request.method == 'POST':
        form = ImportSetForm(request.POST, request.FILES)
        if form.is_valid():
            from .set_importer import import_set_from_file, SetImportError
            try:
                summary = import_set_from_file(
                    form.cleaned_data['set_file'], form.cleaned_data['set_name'], user)
                qset = summary['question_set']
                message = ('Imported "{0}": {1} tossups, {2} bonuses, {3} comments.'
                           .format(qset.name, summary['tossups'], summary['bonuses'],
                                   summary['comments']))
                if summary.get('users_created'):
                    message += ' Created {0} placeholder commenter account(s).'.format(summary['users_created'])
                if summary.get('legacy_authors'):
                    message += ' Created {0} legacy author account(s) for authors with no current account.'.format(summary['legacy_authors'])
                message += ' Search indexing is running in the background and will finish shortly.'
                message_class = 'alert-box success'
            except SetImportError as ex:
                message = str(ex)
                message_class = 'alert-box alert'
            except Exception as ex:
                message = 'Import failed: {0}'.format(ex)
                message_class = 'alert-box alert'
    else:
        form = ImportSetForm()

    return render(request, 'import_set.html',
                  {'form': form, 'user': user, 'summary': summary,
                   'message': message, 'message_class': message_class})

@login_required
def import_packets(request):
    """Admin-only: create a new tournament (question set) from uploaded packet
    files (.docx or .pdf), one packet per file."""
    user = request.user.writer

    if not request.user.is_superuser:
        messages.error(request, 'Only the admin account may import packets.')
        return HttpResponseRedirect('/failure.html/')

    message = ''
    message_class = ''
    summary = None

    if request.method == 'POST':
        form = ImportPacketsForm(request.POST, request.FILES)
        if form.is_valid():
            from .packet_set_importer import (import_packets_from_files,
                                              import_packets_into_set, PacketImportError)
            files = form.cleaned_data['packet_files']
            target = form.cleaned_data.get('target_set')
            try:
                if target is not None:
                    summary = import_packets_into_set(files, target, user)
                    qset = summary['question_set']
                    message = ('Added {0} packet(s) to "{1}": {2} tossups, {3} bonuses.'
                               .format(len(summary['packets']), qset.name,
                                       summary['tossups'], summary['bonuses']))
                else:
                    summary = import_packets_from_files(
                        files, form.cleaned_data['set_name'], user)
                    qset = summary['question_set']
                    message = ('Imported "{0}" from {1} packet(s): {2} tossups, {3} bonuses.'
                               .format(qset.name, len(summary['packets']),
                                       summary['tossups'], summary['bonuses']))
                if summary['errors']:
                    message += ' {0} question(s) could not be parsed (see below).'.format(len(summary['errors']))
                message_class = 'alert-box success'
            except PacketImportError as ex:
                message = str(ex)
                message_class = 'alert-box alert'
            except Exception as ex:
                message = 'Import failed: {0}'.format(ex)
                message_class = 'alert-box alert'
        else:
            errs = form.errors.get('__all__')
            message = errs[0] if errs else 'Please choose files and a destination (new name or existing set).'
            message_class = 'alert-box alert'
    else:
        form = ImportPacketsForm()

    return render(request, 'import_packets.html',
                  {'form': form, 'user': user, 'summary': summary,
                   'message': message, 'message_class': message_class})

def packet(request):
    if request.user.is_authenticated:
        player = request.user.get_profile()
        packets = player.packet_set.filter(date_submitted=None)

        print('packets: ', packets)

        return render(request, 'packetview.html',
                                  {'packet_list': packets})

    else:
        return HttpResponseRedirect('/accounts/login/')

@login_required
def create_question_set (request):
    user = request.user.writer

    if not _account_can_create(request.user):
        return render(request, 'failure.html',
                      {'message': _ACCOUNT_TOO_NEW_MSG, 'message_class': 'alert-box alert'})

    if request.method == 'POST':
        form = QuestionSetForm(data=request.POST)
        if form.is_valid():
            # for the moment, just use the default ACF Distribution
            #dist = Distribution.objects.get(id=1)
            question_set = form.save(commit=False)
            question_set.owner = user
            question_set.editors = []
            question_set.editors.append(user)
            #question_set.distribution = dist
            question_set.save()
            form.save_m2m()
            user.question_set_editor.add(question_set)
            user.save()

            dist = question_set.distribution
            dist_entries = dist.distributionentry_set.all()
            for entry in dist_entries:
                set_wide_entry = SetWideDistributionEntry()
                set_wide_entry.num_bonuses = question_set.num_packets * entry.min_tossups
                set_wide_entry.num_tossups = question_set.num_packets * entry.min_bonuses
                set_wide_entry.question_set = question_set
                set_wide_entry.dist_entry = entry
                set_wide_entry.save()

                tiebreak_entry = TieBreakDistributionEntry()
                tiebreak_entry.num_bonuses = 1
                tiebreak_entry.num_tossups = 1
                tiebreak_entry.question_set = question_set
                tiebreak_entry.dist_entry = entry
                tiebreak_entry.save()

            set_distro_formset = create_set_distro_formset(question_set)
            tiebreak_formset = create_tiebreak_formset(question_set)
            comment_tab_list = []

            return render(request, 'edit_question_set.html',
                                      {'message': 'Your question set has been successfully created!',
                                       'message_class': 'alert-box success',
                                       'qset': question_set,
                                       'user': user,
                                       'form': form,
                                       'set_distro_formset': set_distro_formset,
                                       'tiebreak_formset': tiebreak_formset,
                                       'editors': [ed for ed in question_set.editor.all() if ed != question_set.owner],
                                       'writers': question_set.writer.all(),
                                       'tossups': Tossup.objects.filter(question_set=question_set),
                                       'bonuses': Bonus.objects.filter(question_set=question_set),
                                       'comment_tab_list': comment_tab_list,
                                       'packets': sorted_packets(question_set),})
        else:
            print(form.errors)
            distributions = Distribution.objects.all()
            return render(request, 'create_question_set.html',
                                      {'message': 'There was an error in creating your question set!',
                                       'message_class': 'alert-box warning',
                                       'form': form,
                                       'distributions': distributions,
                                       'user': user})
    else:
        form = QuestionSetForm()
        distributions = Distribution.objects.all()

    return render(request, 'create_question_set.html',
                              {'form': form,
                               'distributions': distributions,
                               'user': user})

@login_required
def edit_question_set(request, qset_id):
    read_only = False
    message = ''
    tossups = []
    bonuses = []
    
    qset = QuestionSet.objects.get(id=qset_id)
    qset_editors = qset.editor.all()
    qset_writers = qset.writer.all()
    user = request.user.writer
    # Membership provenance: which members come from a role group (and which
    # groups), so the writers/editors list can show it. group_granted_ids are
    # those present ONLY via a group (no direct assignment) — they can't be
    # removed here; you remove them from the group instead.
    member_groups = {}
    for _a in qset.role_group_assignments.select_related('role_group'):
        for _w_id in _a.role_group.members.values_list('id', flat=True):
            member_groups.setdefault(_w_id, []).append(_a.role_group.name)
    group_granted_ids = set(GroupRoleGrant.objects.filter(question_set=qset)
                            .values_list('writer_id', flat=True))
    set_status = {}
    set_distro_formset = None
    tiebreak_formset = None
    writer_stats = {}

    total_tu_req = 0
    total_bs_req = 0
    total_tu_written = 0
    total_bs_written = 0
    comment_tab_list = []
    tu_needed = 0
    bs_needed = 0
    set_pct_complete = 0

    role = get_role_no_owner(user, qset)

    if not qset.is_owner(user) and user not in qset_editors and user not in qset_writers:
        messages.error(request, 'You are not authorized to view information about this tournament!')
        return HttpResponseRedirect('/failure.html/')

    new_activity = _new_activity_count(user, qset)

    if request.method == 'POST':
        if (qset.is_owner(user) or user in qset_editors):
            form = QuestionSetForm(data=request.POST)
            if form.is_valid():
                qset = QuestionSet.objects.get(id=qset_id)
                qset.name = form.cleaned_data['name']
                qset.date = form.cleaned_data['date']
                qset.distribution = form.cleaned_data['distribution']
                qset.num_packets = form.cleaned_data['num_packets']
                qset.char_count_ignores_pronunciation_guides = form.cleaned_data['char_count_ignores_pronunciation_guides']
                qset.tossups_only = form.cleaned_data['tossups_only']
                qset.public = form.cleaned_data['public']
                qset.max_acf_tossup_length = form.cleaned_data['max_acf_tossup_length']
                qset.max_acf_bonus_length = form.cleaned_data['max_acf_bonus_length']
                qset.save()
                cache.clear()

                tossups, tossup_dict, bonuses, bonus_dict = get_tossup_and_bonuses_in_set(qset, question_limit=30, preview_only=True)

                if qset.is_owner(user):
                    read_only = False
                else:
                    read_only = True

                set_status, total_tu_req, total_bs_req, tu_needed, bs_needed, set_pct_complete = get_questions_remaining(qset)
                writer_stats = get_writer_questions_remaining(qset, total_tu_req, total_bs_req)
                                                                
                comment_tab_list = get_comment_tab_list(tossup_dict, bonus_dict)

                return render(request, 'edit_question_set.html',
                                          {'form': form,
                                           'qset': qset,
                                           'user': user,
                                           'editors': [ed for ed in qset_editors if ed != qset.owner],
                                           'writers': qset.writer.all(),
                                           'writer_stats': writer_stats,
                                           'upload_form': QuestionUploadForm(),
                                           'set_status': set_status,
                                           'set_pct_complete': '{0:0.2f}%'.format(set_pct_complete),
                                           'set_pct_progress_bar': '{0:0.0f}%'.format(set_pct_complete),
                                           'tu_needed': tu_needed,
                                           'bs_needed': bs_needed,
                                           'tossups': tossups,
                                           'bonuses': bonuses,
                                           'packets': sorted_packets(qset),
                                           'comment_list': comment_tab_list,
                                           'role': role,
                                           'new_activity': new_activity,
                                           'member_groups': member_groups,
                                           'group_granted_ids': group_granted_ids,
                                           'message': 'Your changes have been successfully saved.',
                                           'message_class': 'alert-success'})
            else:
                # Form invalid: still populate question data so the page isn't
                # blank (and so an empty render isn't cached).
                tossups, tossup_dict, bonuses, bonus_dict = get_tossup_and_bonuses_in_set(qset, question_limit=30, preview_only=True)
                set_status, total_tu_req, total_bs_req, tu_needed, bs_needed, set_pct_complete = get_questions_remaining(qset)
                writer_stats = get_writer_questions_remaining(qset, total_tu_req, total_bs_req)
                comment_tab_list = get_comment_tab_list(tossup_dict, bonus_dict)
                read_only = not (qset.is_owner(user) or user in qset_editors)
        else:
            return render(request, 'failure.html', {'message': 'You are not authorized to change this set!', 'message_class': 'alert-box alert'})
    else:
        print("Begin edit_question_set get", time.strftime("%H:%M:%S"))
        if user not in qset_editors and not qset.is_owner(user) and user not in qset.writer.all():
            # Just redirect to main in this case of no permissions
            # TODO: a better story
            return HttpResponseRedirect('/main.html')

        tossups, tossup_dict, bonuses, bonus_dict = get_tossup_and_bonuses_in_set(qset, question_limit=30, preview_only=True)
        
        if user not in qset_editors and not qset.is_owner(user):
            form = QuestionSetForm(instance=qset, read_only=True)
            read_only = True
            message = ''
        else:
            if qset.is_owner(user):
                read_only = False
            elif user in qset.writer.all() or user in qset.editor.all():
                read_only = True
            form = QuestionSetForm(instance=qset)

        set_status, total_tu_req, total_bs_req, tu_needed, bs_needed, set_pct_complete = get_questions_remaining(qset)
        writer_stats = get_writer_questions_remaining(qset, total_tu_req, total_bs_req)
                                                                
        comment_tab_list = get_comment_tab_list(tossup_dict, bonus_dict)                    

    print("End edit_question_set get", time.strftime("%H:%M:%S"))
        
    return render(request, 'edit_question_set.html',
                              {'form': form,
                               'user': user,
                               'editors': [ed for ed in qset_editors if ed != qset.owner],
                               'writers': [wr for wr in qset_writers if wr != qset.owner],
                               'writer_stats': writer_stats,
                               'set_status': set_status,
                               'set_pct_complete': '{0:0.2f}%'.format(set_pct_complete),
                               'set_pct_progress_bar': '{0:0.0f}%'.format(set_pct_complete),
                               'tu_needed': tu_needed,
                               'bs_needed': bs_needed,
                               'upload_form': QuestionUploadForm(),
                               'tossups': tossups,
                               'bonuses': bonuses,
                               'packets': sorted_packets(qset),
                               'comment_tab_list': comment_tab_list,
                               'qset': qset,
                               'role': role,
                               'new_activity': new_activity,
                               'read_only': read_only,
                               'all_role_groups': RoleGroup.objects.all().order_by('name'),
                               'attached_role_groups': list(
                                   qset.role_group_assignments.select_related('role_group')),
                               'member_groups': member_groups,
                               'group_granted_ids': group_granted_ids,
                               'message': message})

@login_required
def categories(request, qset_id, category_id):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    qset_editors = qset.editor.all()
    qset_writers = qset.writer.all()

    category_object = DistributionEntry.objects.get(id=category_id)

    entry = qset.setwidedistributionentry_set.get(dist_entry=category_object)
    tu_required = entry.num_tossups
    bs_required = entry.num_bonuses
    tu_written = qset.tossup_set.filter(category=entry.dist_entry).count()
    bs_written = qset.bonus_set.filter(category=entry.dist_entry).count()

    category_status =   {'tu_req': tu_required,
                         'tu_in_cat': tu_written,
                         'bs_req': bs_required,
                         'bs_in_cat': bs_written
                         }

    message = category_object.category
    tossups = []
    bonuses = []
    if user not in qset_editors and not qset.is_owner(user) and user not in qset.writer.all():
        message = 'You are not authorized to view this set'
    else:
        tossups = list(Tossup.objects.filter(question_set=qset).filter(category=category_id)
                       .select_related(*QUESTION_LIST_RELATED))
        bonuses = list(Bonus.objects.filter(question_set=qset).filter(category=category_id)
                       .select_related(*QUESTION_LIST_RELATED))
        attach_question_comments({t.id: t for t in tossups}, {b.id: b for b in bonuses})

    return render(request, 'categories.html',
        {
        'user': user,
        'tossups': tossups,
        'bonuses': bonuses,
        'category_status': category_status,
        'qset': qset,
        'message': message,
        'category': category_object})

@login_required
def top_category(request, qset_id, category_name):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)

    if user not in qset.editor.all() and not qset.is_owner(user) and user not in qset.writer.all():
        return render(request, 'failure.html', {'message': 'You are not authorized to view this set'})

    role = get_role_no_owner(user, qset)

    entries = DistributionEntry.objects.filter(
        setwidedistributionentry__question_set=qset,
        category=category_name
    ).distinct().order_by('subcategory')

    sub_categories = []
    for entry in entries:
        swide = qset.setwidedistributionentry_set.filter(dist_entry=entry).first()
        if swide:
            tu_req = swide.num_tossups
            bs_req = swide.num_bonuses
            tu_written = qset.tossup_set.filter(category=entry).count()
            bs_written = qset.bonus_set.filter(category=entry).count()
            sub_categories.append({
                'entry': entry,
                'tu_req': tu_req,
                'tu_written': tu_written,
                'bs_req': bs_req,
                'bs_written': bs_written,
            })

    return render(request, 'top_category.html', {
        'user': user,
        'qset': qset,
        'category_name': category_name,
        'sub_categories': sub_categories,
        'role': role,
    })

@login_required
def category_document(request, qset_id, category_id=None, category_name=None):
    """Read-only document view of every question in a category or subcategory
    (e.g. all of History - American, or all of History)."""
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    if not (qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all()):
        return render(request, 'failure.html',
                      {'message': 'You are not authorized to view this set',
                       'message_class': 'alert-box alert'})

    if category_id is not None:
        cat = DistributionEntry.objects.get(id=category_id)
        title = str(cat)
        tu_qs = Tossup.objects.filter(question_set=qset, category=cat)
        bs_qs = Bonus.objects.filter(question_set=qset, category=cat)
    else:
        title = category_name
        tu_qs = Tossup.objects.filter(question_set=qset, category__category=category_name)
        bs_qs = Bonus.objects.filter(question_set=qset, category__category=category_name)

    related = ('category', 'author__user', 'packet')
    order = ('packet__sort_order', 'packet__id', 'question_number', 'id')
    tu_qs = tu_qs.select_related(*related).order_by(*order)
    bs_qs = bs_qs.select_related(*related).order_by(*order)

    def label(writer):
        if writer is None:
            return ''
        return '{0} {1}'.format(writer.user.first_name, writer.user.last_name).strip() or writer.user.username

    def item(q, qtype):
        return {
            'id': q.id, 'qtype': qtype, 'html': q.to_html(),
            'category': str(q.category) if q.category else '',
            'packet': q.packet.packet_name if q.packet else 'Unpacketized',
            'number': q.question_number or '',
            'author': label(q.author),
            'edit_url': '/edit_{0}/{1}/'.format(qtype, q.id),
        }

    tossups = [item(t, 'tossup') for t in tu_qs]
    bonuses = [item(b, 'bonus') for b in bs_qs]
    return render(request, 'category_document.html', {
        'qset': qset, 'title': title, 'tossups': tossups, 'bonuses': bonuses,
        'count': len(tossups) + len(bonuses), 'user': user,
        'category_id': category_id, 'category_name': category_name})


def _dup_fingerprint(qset):
    """A cheap signature of the set's question state; changes whenever a
    question is added, removed, or edited, so the cached report auto-refreshes."""
    import hashlib
    from django.db.models import Count, Max
    t = Tossup.objects.filter(question_set=qset).aggregate(n=Count('id'), m=Max('last_changed_date'))
    b = Bonus.objects.filter(question_set=qset).aggregate(n=Count('id'), m=Max('last_changed_date'))
    raw = '{0}-{1}-{2}-{3}'.format(t['n'], t['m'], b['n'], b['m'])
    return hashlib.md5(raw.encode('utf-8')).hexdigest()


def _dup_render_answer(raw):
    return get_formatted_question_html(get_primary_answer(raw or ''), True, True, False, False).strip()


def _post_submit_dup_matches(qset, question, qtype):
    """Find duplicate-answer matches for a just-saved question and render their
    answers, for the 'you may have created a duplicate' warning shown after a
    question is submitted."""
    matches = find_answer_matches(qset, question, qtype)
    for m in matches:
        m['answer_html'] = _dup_render_answer(m['answer_raw'])
    return matches


def _dup_answer_html(entry, bonus_map):
    """Rendered answer for an entry. For bonuses, show all three answer lines
    (each cut at the '[') so the whole bonus is identifiable."""
    if entry['type'] == 'bonus':
        bonus = bonus_map.get(entry['id'])
        if bonus is not None:
            parts = [_dup_render_answer(a) for a in (bonus.part1_answer, bonus.part2_answer, bonus.part3_answer)
                     if a and a.strip()]
            if parts:
                return ' / '.join(parts)
    return _dup_render_answer(entry.get('answer_raw'))


def _dup_preview_html(raw_text, term, width=280):
    """A formatted preview window centered on where the repeat (`term`) occurs,
    falling back to the start of the text. Keeps QEMS markup so italics/
    underlines render."""
    raw = (raw_text or '').strip()
    if not raw:
        return ''
    center = 0
    if term:
        i = raw.lower().find(term.lower())
        if i < 0:
            i = strip_markup(raw).lower().find(term.lower())
        if i > 0:
            center = i
    start = max(0, center - width // 3)
    end = min(len(raw), start + width)
    start = max(0, end - width)
    snippet = raw[start:end]
    if start > 0 and ' ' in snippet:
        snippet = snippet.split(' ', 1)[1]
    if end < len(raw) and ' ' in snippet:
        snippet = snippet.rsplit(' ', 1)[0]
    html = get_formatted_question_html(snippet, False, True, False, True)
    return ('&hellip; ' if start > 0 else '') + html + (' &hellip;' if end < len(raw) else '')


@login_required
def duplicate_check(request, qset_id):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)

    if not qset.is_owner(user) and user not in qset.editor.all() and user not in qset.writer.all():
        messages.error(request, 'You are not authorized to view this set.')
        return render(request, 'failure.html',
                      {'message': 'You are not authorized to view this set.',
                       'message_class': 'alert-box alert'})

    # Cache the (expensive) report, keyed by a fingerprint of the question
    # state so it recomputes only when something actually changed.
    cache_key = 'dupcheck:{0}:{1}'.format(qset.id, _dup_fingerprint(qset))
    context = cache.get(cache_key)

    if context is None:
        groups = find_duplicates(qset)
        internal_issues = find_internal_issues(qset)
        topic_groups = find_topic_repeats(qset)

        # Batch-load packets and bonus answers for all referenced questions
        tu_ids, bs_ids = set(), set()
        for group_list in (groups, topic_groups):
            for group in group_list:
                for entry in group['entries']:
                    (bs_ids if entry['type'] == 'bonus' else tu_ids).add(entry['id'])
        bonus_map = {b.id: b for b in Bonus.objects.filter(id__in=bs_ids).select_related('packet')}
        tu_map = {t.id: t for t in Tossup.objects.filter(id__in=tu_ids).select_related('packet')}

        def packet_of(entry):
            obj = tu_map.get(entry['id']) if entry['type'] == 'tossup' else bonus_map.get(entry['id'])
            return obj.packet.packet_name if (obj is not None and obj.packet) else ''

        for group in groups:
            for entry in group['entries']:
                entry['packet'] = packet_of(entry)
                entry['answer_html'] = _dup_answer_html(entry, bonus_map)
                entry['preview_html'] = _dup_preview_html(entry.get('text'), None)
            for pair in group['pairs']:
                pair['similarity_pct'] = int(pair['similarity'] * 100)

        # Topic repeats are already severity-sorted (critical first). A large
        # set can have hundreds; rendering them all is the slow part, so only
        # enrich+render the top N and report how many more exist.
        TOPIC_RENDER_CAP = 150
        topic_total = len(topic_groups)
        topic_render = topic_groups[:TOPIC_RENDER_CAP]
        for group in topic_render:
            term = group.get('label')
            for entry in group['entries']:
                entry['packet'] = packet_of(entry)
                entry['answer_html'] = _dup_answer_html(entry, bonus_map)
                entry['preview_html'] = _dup_preview_html(entry.get('text'), term)

        context = {
            'groups': groups,
            'critical_count': sum(1 for g in groups if g['severity'] == CRITICAL),
            'warning_count': sum(1 for g in groups if g['severity'] == WARNING),
            'info_count': sum(1 for g in groups if g['severity'] == INFO),
            'total_groups': len(groups),
            'total_questions': sum(len(g['entries']) for g in groups),
            'internal_issues': internal_issues,
            'bonus_repeat_count': sum(1 for i in internal_issues if i['issue_type'] == 'bonus_repeat_answer'),
            'clue_reuse_count': sum(1 for i in internal_issues if i['issue_type'] == 'tossup_clue_reuse'),
            'topic_groups': topic_render,
            'topic_total': topic_total,
            'topic_shown': len(topic_render),
            'topic_truncated': topic_total > len(topic_render),
            'topic_critical': sum(1 for g in topic_groups if g['severity'] == CRITICAL),
            'topic_warning': sum(1 for g in topic_groups if g['severity'] == WARNING),
            'topic_info': sum(1 for g in topic_groups if g['severity'] == INFO),
            'has_packets': qset.packet_set.exists(),
        }
        cache.set(cache_key, context, 1800)

    context = dict(context)
    context['qset'] = qset
    context['user'] = user
    return render(request, 'duplicate_check.html', context)

@login_required
def view_all_questions(request, qset_id):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    qset_editors = qset.editor.all()
    qset_writers = qset.writer.all()

    message = ''
    tossups = []
    bonuses = []
    if user not in qset_editors and not qset.is_owner(user) and user not in qset.writer.all():
        message = 'You are not authorized to view this set'
        return render(request, 'failure.html',
                                 {'message': message,
                                  'message_class': 'alert-box alert'})        
    else:
        tossups, tossup_dict, bonuses, bonus_dict = get_tossup_and_bonuses_in_set(qset, question_limit=10000, preview_only=True)
            
    return render(request, 'view_all_questions.html',
        {
        'user': user,
        'tossups': tossups,
        'bonuses': bonuses,
        'qset': qset,
        'message': message})	

@login_required
def view_all_comments(request, qset_id):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    qset_editors = qset.editor.all()
    qset_writers = qset.writer.all()

    message = ''
    tossups = []
    bonuses = []
    if user not in qset_editors and not qset.is_owner(user) and user not in qset.writer.all():
        message = 'You are not authorized to view this set'
        return render(request, 'failure.html',
                                 {'message': message,
                                  'message_class': 'alert-box alert'})        
    else:
        tossups, tossup_dict, bonuses, bonus_dict = get_tossup_and_bonuses_in_set(qset, question_limit=10000, preview_only=True)
        comment_tab_list = get_comment_tab_list(tossup_dict, bonus_dict, comment_limit=10000)
            
    return render(request, 'view_all_comments.html',
        {
        'user': user,
        'comment_tab_list': comment_tab_list,
        'qset': qset,
        'message': message})	

@login_required
def question_set_distribution(request, qset_id):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    qset_editors = qset.editor.all()
    qset_writers = qset.writer.all()
    set_distro_formset = []
    tiebreak_formset = []
    read_only = True

    message = ''
    if user not in qset_editors and not qset.is_owner(user) and user not in qset.writer.all():
        message = 'You are not authorized to view this set'
        return render(request, 'failure.html',
                                 {'message': message,
                                  'message_class': 'alert-box alert'})                
    elif qset.is_owner(user):
        set_distro_formset = create_set_distro_formset(qset)
        tiebreak_formset = create_tiebreak_formset(qset)    
        read_only = False
    else:
        set_distro_formset = create_set_distro_formset(qset)
        tiebreak_formset = create_tiebreak_formset(qset)        
            
    return render(request, 'question_set_distribution.html',
        {
        'user': user,
        'set_distro_formset': set_distro_formset,
        'tiebreak_formset': tiebreak_formset,
        'qset': qset,
        'message': message,
        'read_only': read_only})	

@login_required
def edit_set_distribution(request, qset_id):

    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)

    if request.method == 'POST':

        DistributionEntryFormset = formset_factory(SetWideDistributionEntryForm, can_delete=False, extra=0)
        formset = DistributionEntryFormset(data=request.POST, prefix='distentry')

        if formset.is_valid() and qset.is_owner(user):
            for dist_form in formset.forms:
                entry_id = int(dist_form.cleaned_data['entry_id'])
                num_tossups = int(dist_form.cleaned_data['num_tossups'])
                num_bonuses = int(dist_form.cleaned_data['num_bonuses'])

                entry = SetWideDistributionEntry.objects.get(id=entry_id)
                entry.num_tossups = num_tossups
                entry.num_bonuses = num_bonuses
                entry.save()

            return HttpResponseRedirect('/question_set_distribution/{0}'.format(qset_id))
        else:
            return render(request, 'failure.html',
                                     {'message': 'Something went wrong. We\'re working on it.',
                                      'message_class': 'alert-box alert'})
    elif request.method == 'GET':
        if qset.is_owner(user):
            return render(request, 'view_all_questions.html',
                {
                'user': user,
                'tossups': tossups,
                'bonuses': bonuses,
                'qset': qset,
                'message': message})	
            

@login_required
def edit_set_tiebreak(request, qset_id):

    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)

    if request.method == 'POST':

        TiebreakDistributionEntryFormset = formset_factory(TieBreakDistributionEntryForm, can_delete=False, extra=0)
        formset = TiebreakDistributionEntryFormset(data=request.POST, prefix='tiebreak')

        if formset.is_valid() and qset.is_owner(user):
            for dist_form in formset.forms:
                entry_id = int(dist_form.cleaned_data['entry_id'])
                num_tossups = int(dist_form.cleaned_data['num_tossups'])
                num_bonuses = int(dist_form.cleaned_data['num_bonuses'])

                entry = TieBreakDistributionEntry.objects.get(id=entry_id)
                entry.num_tossups = num_tossups
                entry.num_bonuses = num_bonuses
                entry.save()

            return HttpResponseRedirect('/question_set_distribution/{0}'.format(qset_id))
        else:
            return render(request, 'failure.html',
                                     {'message': 'Something went wrong. We\'re working on it.',
                                      'message_class': 'alert-box alert'})

@login_required
def find_editor(request):
    pass

@login_required
def add_editor(request, qset_id):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    message = ''

    if request.method == 'GET':
        if qset.is_owner(user):
            current_editors = qset.editor.all()

            available_editors = [writer for writer in Writer.objects.all().order_by('user__last_name', 'user__first_name', 'user__username') #exclude(is_active=False)
                                 if writer not in current_editors and
                                    not qset.is_owner(writer) and writer.id != 1
                                    and writer.user.is_active]
        else:
            available_editors = []
            return render(request, 'failure.html',
                                     {'message': 'You are not authorized to make changes to this tournament!',
                                      'message_class': 'alert-box alert'})

        return render(request, 'add_editor.html',
                                 {'qset': qset,
                                  'available_editors': available_editors,
                                  'message': message,
                                  'user': user})


    elif request.method == 'POST':
        if qset.is_owner(user):
            editors_to_add = request.POST.getlist('editors_to_add')
            # do some basic validation here
            if all([x.isdigit() for x in editors_to_add]):
                for editor_id in editors_to_add:
                    editor = Writer.objects.get(id=editor_id)
                    qset.editor.add(editor)
                    _notify_added_to_set(editor, qset, 'editor', user)
                    # A direct add takes ownership of the membership (so it
                    # survives role-group changes).
                    GroupRoleGrant.objects.filter(
                        question_set=qset, writer=editor, role='editor').delete()

                    # Don't have someone be both a writer and editor--delete them
                    try:
                        writer = qset.writer.get(id=editor_id)
                        if (writer is not None):
                            qset.writer.remove(writer)
                    except:
                        print("No writer to delete") # TODO: Come up with a better way of handling this

                qset.save()
                cache.clear()
                set_editors = qset.editor.all()
                available_editors = [writer for writer in Writer.objects.all().order_by('user__last_name', 'user__first_name', 'user__username') #exclude(is_active=False)
                                     if writer not in set_editors and
                                        not qset.is_owner(writer) and writer.id != 1
                                        and writer.user.is_active]
            else:
                message = 'Invalid data entered!'
                available_editors = []
        else:
            available_editors = []
            return render(request, 'failure.html',
                                     {'message': 'You are not authorized to make changes to this tournament!',
                                      'message_class': 'alert-box alert'})

        return render(request, 'add_editor.html',
                                 {'qset': qset,
                                  'available_editors': available_editors,
                                  'message': message,
                                  'user': user})

@login_required
def add_co_owner(request, qset_id):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    message = ''

    def get_available():
        current_owners = qset.all_owners()
        return [writer for writer in Writer.objects.all().order_by('user__last_name', 'user__first_name', 'user__username')
                if writer not in current_owners and writer.id != 1 and writer.user.is_active]

    if not qset.is_owner(user):
        return render(request, 'failure.html',
                                 {'message': 'You are not authorized to make changes to this tournament!',
                                  'message_class': 'alert-box alert'})

    if request.method == 'POST':
        co_owners_to_add = request.POST.getlist('co_owners_to_add')
        if all([x.isdigit() for x in co_owners_to_add]):
            for co_owner_id in co_owners_to_add:
                co_owner = Writer.objects.get(id=co_owner_id)
                qset.co_owners.add(co_owner)
                # Co-owners get full editor privileges as well
                qset.editor.add(co_owner)
                _notify_added_to_set(co_owner, qset, 'co-owner', user)

                # Don't have someone be both a writer and a co-owner--remove them as writer
                try:
                    writer = qset.writer.get(id=co_owner_id)
                    if writer is not None:
                        qset.writer.remove(writer)
                except:
                    pass

            qset.save()
            cache.clear()
            message = 'Co-owner(s) added'
        else:
            message = 'Invalid data entered!'

    return render(request, 'add_co_owner.html',
                             {'qset': qset,
                              'available_co_owners': get_available(),
                              'message': message,
                              'user': user})

@login_required
def delete_co_owner(request):
    user = request.user.writer
    message = ''
    message_class = ''

    if request.method == 'POST':
        qset_id = request.POST['qset_id']
        qset = QuestionSet.objects.get(id=qset_id)
        co_owner_id = request.POST['co_owner_id']
        if qset.is_owner(user):
            co_owner = qset.co_owners.get(id=co_owner_id)
            qset.co_owners.remove(co_owner)
            # Also drop the editor privilege that came with co-ownership
            qset.editor.remove(co_owner)
            cache.clear()
            message = 'Co-owner removed'
            message_class = 'alert-box success'
        else:
            message = 'You are not authorized to remove co-owners from this set!'
            message_class = 'alert-box warning'

    return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))

@login_required
def add_writer(request, qset_id):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    message = ''
    message_class = ''

    if request.method == 'GET':
        if qset.is_owner(user):
            set_writers = Writer.objects.filter(Q(question_set_writer=qset) | Q(question_set_editor=qset)).distinct().order_by('user__last_name', 'user__first_name', 'user__username')
            available_writers = [writer for writer in Writer.objects.all().order_by('user__last_name', 'user__first_name', 'user__username') #exclude(is_active=False)
                                 if writer not in set_writers and
                                    not qset.is_owner(writer) and writer.id != 1
                                    and writer.user.is_active]
        else:
            available_writers = []
            return render(request, 'failure.html',
                                     {'message': 'You are not authorized to make changes to this tournament!',
                                      'message_class': 'alert-box alert'})

        return render(request, 'add_writer.html',
                                 {'qset': qset,
                                  'available_writers': available_writers,
                                  'message': message,
                                  'user': user})


    elif request.method == 'POST':
        if qset.is_owner(user):
            writers_to_add = request.POST.getlist('writers_to_add')
            # do some basic validation here
            if all([x.isdigit() for x in writers_to_add]):
                for writer_id in writers_to_add:
                    writer = Writer.objects.get(id=writer_id)
                    qset.writer.add(writer)
                    _notify_added_to_set(writer, qset, 'writer', user)
                    GroupRoleGrant.objects.filter(
                        question_set=qset, writer=writer, role='writer').delete()
                qset.save()
                cache.clear()
                set_writers = Writer.objects.filter(Q(question_set_writer=qset) | Q(question_set_editor=qset)).distinct().order_by('user__last_name', 'user__first_name', 'user__username')
                available_writers = [writer for writer in Writer.objects.all().order_by('user__last_name', 'user__first_name', 'user__username') #exclude(is_active=False)
                                     if writer not in set_writers and
                                        not qset.is_owner(writer) and writer.id != 1
                                        and writer.user.is_active]
            else:
                message = 'Invalid data entered!'
                available_writers = []
        else:
            available_writers = []
            return render(request, 'failure.html',
                                     {'message': 'You are not authorized to make changes to this tournament!',
                                      'message_class': 'alert-box alert'})

        return render(request, 'add_writer.html',
                                 {'qset': qset,
                                  'available_writers': available_writers,
                                  'message': message,
                                  'message_class': message_class,
                                  'user': user})


@login_required
def role_groups(request):
    """Create and manage role groups (named groups of writers). Members added
    here propagate to every set the group is attached to."""
    user = request.user.writer
    message = message_class = ''

    if request.method == 'POST':
        action = request.POST.get('action', '')
        if action == 'create':
            if not _account_can_create(request.user):
                message, message_class = ("New accounts can't create role groups until 2 days "
                                          "after sign-up. Please try again later.", 'alert-box warning')
            else:
                name = (request.POST.get('name') or '').strip()
                if not name:
                    message, message_class = 'Enter a group name.', 'alert-box warning'
                elif RoleGroup.objects.filter(name__iexact=name).exists():
                    message, message_class = 'A group with that name already exists.', 'alert-box warning'
                else:
                    RoleGroup.objects.create(name=name, created_by=user)
                    message, message_class = 'Group "{0}" created.'.format(name), 'alert-box success'
        elif action == 'request_join':
            try:
                group = RoleGroup.objects.get(id=int(request.POST.get('group_id', 0)))
            except (ValueError, RoleGroup.DoesNotExist):
                group = None
            if group is None:
                message, message_class = 'Group not found.', 'alert-box warning'
            elif group.can_manage(user) or group.members.filter(id=user.id).exists():
                message, message_class = 'You are already part of this group.', 'alert-box warning'
            elif _notify_group_join_request(group, user):
                message, message_class = ('Your request to join "{0}" was emailed to the group '
                                          'owner.'.format(group.name), 'alert-box success')
            else:
                message, message_class = ('The group owner has no email on file, so the request '
                                          'could not be sent.', 'alert-box warning')
        else:
            try:
                group = RoleGroup.objects.get(id=int(request.POST.get('group_id', 0)))
            except (ValueError, RoleGroup.DoesNotExist):
                group = None
            if group is None or not group.can_manage(user):
                message, message_class = 'You can only manage groups you created.', 'alert-box alert'
            elif action == 'delete':
                sets = [a.question_set for a in group.set_assignments.all()]
                group.delete()
                for qs in sets:
                    reconcile_group_roles(qs)
                message, message_class = 'Group deleted.', 'alert-box success'
            elif action == 'add_member':
                uname = (request.POST.get('username') or '').strip()
                try:
                    w = Writer.objects.get(user__username__iexact=uname)
                    if group.members.filter(id=w.id).exists():
                        message, message_class = '{0} is already a member.'.format(uname), 'alert-box warning'
                    else:
                        group.members.add(w)
                        reconcile_group(group)
                        _notify_added_to_group(w, group, user)
                        message, message_class = 'Added {0}.'.format(uname), 'alert-box success'
                except Writer.DoesNotExist:
                    message, message_class = 'No user named "{0}".'.format(uname), 'alert-box warning'
            elif action == 'remove_member':
                try:
                    w = Writer.objects.get(id=int(request.POST.get('writer_id', 0)))
                    group.members.remove(w)
                    reconcile_group(group)
                    message, message_class = 'Member removed.', 'alert-box success'
                except (ValueError, Writer.DoesNotExist):
                    pass

    groups = []
    for g in RoleGroup.objects.all().order_by('name').prefetch_related('members__user', 'set_assignments__question_set'):
        member_list = list(g.members.all().order_by('user__username'))
        is_member = any(m.id == user.id for m in member_list)
        can_manage = g.can_manage(user)
        # Membership is private: only members (and managers) see who's in a group.
        groups.append({'group': g,
                       'members': member_list if (is_member or can_manage) else None,
                       'member_count': len(member_list),
                       'assignments': list(g.set_assignments.all()),
                       'can_manage': can_manage,
                       'is_member': is_member})
    return render(request, 'role_groups.html',
                  {'groups': groups, 'message': message, 'message_class': message_class, 'user': user})


@login_required
def attach_role_group(request, qset_id):
    """Attach a role group to a set with a role (owner only); members gain the role."""
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    if request.method == 'POST' and qset.is_owner(user):
        role = request.POST.get('role', 'writer')
        role = role if role in ('editor', 'writer') else 'writer'
        try:
            group = RoleGroup.objects.get(id=int(request.POST.get('role_group_id', 0)))
            SetRoleGroupAssignment.objects.update_or_create(
                question_set=qset, role_group=group, defaults={'role': role})
            reconcile_group_roles(qset)
            cache.clear()
        except (ValueError, RoleGroup.DoesNotExist):
            pass
    return HttpResponseRedirect('/edit_question_set/{0}/#editors'.format(qset_id))


@login_required
def detach_role_group(request, qset_id):
    """Remove a role group from a set (owner only); group-granted members lose the role."""
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    if request.method == 'POST' and qset.is_owner(user):
        try:
            SetRoleGroupAssignment.objects.filter(
                question_set=qset, role_group_id=int(request.POST.get('role_group_id', 0))).delete()
            reconcile_group_roles(qset)
            cache.clear()
        except ValueError:
            pass
    return HttpResponseRedirect('/edit_question_set/{0}/#editors'.format(qset_id))


@login_required
def edit_packet(request, packet_id):
    user = request.user.writer
    packet = Packet.objects.get(id=packet_id)
    qset = packet.question_set
    message = ''
    message_class = ''
    read_only = True
    tossup_status = []
    bonus_status = []
    can_rename = qset.is_owner(user) or user in qset.editor.all()

    if request.method == 'POST' and 'packet_name' in request.POST:
        if can_rename:
            new_name = (request.POST.get('packet_name') or '').strip()[:200]
            if not new_name:
                message, message_class = 'Packet name cannot be empty.', 'alert-box warning'
            elif Packet.objects.filter(question_set=qset, packet_name=new_name).exclude(id=packet.id).exists():
                message = 'A packet named "{0}" already exists in this set.'.format(new_name)
                message_class = 'alert-box warning'
            elif new_name != packet.packet_name:
                packet.packet_name = new_name
                packet.save(update_fields=['packet_name'])
                cache.clear()
                message, message_class = 'Packet renamed to "{0}".'.format(new_name), 'alert-box success'
        else:
            message = 'Only an owner or editor can rename packets.'
            message_class = 'alert-box alert'

    if request.method in ('GET', 'POST'):
        if qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all():
            tossups = packet.tossup_set.order_by('question_number').all()
            bonuses = packet.bonus_set.order_by('question_number').all()
            if user not in qset.writer.all():
                read_only = False

            # Per-packet requirements, shown per top-level category AND broken
            # out by subcategory so it's clear which subcategories are needed.
            # Use the packetization quota for a path when defined, otherwise the
            # set-wide total for that path divided by the packet count.
            num_packets = max(qset.num_packets, 1)
            quota_by_path = {e.path: e for e in PacketizationEntry.objects.filter(question_set=qset)}

            def _req(path, total, attr):
                quota = quota_by_path.get(path)
                if quota is not None and getattr(quota, attr) is not None:
                    return float(getattr(quota, attr))
                return round(total / float(num_packets), 1)

            entries = list(qset.setwidedistributionentry_set.select_related('dist_entry')
                           .order_by('dist_entry__category', 'dist_entry__subcategory'))
            by_top = {}
            for swde in entries:
                by_top.setdefault(swde.dist_entry.category, []).append(swde)

            tossup_status = []
            bonus_status = []
            for top, swdes in by_top.items():
                tu_total = sum(s.num_tossups or 0 for s in swdes)
                bs_total = sum(s.num_bonuses or 0 for s in swdes)
                tu_in_top = Tossup.objects.filter(packet=packet, category__category=top).count()
                bs_in_top = Bonus.objects.filter(packet=packet, category__category=top).count()
                tossup_status.append({'label': top, 'is_sub': False,
                                      'tu_req': _req(top, tu_total, 'min_tossups'), 'tu_in_cat': tu_in_top})
                bonus_status.append({'label': top, 'is_sub': False,
                                     'bs_req': _req(top, bs_total, 'min_bonuses'), 'bs_in_cat': bs_in_top})
                # Subcategory detail rows (only when the category has subcategories).
                for swde in swdes:
                    de = swde.dist_entry
                    if not de.subcategory:
                        continue
                    path = '{0} - {1}'.format(de.category, de.subcategory)
                    tossup_status.append({
                        'label': de.subcategory, 'is_sub': True,
                        'tu_req': _req(path, swde.num_tossups or 0, 'min_tossups'),
                        'tu_in_cat': Tossup.objects.filter(packet=packet, category=de).count()})
                    bonus_status.append({
                        'label': de.subcategory, 'is_sub': True,
                        'bs_req': _req(path, swde.num_bonuses or 0, 'min_bonuses'),
                        'bs_in_cat': Bonus.objects.filter(packet=packet, category=de).count()})


        else:
            message = 'You are not authorized to view or edit this packet!'
            message_class = 'alert-box alert'
            tossups = None
            bonuses = None

    return render(request, 'edit_packet.html',
        {'qset': qset,
         'packet': packet,
         'message': message,
         'message_class': message_class,
         'tossups': tossups,
         'bonuses': bonuses,
         'tossup_status': tossup_status,
         'bonus_status': bonus_status,
         'role': get_role_no_owner(user, qset),
         'read_only': read_only,
         'can_rename': can_rename,
         'user': user})

@login_required
def add_tossups(request, qset_id, packet_id=None):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    message = ''
    message_class = ''
    tossup = None
    read_only = True
    question_type_id = []
    tossup_form = []

    if (QuestionType.objects.exists()):
        question_type_id = QuestionType.objects.get(question_type=ACF_STYLE_TOSSUP)

    if request.method == 'GET':
        if user in qset.editor.all() or user in qset.writer.all() or qset.is_owner(user):
            if user in qset.writer.all() and user not in qset.editor.all() and not qset.is_owner(user):
                tossup_form = TossupForm(qset_id=qset.id, packet_id=packet_id, role='writer', writer=user.user.username, initial={'question_type': question_type_id})
            else:
                tossup_form = TossupForm(qset_id=qset.id, packet_id=packet_id, writer=user.user.username, initial={'question_type': question_type_id})
            read_only = False
        else:
            tossup_form = []
            message = 'You are not authorized to add questions to this tournament!'
            message_class = 'alert-box warning'
            read_only = True

        return render(request, 'add_tossups.html',
            {'form': tossup_form,
             'message': message,
             'message_class': message_class,
             'read_only': read_only,
             'user': user,
             'qset': qset})

    elif request.method == 'POST':
        if user in qset.editor.all() or user in qset.writer.all() or qset.is_owner(user):
            read_only = False

            # The user may have set the packet ID through the POST body, so check for it there
            if packet_id == None and 'packet' in request.POST and request.POST['packet'] != '':
                packet_id = int(request.POST['packet'])
            tossup_form = TossupForm(request.POST, qset_id=qset.id, packet_id=packet_id, writer=user.user.username)

            if tossup_form.is_valid():
                tossup = tossup_form.save(commit=False)
                if (tossup.author is None):
                    tossup.author = user
                tossup.question_set = qset
                tossup.tossup_text = strip_markup(tossup.tossup_text)
                tossup.tossup_answer = strip_markup(tossup.tossup_answer)
                tossup.locked = False

                try:
                    tossup.is_valid()

                    if packet_id is None or packet_id == '':
                        # If the tossup doesn't have a packet, set its number to be the magic number
                        # of 999, meaning that it's unassigned.  Can't assign -1 because this is outside
                        # of the legal range of tossup numbers and it ends up getting set to 1 for some
                        # reason, except in the case where there are no packets in the system in which
                        # case there's an error adding the question
                        tossup.question_number = 999
                    else:
                        tossup.packet_id = packet_id
                        tossup.question_number = Tossup.objects.filter(packet_id=packet_id).count() + 1

                    tossup.save_question(edit_type=QUESTION_CREATE, changer=user)
                    cache.clear()
                    message = 'Your tossup has been added to the set.'
                    message_class = 'alert-box info radius'

                    # In the success case, don't return the whole tossup object so as to clear the fields
                    return render(request, 'add_tossups.html',
                             {'form': TossupForm(qset_id=qset.id, packet_id=packet_id, initial={'question_type': question_type_id}, writer=user.user.username),
                             'message': message,
                             'message_class': message_class,
                             'tossup' : None,
                             'tossup_id': tossup.id,
                             'dup_matches': _post_submit_dup_matches(qset, tossup, 'tossup'),
                             'read_only': read_only,
                             'user': user,
                             'qset': qset})

                except InvalidTossup as ex:
                    message = str(ex)
                    message_class = 'alert-box warning'

            else:
                message = 'Problem adding a tossup.  Make sure that all required fields are filled out!'
                message_class = 'alert-box warning'

        else:
            tossup = None
            message = 'You are not authorized to add questions to this tournament!'
            message_class = 'alert-box warning'
            tossup_form = []
            read_only = True
            
        if (tossup_form is None):
            tossup_form = TossupForm(qset_id=qset.id, packet_id=packet_id, initial={'question_type': question_type_id})

        # In the error case, return the whole tossup object so you can edit it
        return render(request, 'add_tossups.html',
                 {'form': tossup_form,
                 'message': message,
                 'message_class': message_class,
                 'tossup' : tossup,
                 'tossup_id': None,
                 'read_only': read_only,
                 'user': user,
                 'qset': qset})

    else:
        return render(request, 'failure.html',
            {'message': 'The request cannot be completed as specified',
             'message_class': 'alert-box alert'})

@login_required
def add_bonuses(request, qset_id, bonus_type, packet_id=None):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    message = ''
    message_class = ''
    read_only = True
    role = get_role_no_owner(user, qset)
    question_type_id = []
    bonus_form = []

    if (QuestionType.objects.exists()):
        if (bonus_type == VHSL_BONUS):
            question_type_id = QuestionType.objects.get(question_type=VHSL_BONUS)
        elif (bonus_type == ACF_STYLE_BONUS):
            question_type_id = QuestionType.objects.get(question_type=ACF_STYLE_BONUS)
        else:
            return render(request, 'failure.html',
                {'message': 'The request cannot be completed as specified.  Bonus type is invalid.',
                 'message_class': 'alert-box alert'})

    if request.method == 'GET':
        if user in qset.editor.all() or user in qset.writer.all() or qset.is_owner(user):
            form = BonusForm(qset_id=qset.id, packet_id=packet_id, role=role, initial={'question_type': question_type_id}, writer=user.user.username, question_type=bonus_type)
            read_only = False
        else:
            form = None
            message = 'You are not authorized to add questions to this tournament!'
            message_class = 'alert-box warning'
            read_only = True

        return render(request, 'add_bonuses.html',
            {'form': form,
             'message': message,
             'message_class': message_class,
             'read_only': read_only,
             'question_type': bonus_type,
             'user': user,
             'qset': qset})

    elif request.method == 'POST':
        bonus = None
        if user in qset.editor.all() or user in qset.writer.all() or qset.is_owner(user):
            bonus_form = BonusForm(request.POST, qset_id=qset.id, packet_id=packet_id, initial={'question_type': question_type_id}, writer=user.user.username, question_type=bonus_type)
            read_only = False

            if bonus_form.is_valid():
                bonus = bonus_form.save(commit=False)
                if (bonus.author is None):
                    bonus.author = user

                bonus.question_set = qset
                bonus.leadin = strip_markup(bonus.leadin)
                bonus.part1_text = strip_markup(bonus.part1_text)
                bonus.part1_answer = strip_markup(bonus.part1_answer)
                bonus.part2_text = strip_markup(bonus.part2_text)
                bonus.part2_answer = strip_markup(bonus.part2_answer)
                bonus.part3_text = strip_markup(bonus.part3_text)
                bonus.part3_answer = strip_markup(bonus.part3_answer)
                bonus.locked = False

                if packet_id is None or packet_id == '':
                    # If the bonus doesn't have a packet, set its number to be the magic number
                    # of 999, meaning that it's unassigned.  Can't assign -1 because this is outside
                    # of the legal range of bonus numbers and it ends up getting set to 1 for some
                    # reason, except in the case where there are no packets in the system in which
                    # case there's an error adding the question
                    bonus.question_number = 999
                else:
                    bonus.packet_id = packet_id
                    bonus.question_number = Bonus.objects.filter(packet_id=packet_id).count() + 1

                try:
                    bonus.is_valid()
                    bonus.save_question(edit_type=QUESTION_CREATE, changer=user)
                    cache.clear()
                    message = 'Your bonus has been added to the set.'
                    message_class = 'alert-box success'

                    # On success case, don't return the full bonus so that field gets cleared
                    return render(request, 'add_bonuses.html',
                             {'form': BonusForm(qset_id=qset.id, packet_id=packet_id, initial={'question_type': question_type_id}, writer=user.user.username, question_type=bonus_type),
                             'message': message,
                             'message_class': message_class,
                             'bonus': None,
                             'bonus_id': bonus.id,
                             'dup_matches': _post_submit_dup_matches(qset, bonus, 'bonus'),
                             'read_only': read_only,
                             'question_type': bonus_type,
                             'user': user,
                             'qset': qset})

                except InvalidBonus as ex:
                    message = str(ex)
                    message_class = 'alert-box alert'

            else:
                message = 'There was an error with the form: ' + str(bonus_form.errors)
                message_class = 'alert-box alert'

            read_only = False
        else:
            message = 'You are not authorized to add questions to this tournament!'
            message_class = 'alert-box alert'
            bonus_form = []
            bonus = None
            read_only = True

        if (bonus_form is None):
            bonus_form = BonusForm(qset_id=qset.id, packet_id=packet_id, initial={'question_type': question_type_id}, writer=user.user.username)

        return render(request, 'add_bonuses.html',
                 {'form': bonus_form,
                 'message': message,
                 'message_class': message_class,
                 'bonus': bonus,
                 'bonus_id': None,
                 'read_only': read_only,
                 'question_type': bonus_type,
                 'user': user,
                 'qset': qset})

    else:
        return render(request, 'failure.html',
            {'message': 'The request cannot be completed as specified',
             'message_class': 'alert-box alert'})


#########################################################################
# Suggested edits (track-changes style proposals on questions)
#########################################################################

SUGGESTABLE_FIELDS = {
    'tossup': [('tossup_text', 'Tossup Text'), ('tossup_answer', 'Answer')],
    'bonus': [('leadin', 'Leadin'),
              ('part1_text', 'Part 1'), ('part1_answer', 'Part 1 Answer'),
              ('part2_text', 'Part 2'), ('part2_answer', 'Part 2 Answer'),
              ('part3_text', 'Part 3'), ('part3_answer', 'Part 3 Answer')],
}


def _is_set_member(user, qset):
    return qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all()


def _can_review_suggestions(user, question, qset):
    """The question's author, or any editor/owner, may accept/reject suggestions."""
    return (getattr(question, 'author_id', None) == user.id
            or qset.is_owner(user) or user in qset.editor.all())


def _pending_suggestions(question, qtype):
    if question is None:
        return []
    return list(SuggestedEdit.objects.filter(
        question_type=qtype, question_id=question.id, status='pending')
        .select_related('suggested_by__user'))


def _suggestion_render_ctx(user, question, qtype, qset):
    if question is None:
        return {'pending_suggestions': [], 'can_review_suggestions': False,
                'suggest_fields': [], 'can_suggest': False}
    return {
        'pending_suggestions': _pending_suggestions(question, qtype),
        'can_review_suggestions': _can_review_suggestions(user, question, qset),
        'suggest_fields': [{'name': f, 'label': lbl, 'value': getattr(question, f, '') or ''}
                           for f, lbl in SUGGESTABLE_FIELDS.get(qtype, [])],
        'can_suggest': _is_set_member(user, qset),
    }


def _apply_suggestion(question, suggestion, user):
    setattr(question, suggestion.field, suggestion.new_value)
    question.save_question(edit_type=QUESTION_CHANGE, changer=user)
    suggestion.status = 'accepted'
    suggestion.resolved_by = user
    suggestion.resolved_date = timezone.now()
    suggestion.save()
    # Other pending suggestions for the same field were diffed against the old
    # text, so they no longer apply cleanly — mark them superseded.
    SuggestedEdit.objects.filter(
        question_type=suggestion.question_type, question_id=suggestion.question_id,
        field=suggestion.field, status='pending').exclude(id=suggestion.id).update(
        status='superseded', resolved_by=user, resolved_date=timezone.now())


def _record_suggestions(qset, question, qtype, user, proposed, note=''):
    """Create a pending SuggestedEdit for each content field whose proposed value
    differs from the current one. `proposed` maps field name -> raw new value."""
    created = 0
    for f, lbl in SUGGESTABLE_FIELDS.get(qtype, []):
        if f not in proposed:
            continue
        new_val = strip_markup(proposed.get(f, '') or '')
        old_val = getattr(question, f, '') or ''
        if (new_val or '').strip() != (old_val or '').strip():
            SuggestedEdit.objects.create(
                question_set=qset, question_type=qtype, question_id=question.id,
                field=f, field_label=lbl, old_value=old_val, new_value=new_val,
                note=note, suggested_by=user)
            created += 1
    return created


@login_required
def suggest_edit(request):
    """Any set member proposes changes to a question's fields (track-changes).
    Used by the standalone suggest form; the edit pages now post the main form
    with save_as_suggestion instead."""
    user = request.user.writer
    if request.method != 'POST':
        return HttpResponseRedirect('/')
    qtype = request.POST.get('question_type', '')
    question = _style_question(qtype, request.POST.get('question_id', ''))
    if question is None:
        return render(request, 'failure.html',
                      {'message': 'No such question.', 'message_class': 'alert-box alert'})
    qset = question.question_set
    if not _is_set_member(user, qset):
        return render(request, 'failure.html',
                      {'message': 'You must be a member of this set to suggest changes.',
                       'message_class': 'alert-box alert'})
    note = (request.POST.get('note') or '').strip()[:255]
    proposed = {f: request.POST.get('field_' + f, '')
                for f, _ in SUGGESTABLE_FIELDS.get(qtype, []) if ('field_' + f) in request.POST}
    created = _record_suggestions(qset, question, qtype, user, proposed, note)
    return HttpResponseRedirect('/edit_{0}/{1}/?suggested={2}#suggested-changes'.format(
        qtype, question.id, created))


@login_required
def resolve_suggestion(request):
    """The question's author or an editor accepts/rejects a single suggestion."""
    user = request.user.writer
    if request.method != 'POST':
        return HttpResponse(json.dumps({'ok': False, 'error': 'POST required'}))
    try:
        s = SuggestedEdit.objects.get(id=int(request.POST['suggestion_id']))
    except (KeyError, ValueError, SuggestedEdit.DoesNotExist):
        return HttpResponse(json.dumps({'ok': False, 'error': 'Not found'}))
    question = _style_question(s.question_type, s.question_id)
    if question is None or not _can_review_suggestions(user, question, s.question_set):
        return HttpResponse(json.dumps({'ok': False, 'error': 'Not authorized'}))
    if s.status != 'pending':
        return HttpResponse(json.dumps({'ok': False, 'error': 'Already resolved'}))
    action = request.POST.get('action', '')
    if action == 'accept':
        _apply_suggestion(question, s, user)
    elif action == 'reject':
        s.status = 'rejected'
        s.resolved_by = user
        s.resolved_date = timezone.now()
        s.save()
    else:
        return HttpResponse(json.dumps({'ok': False, 'error': 'Bad action'}))
    cache.clear()
    return HttpResponse(json.dumps({'ok': True}))


@login_required
def resolve_all_suggestions(request):
    """Accept or reject every pending suggestion on a question at once."""
    user = request.user.writer
    if request.method != 'POST':
        return HttpResponse(json.dumps({'ok': False, 'error': 'POST required'}))
    qtype = request.POST.get('question_type', '')
    question = _style_question(qtype, request.POST.get('question_id', ''))
    if question is None or not _can_review_suggestions(user, question, question.question_set):
        return HttpResponse(json.dumps({'ok': False, 'error': 'Not authorized'}))
    action = request.POST.get('action', '')
    pending = SuggestedEdit.objects.filter(
        question_type=qtype, question_id=question.id, status='pending')
    if action == 'accept':
        for s in list(pending):
            s.refresh_from_db()
            if s.status == 'pending':
                _apply_suggestion(question, s, user)
    elif action == 'reject':
        pending.update(status='rejected', resolved_by=user, resolved_date=timezone.now())
    else:
        return HttpResponse(json.dumps({'ok': False, 'error': 'Bad action'}))
    cache.clear()
    return HttpResponse(json.dumps({'ok': True}))


@login_required
def edit_tossup(request, tossup_id):
    user = request.user.writer
    tossup = Tossup.objects.get(id=tossup_id)
    tossup_length = tossup.character_count()
    qset = tossup.question_set
    packet = tossup.packet
    message = ''
    message_class = ''
    read_only = True
    dup_matches = []
    role = get_role_no_owner(user, qset)

    if request.method == 'GET':
        if user == tossup.author or qset.is_owner(user) or user in qset.editor.all():
            form = TossupForm(instance=tossup, qset_id=qset.id, role=role)
            if user == tossup.author and not qset.is_owner(user) and not user in qset.editor.all() and tossup.locked:
                read_only = True
                message = 'This tossup has been locked by an editor. It cannot be changed except by another editor.'
                message_class = 'alert-box warning'
            else:
                read_only = False

        elif _is_set_member(user, qset):
            # Members who can't edit directly can still propose suggested changes,
            # so render the form (the template hides the direct Save button).
            read_only = True
            form = TossupForm(instance=tossup, qset_id=qset.id, role=role)
        else:
            read_only = True
            tossup = None
            form = None
            message = 'You are not authorized to view or edit this question!'
            message_class = 'alert-box alert'

        if request.GET.get('suggested') is not None and tossup is not None:
            if request.GET.get('suggested') != '0':
                message = '{0} change(s) saved as suggestions for the author/editors to review.'.format(request.GET.get('suggested'))
                message_class = 'alert-box success'
            else:
                message = 'No changes were detected, so nothing was suggested.'
                message_class = 'alert-box warning'

        return render(request, 'edit_tossup.html',
            {'tossup': tossup,
             'tossup_length': tossup_length,
             'form': form,
             'qset': qset,
             'packet': packet,
             'available_tags': build_tag_checkboxes(qset, tossup, tossup.category if tossup else None),
             'message': message,
             'message_class': message_class,
             'read_only': read_only,
             'role': role,
             'playtest': _question_buzz_data(tossup, 'tossup'),
             'discord_threads': tossup.discord_threads.order_by('created_date'),
             'user': user,
             **_suggestion_render_ctx(user, tossup, 'tossup', qset)})

    elif request.method == 'POST':
        print("start post for edit tossup")
        if 'save_as_suggestion' in request.POST:
            if not _is_set_member(user, qset):
                return render(request, 'failure.html',
                              {'message': 'You must be a member of this set to suggest changes.',
                               'message_class': 'alert-box alert'})
            proposed = {f: request.POST.get(f, '') for f, _ in SUGGESTABLE_FIELDS['tossup']}
            note = (request.POST.get('suggestion_note') or '').strip()[:255]
            created = _record_suggestions(qset, tossup, 'tossup', user, proposed, note)
            return HttpResponseRedirect(
                '/edit_tossup/{0}/?suggested={1}#suggested-changes'.format(tossup.id, created))

        if user == tossup.author or qset.is_owner(user) or user in qset.editor.all():
            form = TossupForm(request.POST, qset_id=qset.id, role=role)
            can_change = True
            if tossup.locked and not (qset.is_owner(user) or user in qset.editor.all()):
                can_change = False

            if form.is_valid() and can_change:
                read_only = False

                is_tossup_already_edited = tossup.edited
                is_tossup_already_proofread = tossup.proofread
                is_tossup_already_read_carefully = tossup.read_carefully

                tossup.tossup_text = strip_markup(form.cleaned_data['tossup_text'])
                tossup.tossup_answer = strip_markup(form.cleaned_data['tossup_answer'])
                tossup.category = form.cleaned_data['category']
                tossup.packet = form.cleaned_data['packet']
                tossup.locked = form.cleaned_data['locked']
                tossup.edited = form.cleaned_data['edited']
                tossup.proofread = form.cleaned_data['proofread']
                tossup.read_carefully = form.cleaned_data['read_carefully']
                tossup.question_type = form.cleaned_data['question_type']
                tossup.author = form.cleaned_data['author']
                print("trying to save tossup")

                try:
                    tossup.is_valid()
                    change_type = QUESTION_CHANGE
                    if (not is_tossup_already_edited and tossup.edited == True):
                        change_type = QUESTION_EDIT

                    if (not is_tossup_already_proofread and tossup.proofread == True):
                        change_type = QUESTION_PROOFREAD

                    if (not is_tossup_already_read_carefully and tossup.read_carefully == True):
                        change_type = QUESTION_READ_CAREFULLY

                    tossup.save_question(edit_type=change_type, changer=user)
                    save_tag_selection(request, qset, tossup, tossup.category, is_tossup=True)
                    tossup_length = tossup.character_count()
                    cache.clear()
                    print("Tossup saved")
                    message = 'Your changes have been saved!'
                    message_class = 'alert-box success'
                    dup_matches = _post_submit_dup_matches(qset, tossup, 'tossup')

                except InvalidTossup as ex:
                    message = str(ex)
                    message_class = 'alert-box warning'

            elif form.is_valid() and not can_change:
                message = 'This tossup is locked and can only be changed by an editor!'
                message_class = 'alert-box warning'
                read_only = True
            else:
                message = 'There was an error with the form: ' + str(form.errors)
                message_class = 'alert-box warning'


        elif user in qset.writer.all():
            read_only = True
            form = None
            message = 'You are only authorized to view, not to edit, this question!'
            message_class = 'alert-box warning'
        else:
            read_only = True
            tossup = None
            message = 'You are not authorized to view or edit this question!'
            message_class = 'alert-box alert'

        return render(request, 'edit_tossup.html',
            {'tossup': tossup,
             'tossup_length': tossup_length,
             'form': form,
             'role': role,
             'qset': qset,
             'packet': packet,
             'available_tags': build_tag_checkboxes(qset, tossup, tossup.category if tossup else None),
             'message': message,
             'message_class': message_class,
             'dup_matches': dup_matches,
             'read_only': read_only,
             'playtest': _question_buzz_data(tossup, 'tossup'),
             'discord_threads': tossup.discord_threads.order_by('created_date'),
             'user': user,
             **_suggestion_render_ctx(user, tossup, 'tossup', qset)})

@login_required
def edit_bonus(request, bonus_id):
    user = request.user.writer
    bonus = Bonus.objects.get(id=bonus_id)
    char_count = bonus.character_count()
    qset = bonus.question_set
    packet = bonus.packet
    message = ''
    message_class = ''
    read_only = True
    dup_matches = []
    role = get_role_no_owner(user, qset)

    question_type = ACF_STYLE_BONUS
    if (bonus.question_type is not None):
        question_type = bonus.question_type.question_type
        
    if request.method == 'GET':
        if user == bonus.author or qset.is_owner(user) or user in qset.editor.all():
            form = BonusForm(instance=bonus, qset_id=qset.id, role=role, question_type=question_type)
            if user == bonus.author and not qset.is_owner(user) and not user in qset.editor.all() and bonus.locked:
                read_only = True
                message = 'This bonus has been locked by an editor. It cannot be changed except by another editor.'
                message_class = 'alert-box warning'
            else:
                read_only = False

        elif _is_set_member(user, qset):
            # Members who can't edit directly can still propose suggested changes.
            read_only = True
            form = BonusForm(instance=bonus, qset_id=qset.id, role=role, question_type=question_type)
        else:
            read_only = True
            bonus = None
            form = None
            message = 'You are not authorized to view or edit this question!'
            message_class = 'alert-box alert'

        if request.GET.get('suggested') is not None and bonus is not None:
            if request.GET.get('suggested') != '0':
                message = '{0} change(s) saved as suggestions for the author/editors to review.'.format(request.GET.get('suggested'))
                message_class = 'alert-box success'
            else:
                message = 'No changes were detected, so nothing was suggested.'
                message_class = 'alert-box warning'

        return render(request, 'edit_bonus.html',
            {'bonus': bonus,
             'char_count': char_count,
             'question_type': question_type,
             'form': form,
             'qset': qset,
             'packet': packet,
             'available_tags': build_tag_checkboxes(qset, bonus, bonus.category if bonus else None),
             'message': message,
             'message_class': message_class,
             'read_only': read_only,
             'role': role,
             'playtest': _question_buzz_data(bonus, 'bonus'),
             'discord_threads': bonus.discord_threads.order_by('created_date'),
             'user': user,
             **_suggestion_render_ctx(user, bonus, 'bonus', qset)})

    elif request.method == 'POST':
        if 'save_as_suggestion' in request.POST:
            if not _is_set_member(user, qset):
                return render(request, 'failure.html',
                              {'message': 'You must be a member of this set to suggest changes.',
                               'message_class': 'alert-box alert'})
            proposed = {f: request.POST.get(f, '') for f, _ in SUGGESTABLE_FIELDS['bonus']}
            note = (request.POST.get('suggestion_note') or '').strip()[:255]
            created = _record_suggestions(qset, bonus, 'bonus', user, proposed, note)
            return HttpResponseRedirect(
                '/edit_bonus/{0}/?suggested={1}#suggested-changes'.format(bonus.id, created))

        if user == bonus.author or qset.is_owner(user) or user in qset.editor.all():
            form = BonusForm(request.POST, qset_id=qset.id, role=role, question_type=question_type)

            can_change = True
            if bonus.locked and not (qset.is_owner(user) or user in qset.editor.all()):
                can_change = False

            if form.is_valid() and can_change:
                is_bonus_already_edited = bonus.edited
                is_bonus_already_proofread = bonus.proofread
                is_bonus_already_read_carefully = bonus.read_carefully

                bonus.leadin = strip_markup(form.cleaned_data['leadin'])
                bonus.part1_text = strip_markup(form.cleaned_data['part1_text'])
                bonus.part1_answer = strip_markup(form.cleaned_data['part1_answer'])
                bonus.part2_text = strip_markup(form.cleaned_data['part2_text'])
                bonus.part2_answer = strip_markup(form.cleaned_data['part2_answer'])
                bonus.part3_text = strip_markup(form.cleaned_data['part3_text'])
                bonus.part3_answer = strip_markup(form.cleaned_data['part3_answer'])
                bonus.part1_difficulty = form.cleaned_data.get('part1_difficulty', '')
                bonus.part2_difficulty = form.cleaned_data.get('part2_difficulty', '')
                bonus.part3_difficulty = form.cleaned_data.get('part3_difficulty', '')
                bonus.category = form.cleaned_data['category']
                bonus.packet = form.cleaned_data['packet']
                bonus.locked = form.cleaned_data['locked']
                bonus.edited = form.cleaned_data['edited']
                bonus.proofread = form.cleaned_data['proofread']
                bonus.read_carefully = form.cleaned_data['read_carefully']
                bonus.question_type = form.cleaned_data['question_type']
                bonus.author = form.cleaned_data['author']

                try:
                    bonus.is_valid()
                    change_type = QUESTION_CHANGE
                    if (not is_bonus_already_edited and bonus.edited):
                        change_type = QUESTION_EDIT

                    if (not is_bonus_already_proofread and bonus.proofread):
                        change_type = QUESTION_PROOFREAD
                        
                    if (not is_bonus_already_read_carefully and bonus.read_carefully):
                        change_type = QUESTION_READ_CAREFULLY

                    bonus.save_question(edit_type=change_type, changer=user)
                    save_tag_selection(request, qset, bonus, bonus.category, is_tossup=False)
                    char_count = bonus.character_count()
                    cache.clear()

                    message = 'Your changes have been saved!'
                    message_class = 'alert-box success'
                    read_only = False
                    dup_matches = _post_submit_dup_matches(qset, bonus, 'bonus')
                except InvalidBonus as ex:
                    message = str(ex)
                    message_class = 'alert-box warning'
                    read_only = False

            elif form.is_valid() and not can_change:
                message = 'This bonus is locked and can only be changed by an editor!'
                message_class = 'alert-box warning'
                read_only = True
            else:
                message = 'There was an error with the form: ' + str(form.errors)
                message_class = 'alert-box warning'

        elif user in qset.writer.all():
            form = None
            read_only = True
            message = 'You are only authorized to view, not to edit, this question!'
            message_class = 'alert-box warning'
        else:
            form = None
            bonus = None
            read_only = True
            message = 'You are not authorized to view or edit this question!'
            message_class = 'alert-box alert'

        return render(request, 'edit_bonus.html',
            {'bonus': bonus,
             'char_count': char_count,
             'question_type': question_type,
             'form': form,
             'qset': qset,
             'packet': packet,
             'available_tags': build_tag_checkboxes(qset, bonus, bonus.category if bonus else None),
             'message': message,
             'message_class': message_class,
             'dup_matches': dup_matches,
             'read_only': read_only,
             'role': role,
             'playtest': _question_buzz_data(bonus, 'bonus'),
             'discord_threads': bonus.discord_threads.order_by('created_date'),
             'user': user,
             **_suggestion_render_ctx(user, bonus, 'bonus', qset)})

@login_required
def delete_tossup(request):
    user = request.user.writer
    message = ''
    message_class = ''
    read_only = True

    if request.method == 'POST':
        tossup_id = int(request.POST['tossup_id'])
        tossup = Tossup.objects.get(id=tossup_id)
        qset = tossup.question_set
        if user == tossup.author or qset.is_owner(user) or user in qset.editor.all():
            tossup.delete()
            message = 'Tossup deleted'
            message_class = 'alert-box success'
            read_only = False
        else:
            message = 'You are not authorized to delete questions from this set!'
            message_class = 'alert-box warning'

    return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))

@login_required
def delete_bonus(request):
    user = request.user.writer
    message = ''
    message_class = ''
    read_only = True

    if request.method == 'POST':
        bonus_id = int(request.POST['bonus_id'])
        bonus = Bonus.objects.get(id=bonus_id)
        qset = bonus.question_set
        if user == bonus.author or qset.is_owner(user) or user in qset.editor.all():
            bonus.delete()
            cache.clear()
            message = 'Bonus deleted'
            message_class = 'alert-box success'
            read_only = False
        else:
            message = 'You are not authorized to delete questions from this set!'
            message_class = 'alert-box warning'

    return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))

@login_required
def delete_writer(request):
    user = request.user.writer
    message = ''
    message_class = ''
    read_only = True

    if request.method == 'POST':
        qset_id = request.POST['qset_id']
        qset = QuestionSet.objects.get(id=qset_id)
        writer_id = request.POST['writer_id']
        writer = qset.writer.get(id=writer_id)
        role = get_role_no_owner(user, qset)
        if role == "editor":
            qset.writer.remove(writer)
            # If they're still in a role group attached as writer, the group
            # re-grants access; otherwise they're fully removed.
            reconcile_group_roles(qset)
            cache.clear()
            message = 'Writer removed'
            message_class = 'alert-box success'
        else:
            message = 'You are not authorized to remove writers from this set!'
            message_class = 'alert-box warning'

    return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))

@login_required
def delete_editor(request):
    user = request.user.writer
    message = ''
    message_class = ''
    read_only = True

    if request.method == 'POST':
        qset_id = request.POST['qset_id']
        qset = QuestionSet.objects.get(id=qset_id)
        editor_id = request.POST['editor_id']
        editor = qset.editor.get(id=editor_id)
        role = get_role_no_owner(user, qset)
        if role == "editor":
            qset.editor.remove(editor)
            reconcile_group_roles(qset)
            cache.clear()
            message = 'Editor removed'
            message_class = 'alert-box success'
        else:
            message = 'You are not authorized to remove editors from this set!'
            message_class = 'alert-box warning'

    return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))

@login_required
def delete_set(request):
    user = request.user.writer
    message = ''
    message_class = ''
    read_only = True

    print("In editor removed")
    if request.method == 'POST':
        qset_id = request.POST['qset_id']
        qset = QuestionSet.objects.get(id=qset_id)
        if qset.is_owner(user) or user in qset.editor.all():
            from .set_importer import delete_question_set
            delete_question_set(qset)
            cache.clear()
            message = 'Set deleted'
            message_class = 'alert-box success'
        else:
            message = 'You are not authorized to delete this set!'
            message_class = 'alert-box warning'

    return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))

@login_required
def delete_comment(request):
    user = request.user.writer
    message = ''
    message_class = ''
    read_only = True

    if request.method == 'POST':
        qset_id = request.POST['qset_id']
        qset = QuestionSet.objects.get(id=qset_id)
        qset_editors = qset.editor.all()
        comment_id = request.POST['comment_id']
        comment = Comment.objects.get(id=comment_id)

        if (comment is None):
            message = 'Error retrieving comment.'
            message_class = 'alert-box warning'
        else:
            if user in qset_editors:
                comment.is_removed = True
                comment.save()
                cache.clear()
                message = 'Comment removed'
                message_class = 'alert-box success'
            else:
                message = 'You are not authorized to remove comments from this set!'
                message_class = 'alert-box warning'

    return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))

@login_required
def post_comment(request):
    """Create a top-level comment on a tossup, bonus, or packet via AJAX.

    Bypasses django_comments' security form (whose timestamp expires after a
    couple of hours, producing a 400 on a tab left open too long). The Comment
    post_save signals still fire, so @mentions and notification emails work.
    """
    if request.method != 'POST':
        return HttpResponse(json.dumps({'success': False, 'message': 'Invalid request'}))
    target_type = request.POST.get('target_type')
    target_id = request.POST.get('target_id')
    text = (request.POST.get('comment_text') or '').strip()
    models_by_type = {'tossup': Tossup, 'bonus': Bonus, 'packet': Packet}
    if target_type not in models_by_type or not target_id or not text:
        return HttpResponse(json.dumps({'success': False, 'message': 'Missing or invalid fields.'}))
    try:
        obj = models_by_type[target_type].objects.select_related('question_set').get(id=target_id)
    except (Tossup.DoesNotExist, Bonus.DoesNotExist, Packet.DoesNotExist):
        return HttpResponse(json.dumps({'success': False, 'message': 'Target not found.'}))

    qset = obj.question_set
    user = request.user.writer
    if user not in qset.writer.all() and user not in qset.editor.all() and not qset.is_owner(user):
        return HttpResponse(json.dumps({'success': False, 'message': 'You are not authorized to comment on this set.'}))

    from django.contrib.sites.models import Site
    Comment.objects.create(
        content_type=ContentType.objects.get_for_model(obj),
        object_pk=str(obj.id), site=Site.objects.get_current(),
        user=request.user, comment=text, is_public=True, is_removed=False)
    cache.clear()
    return HttpResponse(json.dumps({'success': True, 'message': 'Comment posted.'}))


@login_required
def reply_to_comment(request):
    message = ''
    message_class = ''

    if request.method == 'POST':
        parent_id = request.POST.get('parent_id')
        comment_text = request.POST.get('comment_text', '').strip()
        qset_id = request.POST.get('qset_id')

        if not parent_id or not comment_text or not qset_id:
            message = 'Missing required fields.'
            message_class = 'alert-box warning'
            return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))

        try:
            parent_comment = Comment.objects.get(id=parent_id)
            qset = QuestionSet.objects.get(id=qset_id)
        except (Comment.DoesNotExist, QuestionSet.DoesNotExist):
            message = 'Comment or question set not found.'
            message_class = 'alert-box warning'
            return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))

        user = request.user.writer
        qset_writers = qset.writer.all()
        qset_editors = qset.editor.all()

        if user not in qset_writers and user not in qset_editors and not qset.is_owner(user):
            message = 'You are not authorized to comment on this set.'
            message_class = 'alert-box warning'
            return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))

        from django.contrib.sites.models import Site
        new_comment = Comment(
            content_type=parent_comment.content_type,
            object_pk=parent_comment.object_pk,
            site=Site.objects.get_current(),
            user=request.user,
            comment=comment_text,
        )
        new_comment.save()

        CommentReply.objects.create(comment=new_comment, parent=parent_comment)

        message = 'Reply posted.'
        message_class = 'alert-box success'

    return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))


@login_required
def add_anchored_comment(request):
    message = ''
    message_class = ''

    if request.method == 'POST':
        question_type = request.POST.get('question_type')
        question_id = request.POST.get('question_id')
        comment_text = request.POST.get('comment_text', '').strip()
        selected_text = request.POST.get('selected_text', '').strip()
        prefix = request.POST.get('prefix', '')
        suffix = request.POST.get('suffix', '')

        if not question_id or not comment_text or not selected_text or question_type not in ('tossup', 'bonus'):
            message = 'Missing required fields.'
            message_class = 'alert-box warning'
            return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))

        try:
            if question_type == 'tossup':
                question = Tossup.objects.get(id=question_id)
            else:
                question = Bonus.objects.get(id=question_id)
        except (Tossup.DoesNotExist, Bonus.DoesNotExist):
            message = 'Question not found.'
            message_class = 'alert-box warning'
            return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))

        qset = question.question_set
        user = request.user.writer

        if user not in qset.writer.all() and user not in qset.editor.all() and not qset.is_owner(user):
            message = 'You are not authorized to comment on this set.'
            message_class = 'alert-box warning'
            return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))

        from django.contrib.sites.models import Site
        new_comment = Comment(
            content_type=ContentType.objects.get_for_model(question),
            object_pk=str(question.id),
            site=Site.objects.get_current(),
            user=request.user,
            comment=comment_text,
        )
        new_comment.save()

        CommentAnchor.objects.create(
            comment=new_comment,
            selected_text=selected_text[:1000],
            prefix=prefix[:100],
            suffix=suffix[:100],
        )

        message = 'Comment posted.'
        message_class = 'alert-box success'

    return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))


@login_required
def add_question_comment(request):
    """Add a plain (non-anchored) comment to a tossup or bonus. Used by the
    document viewer's inline comment boxes."""
    message = ''
    message_class = ''

    if request.method == 'POST':
        question_type = request.POST.get('question_type')
        question_id = request.POST.get('question_id')
        comment_text = request.POST.get('comment_text', '').strip()

        if not question_id or not comment_text or question_type not in ('tossup', 'bonus'):
            return HttpResponse(json.dumps({'message': 'Missing required fields.', 'message_class': 'alert-box warning'}))

        try:
            if question_type == 'tossup':
                question = Tossup.objects.get(id=question_id)
            else:
                question = Bonus.objects.get(id=question_id)
        except (Tossup.DoesNotExist, Bonus.DoesNotExist):
            return HttpResponse(json.dumps({'message': 'Question not found.', 'message_class': 'alert-box warning'}))

        qset = question.question_set
        user = request.user.writer
        if user not in qset.writer.all() and user not in qset.editor.all() and not qset.is_owner(user):
            return HttpResponse(json.dumps({'message': 'You are not authorized to comment on this set.', 'message_class': 'alert-box warning'}))

        from django.contrib.sites.models import Site
        Comment.objects.create(
            content_type=ContentType.objects.get_for_model(question),
            object_pk=str(question.id),
            site=Site.objects.get_current(),
            user=request.user,
            comment=comment_text,
            is_public=True,
            is_removed=False,
        )
        message = 'Comment posted.'
        message_class = 'alert-box success'

    return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))


@login_required
def delete_all_comments(request):
    user = request.user.writer
    message = ''
    message_class = ''
    read_only = True

    tossup_content_type_id = ContentType.objects.get_for_model(Tossup).id
    bonus_content_type_id = ContentType.objects.get_for_model(Bonus).id

    if request.method == 'POST':
        qset_id = request.POST['qset_id']
        qset = QuestionSet.objects.get(id=qset_id)
        qset_editors = qset.editor.all()
        question_type = request.POST['question_type']
        question_id = request.POST['question_id']

        if (question_type == 'tossup'):
            comment_list = Comment.objects.filter(content_type_id=tossup_content_type_id).filter(object_pk=question_id).order_by('submit_date')
        else:
            comment_list = Comment.objects.filter(content_type_id=bonus_content_type_id).filter(object_pk=question_id).order_by('submit_date')

        if (comment_list is None):
            message = 'Error retrieving comments.'
            message_class = 'alert-box warning'
        else:
            if user in qset_editors:
                comment_list.update(is_removed=True)
                cache.clear()

                message = 'Comments removed'
                message_class = 'alert-box success'
            else:
                message = 'You are not authorized to remove comments from this set!'
                message_class = 'alert-box warning'

    return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))

@login_required
def add_packets(request, qset_id):

    qset = QuestionSet.objects.get(id=qset_id)
    user = request.user.writer
    message = ''
    message_class = ''

    if qset.is_owner(user):
        if request.method == 'GET':
            form = NewPacketsForm()
        elif request.method == 'POST':
            form = NewPacketsForm(data=request.POST)
            if form.is_valid():
                packet_name = form.cleaned_data['packet_name']
                name_base = form.cleaned_data['name_base']
                num_packets = form.cleaned_data['num_packets']
                if packet_name and len(packet_name.strip()) > 0 and (name_base is None or num_packets is None):
                    if Packet.objects.filter(question_set=qset, packet_name=packet_name).exists():
                        message = 'The packet name "{0}" arleady exists.'.format(packet_name)
                        message_class = 'alert-box warning'
                    else:
                        new_packet = Packet()
                        new_packet.packet_name = packet_name
                        new_packet.created_by = user
                        new_packet.question_set = qset
                        new_packet.save()
                        cache.clear()
                        message = 'Your packet named {0} has been created.'.format(packet_name)
                        message_class = 'alert-box success'

                elif name_base and len(name_base.strip()) > 0 and num_packets is not None:
                    create_all_failed = False
                    for i in range(1, num_packets + 1):
                        new_packet = Packet()
                        packet_name = '{0!s} {1:02}'.format(name_base, i)
                        if Packet.objects.filter(question_set=qset, packet_name=packet_name).exists():
                            message = 'The packet name "{0}" arleady exists.'.format(packet_name)
                            message_class = 'alert-box warning'
                            create_all_failed = True
                            break
                        new_packet.packet_name = packet_name
                        new_packet.created_by = user
                        new_packet.question_set = qset
                        new_packet.save()
                        cache.clear()
                    if not create_all_failed:
                        message = 'Your {0} packet(s) with the base name {1} have been created.'.format(num_packets, name_base)
                        message_class = 'alert-box success'
                else:
                    message = 'You must enter either the name for an individual packet or a base name and the number of packets to create!'
                    message_class = 'alert-box warning'

            else:
                message = 'Invalid information entered into form!'
                message_class = 'alert-box alert'
        else:
            message = 'Invalid method!'
            message_class = 'alert-box alert'
            form = None

    else:
        message = 'You are not authorized to add packets to this set!'
        message_class = 'alert-box alert'
        form = None

    return render(request, 'add_packets.html',
                             {'message': message,
                              'message_class': message_class,
                              'form': form,
                              'qset': qset,
                              'user': user})

@login_required
def delete_packet(request):
    user = request.user.writer
    message = ''
    message_class = ''
    read_only = True

    if request.method == 'POST':
        packet_id = int(request.POST['packet_id'])
        packet = Packet.objects.get(id=packet_id)
        qset = packet.question_set
        if qset.is_owner(user):
            packet.delete()
            cache.clear()
            message = 'Packet deleted'
            message_class = 'alert-box success'
            read_only = False
        else:
            message = 'You are not authorized to delete packets from this set!'
            message_class = 'alert-box warning'

    return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))

@login_required
def get_unassigned_tossups(request):
    user = request.user.writer
    qset_id = request.GET['qset_id']
    message = ''
    message_class = ''
    data = []

    try:
        qset = QuestionSet.objects.get(id=qset_id)

        if request.method == 'GET':
            if qset.is_owner(user):
                available_tossups = Tossup.objects.filter(question_set=qset, packet=None)
                for tu in available_tossups:
                    data.append(tu.to_json())
            else:
                available_tossups = []
                message = 'Only the set owner has the power to add questions to it!'
                message_class = 'alert-box alert'

        else:
            message = 'Invalid request!'
            message_class = 'alert-box alert'
    except Exception as ex:
        print(ex)
        message = 'Unable to retrieve question set; qset_id either missing or incorrect!'
        message_class = 'alert-box alert'

    return HttpResponse(json.dumps(data))

@login_required
def get_unassigned_bonuses(request):
    user = request.user.writer
    qset_id = request.GET['qset_id']
    message = ''
    message_class = ''
    data = []

    try:
        qset = QuestionSet.objects.get(id=qset_id)

        if request.method == 'GET':
            if qset.is_owner(user):
                available_bonuses= Bonus.objects.filter(question_set=qset, packet=None)
                for bs in available_bonuses:
                    data.append(bs.to_json())
            else:
                available_tossups = []
                message = 'Only the set owner has the power to add questions to it!'
                message_class = 'alert-box alert'

        else:
            message = 'Invalid request!'
            message_class = 'alert-box alert'
    except Exception as ex:
        print(ex)
        message = 'Unable to retrieve question set; qset_id either missing or incorrect!'
        message_class = 'alert-box alert'

    return HttpResponse(json.dumps(data))

@login_required
def assign_tossups_to_packet(request):

    user = request.user.writer
    packet_id = int(request.POST['packet_id'])
    tossup_ids = request.POST.getlist('tossup_ids[]')
    packet = Packet.objects.get(id=packet_id)
    qset = packet.question_set
    message = ''
    message_class = ''


    if request.method == 'POST':
        if qset.is_owner(user):
            for tu_id in tossup_ids:
                # Only questions already in this packet's set may be assigned,
                # so foreign question IDs can't be pulled across sets
                tossup = Tossup.objects.filter(id=tu_id, question_set=qset).first()
                if tossup is None:
                    continue
                tossup.packet = packet
                # Potential race condition?
                tossup.question_number = Tossup.objects.filter(packet_id=packet_id).count() + 1
                message = 'Your tossups have been added to the set!'
                message_class = 'alert-box success'
                tossup.save()
                cache.clear()
        else:
            message = 'Only the set owner is authorized to add questions to the set!'
            message_class = 'alert-box warning'

    else:
        message = 'Invalid request!'
        message_class = 'alert-box alert'

    return HttpResponse(json.dumps({'message': message,
                                    'message_class': message_class}))

@login_required
def assign_bonuses_to_packet(request):

    user = request.user.writer
    packet_id = int(request.POST['packet_id'])
    bonus_ids = request.POST.getlist('bonus_ids[]')
    packet = Packet.objects.get(id=packet_id)
    qset = packet.question_set
    message = ''
    message_class = ''

    if request.method == 'POST':
        if qset.is_owner(user):
            for bs_id in bonus_ids:
                # Only questions already in this packet's set may be assigned,
                # so foreign question IDs can't be pulled across sets
                bonus = Bonus.objects.filter(id=bs_id, question_set=qset).first()
                if bonus is None:
                    continue
                bonus.packet = packet
                bonus.question_number = Bonus.objects.filter(packet_id=packet_id).count() + 1
                message = 'Your bonuses have been added to the set!'
                message_class = 'alert-box success'
                bonus.save()
                cache.clear()
        else:
            message = 'Only the set owner is authorized to add questions to the set!'
            message_class = 'alert-box warning'

    else:
        message = 'Invalid request!'
        message_class = 'alert-box alert'

    return HttpResponse(json.dumps({'message': message,
                                    'message_class': message_class}))

@login_required
def change_question_order(request):

    user = request.user.writer
    packet_id = int(request.POST['packet_id'])
    num_questions = int(request.POST['num_questions'])
    question_type = request.POST['question_type']
    packet = Packet.objects.get(id=packet_id)
    qset = packet.question_set

    if request.method == 'POST':
        if qset.is_owner(user):
            try:
                for i in range(num_questions):
                    id_key = 'order_data[{0}][id]'.format(i)
                    order_key = 'order_data[{0}][order]'.format(i)
                    id = int(request.POST[id_key])
                    order = int(request.POST[order_key])
                    if question_type == 'tossup':
                        question = Tossup.objects.get(id=id, question_set=qset)
                    elif question_type == 'bonus':
                        question = Bonus.objects.get(id=id, question_set=qset)
                    question.question_number = order
                    question.save()
                    cache.clear()
                message = ''
                message_class = ''

            except Exception as ex:
                print(ex)
                message = 'Something went terribly wrong!'
                message_class = 'alert-box alert'

        else:
            message = 'Only the owner of the set is allowed to change the order of questions!'
            message_class = 'alert-box warning'
    else:
        message = 'Invalid request!'
        message_class = 'alert-box warning'

    return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))

# @login_required
# def change_tossup_order(request):
#     # packet_id, old_index, new_index
#     packet_id = int(request.POST['packet_id'])
#     user = request.user.writer
#     packet = Packet.objects.get(id=packet_id)
#     qset = packet.question_set
#
#     old_index = int(request.POST['old_index'])
#     new_index = int(request.POST['new_index'])
#
#     if request.method == 'POST':
#         if qset.is_owner(user):
#             change_question_order(packet, int(old_index), int(new_index), Tossup)
#             message = ''
#             message_class = ''
#         else:
#             message = 'Only the set owner is authorized to change question order'
#             message_class = 'alert-box warning'
#     else:
#         message = 'Invalid request!'
#         message_class = 'alert-box alert'
#
#     return HttpResponse(json.dumps({'message': message,
#                                     'message_class': message_class}))


# @login_required
# def change_bonus_order(request):
#     user = request.user.writer
#     packet = Packet.objects.get(id=packet_id)
#     qset = packet.question_set
#
#     if request.method == 'POST':
#         if qset.is_owner(user):
#             change_question_order(packet, int(old_index), int(new_index), Bonus)
#             message = ''
#             message_class = ''
#         else:
#             message = 'Only the set owner is authorized to change question order'
#             message_class = 'alert-box warning'
#     else:
#         message = 'Invalid request!'
#         message_class = 'alert-box alert'
#     return HttpResponse(json.dumps({'message': message,
#                                     'message_class': message_class}))
#
# # Not a URL action, just a helper method. old_index and new_index should be integers
# def change_question_order(packet, old_index, new_index, model_class):
#     if old_index != new_index and old_index >= 0 and new_index >= 0:
#         # If oldIndex < newIndex, decrease question_number for questions [oldIndex + 1, newIndex]
#         # Otherwise, increase question_number for questions [newIndex, oldIndex - 1]
#         lowerIndex = old_index + 1 if old_index < new_index else new_index
#         higherIndex = old_index - 1 if old_index > new_index else new_index
#
#         selected_question = model_class.objects.get(packet=packet, question_number=old_index)
#         selected_id = selected_question.id
#         reordered_questions = model_class.objects.filter(packet=packet, question_number__range=(lowerIndex, higherIndex))
#         # This prevents a race condition where selected_question's question_number is set to something in the range
#         # before this QuerySet is evaluated.
#         reordered_questions = reordered_questions.exclude(id=selected_id)
#         selected_question.question_number = new_index
#         selected_question.save()
#
#         direction = -1 if old_index < new_index else 1
#         for question in reordered_questions:
#             question.question_number += direction
#             question.save()

def _account_can_create(user):
    """Anti-spam: a new account can't create question sets or distributions
    until 2 days after it was created."""
    from datetime import timedelta
    return (timezone.now() - user.date_joined) >= timedelta(days=2)


_ACCOUNT_TOO_NEW_MSG = ('New accounts can\'t create question sets or distributions '
                        'until 2 days after sign-up. Please try again later.')


def _writer_distribution_ids(writer):
    """Distribution ids for the sets a writer belongs to (owner, co-owner,
    editor or writer) — the distributions they're allowed to see and edit."""
    sets = (QuestionSet.objects.filter(owner=writer)
            | writer.question_set_editor.all()
            | writer.question_set_writer.all()
            | writer.co_owned_sets.all())
    ids = set(sets.values_list('distribution_id', flat=True))
    # Also include distributions this writer created but hasn't yet attached
    # to one of their sets.
    ids |= set(Distribution.objects.filter(created_by=writer).values_list('id', flat=True))
    ids.discard(None)
    return ids


@login_required
def distributions (request):
    # Only show distributions for sets this user is part of.
    user = request.user.writer
    dists = Distribution.objects.filter(id__in=_writer_distribution_ids(user))

    return render(request, 'distributions.html',
                             {'dists': dists,
                              'user': user})

@login_required
def clone_distribution(request, dist_id):
    if request.method != 'POST':
        return HttpResponseRedirect('/distributions/')

    user = request.user.writer
    if int(dist_id) not in _writer_distribution_ids(user):
        return render(request, 'failure.html',
                      {'message': 'You can only clone distributions for your own sets.',
                       'message_class': 'alert-box alert'})
    if not _account_can_create(request.user):
        return render(request, 'failure.html',
                      {'message': _ACCOUNT_TOO_NEW_MSG, 'message_class': 'alert-box alert'})

    source = Distribution.objects.get(id=dist_id)
    new_dist = Distribution()
    new_dist.name = source.name + ' (Copy)'
    new_dist.created_by = user
    new_dist.acf_tossup_per_period_count = source.acf_tossup_per_period_count
    new_dist.acf_bonus_per_period_count = source.acf_bonus_per_period_count
    new_dist.vhsl_bonus_per_period_count = source.vhsl_bonus_per_period_count
    new_dist.save()

    for entry in DistributionEntry.objects.filter(distribution=source):
        DistributionEntry.objects.create(
            distribution=new_dist,
            category=entry.category,
            subcategory=entry.subcategory,
            min_tossups=entry.min_tossups,
            min_bonuses=entry.min_bonuses,
            max_tossups=entry.max_tossups,
            max_bonuses=entry.max_bonuses,
        )

    return HttpResponseRedirect('/edit_distribution/' + str(new_dist.id) + '/')

@login_required
def edit_distribution(request, dist_id=None):

    data = []
    message = ''
    message_class = ''

    user = request.user.writer
    if dist_id is None:
        # Creating a brand-new distribution: gate by account age.
        if not _account_can_create(request.user):
            return render(request, 'failure.html',
                          {'message': _ACCOUNT_TOO_NEW_MSG, 'message_class': 'alert-box alert'})
    elif int(dist_id) not in _writer_distribution_ids(user):
        return render(request, 'failure.html',
                      {'message': 'You can only view or edit distributions for your own sets.',
                       'message_class': 'alert-box alert'})

    if request.user.is_authenticated:
        DistributionEntryFormset = formset_factory(DistributionEntryForm, can_delete=True)
        if request.method == 'POST':
            # no dist_id supplied means new dist
            if dist_id is None:
                formset = DistributionEntryFormset(data=request.POST, prefix='distentry')
                dist_form = DistributionForm(data=request.POST)
                if dist_form.is_valid() and formset.is_valid():
                    new_dist = Distribution()
                    new_dist.name = dist_form.cleaned_data['name']
                    new_dist.created_by = user
                    new_dist.save()

                    for form in formset:
                        if form.cleaned_data != {}:
                            new_entry = DistributionEntry()
                            new_entry.category = form.cleaned_data['category']
                            new_entry.subcategory = form.cleaned_data['subcategory']
                            new_entry.min_bonuses = form.cleaned_data['min_bonuses']
                            new_entry.min_tossups = form.cleaned_data['min_tossups']
                            new_entry.max_bonuses = form.cleaned_data['max_bonuses']
                            new_entry.max_tossups = form.cleaned_data['max_tossups']
                            if new_entry.min_bonuses > new_entry.max_bonuses:
                                new_entry.min_bonuses = new_entry.max_bonuses
                                #TODO: display the message
                                message = 'Minimum bonuses for ' + new_entry.category + ' - ' + new_entry.subcategory +\
                                          ' was higher than maximum bonuses and has been set to maximum bonuses.'
                                message_class = 'alert-box warning'
                            if new_entry.min_tossups > new_entry.max_tossups:
                                new_entry.min_tossups = new_entry.max_tossups
                                #TODO: display the message
                                message = 'Minimum tossups for ' + new_entry.category + ' - ' + new_entry.subcategory +\
                                          ' was higher than maximum tossups and has been set to maximum tossups.'
                                message_class = 'alert-box warning'

                            new_entry.distribution = new_dist
                            new_entry.save()

                    return HttpResponseRedirect('/edit_distribution/' + str(new_dist.id))

            else:
                dist_form = DistributionForm(data=request.POST)
                #print dist_form.is_valid()
                #print formset.is_valid()
                #print formset.errors
                if 'add_row' in request.POST:
                    distentry_post = request.POST.copy()
                    #TODO: grab a value from an input
                    num_rows = 1
                    distentry_post['distentry-TOTAL_FORMS'] = int(distentry_post['distentry-TOTAL_FORMS']) + num_rows
                    formset = DistributionEntryFormset(data=distentry_post, prefix='distentry')
                else:
                    formset = DistributionEntryFormset(data=request.POST, prefix='distentry')
                    if dist_form.is_valid() and formset.is_valid():
                        dist = Distribution.objects.get(id=dist_id)
                        dist.name = dist_form.cleaned_data['name']
                        dist.save()

                        qsets = dist.questionset_set.all()
                        for form in formset:
                            if form.cleaned_data != {}:
                                if form.cleaned_data['entry_id'] is not None:
                                    entry_id = int(form.cleaned_data['entry_id'])
                                    entry = DistributionEntry.objects.get(id=entry_id)
                                    if form.cleaned_data['DELETE']:
                                        entry.delete()
                                        entry = None
                                    else:
                                        entry.category = form.cleaned_data['category']
                                        entry.subcategory = form.cleaned_data['subcategory']
                                        entry.min_bonuses = form.cleaned_data['min_bonuses']
                                        entry.min_tossups = form.cleaned_data['min_tossups']
                                        entry.max_bonuses = form.cleaned_data['max_bonuses']
                                        entry.max_tossups = form.cleaned_data['max_tossups']
                                        if entry.min_bonuses > entry.max_bonuses:
                                            entry.min_bonuses = entry.max_bonuses
                                            message = 'Minimum bonuses for ' + entry.category + ' - ' + entry.subcategory +\
                                                      ' was higher than maximum bonuses and has been set to maximum bonuses.'
                                            message_class = 'alert-box warning'
                                        if entry.min_tossups > entry.max_tossups:
                                            entry.min_tossups = entry.max_tossups
                                            message = 'Minimum tossups for ' + entry.category + ' - ' + entry.subcategory +\
                                                      ' was higher than maximum tossups and has been set to maximum tossups.'
                                            message_class = 'alert-box warning'

                                        entry.save()
                                else:
                                    entry = form.save(commit=False)
                                    entry.distribution = dist
                                    entry.save()

                                if entry is not None:
                                    for qset in qsets:
                                        set_wide_entry = qset.setwidedistributionentry_set.filter(dist_entry=entry)
                                        print(set_wide_entry)
                                        if set_wide_entry.count() == 0:
                                            new_set_wide_entry = SetWideDistributionEntry()
                                            new_set_wide_entry.dist_entry = entry
                                            new_set_wide_entry.question_set = qset
                                            new_set_wide_entry.num_tossups = qset.num_packets * entry.min_tossups
                                            new_set_wide_entry.num_bonuses = qset.num_packets * entry.min_bonuses
                                            new_set_wide_entry.save()

                        entries = dist.distributionentry_set.all()
                        initial_data = []
                        for entry in entries:
                            initial_data.append({'entry_id': entry.id,
                                                 'category': entry.category,
                                                 'subcategory': entry.subcategory,
                                                 'min_tossups': entry.min_tossups,
                                                 'min_bonuses': entry.min_bonuses,
                                                 'max_tossups': entry.max_tossups,
                                                 'max_bonuses': entry.max_bonuses})
                        formset = DistributionEntryFormset(initial=initial_data, prefix='distentry')

                    else:
                        dist = Distribution.objects.get(id=dist_id)
                        dist_form = DistributionForm(instance=dist)
                        formset = DistributionEntryFormset(data=request.POST, prefix='distentry')

            return render(request, 'edit_distribution.html',
                                     {'form': dist_form,
                                      'formset': formset,
                                      'message': message,
                                      'message_class': message_class,
                                      'user': request.user.writer})
        else:
            if dist_id is not None:
                dist = Distribution.objects.get(id=dist_id)
                entries = dist.distributionentry_set.all()
                initial_data = []
                for entry in entries:
                    initial_data.append({'entry_id': entry.id,
                                         'category': entry.category,
                                         'subcategory': entry.subcategory,
                                         'min_tossups': entry.min_tossups,
                                         'min_bonuses': entry.min_bonuses,
                                         'max_tossups': entry.max_tossups,
                                         'max_bonuses': entry.max_bonuses})
                dist_form = DistributionForm(instance=dist)
                formset = DistributionEntryFormset(initial=initial_data, prefix='distentry')
            else:
                dist_form = DistributionForm()
                formset = DistributionEntryFormset(prefix='distentry')

            return render(request, 'edit_distribution.html',
                                     {'form': dist_form,
                                      'formset': formset,
                                      'message': message,
                                      'message_class': message_class,
                                      'user': request.user.writer})

@login_required()
def edit_tiebreak(request, dist_id=None):

    user = request.user.writer
    data = []


    TiebreakDistributionEntryFormset = formset_factory(TieBreakDistributionEntryForm, can_delete=True)
    if request.method == 'POST':
        # no dist_id supplied means new dist
        if dist_id is None:
            formset = TiebreakDistributionEntryFormset(data=request.POST, prefix='tiebreak')
            dist_form = TieBreakDistributionForm(data=request.POST)
            if dist_form.is_valid() and formset.is_valid():
                new_dist = TieBreakDistribution()
                new_dist.name = dist_form.cleaned_data['name']
                new_dist.save()

                for form in formset:
                    if form.cleaned_data != {}:
                        new_entry = DistributionEntry()
                        new_entry.category = form.cleaned_data['category']
                        new_entry.subcategory = form.cleaned_data['subcategory']
                        new_entry.bonuses = form.cleaned_data['num_bonuses']
                        new_entry.tossups = form.cleaned_data['num_tossups']
                        new_entry.distribution = new_dist
                        new_entry.save()

                return HttpResponseRedirect('/edit_tiebreak/' + str(new_dist.id))
        else:
            formset = TiebreakDistributionEntryFormset(data=request.POST, prefix='tiebreak')
            dist_form = TieBreakDistributionForm(data=request.POST)
            print(dist_form.is_valid())
            print(formset.is_valid())
            print(formset.errors)
            if dist_form.is_valid() and formset.is_valid():

                dist = TieBreakDistribution.objects.get(id=dist_id)
                dist.name = dist_form.cleaned_data['name']
                qsets = dist.questionset_set.all()
                for form in formset:
                    if form.cleaned_data != {}:
                        if form.cleaned_data['entry_id'] is not None:
                            entry_id = int(form.cleaned_data['entry_id'])
                            entry = DistributionEntry.objects.get(id=entry_id)
                            if form.cleaned_data['DELETE']:
                                entry.delete()
                                entry = None
                            else:
                                entry.category = form.cleaned_data['category']
                                entry.subcategory = form.cleaned_data['subcategory']
                                entry.bonuses = form.cleaned_data['num_bonuses']
                                entry.tossups = form.cleaned_data['num_tossups']
                                entry.save()
                        else:
                            entry = form.save(commit=False)
                            entry.distribution = dist
                            entry.save()

                        if entry is not None:
                            for qset in qsets:
                                set_wide_entry = qset.tiebreakdistributionentry_set.filter(dist_entry=entry)
                                print(set_wide_entry)
                                if set_wide_entry.count() == 0:
                                    print('here')
                                    new_set_wide_entry = DistributionEntry()
                                    new_set_wide_entry.dist_entry = entry
                                    new_set_wide_entry.question_set = qset
                                    new_set_wide_entry.num_tossups = qset.num_packets * entry.min_tossups
                                    new_set_wide_entry.num_bonuses = qset.num_packets * entry.min_bonuses
                                    new_set_wide_entry.save()

                entries = dist.distributionentry_set.all()
                initial_data = []
                for entry in entries:
                    initial_data.append({'entry_id': entry.id,
                                         'category': entry.category,
                                         'subcategory': entry.subcategory,
                                         'num_bonuses': entry.min_bonuses,
                                         'num_tossups': entry.max_tossups,})
                formset = TiebreakDistributionEntryFormset(initial=initial_data, prefix='tiebreak')

            else:
                dist = Distribution.objects.get(id=dist_id)
                dist_form = DistributionForm(instance=dist)
                formset = TiebreakDistributionEntryFormset(data=request.POST, prefix='tiebreak')

        return render(request, 'edit_tiebreak.html',
                                  {'form': dist_form,
                                   'formset': formset})

    else:
        if dist_id is not None:
            dist = TieBreakDistribution.objects.get(id=dist_id)
            entries = dist.distributionentry_set.all()
            initial_data = []
            for entry in entries:
                initial_data.append({'entry_id': entry.id,
                                     'category': entry.category,
                                     'subcategory': entry.subcategory,
                                     'num_tossups': entry.min_tossups,
                                     'num_bonuses': entry.min_bonuses,})
            dist_form = TieBreakDistributionForm(instance=dist)
            formset = TiebreakDistributionEntryFormset(initial=initial_data, prefix='tiebreak')
        else:
            dist_form = TieBreakDistributionForm()
            formset = TiebreakDistributionEntryFormset(prefix='tiebreak')

        return render(request, 'edit_tiebreak.html',
        {'form': dist_form,
        'formset': formset,})


@login_required
def add_comment(request):

    user = request.user.writer
    qset_id = request.POST['qset-id']
    qset = QuestionSet.objects.get(id=qset_id)

    if request.method == 'POST':

        comment_text = request.POST['comment-text']
        cache.clear()
        print(comment_text)


@login_required
def upload_questions(request, qset_id):
    qset = QuestionSet.objects.get(id=qset_id)
    user = request.user.writer

    if request.method == 'POST':
        if (qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all()):
            form = QuestionUploadForm(request.POST, request.FILES)
            if form.is_valid():
                uploaded_tossups, uploaded_bonuses = parse_uploaded_packet(request.FILES['questions_file'])
                cache.clear()

                return render(request, 'upload_preview.html',
                {'tossups': uploaded_tossups,
                'bonuses': uploaded_bonuses,
                'message': mark_safe('Please verify that this data is correct. Hitting "Submit" will upload these questions '\
                'If you see any mistakes in the submissions, please correct them in the <strong><em>original file</em></strong> and reupload.'),
                'message_class': 'alert-box warning',
                'qset': qset})
            else:
                messages.error(request, form.questions_file.errors)
                return HttpResponseRedirect('/edit_question_set/{0}'.format(qset_id))
        else:
            messages.error(request, 'You do not have permission to upload ')

@login_required
def type_questions(request, qset_id=None):
    if qset_id is not None:
        qset = QuestionSet.objects.get(id=qset_id)
    else:
        qset = QuestionSet.objects.get(id=request.POST['qset_id'])

    user = request.user.writer

    if request.method == 'POST':
        if (qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all()):
            form = TypeQuestionsForm(request.POST)
            if form.is_valid():
                question_data = request.POST['questions'].splitlines()
                tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(question_data, qset)

                return render(request, 'type_questions_preview.html',
                                         {'tossups': tossups,
                                          'bonuses': bonuses,
                                          'tossup_errors': tossup_errors,
                                          'bonus_errors': bonus_errors,
                                          'message': 'Please verify that these questions have been correctly parsed. Hitting "Submit" will '\
                                          'commit these questions to the database. If you see any mistakes, hit "Cancel" and correct your mistakes.',
                                          'qset': qset,
                                          'user': user})
            else:
                question_data = request.POST['questions']
                tossups, bonuses = parse_packet_data(question_data, qset)
                messages.error(request, form.questions.errors)

        else:
            tossups = None
            bonuses = None
            messages.error(request, 'You do not have permission to add questions to this set')
            return render(request, 'type_questions.html',
                                     {'tossups': tossups,
                                      'bonuses': bonuses,
                                      'qset': qset,
                                      'user': user})
    elif request.method == 'GET':
        if (qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all()):
            dist_entries = qset.setwidedistributionentry_set.all().order_by('dist_entry__category', 'dist_entry__subcategory')

            form = TypeQuestionsForm(request.POST)
            return render(request, 'type_questions.html',
                                     {'user': user,
                                      'qset': qset,
                                      'form': form,
                                      'dist_entries': dist_entries})
        else:
            messages.error(request, 'You do not have permission to add questions to this set')
            return render(request, 'type_questions.html',
                                     {'qset': qset,
                                      'user': user})

@login_required
def type_questions_edit(request, question_type, question_id):
    user = request.user.writer
    
    if (question_type == "tossup"):
        question = Tossup.objects.get(id=question_id)
    elif (question_type == "bonus"):
        question = Bonus.objects.get(id=question_id)
    
    qset = question.question_set
    packet = question.packet
    message = ''
    message_class = ''
    read_only = True
    role = get_role_no_owner(user, qset)

    if request.method == 'POST':
        if (qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all()):
            form = TypeQuestionsForm(request.POST)
            if form.is_valid():
                question_data = request.POST['questions'].splitlines()
                tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(question_data, qset)

                return render(request, 'type_questions_edit_preview.html',
                                         {'tossups': tossups,
                                          'bonuses': bonuses,
                                          'tossup_errors': tossup_errors,
                                          'bonus_errors': bonus_errors,
                                          'message': 'Please verify that these questions have been correctly parsed. Hitting "Submit" will '\
                                          'commit these questions to the database. If you see any mistakes, hit "Cancel" and correct your mistakes.',
                                          'qset': qset,
                                          'user': user})
            else:
                question_data = request.POST['questions']
                tossups, bonuses = parse_packet_data(question_data, qset)
                messages.error(request, form.questions.errors)

        else:
            tossups = None
            bonuses = None
            messages.error(request, 'You do not have permission to edit this question')
            return render(request, 'type_questions_edit.html',
                                     {'tossups': tossups,
                                      'bonuses': bonuses,
                                      'qset': qset,
                                      'user': user})
    elif request.method == 'GET':
        if (qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all()):
            dist_entries = qset.setwidedistributionentry_set.all().order_by('dist_entry__category', 'dist_entry__subcategory')

            form = TypeQuestionsForm(request.POST)
            return render(request, 'type_questions_edit.html',
                                     {'user': user,
                                      'qset': qset,
                                      'form': form,
                                      'dist_entries': dist_entries})
        else:
            messages.error(request, 'You do not have permission to edit this question')
            return render(request, 'type_questions_edit.html',
                                     {'qset': qset,
                                      'user': user})

@login_required
def complete_upload(request):
    user = request.user.writer
    if request.method == 'POST':
        qset_id = request.POST['qset-id']
        qset = QuestionSet.objects.get(id=qset_id)

        if not (qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all()):
            messages.error(request, 'You are not authorized to add questions to this set!')
            return HttpResponseRedirect('/failure.html/')

        num_tossups = int(request.POST['num-tossups'])
        num_bonuses = int(request.POST['num-bonuses'])
        categories = DistributionEntry.objects.filter(distribution=qset.distribution)
        questionTypes = QuestionType.objects.all()

        new_tossups = []
        new_bonuses = []

        for tu_num in range(num_tossups):
            data="UTF-8 DATA"
            tu_text_name = 'tossup-text-{0}'.format(tu_num)
            tu_ans_name = 'tossup-answer-{0}'.format(tu_num)
            tu_cat_name = 'tossup-category-{0}'.format(tu_num)
            tu_type_name = 'tossup-type-{0}'.format(tu_num)

            tu_text = strip_markup(request.POST[tu_text_name])
            tu_ans = strip_markup(request.POST[tu_ans_name])
            tu_cat = request.POST[tu_cat_name]
            tu_type = request.POST[tu_type_name]

            new_tossup = Tossup()
            new_tossup.tossup_text = tu_text
            new_tossup.tossup_answer = tu_ans
            new_tossup.author = user
            new_tossup.question_set = qset

            for category in categories:
                formattedCategory = category.category + " - " + category.subcategory
                if (formattedCategory == tu_cat):
                    new_tossup.category = category
                    break

            for questionType in questionTypes:
                if (str(questionType) == tu_type):
                    new_tossup.question_type = questionType
                    break

            new_tossup.locked = False
            new_tossup.edited = False

            new_tossup.save_question(edit_type=QUESTION_CREATE, changer=user)
            new_tossups.append(new_tossup)

        for bs_num in range(num_bonuses):
            bs_leadin_name = 'bonus-leadin-{0}'.format(bs_num)

            bs_part1_name = 'bonus-part1-{0}'.format(bs_num)
            bs_ans1_name = 'bonus-answer1-{0}'.format(bs_num)
            bs_part2_name = 'bonus-part2-{0}'.format(bs_num)
            bs_ans2_name = 'bonus-answer2-{0}'.format(bs_num)
            bs_part3_name = 'bonus-part3-{0}'.format(bs_num)
            bs_ans3_name = 'bonus-answer3-{0}'.format(bs_num)
            bs_cat_name = 'bonus-category-{0}'.format(bs_num)
            bs_type_name = 'bonus-type-{0}'.format(bs_num)
            bs_type = request.POST[bs_type_name]

            new_bonus = Bonus()
            new_bonus.question_set = qset
            new_bonus.author = user
            new_bonus.edited = False
            new_bonus.locked = False
            new_bonus.leadin = strip_markup(request.POST[bs_leadin_name])
            new_bonus.part1_text = strip_markup(request.POST[bs_part1_name])
            new_bonus.part1_answer = strip_markup(request.POST[bs_ans1_name])
            new_bonus.part2_text = strip_markup(request.POST[bs_part2_name])
            new_bonus.part2_answer = strip_markup(request.POST[bs_ans2_name])
            new_bonus.part3_text = strip_markup(request.POST[bs_part3_name])
            new_bonus.part3_answer = strip_markup(request.POST[bs_ans3_name])
            new_bonus.part1_difficulty = request.POST.get('bonus-difficulty1-{0}'.format(bs_num), '')
            new_bonus.part2_difficulty = request.POST.get('bonus-difficulty2-{0}'.format(bs_num), '')
            new_bonus.part3_difficulty = request.POST.get('bonus-difficulty3-{0}'.format(bs_num), '')

            bonus_cat = request.POST[bs_cat_name]
            for category in categories:
                formattedCategory = category.category + " - " + category.subcategory
                if (formattedCategory == bonus_cat):
                    new_bonus.category = category
                    break

            for questionType in questionTypes:
                if (str(questionType) == bs_type):
                    new_bonus.question_type = questionType
                    break

            new_bonus.save_question(edit_type=QUESTION_CREATE, changer=user)
            new_bonuses.append(new_bonus)

        cache.clear()
        messages.success(request, 'Your questions have been uploaded.', extra_tags='alert-box success')        
        for tossup in new_tossups:
            messages.success(request, u'View your tossup on <a href="/edit_tossup/{0}">{1}.</a>'.format(tossup.id, get_answer_no_formatting(tossup.tossup_answer)), extra_tags='safe alert-box info')

        for bonus in new_bonuses:
            messages.success(request, u'View your bonus on <a href="/edit_bonus/{0}">{1}.</a>'.format(bonus.id, get_answer_no_formatting(bonus.part1_answer)), extra_tags='safe alert-box info')

        return HttpResponseRedirect('/edit_question_set/{0}'.format(qset_id))

    else:
        messages.error(request, 'Invalid request!')
        return render(request, 'failure.html')

@login_required
def settings(request):

    if request.method == 'GET':
        return render(request, 'settings.html', {})

    else:
        messages.error(request, 'Invalid request!')
        return render(request, 'failure.html', {})

@login_required
def profile(request):

    user = request.user
    writer = Writer.objects.get(user=user)

    if request.method == 'GET':
        initial_data = {'username': user.username,
                        'first_name': user.first_name,
                        'last_name': user.last_name,
                        'email': user.email,
                        'send_mail_on_comments': writer.send_mail_on_comments}

        form = WriterChangeForm(initial=initial_data)

    elif request.method == 'POST':

        print(request.POST)
        form = WriterChangeForm(request.POST)

        if form.is_valid():
            user.username = form.cleaned_data['username']
            user.first_name = form.cleaned_data['first_name']
            user.last_name = form.cleaned_data['last_name']
            user.email = form.cleaned_data['email']
            writer.send_mail_on_comments = form.cleaned_data['send_mail_on_comments']
            user.save()
            writer.save()

    return render(request, 'profile.html',
            {'form': form,
             'user': request.user.writer})

@login_required()
def search(request, passed_qset_id=None):

    user = request.user.writer

    passed_q_set = None
    if passed_qset_id is not None:
        passed_q_set = QuestionSet.objects.get(id=passed_qset_id)

    question_sets = QuestionSet.objects.filter(Q(writer=user) | Q(editor=user) | Q(owner=user)).distinct()

    if request.method == 'GET':
        all_categories = [(cat.category, cat.subcategory) for cat in DistributionEntry.objects.all()]
        categories = []
        for cat in all_categories:
            if cat not in categories:
                categories.append(cat)

        if request.GET.dict() == {}:

            q_set = passed_q_set

            return render(request, 'search/search.html',
                                      {'user': user,
                                       'categories': categories,
                                       'q_sets': question_sets,
                                       'selected_qset': q_set,
                                       'tossups_selected': 'checked',
                                       'bonuses_selected': 'checked',
                                       'search_all_selected' :'unchecked',
                                       'passed_q_set': passed_q_set})

        else:
            query = request.GET.get('q')
            search_models = request.GET.getlist('models')
            qset_id = int(request.GET.get('qset'))
            qset = QuestionSet.objects.get(id=qset_id)
            search_category = request.GET.get('category')
            tossups_selected = "unchecked"
            bonuses_selected = "unchecked"
            search_all_selected = "unchecked"
            if 'qsub.tossup' in search_models:
                tossups_selected = "checked"
            if 'qsub.bonus' in search_models:
                bonuses_selected = "checked"
            if 'qsub.search_all' in search_models:
                search_all_selected = "checked"

            if user in qset.writer.all() or user in qset.editor.all() or qset.is_owner(user):
                # Postgres full-text search (icontains on SQLite) over the
                # maintained search fields, scoped to the relevant set(s).
                if search_all_selected == 'checked':
                    set_ids = list(question_sets.values_list('id', flat=True))
                else:
                    set_ids = [qset.id]

                questions = []
                if 'qsub.tossup' in search_models:
                    tu_qs = Tossup.objects.filter(question_set_id__in=set_ids).select_related(*QUESTION_LIST_RELATED)
                    questions += list(fulltext_filter(tu_qs, query))
                if 'qsub.bonus' in search_models:
                    bs_qs = Bonus.objects.filter(question_set_id__in=set_ids).select_related(*QUESTION_LIST_RELATED)
                    questions += list(fulltext_filter(bs_qs, query))

                if search_category and search_category != 'All':
                    questions = [q for q in questions if str(q.category) == search_category]

                result = questions
                message = ''
                message_class = ''

            else:
                result = []
                message = 'You are not authorized to view questions from this set.'
                message_class = 'alert-box alert'

            return render(request, 'search/search.html',
                                      {'user': user,
                                       'categories': categories,
                                       'q_sets': question_sets,
                                       'result': result,
                                       'search_term': query,
                                       'search_category': search_category,
                                       'selected_qset': qset,
                                       'tossups_selected': tossups_selected,
                                       'bonuses_selected': bonuses_selected,
                                       'search_all_selected': search_all_selected,
                                       'passed_q_set': passed_q_set,
                                       'message': message,
                                       'message_class': message_class})

def _quick_search_scope(user, qset_param):
    """(scope_ids, selected_value) for quick search. `scope_ids` is the set of
    QuestionSet ids the user may search; `selected_value` is 'all' or the chosen
    id as a string. Only sets the user can access are ever included."""
    accessible = list(QuestionSet.objects.filter(
        Q(writer=user) | Q(editor=user) | Q(owner=user) | Q(co_owners=user)
    ).distinct().values_list('id', flat=True))
    if qset_param and qset_param != 'all':
        try:
            sid = int(qset_param)
        except (TypeError, ValueError):
            sid = None
        if sid in set(accessible):
            return [sid], str(sid)
    return accessible, 'all'


@login_required
def quick_search(request, passed_qset_id=None):
    """Fast type-ahead search over answer lines, filterable by category, packet,
    and writer. This view only renders the page shell and the filter facets for
    the chosen scope; results stream in from quick_search_results as the user
    types."""
    user = request.user.writer
    qset_param = request.GET.get('qset') or (str(passed_qset_id) if passed_qset_id else 'all')
    scope_ids, selected = _quick_search_scope(user, qset_param)
    multi = len(scope_ids) != 1  # spanning multiple sets: qualify facets by set

    accessible = QuestionSet.objects.filter(
        Q(writer=user) | Q(editor=user) | Q(owner=user) | Q(co_owners=user)
    ).distinct().order_by('-date', 'name')

    packets = (Packet.objects.filter(question_set_id__in=scope_ids)
               .select_related('question_set')
               .order_by('question_set__name', 'sort_order', 'packet_name'))
    packet_facets = [{
        'id': p.id,
        'name': '{0} — {1}'.format(p.question_set.name, p.packet_name) if multi else p.packet_name,
    } for p in packets]

    writers = (Writer.objects.filter(
        Q(tossup__question_set_id__in=scope_ids) | Q(bonus__question_set_id__in=scope_ids))
        .select_related('user').distinct())
    writer_facets = sorted(
        ({'id': w.id, 'name': w.get_real_name().strip() or w.user.username} for w in writers),
        key=lambda w: w['name'].lower())

    dist_ids = list(QuestionSet.objects.filter(id__in=scope_ids).values_list('distribution_id', flat=True))
    seen, category_facets = set(), []
    for cat, sub in DistributionEntry.objects.filter(
            distribution_id__in=dist_ids).values_list('category', 'subcategory'):
        label = '{0} - {1}'.format(cat, sub)
        if label not in seen:
            seen.add(label)
            category_facets.append(label)
    category_facets.sort()

    passed_q_set = QuestionSet.objects.filter(id__in=scope_ids).first() if len(scope_ids) == 1 else None

    return render(request, 'quick_search.html', {
        'user': user,
        'q_sets': accessible,
        'selected_qset': selected,
        'passed_q_set': passed_q_set,
        'packet_facets': packet_facets,
        'writer_facets': writer_facets,
        'category_facets': category_facets,
        'multi_set': multi,
    })


@login_required
def quick_search_results(request):
    """JSON answer-line search for the quick-search page. Matches the maintained
    (markup-stripped) search_question_answers field with optional category /
    packet / writer filters, and returns each hit's rendered content for inline
    preview plus an edit URL to open it in a new tab."""
    user = request.user.writer
    scope_ids, _ = _quick_search_scope(user, request.GET.get('qset', 'all'))
    if not scope_ids:
        return JsonResponse({'results': [], 'truncated': False})

    norm_q = strip_special_chars((request.GET.get('q') or '').strip())
    category = (request.GET.get('category') or '').strip()
    packet_id = request.GET.get('packet') or ''
    writer_id = request.GET.get('writer') or ''
    types = request.GET.getlist('types') or ['tossup', 'bonus']

    # Don't dump the whole scope when nothing is specified.
    if not (norm_q or category or packet_id or writer_id):
        return JsonResponse({'results': [], 'truncated': False, 'empty': True})

    cat_ids = None
    if category:
        dist_ids = list(QuestionSet.objects.filter(id__in=scope_ids).values_list('distribution_id', flat=True))
        cat_ids = [d.id for d in DistributionEntry.objects.filter(distribution_id__in=dist_ids)
                   if str(d) == category]
        if not cat_ids:
            return JsonResponse({'results': [], 'truncated': False})

    LIMIT = 50
    multi = len(scope_ids) != 1

    def scoped(qs):
        qs = qs.filter(question_set_id__in=scope_ids)
        if norm_q:
            qs = qs.filter(search_question_answers__icontains=norm_q)
        if cat_ids is not None:
            qs = qs.filter(category_id__in=cat_ids)
        if packet_id:
            qs = qs.filter(packet_id=packet_id)
        if writer_id:
            qs = qs.filter(author_id=writer_id)
        return (qs.select_related('packet', 'category', 'author', 'author__user', 'question_set')
                  .order_by('question_set__name', 'packet__sort_order', 'packet__packet_name', 'question_number'))

    rows = []
    if 'tossup' in types:
        rows += [('tossup', t) for t in scoped(Tossup.objects.all())[:LIMIT + 1]]
    if 'bonus' in types:
        rows += [('bonus', b) for b in scoped(Bonus.objects.all())[:LIMIT + 1]]
    truncated = len(rows) > LIMIT
    rows = rows[:LIMIT]

    def answer_preview(qtype, q):
        if qtype == 'tossup':
            return _grid_answer_preview(q.tossup_answer, 90)
        return ' / '.join(filter(None, [
            _grid_answer_preview(q.part1_answer, 30),
            _grid_answer_preview(q.part2_answer, 30),
            _grid_answer_preview(q.part3_answer, 30)]))

    data = []
    for qtype, q in rows:
        pkt = q.packet.packet_name if q.packet_id else '(unpacketized)'
        loc = '{0} #{1}'.format(pkt, q.question_number) if q.question_number else pkt
        if multi:
            loc = '{0} · {1}'.format(q.question_set.name, loc)
        data.append({
            'type': qtype,
            'id': q.id,
            'answer': answer_preview(qtype, q),
            'category': str(q.category) if q.category_id else '',
            'location': loc,
            'author': (q.author.get_real_name().strip() or q.author.user.username) if q.author_id else '',
            'edit_url': '/edit_{0}/{1}/'.format(qtype, q.id),
            'content': q.to_html(),
        })

    return JsonResponse({'results': data, 'truncated': truncated})


@login_required
def logout_view(request):
    logout(request)
    return HttpResponseRedirect("/main/")

def forgot_username(request):
    sent = False
    if request.method == 'POST':
        email = request.POST.get('email', '').strip()
        if email:
            from django.contrib.auth.models import User
            from django.core.mail import send_mail
            from django.conf import settings as django_settings
            users = User.objects.filter(email__iexact=email)
            if users.exists():
                usernames = ', '.join(u.username for u in users)
                send_mail(
                    'Your QEMS2 Username',
                    f'Your username is: {usernames}\n\nYou can sign in at {request.build_absolute_uri("/accounts/login/")}',
                    django_settings.DEFAULT_FROM_EMAIL,
                    [email],
                    fail_silently=False,
                )
        # Always show success to prevent email enumeration
        sent = True
    return render(request, 'account/forgot_username.html', {'sent': sent})

@login_required
def move_tossup(request, q_set_id, tossup_id):
    user = request.user.writer
    q_set = QuestionSet.objects.get(id=q_set_id)
    role = get_role_no_owner(user, q_set)
    
    tossup = Tossup.objects.get(id=tossup_id)
    if (tossup is None or tossup.question_set != q_set):
        message = 'Invalid tossup'
        message_class = 'alert-box alert'
        tossup = None

    move_sets = user.question_set_editor.exclude(id=q_set_id)

    if request.method == 'GET':
        if (role == "editor"):
            if (tossup is not None):
                form = MoveTossupForm(move_sets=move_sets)

                message = ''
                message_class = ''

                return render(request, 'move_tossup.html',
                                    {'user': user,
                                     'q_set': q_set,
                                     'form': form,
                                     'tossup': tossup,
                                     'message': message,
                                     'message_class': message_class})
            else:
                form = []
                return render(request, 'move_tossup.html',
                                    {'user': user,
                                     'q_set': q_set,
                                     'form': form,
                                     'tossup': tossup,
                                     'message': message,
                                     'message_class': message_class})
        else:
            form = []
            message = 'You do not have permissions to move this question.'
            message_class = 'alert-box alert'
            q_set = []
            return render(request, 'move_tossup.html',
                                {'user': user,
                                 'q_set': q_set,
                                 'tossup': None,
                                 'form': form,
                                 'message': message,
                                 'message_class': message_class})

    else:
        # Update the question set for this tossup
        if (role == 'editor'):
            form = MoveTossupForm(request.POST, move_sets=move_sets)

            if form.is_valid():
                dest_qset_id = request.POST["move_sets"]
                dest_qset = QuestionSet.objects.get(id=dest_qset_id)

                if (tossup is not None and dest_qset is not None):
                    tossup.question_set = dest_qset
                    tossup.packet = None

                    tossup.save()
                    cache.clear()
                    message = "Successfully moved tossup to " + str(dest_qset)
                    message_class = 'alert-box success'
                    return render(request, 'move_tossup_success.html',
                                        {'user': user,
                                         'q_set': q_set,
                                         'dest_q_set': dest_qset,
                                         'tossup': tossup,
                                         'message': message,
                                         'message_class': message_class})
                else:
                    message = 'There was an error with your submission.  Hit the back button and make sure you selected a valid question set to move to.'
                    message_class = 'alert-box warning'

                    return render(request, 'move_tossup.html',
                                        {'user': user,
                                         'q_set': q_set,
                                         'form': form,
                                         'tossup': tossup,
                                         'message': message,
                                         'message_class': message_class})
            else:
                message = 'There was an error moving your question.  Hit the back button and make sure you selected a valid question set to move to.'
                message_class = 'alert-box warning'
                q_set = []
                return render(request, 'move_tossup.html',
                                    {'user': user,
                                     'q_set': q_set,
                                     'tossup': None,
                                     'form': form,
                                     'message': message,
                                     'message_class': message_class})

        else:
            message = 'You do not have permissions to move this question.'
            message_class = 'alert-box alert'
            q_set = []
            form = []
            return render(request, 'move_tossup.html',
                                {'user': user,
                                 'q_set': q_set,
                                 'tossup': None,
                                 'form': form,
                                 'message': message,
                                 'message_class': message_class})

@login_required
def move_bonus(request, q_set_id, bonus_id):
    user = request.user.writer
    q_set = QuestionSet.objects.get(id=q_set_id)
    role = get_role_no_owner(user, q_set)

    bonus = Bonus.objects.get(id=bonus_id)
    if (bonus is None or bonus.question_set != q_set):
        message = 'Invalid bonus'
        message_class = 'alert-box alert'
        bonus = None

    move_sets = user.question_set_editor.exclude(id=q_set_id)

    if request.method == 'GET':
        if (role == 'editor'):
            if (bonus is not None):
                form = MoveBonusForm(move_sets=move_sets)

                message = ''
                message_class = ''

                return render(request, 'move_bonus.html',
                                    {'user': user,
                                     'q_set': q_set,
                                     'form': form,
                                     'bonus': bonus,
                                     'message': message,
                                     'message_class': message_class})
            else:
                form = []
                return render(request, 'move_bonus.html',
                                    {'user': user,
                                     'q_set': q_set,
                                     'form': form,
                                     'bonus': bonus,
                                     'message': message,
                                     'message_class': message_class})
        else:
            form = []
            message = 'You do not have permissions to move this question.'
            message_class = 'alert-box alert'
            q_set = []
            return render(request, 'move_bonus.html',
                                {'user': user,
                                 'q_set': q_set,
                                 'bonus': None,
                                 'form': form,
                                 'message': message,
                                 'message_class': message_class})

    else:
        # Update the question set for this bonus
        if (role == 'editor'):
            form = MoveBonusForm(request.POST, move_sets=move_sets)
            if form.is_valid():
                dest_qset_id = request.POST["move_sets"]
                dest_qset = QuestionSet.objects.get(id=dest_qset_id)

                if (bonus is not None and dest_qset is not None):
                    bonus.question_set = dest_qset
                    bonus.packet = None

                    bonus.save()
                    cache.clear()
                    return render(request, 'move_bonus_success.html',
                                        {'user': user,
                                         'q_set': q_set,
                                         'dest_q_set': dest_qset,
                                         'bonus': bonus})
                else:
                    message = 'There was an error with your submission.  Hit the back button and make sure you selected a valid question set to move to.'
                    message_class = 'alert-box warning'

                    return render(request, 'move_bonus.html',
                                        {'user': user,
                                         'q_set': q_set,
                                         'form': form,
                                         'bonus': bonus,
                                         'message': message,
                                         'message_class': message_class})
            else:
                message = 'There was an error moving your question.  Hit the back button and make sure you selected a valid question set to move to.'
                message_class = 'alert-box warning'
                q_set = []
                return render(request, 'move_bonus.html',
                                    {'user': user,
                                     'q_set': q_set,
                                     'bonus': None,
                                     'form': form,
                                     'message': message,
                                     'message_class': message_class})

        else:
            message = 'You do not have permissions to move this question.'
            message_class = 'alert-box alert'
            q_set = []
            form = []
            return render(request, 'move_bonus.html',
                                {'user': user,
                                 'q_set': q_set,
                                 'bonus': None,
                                 'form': form,
                                 'message': message,
                                 'message_class': message_class})

def add_qems_formatted_runs(paragraph, text, bold=False, is_answer=False):
    """Convert QEMS markup to python-docx runs on a paragraph.

    Markup rules (mirroring get_formatted_question_html in utils.py):
      _text_   → bold + underline (answer line)
      __text__ → underline only (prompt)
      ~text~   → italic
      (text)   → bold (pronunciation guide)
      (*)      → bold (power mark)
      \\s / \\S  → subscript / superscript (approximated with smaller font)
    """
    if text is None:
        return paragraph

    # Stored question text is HTML-escaped (e.g. apostrophes as &#x27; from the
    # YAPP import). Word gets raw characters, not a browser, so decode entities
    # first or they'd render literally as "&#x27;".
    text = html.unescape(text)

    allow_underlines = True
    allow_parens = True
    allow_powers = not is_answer  # powers only in question text

    italics_flag = False
    parens_flag = False
    underline_flag = False
    prompt_flag = False
    power_flag = False
    power_index = -1
    sub_flag = False
    super_flag = False
    bold_flag = False
    need_restore_italics = False

    if allow_powers:
        power_index = text.find("(*)")
        if power_index > -1:
            power_flag = True

    buf = ""
    # Current formatting state
    cur_bold = bold or power_flag
    cur_italic = False
    cur_underline = False
    cur_sub = False
    cur_super = False

    def flush(b=cur_bold, i=cur_italic, u=cur_underline, sub=cur_sub, sup=cur_super):
        nonlocal buf
        if buf:
            run = paragraph.add_run(buf)
            run.bold = b
            run.italic = i
            run.underline = u
            if sub:
                run.font.subscript = True
            if sup:
                run.font.superscript = True
            buf = ""

    def current_state():
        b = bold or power_flag or parens_flag or underline_flag or bold_flag
        i = italics_flag
        u = underline_flag or prompt_flag
        return b, i, u, sub_flag, super_flag

    index = 0
    prev = ""
    prev2 = ""

    while index < len(text):
        c = text[index]
        next_c = text[index + 1] if index < len(text) - 1 else ""

        new_bold, new_italic, new_underline, new_sub, new_sup = current_state()

        # Power mark
        if index == power_index and power_flag:
            flush(new_bold, new_italic, new_underline, new_sub, new_sup)
            run = paragraph.add_run("(*)")
            run.bold = True
            power_flag = False
            index += 3
            prev2, prev = prev, ")"
            continue

        # Tildes → italic toggle
        if c == "~":
            flush(new_bold, new_italic, new_underline, new_sub, new_sup)
            italics_flag = not italics_flag
            index += 1
            prev2, prev = prev, c
            continue

        # Open paren (pronunciation guide)
        if c == "(" and allow_parens and prev != "\\":
            flush(new_bold, new_italic, new_underline, new_sub, new_sup)
            if italics_flag:
                need_restore_italics = True
                italics_flag = False
            if not power_flag:
                parens_flag = True
            buf = "("
            index += 1
            prev2, prev = prev, c
            continue

        # Escaped open paren
        if c == "(" and allow_parens and prev == "\\" and prev2 != "\\":
            # Remove the backslash from buffer
            if buf.endswith("\\"):
                buf = buf[:-1]
            buf += c
            index += 1
            prev2, prev = prev, c
            continue

        # Close paren
        if c == ")" and allow_parens and prev != "\\":
            buf += ")"
            new_b, new_i, new_u, new_sub2, new_sup2 = current_state()
            flush(new_b, new_i, new_u, new_sub2, new_sup2)
            if not power_flag:
                parens_flag = False
            if need_restore_italics:
                italics_flag = True
                need_restore_italics = False
            index += 1
            prev2, prev = prev, c
            continue

        # Escaped close paren
        if c == ")" and allow_parens and prev == "\\":
            if buf.endswith("\\"):
                buf = buf[:-1]
            buf += c
            index += 1
            prev2, prev = prev, c
            continue

        # Subscript toggle: \s
        if c == "s" and prev == "\\" and prev2 != "\\" and not super_flag:
            if buf.endswith("\\"):
                buf = buf[:-1]
            flush(*current_state())
            sub_flag = not sub_flag
            index += 1
            prev2, prev = prev, c
            continue

        # Superscript toggle: \S
        if c == "S" and prev == "\\" and prev2 != "\\" and not sub_flag:
            if buf.endswith("\\"):
                buf = buf[:-1]
            flush(*current_state())
            super_flag = not super_flag
            index += 1
            prev2, prev = prev, c
            continue

        # Bold-only toggle: \B
        if c == "B" and prev == "\\" and prev2 != "\\":
            if buf.endswith("\\"):
                buf = buf[:-1]
            flush(*current_state())
            bold_flag = not bold_flag
            index += 1
            prev2, prev = prev, c
            continue

        # Underline markup
        if c == "_" and allow_underlines:
            if next_c == "_":
                # Double underscore → prompt (underline only)
                flush(*current_state())
                prompt_flag = not prompt_flag
                index += 2
                prev2, prev = "_", "_"
                continue
            else:
                # Single underscore → answer line (bold + underline)
                flush(*current_state())
                underline_flag = not underline_flag
                index += 1
                prev2, prev = prev, c
                continue

        # Regular character
        buf += c
        index += 1
        prev2, prev = prev, c

    # Flush remaining buffer
    flush(*current_state())
    return paragraph


@login_required
def export_question_set(request, qset_id, output_format):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    role = get_role_no_owner(user, qset)

    tossup_content_type_id = ContentType.objects.get_for_model(Tossup).id
    bonus_content_type_id = ContentType.objects.get_for_model(Bonus).id

    def safe_text(val):
        """Safely convert a value to string and remove newlines."""
        if val is None:
            return ""
        return remove_new_lines(str(val))

    def safe_name(writer_obj):
        """Safely get a writer's real name, handling None."""
        if writer_obj is None:
            return ""
        try:
            return writer_obj.get_real_name()
        except Exception:
            return str(writer_obj)

    def safe_category(cat_obj):
        """Safely convert a category to string."""
        if cat_obj is None:
            return ""
        return str(cat_obj)

    def safe_packet(packet_obj):
        """Safely convert a packet to string."""
        if packet_obj is None:
            return ""
        return str(packet_obj)

    def build_comment_string(content_type_id, object_id):
        """Build a comment string including threaded replies."""
        comment_list = Comment.objects.filter(
            content_type_id=content_type_id,
            object_pk=object_id,
            is_removed=False
        ).order_by('submit_date')

        # Get reply mappings
        reply_comment_ids = set()
        parent_map = {}
        for cr in CommentReply.objects.filter(comment__in=comment_list):
            reply_comment_ids.add(cr.comment_id)
            parent_map[cr.comment_id] = cr.parent_id

        # Build replies dict: parent_id -> [comments]
        replies = {}
        top_level = []
        for comment in comment_list:
            if comment.id in reply_comment_ids:
                parent_id = parent_map[comment.id]
                replies.setdefault(parent_id, []).append(comment)
            else:
                top_level.append(comment)

        parts = []
        for comment in top_level:
            parts.append(str(comment.user) + ": " + (comment.comment or ""))
            for reply in replies.get(comment.id, []):
                parts.append("  > " + str(reply.user) + ": " + (reply.comment or ""))

        return "||".join(parts)

    if request.method == 'GET':
        if (role == 'editor'):
            if (output_format in ("csv", "tsv")):
                tossups = Tossup.objects.filter(question_set=qset)
                bonuses = Bonus.objects.filter(question_set=qset)

                if output_format == "tsv":
                    response = HttpResponse(content_type='text/tab-separated-values')
                    response['Content-Disposition'] = 'attachment; filename="packet2.tsv"'
                    csv_writer = unicodecsv.writer(response, encoding='utf-8', delimiter='\t')
                else:
                    response = HttpResponse(content_type='text/csv')
                    response['Content-Disposition'] = 'attachment; filename="packet2.csv"'
                    csv_writer = unicodecsv.writer(response, encoding='utf-8', quoting=csv.QUOTE_ALL)

                csv_writer.writerow(["Tossup Question", "Answer", "Category", "Author", "Edited", "Packet", "Question Number", "Comments","Id", "Editor", "Proofreader", "Read Carefully"])
                for tossup in tossups:
                    comment_string = build_comment_string(tossup_content_type_id, tossup.id)

                    editor_name = ""
                    if tossup.edited and tossup.editor is not None:
                        editor_name = safe_name(tossup.editor)

                    proofreader_name = ""
                    if tossup.proofread and tossup.proofreader is not None:
                        proofreader_name = safe_name(tossup.proofreader)

                    csv_writer.writerow([safe_text(tossup.tossup_text), safe_text(tossup.tossup_answer), safe_category(tossup.category), safe_name(tossup.author), tossup.edited, safe_packet(tossup.packet), tossup.question_number, safe_text(comment_string), tossup.id, editor_name, proofreader_name, tossup.read_carefully])

                csv_writer.writerow([])

                csv_writer.writerow(["Bonus Leadin", "Bonus Part 1", "Bonus Answer 1", "Part 1 Difficulty", "Bonus Part 2", "Bonus Answer 2", "Part 2 Difficulty", "Bonus Part 3", "Bonus Answer 3", "Part 3 Difficulty", "Category", "Author", "Edited", "Packet", "Question Number", "Comments", "Id", "Editor", "Proofreader", "Read Carefully"])
                for bonus in bonuses:
                    comment_string = build_comment_string(bonus_content_type_id, bonus.id)

                    editor_name = ""
                    if bonus.edited and bonus.editor is not None:
                        editor_name = safe_name(bonus.editor)

                    proofreader_name = ""
                    if bonus.proofread and bonus.proofreader is not None:
                        proofreader_name = safe_name(bonus.proofreader)

                    csv_writer.writerow([safe_text(bonus.leadin), safe_text(bonus.part1_text), safe_text(bonus.part1_answer), bonus.part1_difficulty, safe_text(bonus.part2_text), safe_text(bonus.part2_answer), bonus.part2_difficulty, safe_text(bonus.part3_text), safe_text(bonus.part3_answer), bonus.part3_difficulty, safe_category(bonus.category), safe_name(bonus.author), bonus.edited, safe_packet(bonus.packet), bonus.question_number, safe_text(comment_string), bonus.id, editor_name, proofreader_name, bonus.read_carefully])

                csv_writer.writerow([])
                entries = qset.setwidedistributionentry_set.all()
                csv_writer.writerow(["Category", "Subcategory", "Total Tossups", "Total Bonuses"])
                for entry in entries:
                    csv_writer.writerow([entry.dist_entry.category, entry.dist_entry.subcategory, entry.num_tossups, entry.num_bonuses])

                return response
            elif output_format in ("docx", "docx-by-category", "docx-packetized"):
                tu_comment_ct = ContentType.objects.get_for_model(Tossup)
                bs_comment_ct = ContentType.objects.get_for_model(Bonus)

                def question_meta(q):
                    """Attribution line in the standard QEMS packet format:
                    ``<Author, Category - Subcategory> ~Id~ <Editor: Name>``."""
                    # get_real_name() pads with spaces and is blank when a writer
                    # has no name, so strip before deciding what to include.
                    author = html.unescape(safe_name(q.author)).strip()
                    cat = html.unescape(safe_category(q.category)).strip()
                    if author and cat:
                        head = '<{0}, {1}>'.format(author, cat)
                    elif author:
                        head = '<{0}>'.format(author)
                    elif cat:
                        head = '<{0}>'.format(cat)
                    else:
                        head = ''
                    meta = '{0} ~{1}~'.format(head, q.id).strip()
                    editor = html.unescape(safe_name(q.editor)).strip() if q.editor else ''
                    if q.edited and editor:
                        meta += ' <Editor: {0}>'.format(editor)
                    return meta

                def _initials(name):
                    parts = (name or '').split()
                    return ''.join(p[0] for p in parts[:2]).upper() or 'QC'

                def open_comment_threads(question, ct):
                    """Open (non-removed, unresolved) comments on a question, as
                    (author, text) pairs with any replies folded into the text.
                    Resolved threads and replies under them are skipped."""
                    comments = list(Comment.objects.filter(
                        content_type=ct, object_pk=str(question.id), is_removed=False
                    ).select_related('user').order_by('submit_date'))
                    if not comments:
                        return []
                    ids = [c.id for c in comments]
                    resolved = set(CommentResolution.objects.filter(
                        comment_id__in=ids, resolved=True).values_list('comment_id', flat=True))
                    parent_of = dict(CommentReply.objects.filter(
                        comment_id__in=ids).values_list('comment_id', 'parent_id'))

                    def author_of(c):
                        if c.user_id and c.user:
                            return c.user.get_username()
                        return c.user_name or 'Anonymous'

                    replies = defaultdict(list)
                    tops = []
                    for c in comments:
                        pid = parent_of.get(c.id)
                        (replies[pid].append(c) if pid is not None else tops.append(c))

                    out = []
                    for c in tops:
                        if c.id in resolved:
                            continue
                        text = strip_markup(c.comment or '').strip()
                        for r in replies.get(c.id, []):
                            text += '\n↳ {0}: {1}'.format(
                                author_of(r), strip_markup(r.comment or '').strip())
                        out.append((author_of(c), text))
                    return out

                def attach_open_comments(document, paragraph, question, ct):
                    """Add each open comment on the question as a Word comment
                    anchored to the question's paragraph."""
                    if not paragraph.runs:
                        return
                    for author, text in open_comment_threads(question, ct):
                        if not text:
                            continue
                        document.add_comment(paragraph.runs, text=text,
                                             author=author or 'QEMS', initials=_initials(author))

                def new_docx():
                    # Match the reference PACE packets: Times New Roman 12pt body,
                    # Word "Narrow" (0.5") margins, black bold centered headings.
                    document = Document()
                    for section in document.sections:
                        section.top_margin = Inches(0.5)
                        section.bottom_margin = Inches(0.5)
                        section.left_margin = Inches(0.5)
                        section.right_margin = Inches(0.5)

                    normal = document.styles['Normal']
                    normal.font.size = Pt(12)
                    normal.font.name = 'Times New Roman'
                    normal.paragraph_format.space_before = Pt(0)
                    normal.paragraph_format.space_after = Pt(0)

                    h1 = document.styles['Heading 1']  # packet/round title
                    h1.font.name = 'Times New Roman'
                    h1.font.size = Pt(16)
                    h1.font.bold = True
                    h1.font.color.rgb = RGBColor(0, 0, 0)
                    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    h1.paragraph_format.space_before = Pt(0)
                    h1.paragraph_format.space_after = Pt(10)

                    h2 = document.styles['Heading 2']  # "Tossups" / "Bonuses"
                    h2.font.name = 'Times New Roman'
                    h2.font.size = Pt(13)
                    h2.font.bold = True
                    h2.font.color.rgb = RGBColor(0, 0, 0)
                    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    h2.paragraph_format.space_before = Pt(12)
                    h2.paragraph_format.space_after = Pt(6)
                    return document

                def _line_break(paragraph):
                    paragraph.add_run().add_break()

                def add_tossup_to_doc(document, tossup, num):
                    # One paragraph per tossup (stem / answer / attribution on
                    # their own lines) so Word keeps the whole question together.
                    p = document.add_paragraph()
                    p.paragraph_format.keep_together = True
                    p.paragraph_format.space_after = Pt(10)
                    p.add_run(f"{num}. ").bold = True
                    add_qems_formatted_runs(p, safe_text(tossup.tossup_text))
                    _line_break(p)
                    p.add_run("ANSWER: ").bold = True
                    add_qems_formatted_runs(p, safe_text(tossup.tossup_answer), is_answer=True)
                    meta = question_meta(tossup)
                    if meta:
                        _line_break(p)
                        p.add_run(meta)
                    attach_open_comments(document, p, tossup, tu_comment_ct)

                def add_bonus_to_doc(document, bonus, num):
                    # One paragraph per bonus (leadin, each part + answer, then
                    # attribution) so the whole bonus stays together.
                    p = document.add_paragraph()
                    p.paragraph_format.keep_together = True
                    p.paragraph_format.space_after = Pt(10)
                    p.add_run(f"{num}. ").bold = True
                    add_qems_formatted_runs(p, safe_text(bonus.leadin))
                    for part_num in range(1, 4):
                        part_text = getattr(bonus, f'part{part_num}_text', None)
                        part_answer = getattr(bonus, f'part{part_num}_answer', None)
                        part_diff = getattr(bonus, f'part{part_num}_difficulty', '')
                        if part_text:
                            diff_tag = part_diff if part_diff else ''
                            _line_break(p)
                            p.add_run(f"[10{diff_tag}] ").bold = True
                            add_qems_formatted_runs(p, safe_text(part_text))
                            _line_break(p)
                            p.add_run("ANSWER: ").bold = True
                            add_qems_formatted_runs(p, safe_text(part_answer), is_answer=True)
                    meta = question_meta(bonus)
                    if meta:
                        _line_break(p)
                        p.add_run(meta)
                    attach_open_comments(document, p, bonus, bs_comment_ct)

                def add_careful_notes_to_doc(document, tossup_qs, bonus_qs):
                    """At the top of a packet, list the answer lines flagged
                    "read answer carefully" so the moderator is warned."""
                    flagged = []
                    for t in tossup_qs:
                        if getattr(t, 'read_carefully', False):
                            flagged.append(('Tossup', t.question_number, t.tossup_answer))
                    for b in bonus_qs:
                        if getattr(b, 'read_carefully', False):
                            ans = ' / '.join(filter(None, [b.part1_answer, b.part2_answer, b.part3_answer]))
                            flagged.append(('Bonus', b.question_number, ans))
                    if not flagged:
                        return
                    head = document.add_paragraph()
                    head.paragraph_format.space_after = Pt(2)
                    head.add_run('Moderator — read these answer lines carefully:').bold = True
                    for label, num, answer in flagged:
                        line = document.add_paragraph()
                        line.paragraph_format.space_after = Pt(0)
                        line.add_run('{0} {1}: '.format(label, num or '?')).bold = True
                        add_qems_formatted_runs(line, safe_text(answer), is_answer=True)
                    document.add_paragraph().paragraph_format.space_after = Pt(8)

                def save_docx_bytes(document):
                    buf = io.BytesIO()
                    document.save(buf)
                    return buf.getvalue()

                def write_all_questions_to_doc(document, tossup_qs, bonus_qs):
                    """Write tossups then bonuses to a document. Accepts a
                    queryset or a list."""
                    if tossup_qs:
                        document.add_heading('Tossups', level=2)
                        for i, tossup in enumerate(tossup_qs, 1):
                            num = tossup.question_number if tossup.question_number else i
                            add_tossup_to_doc(document, tossup, num)
                    if bonus_qs:
                        document.add_heading('Bonuses', level=2)
                        for i, bonus in enumerate(bonus_qs, 1):
                            num = bonus.question_number if bonus.question_number else i
                            add_bonus_to_doc(document, bonus, num)

                if output_format == "docx":
                    document = new_docx()

                    packets = Packet.objects.filter(question_set=qset).order_by('packet_name')
                    for pkt_i, packet in enumerate(packets):
                        if pkt_i > 0:
                            document.add_page_break()
                        document.add_heading('{0} {1}'.format(qset.name, packet.packet_name), level=1)
                        tossups = list(Tossup.objects.filter(
                            packet=packet, question_set=qset
                        ).order_by('question_number'))
                        bonuses = list(Bonus.objects.filter(
                            packet=packet, question_set=qset
                        ).order_by('question_number'))
                        add_careful_notes_to_doc(document, tossups, bonuses)
                        write_all_questions_to_doc(document, tossups, bonuses)

                    # Unpacketed questions
                    unpacketed_tossups = Tossup.objects.filter(
                        packet__isnull=True, question_set=qset
                    ).order_by('question_number')
                    unpacketed_bonuses = Bonus.objects.filter(
                        packet__isnull=True, question_set=qset
                    ).order_by('question_number')
                    if unpacketed_tossups.exists() or unpacketed_bonuses.exists():
                        document.add_heading('Unpacketed Questions', level=1)
                        write_all_questions_to_doc(document, unpacketed_tossups, unpacketed_bonuses)

                    response = HttpResponse(
                        save_docx_bytes(document),
                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                    filename = f"{qset.name}.docx" if qset.name else "questions.docx"
                    response['Content-Disposition'] = f'attachment; filename="{filename}"'
                    return response

                elif output_format == "docx-by-category":
                    # Collect top-level categories
                    all_tossups = Tossup.objects.filter(question_set=qset).select_related('category', 'author', 'editor', 'packet')
                    all_bonuses = Bonus.objects.filter(question_set=qset).select_related('category', 'author', 'editor', 'packet')

                    categories = set()
                    has_uncategorized = False
                    for t in all_tossups:
                        if t.category:
                            categories.add(t.category.category)
                        else:
                            has_uncategorized = True
                    for b in all_bonuses:
                        if b.category:
                            categories.add(b.category.category)
                        else:
                            has_uncategorized = True

                    zip_buf = io.BytesIO()
                    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
                        for cat_name in sorted(categories):
                            document = new_docx()
                            document.add_heading(cat_name, level=1)

                            cat_tossups = all_tossups.filter(
                                category__category=cat_name
                            ).order_by('packet__packet_name', 'question_number')
                            cat_bonuses = all_bonuses.filter(
                                category__category=cat_name
                            ).order_by('packet__packet_name', 'question_number')
                            write_all_questions_to_doc(document, cat_tossups, cat_bonuses)

                            zf.writestr(f"{cat_name}.docx", save_docx_bytes(document))

                        if has_uncategorized:
                            document = new_docx()
                            document.add_heading('Uncategorized', level=1)
                            uncat_tossups = all_tossups.filter(
                                category__isnull=True
                            ).order_by('packet__packet_name', 'question_number')
                            uncat_bonuses = all_bonuses.filter(
                                category__isnull=True
                            ).order_by('packet__packet_name', 'question_number')
                            write_all_questions_to_doc(document, uncat_tossups, uncat_bonuses)
                            zf.writestr("Uncategorized.docx", save_docx_bytes(document))

                    response = HttpResponse(
                        zip_buf.getvalue(),
                        content_type='application/zip'
                    )
                    filename = f"{qset.name} - By Category.zip" if qset.name else "questions-by-category.zip"
                    response['Content-Disposition'] = f'attachment; filename="{filename}"'
                    return response

                else:  # output_format == "docx-packetized"
                    # Export the set's actual packets in their packetized order
                    # (one .docx per packet, questions in question_number order),
                    # not a fresh re-deal. Packets are natural-sorted by name so
                    # "Packet 2" precedes "Packet 10".
                    def _packet_sort_key(pk):
                        nums = re.findall(r'\d+', pk.packet_name or '')
                        return (int(nums[0]) if nums else float('inf'),
                                pk.packet_name or '', pk.id)

                    packets = sorted(Packet.objects.filter(question_set=qset),
                                     key=_packet_sort_key)

                    packet_tus, packet_bos = [], []
                    for packet in packets:
                        packet_tus.append(list(
                            Tossup.objects.filter(packet=packet, question_set=qset)
                            .select_related('category', 'author', 'editor')
                            .order_by('question_number')))
                        packet_bos.append(list(
                            Bonus.objects.filter(packet=packet, question_set=qset)
                            .select_related('category', 'author', 'editor')
                            .order_by('question_number')))

                    # Build answer matrix workbook from the actual packets
                    wb = Workbook()

                    # Tossup answers sheet
                    ws_tu = wb.active
                    ws_tu.title = "Tossup Answers"
                    header_font = Font(bold=True)
                    wrap = Alignment(wrap_text=True, vertical='top')
                    max_tu = max((len(p) for p in packet_tus), default=0)
                    # Header row
                    ws_tu.cell(row=1, column=1, value="Packet").font = header_font
                    for col in range(1, max_tu + 1):
                        ws_tu.cell(row=1, column=col + 1, value=col).font = header_font
                    for pkt_idx, (packet, tus) in enumerate(zip(packets, packet_tus)):
                        row = pkt_idx + 2
                        ws_tu.cell(row=row, column=1, value=packet.packet_name).font = header_font
                        for q_idx, tossup in enumerate(tus):
                            answer = get_answer_no_formatting(
                                get_primary_answer(tossup.tossup_answer)
                            ).strip()
                            cell = ws_tu.cell(row=row, column=q_idx + 2, value=answer)
                            cell.alignment = wrap
                    # Auto-width columns
                    for col_idx in range(1, max_tu + 2):
                        ws_tu.column_dimensions[ws_tu.cell(row=1, column=col_idx).column_letter].width = 18

                    # Bonus answers sheet
                    ws_bo = wb.create_sheet("Bonus Answers")
                    max_bo = max((len(p) for p in packet_bos), default=0)
                    ws_bo.cell(row=1, column=1, value="Packet").font = header_font
                    for col in range(1, max_bo + 1):
                        ws_bo.cell(row=1, column=col + 1, value=col).font = header_font
                    for pkt_idx, (packet, bos) in enumerate(zip(packets, packet_bos)):
                        row = pkt_idx + 2
                        ws_bo.cell(row=row, column=1, value=packet.packet_name).font = header_font
                        for q_idx, bonus in enumerate(bos):
                            answers = []
                            for part_num in range(1, 4):
                                ans = getattr(bonus, f'part{part_num}_answer', None)
                                if ans:
                                    answers.append(get_answer_no_formatting(
                                        get_primary_answer(ans)
                                    ).strip())
                            cell = ws_bo.cell(row=row, column=q_idx + 2,
                                              value=" / ".join(answers))
                            cell.alignment = wrap
                    for col_idx in range(1, max_bo + 2):
                        ws_bo.column_dimensions[ws_bo.cell(row=1, column=col_idx).column_letter].width = 24

                    # Unpacketed questions (if any) go into a trailing document,
                    # kept out of the per-packet answer matrix.
                    unpacketed_tus = list(
                        Tossup.objects.filter(packet__isnull=True, question_set=qset)
                        .select_related('category', 'author', 'editor')
                        .order_by('question_number'))
                    unpacketed_bos = list(
                        Bonus.objects.filter(packet__isnull=True, question_set=qset)
                        .select_related('category', 'author', 'editor')
                        .order_by('question_number'))

                    def _safe_filename(name):
                        return re.sub(r'[\\/:*?"<>|]', '_', name or 'Packet').strip() or 'Packet'

                    # Package everything into a zip
                    zip_buf = io.BytesIO()
                    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
                        used_names = set()
                        for packet, tus, bos in zip(packets, packet_tus, packet_bos):
                            document = new_docx()
                            document.add_heading(
                                '{0} {1}'.format(qset.name, packet.packet_name), level=1)
                            add_careful_notes_to_doc(document, tus, bos)
                            if tus:
                                document.add_heading('Tossups', level=2)
                                for i, tossup in enumerate(tus, 1):
                                    add_tossup_to_doc(
                                        document, tossup, tossup.question_number or i)
                            if bos:
                                document.add_heading('Bonuses', level=2)
                                for i, bonus in enumerate(bos, 1):
                                    add_bonus_to_doc(
                                        document, bonus, bonus.question_number or i)
                            base = _safe_filename(packet.packet_name)
                            fname = base
                            n = 2
                            while fname in used_names:
                                fname = '{0} ({1})'.format(base, n)
                                n += 1
                            used_names.add(fname)
                            zf.writestr(f"{fname}.docx", save_docx_bytes(document))

                        if unpacketed_tus or unpacketed_bos:
                            document = new_docx()
                            document.add_heading(
                                '{0} Unpacketed'.format(qset.name), level=1)
                            if unpacketed_tus:
                                document.add_heading('Tossups', level=2)
                                for i, tossup in enumerate(unpacketed_tus, 1):
                                    add_tossup_to_doc(
                                        document, tossup, tossup.question_number or i)
                            if unpacketed_bos:
                                document.add_heading('Bonuses', level=2)
                                for i, bonus in enumerate(unpacketed_bos, 1):
                                    add_bonus_to_doc(
                                        document, bonus, bonus.question_number or i)
                            zf.writestr("Unpacketed.docx", save_docx_bytes(document))

                        # Add answer matrix
                        xlsx_buf = io.BytesIO()
                        wb.save(xlsx_buf)
                        zf.writestr("Answer Matrix.xlsx", xlsx_buf.getvalue())

                    response = HttpResponse(
                        zip_buf.getvalue(),
                        content_type='application/zip'
                    )
                    filename = f"{qset.name} - Packets.zip" if qset.name else "packets.zip"
                    response['Content-Disposition'] = f'attachment; filename="{filename}"'
                    return response
            elif (output_format == "pdf"):
                # TODO: Experiment with one of those PDF libraries
                message = 'Not supported yet.'
                message_class = 'alert-box alert'
                q_set = []
                tossups = []
                bonuses = []
                return render(request, 'export_question_set.html',
                                    {'user': user,
                                     'q_set': q_set,
                                     'tossups': tossups,
                                     'bonuses': bonuses,
                                     'message': message,
                                     'message_class': message_class})
            else:
                message = 'Unsupported export format.'
                message_class = 'alert-box alert'
                q_set = []
                tossups = []
                bonuses = []
                return render(request, 'export_question_set.html',
                                    {'user': user,
                                     'q_set': q_set,
                                     'tossups': tossups,
                                     'bonuses': bonuses,
                                     'message': message,
                                     'message_class': message_class})

        else:
            message = 'You are not authorized to export questions from this set.'
            message_class = 'alert-box alert'
            q_set = []
            tossups = []
            bonuses = []
            return render(request, 'export_question_set.html',
                                {'user': user,
                                 'q_set': q_set,
                                 'tossups': tossups,
                                 'bonuses': bonuses,
                                 'message': message,
                                 'message_class': message_class})

@login_required
def restore_tossup(request):
    user = request.user.writer

    message = ''
    message_class = ''
    read_only = True

    if request.method == 'POST':
        th_id = request.POST['th_id']
        tossup_history = TossupHistory.objects.get(id=th_id)
        tossup = Tossup.objects.get(question_history=tossup_history.question_history)
        if (tossup_history is None):
            message = 'Invalid tossup history restoration!'
            message_class = 'alert-box warning'
        else:
            qset_id = request.POST['qset_id']
            qset = QuestionSet.objects.get(id=qset_id)
            if user == tossup.author or qset.is_owner(user) or user in qset.editor.all():
                tossup = Tossup.objects.get(question_history=tossup_history.question_history)
                if (tossup is None):
                    message = 'Invalid tossup restoration!'
                    message_class = 'alert-box warning'
                else:
                    tossup.tossup_answer = tossup_history.tossup_answer
                    tossup.tossup_text = tossup_history.tossup_text
                    tossup.save_question(edit_type=QUESTION_RESTORE, changer=user)
                    cache.clear()
                    message = 'Successfully restored question'
                    message_class = 'alert-box success'
            else:
                message = 'You are not authorized to restore this question!'
                message_class = 'alert-box warning'

    return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))

@login_required
def restore_bonus(request):
    user = request.user.writer

    message = ''
    message_class = ''
    read_only = True

    if request.method == 'POST':
        bh_id = request.POST['bh_id']
        bonus_history = BonusHistory.objects.get(id=bh_id)
        bonus = Bonus.objects.get(question_history=bonus_history.question_history)
        if (bonus_history is None):
            message = 'Invalid bonus history restoration!'
            message_class = 'alert-box warning'
        else:
            qset_id = request.POST['qset_id']
            qset = QuestionSet.objects.get(id=qset_id)
            if user == bonus.author or qset.is_owner(user) or user in qset.editor.all():
                bonus = Bonus.objects.get(question_history=bonus_history.question_history)
                if (bonus is None):
                    message = 'Invalid bonus restoration!'
                    message_class = 'alert-box warning'
                else:
                    bonus.question_type = bonus_history.question_type
                    bonus.leadin = bonus_history.leadin
                    bonus.part1_text = bonus_history.part1_text
                    bonus.part1_answer = bonus_history.part1_answer
                    bonus.part2_text = bonus_history.part2_text
                    bonus.part2_answer = bonus_history.part2_answer
                    bonus.part3_text = bonus_history.part3_text
                    bonus.part3_answer = bonus_history.part3_answer
                    bonus.save_question(edit_type=QUESTION_RESTORE, changer=user)
                    cache.clear()
                    message = 'Successfully restored question'
                    message_class = 'alert-box success'
            else:
                message = 'You are not authorized to restore this question!'
                message_class = 'alert-box warning'

    return HttpResponse(json.dumps({'message': message, 'message_class': message_class}))

@login_required
def tossup_history(request, tossup_id):
    user = request.user.writer
    if request.method == 'GET':
        tossup = Tossup.objects.get(id=tossup_id)
        if (tossup is None):
            message = 'Invalid tossup'
            message_class = 'alert-box alert'
            tossup = None
        else:
            q_set = tossup.question_set
            if (q_set is None):
                message = 'Invalid question set'
                message_class = 'alert-box alert'
                tossup = None
            else:
                q_set_writers = Writer.objects.filter(Q(question_set_writer=q_set) | Q(question_set_editor=q_set)).distinct()
                if (user in q_set_writers):
                    tossup_histories, bonus_histories = tossup.get_question_history()
                    tossup_histories = tossup_histories.order_by('-id')
                    bonus_histories = bonus_histories.order_by('-id')
                    message = ''
                    message_class = ''

                    return render(request, 'tossup_history.html',
                                        {'user': user,
                                         'qset': q_set,
                                         'tossup': tossup,
                                         'tossup_histories': tossup_histories,
                                         'bonus_histories': bonus_histories,
                                         'highlight_version': request.GET.get('v', ''),
                                         'message': message,
                                         'message_class': message_class})

                else:
                    message = "You don't have permission to view this question"
                    message_class = 'alert-box alert'
                    tossup = None


    return render(request, 'tossup_history.html',
                        {'user': user,
                         'q_set': q_set,
                         'tossup': tossup,
                         'tossup_histories': [],
                         'bonus_histories': [],
                         'message': message,
                         'message_class': message_class})

@login_required
def bonus_history(request, bonus_id):
    user = request.user.writer
    if request.method == 'GET':
        bonus = Bonus.objects.get(id=bonus_id)
        if (bonus is None):
            message = 'Invalid bonus'
            message_class = 'alert-box alert'
            bonus = None
        else:
            q_set = bonus.question_set
            if (q_set is None):
                message = 'Invalid question set'
                message_class = 'alert-box alert'
                bonus = None
            else:
                q_set_writers = Writer.objects.filter(Q(question_set_writer=q_set) | Q(question_set_editor=q_set)).distinct()
                if (user in q_set_writers):
                    message = ''
                    message_class = ''

                    tossup_histories, bonus_histories = bonus.get_question_history()
                    tossup_histories = tossup_histories.order_by('-id')
                    bonus_histories = bonus_histories.order_by('-id')
                    return render(request, 'bonus_history.html',
                                        {'user': user,
                                         'qset': q_set,
                                         'bonus': bonus,
                                         'tossup_histories': tossup_histories,
                                         'bonus_histories': bonus_histories,
                                         'highlight_version': request.GET.get('v', ''),
                                         'message': message,
                                         'message_class': message_class})

                else:
                    message = "You don't have permission to view this question"
                    message_class = 'alert-box alert'
                    bonus = None


    return render(request, 'bonus_history.html',
                        {'user': user,
                         'q_set': q_set,
                         'bonus': bonus,
                         'tossup_histories': [],
                         'bonus_histories': [],
                         'message': message,
                         'message_class': message_class})

@login_required
def convert_tossup(request):
    user = request.user.writer

    message = ''
    message_class = ''
    redirect_url = None

    if request.method == 'POST':
        tossup_id = request.POST['tossup_id']
        tossup = Tossup.objects.get(id=tossup_id)
        if (tossup is None):
            message = 'Invalid tossup!'
            message_class = 'alert-box warning'
        else:
            qset_id = request.POST['qset_id']
            qset = QuestionSet.objects.get(id=qset_id)
            if user == tossup.author or qset.is_owner(user) or user in qset.editor.all():
                target_type = request.POST['target_type']
                if (target_type == ACF_STYLE_TOSSUP):
                    tossup_to_tossup(tossup, target_type)
                else:
                    result = tossup_to_bonus(tossup, target_type)
                    if result:
                        redirect_url = '/edit_bonus/{}/'.format(result.id)

                message = 'Successfully changed tossup type'
                message_class = 'alert-box success'
                cache.clear()
            else:
                message = 'You are not authorized to change this tossup type!'
                message_class = 'alert-box warning'

    return HttpResponse(json.dumps({'message': message, 'message_class': message_class, 'redirect_url': redirect_url}))

@login_required
def convert_bonus(request):
    user = request.user.writer

    message = ''
    message_class = ''
    redirect_url = None

    if request.method == 'POST':
        bonus_id = request.POST['bonus_id']
        bonus = Bonus.objects.get(id=bonus_id)
        if (bonus is None):
            message = 'Invalid bonus!'
            message_class = 'alert-box warning'
        else:
            qset_id = request.POST['qset_id']
            qset = QuestionSet.objects.get(id=qset_id)
            if user == bonus.author or qset.is_owner(user) or user in qset.editor.all():
                target_type = request.POST['target_type']
                if (target_type == ACF_STYLE_BONUS or target_type == VHSL_BONUS):
                    result = bonus_to_bonus(bonus, target_type)
                    if result:
                        redirect_url = '/edit_bonus/{}/'.format(result.id)
                else:
                    result = bonus_to_tossup(bonus, target_type)
                    if result:
                        redirect_url = '/edit_tossup/{}/'.format(result.id)

                message = 'Successfully changed bonus type'
                message_class = 'alert-box success'
                cache.clear()
            else:
                message = 'You are not authorized to change this bonus type!'
                message_class = 'alert-box warning'

    return HttpResponse(json.dumps({'message': message, 'message_class': message_class, 'redirect_url': redirect_url}))

@login_required
def questions_remaining(request, qset_id):
    message = ''

    qset = QuestionSet.objects.get(id=qset_id)
    user = request.user.writer
    set_status = {}

    total_tu_req = 0
    total_bs_req = 0
    total_tu_written = 0
    total_bs_written = 0
    tu_needed = 0
    bs_needed = 0

    role = get_role_no_owner(user, qset)

    if role == 'none':
        messages.error(request, 'You are not authorized to view information about this tournament!')
        return HttpResponseRedirect('/failure.html/')

    if request.method == 'GET':
        set_status, total_tu_req, total_bs_req, tu_needed, bs_needed, set_pct_complete = get_questions_remaining(qset)

    return render(request, 'questions_remaining.html',
                             {'user': user,
                              'set_status': set_status,
                              'set_pct_complete': '{0:0.2f}%'.format(set_pct_complete),
                              'set_pct_progress_bar': '{0:0.0f}%'.format(set_pct_complete),
                              'tu_needed': tu_needed,
                              'bs_needed': bs_needed,
                              'qset': qset,
                              'message': message})

@login_required
def category_overview(request, qset_id):
    qset = QuestionSet.objects.get(id=qset_id)
    user = request.user.writer

    role = get_role_no_owner(user, qset)

    if role == 'none':
        messages.error(request, 'You are not authorized to view information about this tournament!')
        return HttpResponseRedirect('/failure.html/')

    set_status, total_tu_req, total_bs_req, tu_needed, bs_needed, set_pct_complete = get_questions_remaining(qset)
    overview_rows = get_category_overview(qset)

    return render(request, 'category_overview.html',
                             {'user': user,
                              'overview_rows': overview_rows,
                              'set_pct_complete': '{0:0.2f}%'.format(set_pct_complete),
                              'set_pct_progress_bar': '{0:0.0f}%'.format(set_pct_complete),
                              'tu_needed': tu_needed,
                              'bs_needed': bs_needed,
                              'qset': qset})

@login_required
def bulk_change_set(request, qset_id):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    qset_editors = qset.editor.all()
    qset_writers = qset.writer.all()

    message = ''
    message_class = ''
    tossups = []
    bonuses = []
    role = get_role_no_owner(user, qset)

    if role != 'editor':
        message = 'You are not authorized to make bulk operations on this set'
        return HttpResponseRedirect('/failure.html/')
    else:
        tossups = Tossup.objects.filter(question_set=qset).order_by('-id')
        bonuses = Bonus.objects.filter(question_set=qset).order_by('-id')

    if request.method == 'GET':
        return render(request, 'bulk_change_set.html',
                                 {'user': user,
                                  'tossups': tossups,
                                  'bonuses': bonuses,
                                  'qset': qset,
                                  'message': message,
                                  'message_class': message_class})
    else:
        if ('confirm' in request.POST):
            operation = request.POST['change-type']
            if (operation == "author-step2"):
                return bulk_change_author(request, qset_id)
            elif (operation == "move-step2"):
                return bulk_move_question(request, qset_id)
            elif (operation == "packet-step2"):
                return bulk_change_packet(request, qset_id)

            num_questions_selected = 0
            num_tossups = int(request.POST['num-tossups'])
            num_bonuses = int(request.POST['num-bonuses'])

            change_tossups = []
            change_bonuses = []

            for tu_num in range(num_tossups):
                tu_checked_name = 'tossup-checked-{0}'.format(tu_num)
                tu_id_name = 'tossup-id-{0}'.format(tu_num)

                if (tu_checked_name in request.POST):
                    tu_id = request.POST[tu_id_name]
                    tossup = Tossup.objects.filter(id=tu_id, question_set=qset).first()
                    if tossup is None:
                        continue
                    change_tossups.append(tossup)
                    num_questions_selected += 1

            for bs_num in range(num_bonuses):
                bs_checked_name = 'bonus-checked-{0}'.format(bs_num)
                bs_id_name = 'bonus-id-{0}'.format(bs_num)

                if (bs_checked_name in request.POST):
                    bs_id = request.POST[bs_id_name]
                    bonus = Bonus.objects.filter(id=bs_id, question_set=qset).first()
                    if bonus is None:
                        continue
                    change_bonuses.append(bonus)
                    num_questions_selected += 1

            if (num_questions_selected > 0):
                # Do the actual operation

                if (operation == 'edit'):
                    bulk_edit_questions(True, change_tossups, change_bonuses, qset, user)

                    message = "Successfully edited questions."
                    message_class = 'alert-box success'
                    cache.clear()
                    return render(request, 'bulk_change_set.html',
                                             {'user': user,
                                              'tossups': tossups,
                                              'bonuses': bonuses,
                                              'qset': qset,
                                              'message': message,
                                              'message_class': message_class})
                elif (operation == 'unedit'):
                    bulk_edit_questions(False, change_tossups, change_bonuses, qset, user)

                    message = "Successfully unedited questions."
                    message_class = 'alert-box success'
                    cache.clear()
                    return render(request, 'bulk_change_set.html',
                                             {'user': user,
                                              'tossups': tossups,
                                              'bonuses': bonuses,
                                              'qset': qset,
                                              'message': message,
                                              'message_class': message_class})
                elif (operation == 'packet'):
                    packets = Packet.objects.filter(question_set=qset)
                    
                    cache.clear()
                    return render(request, 'bulk_change_packet.html',
                                             {'user': user,
                                              'tossups': change_tossups,
                                              'bonuses': change_bonuses,
                                              'qset': qset,
                                              'message': message,
                                              'message_class': message_class})
                elif (operation == 'lock'):
                    bulk_lock_questions(True, change_tossups, change_bonuses, qset, user)

                    message = "Successfully locked questions."
                    message_class = 'alert-box success'
                    cache.clear()
                    return render(request, 'bulk_change_set.html',
                                             {'user': user,
                                              'tossups': tossups,
                                              'bonuses': bonuses,
                                              'qset': qset,
                                              'message': message,
                                              'message_class': message_class})
                elif (operation == 'unlock'):
                    bulk_lock_questions(False, change_tossups, change_bonuses, qset, user)

                    message = "Successfully unlocked questions."
                    message_class = 'alert-box success'
                    cache.clear()
                    return render(request, 'bulk_change_set.html',
                                             {'user': user,
                                              'tossups': tossups,
                                              'bonuses': bonuses,
                                              'qset': qset,
                                              'message': message,
                                              'message_class': message_class})
                elif (operation == 'delete'):
                    bulk_delete_questions(change_tossups, change_bonuses, qset, user)
                    message = "Successfully deleted questions."
                    message_class = 'alert-box success'
                    cache.clear()
                    tossups = Tossup.objects.filter(question_set=qset).order_by('-id')
                    bonuses = Bonus.objects.filter(question_set=qset).order_by('-id')

                    return render(request, 'bulk_change_set.html',
                                             {'user': user,
                                              'tossups': tossups,
                                              'bonuses': bonuses,
                                              'qset': qset,
                                              'message': message,
                                              'message_class': message_class})

                elif (operation == 'convert-to-acf-style-tossup'):
                    bulk_convert_to_acf_style_tossup(change_tossups, change_bonuses, qset, user)
                    message = "Successfully converted question type to ACF-style tossups."
                    message_class = 'alert-box success'
                    cache.clear()
                    tossups = Tossup.objects.filter(question_set=qset).order_by('-id')
                    bonuses = Bonus.objects.filter(question_set=qset).order_by('-id')

                    return render(request, 'bulk_change_set.html',
                                             {'user': user,
                                              'tossups': tossups,
                                              'bonuses': bonuses,
                                              'qset': qset,
                                              'message': message,
                                              'message_class': message_class})

                elif (operation == 'convert-to-acf-style-bonus'):
                    bulk_convert_to_acf_style_bonus(change_tossups, change_bonuses, qset, user)
                    message = "Successfully converted question type to ACF-style bonuses."
                    message_class = 'alert-box success'
                    cache.clear()
                    tossups = Tossup.objects.filter(question_set=qset).order_by('-id')
                    bonuses = Bonus.objects.filter(question_set=qset).order_by('-id')

                    return render(request, 'bulk_change_set.html',
                                             {'user': user,
                                              'tossups': tossups,
                                              'bonuses': bonuses,
                                              'qset': qset,
                                              'message': message,
                                              'message_class': message_class})
                elif (operation == 'convert-to-vhsl-bonus'):
                    bulk_convert_to_vhsl_bonus(change_tossups, change_bonuses, qset, user)
                    message = "Successfully converted question type to VHSL bonuses."
                    message_class = 'alert-box success'
                    cache.clear()

                    tossups = Tossup.objects.filter(question_set=qset).order_by('-id')
                    bonuses = Bonus.objects.filter(question_set=qset).order_by('-id')

                    return render(request, 'bulk_change_set.html',
                                             {'user': user,
                                              'tossups': tossups,
                                              'bonuses': bonuses,
                                              'qset': qset,
                                              'message': message,
                                              'message_class': message_class})
                elif (operation == 'move'):
                    new_sets = user.question_set_editor.exclude(id=qset_id)
                    cache.clear()
                    return render(request, 'bulk_move_questions.html',
                                             {'user': user,
                                              'tossups': change_tossups,
                                              'bonuses': change_bonuses,
                                              'qset': qset,
                                              'new_sets': new_sets,
                                              'message': message,
                                              'message_class': message_class})
                elif (operation == 'author'):
                    writers = Writer.objects.filter(Q(question_set_writer=qset) | Q(question_set_editor=qset)).distinct()
                    cache.clear()

                    return render(request, 'bulk_change_author.html',
                                             {'user': user,
                                              'tossups': change_tossups,
                                              'bonuses': change_bonuses,
                                              'qset': qset,
                                              'writers': writers,
                                              'message': message,
                                              'message_class': message_class})

            else:
                message = "Error!  You must select at least one question."
                message_class = 'alert-box warning'
                return render(request, 'bulk_change_set.html',
                                         {'user': user,
                                          'tossups': tossups,
                                          'bonuses': bonuses,
                                          'qset': qset,
                                          'message': message,
                                          'message_class': message_class})
        else:
            message = "You didn't hit the confirm button."
            message_class = 'alert-box warning'
            return render(request, 'bulk_change_set.html',
                                     {'user': user,
                                      'tossups': tossups,
                                      'bonuses': bonuses,
                                      'qset': qset,
                                      'message': message,
                                      'message_class': message_class})

@login_required
def bulk_change_author(request, qset_id):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)

    message = ''
    message_class = ''
    read_only = True

    role = get_role_no_owner(user, qset)

    if role != 'editor':
        message = 'You are not authorized to make bulk operations on this set'
        return HttpResponseRedirect('/failure.html/')

    if request.method == 'POST':
        num_tossups = int(request.POST['num-tossups'])
        num_bonuses = int(request.POST['num-bonuses'])
        new_author_id = request.POST['new-author']
        new_author = Writer.objects.get(id=new_author_id)

        new_author_role = get_role_no_owner(new_author, qset)
        if (new_author_role == 'none'):
            message = 'Could not change author to ' + str(new_author)
            return HttpResponseRedirect('/failure.html/')

        for tu_num in range(num_tossups):
            tu_id_name = 'tossup-id-{0}'.format(tu_num)
            tu_id = request.POST[tu_id_name]
            tossup = Tossup.objects.filter(id=tu_id, question_set=qset).first()
            if tossup is None:
                continue
            tossup.author = new_author
            tossup.save()

        for bs_num in range(num_bonuses):
            bs_id_name = 'bonus-id-{0}'.format(bs_num)
            bs_id = request.POST[bs_id_name]
            bonus = Bonus.objects.filter(id=bs_id, question_set=qset).first()
            if bonus is None:
                continue
            bonus.author = new_author
            bonus.save()

        message = 'Successfully changed author'
        message_class = 'alert-box success'
        cache.clear()

        tossups = Tossup.objects.filter(question_set=qset).order_by('-id')
        bonuses = Bonus.objects.filter(question_set=qset).order_by('-id')

        return render(request, 'bulk_change_set.html',
                                 {'user': user,
                                  'tossups': tossups,
                                  'bonuses': bonuses,
                                  'qset': qset,
                                  'message': message,
                                  'message_class': message_class})

@login_required
def bulk_change_packet(request, qset_id):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)

    message = ''
    message_class = ''
    read_only = True

    role = get_role_no_owner(user, qset)

    if role != 'editor':
        message = 'You are not authorized to make bulk operations on this set'
        return HttpResponseRedirect('/failure.html/')

    if request.method == 'POST':
        num_tossups = int(request.POST['num-tossups'])
        num_bonuses = int(request.POST['num-bonuses'])
        new_packet_id = request.POST['new-packet']
        # The target packet must belong to this set
        new_packet = Packet.objects.filter(id=new_packet_id, question_set=qset).first()
        if new_packet is None:
            message = 'Could not change packet'
            return HttpResponseRedirect('/failure.html/')

        # TODO: We may want to clear the numbers from these questions in the future
        for tu_num in range(num_tossups):
            tu_id_name = 'tossup-id-{0}'.format(tu_num)
            tu_id = request.POST[tu_id_name]
            tossup = Tossup.objects.filter(id=tu_id, question_set=qset).first()
            if tossup is None:
                continue
            tossup.packet = new_packet
            tossup.save()

        for bs_num in range(num_bonuses):
            bs_id_name = 'bonus-id-{0}'.format(bs_num)
            bs_id = request.POST[bs_id_name]
            bonus = Bonus.objects.filter(id=bs_id, question_set=qset).first()
            if bonus is None:
                continue
            bonus.packet = new_packet
            bonus.save()

        message = 'Successfully changed packet'
        message_class = 'alert-box success'
        cache.clear()

        tossups = Tossup.objects.filter(question_set=qset).order_by('-id')
        bonuses = Bonus.objects.filter(question_set=qset).order_by('-id')

        return render(request, 'bulk_change_set.html',
                                 {'user': user,
                                  'tossups': tossups,
                                  'bonuses': bonuses,
                                  'qset': qset,
                                  'message': message,
                                  'message_class': message_class})

@login_required
def bulk_move_question(request, qset_id):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)

    message = ''
    message_class = ''
    read_only = True

    role = get_role_no_owner(user, qset)

    if role != 'editor':
        message = 'You are not authorized to make bulk operations on this set'
        return HttpResponseRedirect('/failure.html/')

    if request.method == 'POST':
        num_tossups = int(request.POST['num-tossups'])
        num_bonuses = int(request.POST['num-bonuses'])
        new_set_id = request.POST['new-set']
        new_set = QuestionSet.objects.get(id=new_set_id)

        new_set_role = get_role_no_owner(user, new_set)
        if (new_set_role != 'editor'):
            message = 'Could not move questions to ' + str(new_set)
            return HttpResponseRedirect('/failure.html/')

        for tu_num in range(num_tossups):
            tu_id_name = 'tossup-id-{0}'.format(tu_num)
            tu_id = request.POST[tu_id_name]
            tossup = Tossup.objects.filter(id=tu_id, question_set=qset).first()
            if tossup is None:
                continue

            tossup.question_set = new_set
            tossup.packet = None

            # It's not guaranteed that these categories exist, so clear them
            tossup.category = None
            tossup.subtype = ''

            tossup.save()

        for bs_num in range(num_bonuses):
            bs_id_name = 'bonus-id-{0}'.format(bs_num)
            bs_id = request.POST[bs_id_name]
            bonus = Bonus.objects.filter(id=bs_id, question_set=qset).first()
            if bonus is None:
                continue

            bonus.question_set = new_set
            bonus.packet = None

            # It's not guaranteed that these categories exist, so clear them
            bonus.category = None
            bonus.subtype = ''

            bonus.save()

        message = 'Successfully moved questions'
        message_class = 'alert-box success'
        cache.clear()

        tossups = Tossup.objects.filter(question_set=qset).order_by('-id')
        bonuses = Bonus.objects.filter(question_set=qset).order_by('-id')

        return render(request, 'bulk_change_set.html',
                                 {'user': user,
                                  'tossups': tossups,
                                  'bonuses': bonuses,
                                  'qset': qset,
                                  'message': message,
                                  'message_class': message_class})

@login_required
def writer_question_set_settings(request, qset_id):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)

    message = ''
    message_class = ''

    role = get_role_no_owner(user, qset)
    if (role == 'none'):
        return render(request, 'failure.html',
            {'message': 'You do not have permissions to this set',
             'message_class': 'alert-box alert'})
    
    # Create the settings if it doesn't exist
    settings = None
    try:
        settings = WriterQuestionSetSettings.objects.get(question_set=qset, writer=user)        
    except:
        settings = WriterQuestionSetSettings(writer=user, question_set=qset)
        settings.save()
        settings.create_per_category_writer_settings()
        
    if request.method == 'POST':
        form = WriterQuestionSetSettingsForm(request.POST)

        PerCategoryWriterSettingsFormset = formset_factory(PerCategoryWriterSettingsForm, can_delete=False, extra=0)
        formset = PerCategoryWriterSettingsFormset(data=request.POST)

        if (form.is_valid() and formset.is_valid()):
            settings.email_on_all_new_comments = form.cleaned_data['email_on_all_new_comments']
            settings.email_on_all_new_questions = form.cleaned_data['email_on_all_new_questions']
            settings.save()
            
            for per_category_form in formset.forms:
                entry_id = int(per_category_form.cleaned_data['entry_id'])
                email_on_new_questions = bool(per_category_form.cleaned_data['email_on_new_questions'])
                email_on_new_comments = bool(per_category_form.cleaned_data['email_on_new_comments'])

                entry = PerCategoryWriterSettings.objects.get(id=entry_id)
                entry.email_on_new_questions = email_on_new_questions
                entry.email_on_new_comments = email_on_new_comments
                entry.save()

            message = 'Your settings have been updated.'
            message_class = 'alert-box success'

            return render(request, 'writer_question_set_settings.html',
                     {'form': form,
                     'formset': formset,
                     'message': message,
                     'message_class': message_class,
                     'user': user,
                     'qset': qset})
            
        else:
            message = 'There was an error saving your settings.'
            message_class = 'alert-box warning'
            return render(request, 'writer_question_set_settings.html',
                     {'form': form,
                     'formset': formset,
                     'message': message,
                     'message_class': message_class,
                     'user': user,
                     'qset': qset})
        
    elif request.method == 'GET':
        entries = settings.percategorywritersettings_set.all()
        initial_data = []
        for entry in entries:
            initial_data.append({
                'entry_id': entry.id,
                'distribution_entry_string': str(entry.distribution_entry),
                'email_on_new_questions': entry.email_on_new_questions,
                'email_on_new_comments': entry.email_on_new_comments})
                
        form = WriterQuestionSetSettingsForm(instance=settings)
        PerCategoryWriterSettingsFormset = formset_factory(PerCategoryWriterSettingsForm, can_delete=False, extra=0)
        formset = PerCategoryWriterSettingsFormset(initial=initial_data)
                
        return render(request, 'writer_question_set_settings.html',
                                 {'form': form,
                                  'formset': formset,
                                  'message': message,
                                  'message_class': message_class,
                                  'user': user,
                                  'qset': qset})

@login_required
def contributor(request, qset_id, writer_id):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    qset_editors = qset.editor.all()
    qset_writers = qset.writer.all()

    writer = Writer.objects.get(id=writer_id)
    
    tossups = []
    bonuses = []

    
    if (writer not in qset_editors and not qset.is_owner(writer) and writer not in qset.writer.all()):
        return render(request, 'failure.html',
            {'message': 'The specified contributor is not in this set',
             'message_class': 'alert-box alert'})
        
    if user not in qset_editors and not qset.is_owner(user) and user not in qset.writer.all():
        return render(request, 'failure.html',
            {'message': 'You are not authorized to view this set',
             'message_class': 'alert-box alert'})

    tossups = Tossup.objects.filter(question_set=qset).filter(author=writer)
    bonuses = Bonus.objects.filter(question_set=qset).filter(author=writer)

    writer_status =   {'tossups_written': tossups.count(),
                         'bonuses_written': bonuses.count()
                         }
            
            
    return render(request, 'contributor.html',
        {
        'user': user,
        'tossups': tossups,
        'bonuses': bonuses,
        'writer_status': writer_status,
        'qset': qset,
        'writer': writer})	


#########################################################################
# Auto-packetization views
#########################################################################

from .packetizer import auto_packetize, build_quota_dict, get_path_parts
from django.db import transaction
from decimal import Decimal, InvalidOperation

def get_packetization_rows(qset):
    """Category tree rows for the packetization page.  Mirrors the category
    overview tree but adds per-packet recommended values derived from the
    set-wide distribution and any previously saved quota entries."""
    num_packets = max(qset.num_packets, 1)
    entries = qset.setwidedistributionentry_set.all().order_by('dist_entry__category', 'dist_entry__subcategory')

    tree = {}
    for entry in entries:
        parts = list(get_path_parts(entry.dist_entry))
        if not parts:
            continue
        leaf_key = tuple(parts)
        tu_total = entry.num_tossups or 0
        bs_total = entry.num_bonuses or 0
        tu_written = qset.tossup_set.filter(category=entry.dist_entry).count()
        bs_written = qset.bonus_set.filter(category=entry.dist_entry).count()

        for i in range(1, len(parts) + 1):
            prefix = tuple(parts[:i])
            if prefix not in tree:
                tree[prefix] = {'tu_total': 0, 'bs_total': 0, 'tu_written': 0, 'bs_written': 0, 'is_leaf': False}
            tree[prefix]['tu_total'] += tu_total
            tree[prefix]['bs_total'] += bs_total
            tree[prefix]['tu_written'] += tu_written
            tree[prefix]['bs_written'] += bs_written
        tree[leaf_key]['is_leaf'] = True

    saved = {e.path: e for e in PacketizationEntry.objects.filter(question_set=qset)}

    rows = []
    for key in sorted(tree.keys()):
        node = tree[key]
        path = ' - '.join(key)
        entry = saved.get(path)
        rec_tu = round(node['tu_total'] / float(num_packets), 1)
        rec_bs = round(node['bs_total'] / float(num_packets), 1)
        is_top = len(key) == 1
        rows.append({
            'path': path,
            'short_name': key[-1],
            'depth': len(key) - 1,
            'padding': (len(key) - 1) * 30,
            'is_top': is_top,
            'is_leaf': node['is_leaf'],
            'tu_total': node['tu_total'],
            'bs_total': node['bs_total'],
            'tu_written': node['tu_written'],
            'bs_written': node['bs_written'],
            'rec_tu': rec_tu,
            'rec_bs': rec_bs,
            'min_tu': entry.min_tossups if entry else (rec_tu if is_top else None),
            'max_tu': entry.max_tossups if entry else (rec_tu if is_top else None),
            'min_bs': entry.min_bonuses if entry else (rec_bs if is_top else None),
            'max_bs': entry.max_bonuses if entry else (rec_bs if is_top else None),
        })
    return rows

def _parse_quota_value(raw):
    raw = (raw or '').strip()
    if raw == '':
        return None
    try:
        value = Decimal(raw)
    except InvalidOperation:
        raise ValueError('"{0}" is not a number'.format(raw))
    if value < 0:
        raise ValueError('Quota values cannot be negative')
    return value

@login_required
def packetize_set(request, qset_id):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    message = ''
    message_class = ''
    report = None

    if not qset.is_owner(user):
        return render(request, 'failure.html',
                                 {'message': 'Only the set owner can packetize a set!',
                                  'message_class': 'alert-box alert'})

    if request.method == 'POST':
        try:
            num_packets = int(request.POST.get('num_packets', qset.num_packets))
            tossups_per_packet = int(request.POST.get('tossups_per_packet', qset.tossups_per_packet))
            bonuses_per_packet = int(request.POST.get('bonuses_per_packet', qset.bonuses_per_packet))
            if num_packets < 1 or tossups_per_packet < 1 or bonuses_per_packet < 1:
                raise ValueError('Packet counts must be positive')

            row_count = int(request.POST.get('row_count', 0))
            new_entries = []
            for i in range(row_count):
                path = request.POST.get('row_{0}_path'.format(i), '').strip()
                if not path:
                    continue
                depth = path.count(' - ')
                min_tu = _parse_quota_value(request.POST.get('row_{0}_min_tu'.format(i)))
                max_tu = _parse_quota_value(request.POST.get('row_{0}_max_tu'.format(i)))
                min_bs = _parse_quota_value(request.POST.get('row_{0}_min_bs'.format(i)))
                max_bs = _parse_quota_value(request.POST.get('row_{0}_max_bs'.format(i)))
                if depth == 0 and (min_tu is None or max_tu is None or min_bs is None or max_bs is None):
                    raise ValueError('Top-level category "{0}" needs minimum and maximum tossups and bonuses'.format(path))
                if (min_tu is not None and max_tu is not None and min_tu > max_tu) or \
                   (min_bs is not None and max_bs is not None and min_bs > max_bs):
                    raise ValueError('Minimum exceeds maximum for "{0}"'.format(path))
                if min_tu is None and max_tu is None and min_bs is None and max_bs is None:
                    continue
                new_entries.append(PacketizationEntry(
                    question_set=qset, path=path, depth=depth,
                    min_tossups=min_tu, max_tossups=max_tu,
                    min_bonuses=min_bs, max_bonuses=max_bs))

            with transaction.atomic():
                qset.num_packets = num_packets
                qset.tossups_per_packet = tossups_per_packet
                qset.bonuses_per_packet = bonuses_per_packet
                qset.save()

                PacketizationEntry.objects.filter(question_set=qset).delete()
                PacketizationEntry.objects.bulk_create(new_entries)

                quotas = build_quota_dict(qset)
                report = auto_packetize(qset, num_packets, tossups_per_packet,
                                        bonuses_per_packet, quotas, created_by=user)
            cache.clear()
            message = 'The set has been packetized into {0} packet(s).'.format(num_packets)
            message_class = 'alert-box success'
        except ValueError as ex:
            message = str(ex)
            message_class = 'alert-box warning'

    rows = get_packetization_rows(qset)
    top_min_tu = sum(float(r['min_tu']) for r in rows if r['is_top'] and r['min_tu'] is not None)
    top_min_bs = sum(float(r['min_bs']) for r in rows if r['is_top'] and r['min_bs'] is not None)

    return render(request, 'packetize_set.html',
                             {'qset': qset,
                              'user': user,
                              'rows': rows,
                              'row_count': len(rows),
                              'top_min_tu': round(top_min_tu, 1),
                              'top_min_bs': round(top_min_bs, 1),
                              'report': report,
                              'message': message,
                              'message_class': message_class})

def _grid_answer_preview(text, limit=45):
    # Decode HTML entities (imported answers store apostrophes as &#x27; etc.);
    # the template re-escapes safely, so the grid shows real characters.
    answer = html.unescape(get_answer_no_formatting(get_primary_answer(text or ''))).strip()
    if len(answer) > limit:
        answer = answer[:limit].rstrip() + '...'
    return answer

@login_required
def packet_grid(request, qset_id):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)

    if not qset.is_owner(user) and user not in qset.editor.all() and user not in qset.writer.all():
        return render(request, 'failure.html',
                                 {'message': 'You are not authorized to view this set!',
                                  'message_class': 'alert-box alert'})

    read_only = not (qset.is_owner(user) or user in qset.editor.all())
    # Natural order (Round 2 before Round 10), extras last, honoring any
    # user-set custom order (Packet.sort_order).
    packets = sorted_packets(qset)

    def build_rows(question_model, preview_func, edit_url):
        qtype = 'tossup' if question_model is Tossup else 'bonus'
        vacancies = {(v.packet_id, v.question_number): v.category
                     for v in PacketSlotVacancy.objects.filter(question_set=qset, question_type=qtype)}
        cells_by_packet = {}
        max_num = 0
        for question in question_model.objects.filter(question_set=qset, packet__in=packets).select_related('category', 'packet'):
            number = question.question_number or 0
            if number <= 0:
                continue
            max_num = max(max_num, number)
            cells_by_packet.setdefault(question.packet_id, {})[number] = {
                'id': question.id,
                'answer': preview_func(question),
                'category': html.unescape(str(question.category)) if question.category else '',
                'edit_url': '{0}{1}/'.format(edit_url, question.id),
                'edited': question.edited,
            }
        rows = []
        for number in range(1, max_num + 1):
            cells = []
            for p in packets:
                cell = cells_by_packet.get(p.id, {}).get(number)
                # Empty slot: carry the packet id so the grid can offer an
                # "add a question here" link (handy for tiebreaker rows), plus
                # the category of whatever was last removed from this slot.
                cells.append(cell if cell is not None else
                             {'empty': True, 'packet_id': p.id,
                              'removed_category': vacancies.get((p.id, number), '')})
            rows.append({'num': number, 'cells': cells})
        return rows

    tossup_rows = build_rows(
        Tossup, lambda t: _grid_answer_preview(t.tossup_answer), '/edit_tossup/')
    bonus_rows = build_rows(
        Bonus, lambda b: ' / '.join(filter(None, [
            _grid_answer_preview(b.part1_answer, 20),
            _grid_answer_preview(b.part2_answer, 20),
            _grid_answer_preview(b.part3_answer, 20)])), '/edit_bonus/')

    def build_unpacketized(question_model, preview_func):
        items = []
        for question in (question_model.objects.filter(question_set=qset, packet=None)
                         .select_related('category').order_by('id')):
            items.append({
                'id': question.id,
                'answer': preview_func(question),
                'category': html.unescape(str(question.category)) if question.category else '',
                'edit_url': '/edit_{0}/{1}/'.format(
                    'tossup' if question_model is Tossup else 'bonus', question.id),
            })
        return items

    unpacketized_tu = build_unpacketized(
        Tossup, lambda t: _grid_answer_preview(t.tossup_answer))
    unpacketized_bs = build_unpacketized(
        Bonus, lambda b: ' / '.join(filter(None, [
            _grid_answer_preview(b.part1_answer, 20),
            _grid_answer_preview(b.part2_answer, 20),
            _grid_answer_preview(b.part3_answer, 20)])))

    return render(request, 'packet_grid.html',
                             {'qset': qset,
                              'user': user,
                              'packets': packets,
                              'tossup_rows': tossup_rows,
                              'bonus_rows': bonus_rows,
                              'tossups_per_packet': qset.tossups_per_packet,
                              'bonuses_per_packet': qset.bonuses_per_packet,
                              'unassigned_tu': len(unpacketized_tu),
                              'unassigned_bs': len(unpacketized_bs),
                              'unpacketized_tu': unpacketized_tu,
                              'unpacketized_bs': unpacketized_bs,
                              'extra_crumb': 'Packet Grid',
                              'extra_crumb_url': '/packet_grid/{0}/'.format(qset.id),
                              'read_only': read_only})

def _log_packet_grid_change(qset, changer, description, prior_states):
    """Record a packet-grid change with the prior state of each affected
    question so it can be undone later."""
    PacketGridLog.objects.create(
        question_set=qset, changer=changer, description=description,
        undo_data=json.dumps(prior_states))


@login_required
def move_packet_question(request):
    user = request.user.writer
    message = ''
    success = False

    if request.method == 'POST':
        try:
            question_type = request.POST['question_type']
            question_id = int(request.POST['question_id'])
            target_packet_id = int(request.POST['target_packet_id'])
            target_number = int(request.POST['target_number'])

            model = Tossup if question_type == 'tossup' else Bonus
            question = model.objects.get(id=question_id)
            qset = question.question_set
            target_packet = Packet.objects.get(id=target_packet_id)

            if target_packet.question_set_id != qset.id:
                message = 'The target packet belongs to a different set!'
            elif not (qset.is_owner(user) or user in qset.editor.all()):
                message = 'You are not authorized to move questions in this set!'
            elif target_number < 1:
                message = 'Invalid question number!'
            else:
                occupant = model.objects.filter(
                    packet=target_packet, question_number=target_number).exclude(id=question.id).first()
                source_packet, source_number = question.packet, question.question_number
                # Capture prior state for undo before mutating.
                prior = [{'qtype': question_type, 'id': question.id,
                          'packet_id': source_packet.id if source_packet else None,
                          'number': source_number}]
                if occupant is not None:
                    prior.append({'qtype': question_type, 'id': occupant.id,
                                  'packet_id': target_packet.id, 'number': target_number})
                question.packet = target_packet
                question.question_number = target_number
                question.save()
                if occupant is not None:
                    occupant.packet = source_packet
                    occupant.question_number = source_number
                    occupant.save()
                src_name = source_packet.packet_name if source_packet else 'Unassigned'
                desc = 'Moved {0} from {1} #{2} to {3} #{4}'.format(
                    question_type, src_name, source_number or '?',
                    target_packet.packet_name, target_number)
                if occupant is not None:
                    desc += ' (swapped)'
                _log_packet_grid_change(qset, user, desc, prior)
                cache.clear()
                success = True
                message = 'Question moved'
        except (KeyError, ValueError):
            message = 'Invalid request!'
        except (Tossup.DoesNotExist, Bonus.DoesNotExist, Packet.DoesNotExist):
            message = 'Question or packet not found!'

    return HttpResponse(json.dumps({'success': success, 'message': message}))


@login_required
def set_packet_order(request):
    """Persist a user-defined packet order (drag-to-reorder on the grid).
    Body: qset_id, packet_ids[] in the desired order."""
    user = request.user.writer
    if request.method != 'POST':
        return HttpResponse(json.dumps({'success': False, 'message': 'Invalid request'}))
    try:
        qset = QuestionSet.objects.get(id=int(request.POST['qset_id']))
    except (KeyError, ValueError, QuestionSet.DoesNotExist):
        return HttpResponse(json.dumps({'success': False, 'message': 'Set not found'}))
    if not (qset.is_owner(user) or user in qset.editor.all()):
        return HttpResponse(json.dumps({'success': False, 'message': 'Not authorized'}))
    ids = request.POST.getlist('packet_ids[]')
    valid = {p.id: p for p in qset.packet_set.all()}
    order = 1
    to_update = []
    for pid in ids:
        try:
            p = valid.get(int(pid))
        except (TypeError, ValueError):
            p = None
        if p is not None:
            p.sort_order = order
            to_update.append(p)
            order += 1
    if to_update:
        Packet.objects.bulk_update(to_update, ['sort_order'])
        cache.clear()
    return HttpResponse(json.dumps({'success': True, 'message': 'Order saved'}))


@login_required
def unassign_packet_question(request):
    """Remove a question from its packet (back to the unpacketized pool)."""
    user = request.user.writer
    message = ''
    success = False
    if request.method == 'POST':
        try:
            question_type = request.POST['question_type']
            question_id = int(request.POST['question_id'])
            model = Tossup if question_type == 'tossup' else Bonus
            question = model.objects.get(id=question_id)
            qset = question.question_set
            if not (qset.is_owner(user) or user in qset.editor.all()):
                message = 'You are not authorized to move questions in this set!'
            elif question.packet_id is None:
                success = True
                message = 'Already unpacketized'
            else:
                prior = [{'qtype': question_type, 'id': question.id,
                          'packet_id': question.packet_id, 'number': question.question_number}]
                src_name = question.packet.packet_name
                src_num = question.question_number
                # Remember the category that used to fill this slot so the empty
                # grid cell can hint what belongs there.
                if src_num:
                    PacketSlotVacancy.objects.update_or_create(
                        packet_id=question.packet_id, question_number=src_num,
                        question_type=question_type,
                        defaults={'question_set': qset,
                                  'category': str(question.category) if question.category_id else ''})
                question.packet = None
                question.question_number = None
                question.save()
                _log_packet_grid_change(
                    qset, user, 'Unpacketized {0} from {1} #{2}'.format(
                        question_type, src_name, src_num or '?'), prior)
                cache.clear()
                success = True
                message = 'Question unpacketized'
        except (KeyError, ValueError):
            message = 'Invalid request!'
        except (Tossup.DoesNotExist, Bonus.DoesNotExist):
            message = 'Question not found!'
    return HttpResponse(json.dumps({'success': success, 'message': message}))


def _create_packets_continuing(qset, count, existing, user):
    """Create `count` new packets, continuing the existing naming scheme
    ("Round 01" -> "Round 11"), and return them. Mirrors _ensure_packets."""
    base = 'Packet'
    if existing:
        m = re.match(r'^(.*?)\s*\d+$', existing[-1].packet_name or '')
        if m and m.group(1).strip():
            base = m.group(1).strip()
    names = set(p.packet_name for p in qset.packet_set.all())
    created = []
    next_num = len(existing) + 1
    while len(created) < count:
        name = '{0} {1:02d}'.format(base, next_num)
        next_num += 1
        if name in names:
            continue
        created.append(Packet.objects.create(
            question_set=qset, packet_name=name, created_by=user))
        names.add(name)
    return created


@login_required
def assign_unpacketized(request):
    """Non-destructively place unpacketized questions: first into empty slots of
    existing packets, then into newly created packets for any overflow. Does not
    move questions that are already assigned."""
    user = request.user.writer
    message = ''
    success = False
    if request.method != 'POST':
        return HttpResponse(json.dumps({'success': False, 'message': 'Invalid request'}))
    try:
        qset = QuestionSet.objects.get(id=int(request.POST['qset_id']))
    except (KeyError, ValueError, QuestionSet.DoesNotExist):
        return HttpResponse(json.dumps({'success': False, 'message': 'Set not found'}))
    if not (qset.is_owner(user) or user in qset.editor.all()):
        return HttpResponse(json.dumps(
            {'success': False, 'message': 'You are not authorized to packetize this set!'}))

    def natural_key(p):
        nums = re.findall(r'\d+', p.packet_name or '')
        return (p.packet_name == EXTRAS_PACKET_NAME,
                int(nums[0]) if nums else float('inf'), p.packet_name or '')

    prior = []
    placed = 0

    def fill_existing(model, per_packet, qtype, packets):
        """Fill empty slots in existing packets; return questions left over."""
        nonlocal placed
        unassigned = list(model.objects.filter(question_set=qset, packet=None).order_by('id'))
        if not unassigned:
            return []
        occupied = defaultdict(set)
        for q in model.objects.filter(question_set=qset, packet__in=packets):
            if q.question_number:
                occupied[q.packet_id].add(q.question_number)
        qi = 0
        for p in packets:
            for n in range(1, per_packet + 1):
                if qi >= len(unassigned):
                    return []
                if n not in occupied[p.id]:
                    q = unassigned[qi]
                    qi += 1
                    prior.append({'qtype': qtype, 'id': q.id, 'packet_id': None, 'number': None})
                    q.packet = p
                    q.question_number = n
                    q.save()
                    placed += 1
        return unassigned[qi:]

    def fill_new(model, per_packet, qtype, leftovers, new_packets):
        nonlocal placed
        qi = 0
        for p in new_packets:
            for n in range(1, per_packet + 1):
                if qi >= len(leftovers):
                    return
                q = leftovers[qi]
                qi += 1
                prior.append({'qtype': qtype, 'id': q.id, 'packet_id': None, 'number': None})
                q.packet = p
                q.question_number = n
                q.save()
                placed += 1

    with transaction.atomic():
        existing = sorted(qset.packet_set.exclude(packet_name=EXTRAS_PACKET_NAME),
                          key=natural_key)
        tu_per = qset.tossups_per_packet or 20
        bs_per = qset.bonuses_per_packet or 20
        rem_tu = fill_existing(Tossup, tu_per, 'tossup', existing)
        rem_bs = fill_existing(Bonus, bs_per, 'bonus', existing)

        n_new = max(math.ceil(len(rem_tu) / tu_per) if rem_tu else 0,
                    math.ceil(len(rem_bs) / bs_per) if rem_bs else 0)
        if n_new:
            new_packets = _create_packets_continuing(qset, n_new, existing, user)
            fill_new(Tossup, tu_per, 'tossup', rem_tu, new_packets)
            fill_new(Bonus, bs_per, 'bonus', rem_bs, new_packets)
            qset.num_packets = max(qset.num_packets or 0, len(existing) + n_new)
            qset.save(update_fields=['num_packets'])

        if prior:
            _log_packet_grid_change(
                qset, user,
                'Assigned {0} unpacketized question(s){1}'.format(
                    placed, ' (+{0} new packet(s))'.format(n_new) if n_new else ''),
                prior)
            cache.clear()
            success = True
            message = 'Placed {0} question(s){1}.'.format(
                placed, ' and created {0} new packet(s)'.format(n_new) if n_new else '')
        else:
            success = True
            message = 'No unpacketized questions to place.'

    return HttpResponse(json.dumps({'success': success, 'message': message}))


def _packet_revision(packet):
    """A short token that changes whenever the packet's composition or any of
    its questions change — used by the document view to detect that someone else
    edited the packet (membership, order, or question text) so it can warn the
    viewer to reload. Cheap: two lightweight queries, no new DB columns."""
    import hashlib
    parts = []
    for qtype, qs in (
        ('t', packet.tossup_set.order_by('question_number')
              .values_list('id', 'question_number', 'last_changed_date')),
        ('b', packet.bonus_set.order_by('question_number')
              .values_list('id', 'question_number', 'last_changed_date'))):
        for qid, num, changed in qs:
            parts.append('{0}{1}:{2}:{3}'.format(qtype, qid, num, changed.isoformat() if changed else ''))
    return hashlib.md5('|'.join(parts).encode('utf-8')).hexdigest()[:16]


@login_required
def packet_revision(request, packet_id):
    """JSON {revision} for the given packet, polled by the document view to
    detect concurrent changes."""
    try:
        packet = Packet.objects.select_related('question_set').get(id=packet_id)
    except Packet.DoesNotExist:
        return HttpResponse(json.dumps({'error': 'not found'}), status=404)
    user = request.user.writer
    qset = packet.question_set
    if not qset.is_owner(user) and user not in qset.editor.all() and user not in qset.writer.all():
        return HttpResponse(json.dumps({'error': 'forbidden'}), status=403)
    return HttpResponse(json.dumps({'revision': _packet_revision(packet)}))


@login_required
def view_packet(request, packet_id):
    user = request.user.writer
    packet = Packet.objects.get(id=packet_id)
    qset = packet.question_set

    if not qset.is_owner(user) and user not in qset.editor.all() and user not in qset.writer.all():
        return render(request, 'failure.html',
                                 {'message': 'You are not authorized to view this packet!',
                                  'message_class': 'alert-box alert'})

    read_only = not (qset.is_owner(user) or user in qset.editor.all())
    order = request.GET.get('order', 'separate')
    if order not in ('separate', 'interleaved'):
        order = 'separate'

    def writer_label(writer):
        if writer is None:
            return ''
        name = '{0} {1}'.format(writer.user.first_name, writer.user.last_name).strip()
        return name or writer.user.username

    # Most recent changer per question, from the question history
    def latest_changers(history_model, questions):
        history_ids = [q.question_history_id for q in questions if q.question_history_id]
        changers = {}
        for h in history_model.objects.filter(question_history_id__in=history_ids) \
                .order_by('change_date').select_related('changer__user'):
            changers[h.question_history_id] = writer_label(h.changer)
        return changers

    def item(question, qtype, edit_url, per_packet, max_length, changers):
        number = question.question_number or 0
        return {
            'id': question.id,
            'qtype': qtype,
            'number': number,
            'html': question.to_html(),
            'category': str(question.category) if question.category else '',
            'edit_url': '{0}{1}/'.format(edit_url, question.id),
            'is_tiebreaker': number > per_packet,
            'author': writer_label(question.author),
            'editor': writer_label(question.editor),
            'edited': question.edited,
            'length': question.character_count(),
            'max_length': max_length,
            'changed_date': question.last_changed_date,
            'changed_by': changers.get(question.question_history_id, ''),
        }

    packet_tossups = list(packet.tossup_set.order_by('question_number')
                          .select_related('category', 'author__user', 'editor__user'))
    packet_bonuses = list(packet.bonus_set.order_by('question_number')
                          .select_related('category', 'author__user', 'editor__user'))
    tossup_changers = latest_changers(TossupHistory, packet_tossups)
    bonus_changers = latest_changers(BonusHistory, packet_bonuses)

    tossups = [item(t, 'tossup', '/edit_tossup/', qset.tossups_per_packet, qset.max_acf_tossup_length, tossup_changers)
               for t in packet_tossups]
    bonuses = [item(b, 'bonus', '/edit_bonus/', qset.bonuses_per_packet,
                    qset.max_vhsl_bonus_length if b.get_bonus_type() == VHSL_BONUS else qset.max_acf_bonus_length,
                    bonus_changers)
               for b in packet_bonuses]

    # Heads-up for the moderator: answer lines flagged "read answer carefully".
    def _careful_answer(q, qtype):
        if qtype == 'tossup':
            return get_formatted_question_html(q.tossup_answer, True, True, False, False)
        parts = [p for p in (q.part1_answer, q.part2_answer, q.part3_answer) if p]
        return ' / '.join(get_formatted_question_html(p, True, True, False, False) for p in parts)

    careful_notes = []
    for t in packet_tossups:
        if t.read_carefully:
            careful_notes.append({'label': 'Tossup', 'number': t.question_number,
                                  'answer': _careful_answer(t, 'tossup')})
    for b in packet_bonuses:
        if b.read_carefully:
            careful_notes.append({'label': 'Bonus', 'number': b.question_number,
                                  'answer': _careful_answer(b, 'bonus')})

    # Raw question data for the client-side Discord copy formatters
    discord_payload = {}
    for t, row in zip(packet_tossups, tossups):
        discord_payload['tossup-{0}'.format(t.id)] = {
            'qtype': 'tossup', 'text': t.tossup_text, 'answer': t.tossup_answer,
            'author': row['author'], 'category': row['category']}
    for b, row in zip(packet_bonuses, bonuses):
        discord_payload['bonus-{0}'.format(b.id)] = {
            'qtype': 'bonus', 'leadin': b.leadin,
            'parts': [{'text': b.part1_text or '', 'answer': b.part1_answer or '', 'diff': b.part1_difficulty or ''},
                      {'text': b.part2_text or '', 'answer': b.part2_answer or '', 'diff': b.part2_difficulty or ''},
                      {'text': b.part3_text or '', 'answer': b.part3_answer or '', 'diff': b.part3_difficulty or ''}],
            'author': row['author'], 'category': row['category']}

    interleaved = []
    if order == 'interleaved':
        for i in range(max(len(tossups), len(bonuses))):
            if i < len(tossups):
                interleaved.append(tossups[i])
            if i < len(bonuses):
                interleaved.append(bonuses[i])

    # Comments for the Google-Docs-style margin
    def attach_comments(items, model):
        ct = ContentType.objects.get_for_model(model)
        comments = list(Comment.objects.filter(
            content_type=ct, object_pk__in=[str(q['id']) for q in items],
            is_removed=False).order_by('submit_date').select_related('user'))
        comment_ids = [c.id for c in comments]
        # Which comments are anchored to a text selection, and which are replies.
        anchored = {a.comment_id: a.selected_text
                    for a in CommentAnchor.objects.filter(comment_id__in=comment_ids)}
        parent_of = {r.comment_id: r.parent_id
                     for r in CommentReply.objects.filter(comment_id__in=comment_ids)}

        def render_comment(c):
            label = ''
            if c.user is not None:
                label = '{0} {1}'.format(c.user.first_name, c.user.last_name).strip()
            if not label:
                label = c.user_name or (c.user.username if c.user else 'unknown')
            return {'id': c.id, 'user': label, 'text': c.comment, 'date': c.submit_date,
                    'anchored': c.id in anchored, 'selection': anchored.get(c.id, ''),
                    'replies': []}

        rendered = {c.id: render_comment(c) for c in comments}
        by_q = {}
        for c in comments:
            by_q.setdefault(int(c.object_pk), []).append(c)
        for q in items:
            qcs = by_q.get(q['id'], [])
            present = {c.id for c in qcs}
            threads = []
            for c in qcs:
                pid = parent_of.get(c.id)
                if pid in present:
                    continue  # nested under its parent below
                node = rendered[c.id]
                node['replies'] = [rendered[r.id] for r in qcs if parent_of.get(r.id) == c.id]
                threads.append(node)
            q['comments'] = threads

    attach_comments(tossups, Tossup)
    attach_comments(bonuses, Bonus)
    comment_count = sum(len(q['comments']) for q in tossups + bonuses)

    siblings = sorted_packets(qset)
    index = next((i for i, p in enumerate(siblings) if p.id == packet.id), 0)
    prev_packet = siblings[index - 1] if index > 0 else None
    next_packet = siblings[index + 1] if index + 1 < len(siblings) else None

    from .audio import VOICE_CHOICES
    return render(request, 'view_packet.html',
                             {'qset': qset,
                              'packet': packet,
                              'order': order,
                              'tossups': tossups,
                              'bonuses': bonuses,
                              'interleaved': interleaved,
                              'prev_packet': prev_packet,
                              'next_packet': next_packet,
                              'packet_index': index + 1,
                              'packet_count': len(siblings),
                              'comment_count': comment_count,
                              'discord_payload': discord_payload,
                              'mp3_voices': VOICE_CHOICES,
                              'packet_revision': _packet_revision(packet),
                              'careful_notes': careful_notes,
                              'role': get_role_no_owner(user, qset),
                              'read_only': read_only,
                              'user': user})


def _packet_mp3_params(request, packet_id):
    """Shared setup for the MP3 endpoints. Returns (packet, qset, tossups,
    bonuses, interleaved, include_answers, voice) or an auth failure response."""
    from .audio import NATURAL_VOICES, DEFAULT_VOICE
    user = request.user.writer
    packet = Packet.objects.get(id=packet_id)
    qset = packet.question_set
    if not qset.is_owner(user) and user not in qset.editor.all() and user not in qset.writer.all():
        return render(request, 'failure.html',
                                 {'message': 'You are not authorized to view this packet!',
                                  'message_class': 'alert-box alert'})
    interleaved = request.GET.get('order') == 'interleaved'
    include_answers = request.GET.get('answers', '1') != '0'
    voice = request.GET.get('voice') or DEFAULT_VOICE
    if voice not in NATURAL_VOICES:
        voice = DEFAULT_VOICE
    tossups = list(packet.tossup_set.order_by('question_number'))
    bonuses = list(packet.bonus_set.order_by('question_number'))
    return packet, qset, tossups, bonuses, interleaved, include_answers, voice


@login_required
def packet_mp3(request, packet_id):
    """Serve an MP3 reading of a packet. If it is not ready yet, kick off a
    background render and show a page that polls until it can download.

    Query params: order=separate|interleaved, answers=1|0.
    """
    parsed = _packet_mp3_params(request, packet_id)
    if isinstance(parsed, HttpResponse):
        return parsed
    packet, qset, tossups, bonuses, interleaved, include_answers, voice = parsed

    from .audio import cache_file, packet_status, start_generation

    if packet_status(packet, tossups, bonuses, interleaved, include_answers, voice) == 'ready':
        order_label = 'interleaved' if interleaved else 'grouped'
        filename = '{0} ({1}).mp3'.format(packet.packet_name, order_label)
        with open(cache_file(packet.id, interleaved, include_answers, voice), 'rb') as f:
            response = HttpResponse(f.read(), content_type='audio/mpeg')
        response['Content-Disposition'] = 'attachment; filename="{0}"'.format(filename)
        return response

    # Not ready: make sure a background render is running, then show a waiting page.
    start_generation(packet, tossups, bonuses, interleaved=interleaved,
                     include_answers=include_answers, voice=voice)
    return render(request, 'packet_mp3_preparing.html',
                             {'packet': packet, 'qset': qset,
                              'order': 'interleaved' if interleaved else 'separate',
                              'answers': '1' if include_answers else '0',
                              'voice': voice},
                             status=202)


@login_required
def packet_mp3_status(request, packet_id):
    """JSON status for a packet MP3 render: ready | running | error | absent."""
    parsed = _packet_mp3_params(request, packet_id)
    if isinstance(parsed, HttpResponse):
        return parsed
    packet, qset, tossups, bonuses, interleaved, include_answers, voice = parsed

    from .audio import packet_status, error_message
    status = packet_status(packet, tossups, bonuses, interleaved, include_answers, voice)
    data = {'status': status}
    if status == 'error':
        data['message'] = error_message(packet.id, interleaved, include_answers, voice) or \
            'generation failed'
    response = JsonResponse(data)
    # Never cache status — each poll must reflect the live generation state.
    response['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    return response


@login_required
def reorder_packet_questions(request):
    user = request.user.writer
    message = ''
    success = False

    if request.method == 'POST':
        try:
            packet = Packet.objects.get(id=int(request.POST['packet_id']))
            qset = packet.question_set
            if not (qset.is_owner(user) or user in qset.editor.all()):
                message = 'You are not authorized to reorder questions in this set!'
            else:
                prior = []
                for key, model in (('tossup_ids[]', Tossup), ('bonus_ids[]', Bonus)):
                    qtype = 'tossup' if model is Tossup else 'bonus'
                    ids = [int(x) for x in request.POST.getlist(key)]
                    if not ids:
                        continue
                    questions = {q.id: q for q in model.objects.filter(packet=packet, id__in=ids)}
                    if len(questions) != len(ids):
                        raise ValueError('Question list does not match the packet')
                    changed = []
                    for number, qid in enumerate(ids, start=1):
                        question = questions[qid]
                        if question.question_number != number:
                            prior.append({'qtype': qtype, 'id': question.id,
                                          'packet_id': packet.id, 'number': question.question_number})
                            question.question_number = number
                            changed.append(question)
                    model.objects.bulk_update(changed, ['question_number'])
                if prior:
                    _log_packet_grid_change(
                        qset, user, 'Reordered questions in {0}'.format(packet.packet_name), prior)
                cache.clear()
                success = True
                message = 'Order saved'
        except (KeyError, ValueError):
            message = 'Invalid request!'
        except Packet.DoesNotExist:
            message = 'Packet not found!'

    return HttpResponse(json.dumps({'success': success, 'message': message}))


@login_required
def swap_candidates(request):
    """JSON list of questions that could go in a slot. With a source
    ``question_id`` it lists swap candidates filtered by category scope ('leaf',
    'sub', 'top') or free-text. Without one (filling an empty slot) it lists
    unpacketized questions, or — with a search — any matching question."""
    user = request.user.writer
    question_type = request.GET.get('question_type')
    if question_type not in ('tossup', 'bonus'):
        return HttpResponse(json.dumps({'error': 'Invalid question type.'}))
    model = Tossup if question_type == 'tossup' else Bonus
    search = request.GET.get('q', '').strip()
    raw_qid = request.GET.get('question_id')
    fill_mode = not raw_qid  # empty-slot mode: find any question to place

    if fill_mode:
        try:
            qset = QuestionSet.objects.get(id=int(request.GET['qset_id']))
        except (KeyError, ValueError, QuestionSet.DoesNotExist):
            return HttpResponse(json.dumps({'error': 'Set not found.'}))
        question = None
    else:
        try:
            question = model.objects.get(id=int(raw_qid))
        except (ValueError, Tossup.DoesNotExist, Bonus.DoesNotExist):
            return HttpResponse(json.dumps({'error': 'Question not found!'}))
        qset = question.question_set

    if not (qset.is_owner(user) or user in qset.editor.all()):
        return HttpResponse(json.dumps({'error': 'You are not authorized to swap questions in this set!'}))

    def text_filter_for(term):
        if question_type == 'tossup':
            return Q(tossup_answer__icontains=term) | Q(tossup_text__icontains=term)
        return (Q(leadin__icontains=term) |
                Q(part1_text__icontains=term) | Q(part1_answer__icontains=term) |
                Q(part2_text__icontains=term) | Q(part2_answer__icontains=term) |
                Q(part3_text__icontains=term) | Q(part3_answer__icontains=term))

    scope = request.GET.get('scope', 'leaf')
    if scope not in ('leaf', 'sub', 'top'):
        scope = 'leaf'

    if fill_mode:
        if search:
            candidates = (model.objects.filter(question_set=qset).filter(text_filter_for(search))
                          .select_related('category', 'packet')
                          .order_by('packet__packet_name', 'question_number')[:200])
            source_label = 'search: "{0}"'.format(search)
        else:
            # Default to the unpacketized pool — the questions you'd most want
            # to drop into an empty slot.
            candidates = (model.objects.filter(question_set=qset, packet=None)
                          .select_related('category', 'packet').order_by('id')[:200])
            source_label = 'unpacketized'
    elif search:
        candidates = (model.objects.filter(question_set=qset).filter(text_filter_for(search))
                      .exclude(id=question.id)
                      .select_related('category', 'packet')
                      .order_by('packet__packet_name', 'question_number')[:200])
        source_label = 'search: "{0}"'.format(search)
    else:
        entry = question.category
        if entry is None:
            return HttpResponse(json.dumps({'error': 'This question has no category. Type to search instead.'}))

        if scope == 'leaf':
            entry_ids = [entry.id]
        else:
            siblings = DistributionEntry.objects.filter(
                distribution=entry.distribution, category=entry.category)
            if scope == 'sub':
                sub_first = (entry.subcategory or '').split(' - ')[0].strip()
                entry_ids = [e.id for e in siblings
                             if (e.subcategory or '').split(' - ')[0].strip() == sub_first]
            else:
                entry_ids = [e.id for e in siblings]

        candidates = model.objects.filter(question_set=qset, category_id__in=entry_ids) \
            .exclude(packet=question.packet) \
            .select_related('category', 'packet') \
            .order_by('packet__packet_name', 'question_number')[:200]
        source_label = str(entry)

    def preview(q):
        if question_type == 'tossup':
            return _grid_answer_preview(q.tossup_answer)
        return ' / '.join(filter(None, [
            _grid_answer_preview(q.part1_answer, 20),
            _grid_answer_preview(q.part2_answer, 20),
            _grid_answer_preview(q.part3_answer, 20)]))

    per_packet = qset.tossups_per_packet if question_type == 'tossup' else qset.bonuses_per_packet
    # Show unpacketized candidates first so they're easy to grab.
    candidates = sorted(candidates, key=lambda q: (q.packet_id is not None,))
    data = [{
        'id': q.id,
        'packet_id': q.packet_id,
        'packet_name': q.packet.packet_name if q.packet_id else '(unpacketized)',
        'number': q.question_number,
        'answer': html.unescape(preview(q)),
        'category': html.unescape(str(q.category)) if q.category else '',
        'is_tiebreaker': (q.question_number or 0) > per_packet if q.packet_id else False,
        'unpacketized': q.packet_id is None,
    } for q in candidates]

    return HttpResponse(json.dumps({'candidates': data, 'source_category': source_label,
                                    'fill_mode': fill_mode,
                                    'source_answer': html.unescape(preview(question)) if question else ''}))


@login_required
def undo_packet_grid_change(request):
    """Undo the most recent (not-yet-undone) packet-grid change for a set,
    restoring each affected question's prior packet and number."""
    user = request.user.writer
    message = ''
    success = False
    if request.method == 'POST':
        try:
            qset = QuestionSet.objects.get(id=int(request.POST['qset_id']))
        except (KeyError, ValueError, QuestionSet.DoesNotExist):
            return HttpResponse(json.dumps({'success': False, 'message': 'Invalid request!'}))
        if not (qset.is_owner(user) or user in qset.editor.all()):
            return HttpResponse(json.dumps({'success': False, 'message': 'You are not authorized to change this set!'}))

        log = PacketGridLog.objects.filter(question_set=qset, undone=False).order_by('-change_date', '-id').first()
        if log is None:
            return HttpResponse(json.dumps({'success': False, 'message': 'Nothing to undo.'}))
        try:
            prior = json.loads(log.undo_data)
        except ValueError:
            prior = []
        for state in prior:
            model = Tossup if state.get('qtype') == 'tossup' else Bonus
            q = model.objects.filter(id=state.get('id'), question_set=qset).first()
            if q is not None:
                q.packet_id = state.get('packet_id')
                q.question_number = state.get('number')
                q.save()
        log.undone = True
        log.save(update_fields=['undone'])
        cache.clear()
        success = True
        message = 'Undid: {0}'.format(log.description)
    return HttpResponse(json.dumps({'success': success, 'message': message}))


@login_required
def packet_grid_log(request, qset_id):
    """Show the history of packet-grid changes for a set."""
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    if not (qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all()):
        return render(request, 'failure.html',
                      {'message': 'You are not authorized to view this set!',
                       'message_class': 'alert-box alert'})
    logs = (PacketGridLog.objects.filter(question_set=qset)
            .select_related('changer__user').order_by('-change_date', '-id')[:500])
    return render(request, 'packet_grid_log.html',
                  {'qset': qset, 'user': user, 'logs': logs,
                   'can_edit': qset.is_owner(user) or user in qset.editor.all()})


def _new_activity_count(user, qset):
    """Count activity (mentions + others' changes to your questions) newer than
    the last time the user viewed their activity feed for this set."""
    seen = ActivitySeen.objects.filter(writer=user, question_set=qset).first()
    last = seen.last_seen if seen else None

    tu_ct = ContentType.objects.get_for_model(Tossup)
    bs_ct = ContentType.objects.get_for_model(Bonus)
    tu_ids = [str(i) for i in qset.tossup_set.values_list('id', flat=True)]
    bs_ids = [str(i) for i in qset.bonus_set.values_list('id', flat=True)]

    mentions = CommentMention.objects.filter(mentioned=user).filter(
        Q(comment__content_type=tu_ct, comment__object_pk__in=tu_ids) |
        Q(comment__content_type=bs_ct, comment__object_pk__in=bs_ids)
    ).exclude(comment__resolution__resolved=True)
    if last is not None:
        mentions = mentions.filter(created_date__gt=last)
    count = mentions.count()

    for model, hist_model in ((Tossup, TossupHistory), (Bonus, BonusHistory)):
        hids = [h for h in model.objects.filter(question_set=qset)
                .filter(Q(author=user) | Q(editor=user))
                .values_list('question_history_id', flat=True) if h]
        if not hids:
            continue
        changes = hist_model.objects.filter(question_history_id__in=hids).exclude(changer=user)
        if last is not None:
            changes = changes.filter(change_date__gt=last)
        count += changes.count()
    return count


@login_required
def activity(request, qset_id):
    """Per-user activity on a set: @mentions of you, plus changes other people
    made to questions you wrote or edited."""
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    if not (qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all()):
        return render(request, 'failure.html',
                      {'message': 'You are not authorized to view this set!',
                       'message_class': 'alert-box alert'})

    tu_ct = ContentType.objects.get_for_model(Tossup)
    bs_ct = ContentType.objects.get_for_model(Bonus)
    tu_ids = set(str(i) for i in qset.tossup_set.values_list('id', flat=True))
    bs_ids = set(str(i) for i in qset.bonus_set.values_list('id', flat=True))

    # --- @mentions of you on this set's questions ---
    mentions = (CommentMention.objects.filter(mentioned=user)
                .filter(Q(comment__content_type=tu_ct, comment__object_pk__in=tu_ids) |
                        Q(comment__content_type=bs_ct, comment__object_pk__in=bs_ids))
                .exclude(comment__resolution__resolved=True)
                .select_related('comment__user', 'comment__content_type')
                .order_by('-comment__submit_date')[:100])
    mention_items = []
    for m in mentions:
        c = m.comment
        is_tu = c.content_type_id == tu_ct.id
        by = str(c.user) if c.user else (c.user_name or 'unknown')
        mention_items.append({
            'date': c.submit_date, 'by': by, 'text': c.comment,
            'qtype': 'tossup' if is_tu else 'bonus',
            'edit_url': '{0}{1}/'.format('/edit_tossup/' if is_tu else '/edit_bonus/', c.object_pk),
        })

    # --- changes by others to questions you authored or edited ---
    change_items = []
    for model, hist_model, edit, preview in (
            (Tossup, TossupHistory, '/edit_tossup/', lambda q: _grid_answer_preview(q.tossup_answer)),
            (Bonus, BonusHistory, '/edit_bonus/', lambda q: _grid_answer_preview(q.part1_answer, 30))):
        mine = (model.objects.filter(question_set=qset)
                .filter(Q(author=user) | Q(editor=user))
                .select_related('category'))
        hist_to_q = {q.question_history_id: q for q in mine if q.question_history_id}
        if not hist_to_q:
            continue
        histories = (hist_model.objects.filter(question_history_id__in=hist_to_q.keys())
                     .exclude(changer=user).select_related('changer__user')
                     .order_by('-change_date')[:100])
        for h in histories:
            q = hist_to_q.get(h.question_history_id)
            if q is None:
                continue
            change_items.append({
                'date': h.change_date,
                'by': str(h.changer) if h.changer else 'unknown',
                'qtype': 'tossup' if model is Tossup else 'bonus',
                'edit_url': '{0}{1}/'.format(edit, q.id),
                'preview': preview(q),
                'role': 'wrote' if q.author_id == user.id else 'edited',
            })
    change_items.sort(key=lambda x: x['date'], reverse=True)
    change_items = change_items[:100]

    # Mark this set's activity as seen (clears the notification badge).
    ActivitySeen.objects.update_or_create(
        writer=user, question_set=qset, defaults={'last_seen': timezone.now()})

    return render(request, 'activity.html',
                  {'qset': qset, 'user': user,
                   'mention_items': mention_items, 'change_items': change_items})


@login_required
def set_members(request, qset_id):
    """JSON list of the set's members (username + real name) for the comment
    @mention autocomplete."""
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    if not (qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all()):
        return HttpResponse(json.dumps({'members': []}))
    by_name = {}
    for w in list(qset.all_owners()) + list(qset.editor.all()) + list(qset.writer.all()):
        u = w.user
        name = '{0} {1}'.format(u.first_name, u.last_name).strip() or u.username
        by_name[u.username] = name
    members = [{'username': k, 'name': v}
               for k, v in sorted(by_name.items(), key=lambda kv: kv[1].lower())]
    return HttpResponse(json.dumps({'members': members}))


@login_required
def resolve_comment(request):
    """Toggle a comment's resolved status. Resolved comments drop off the
    mentioned writers' activity feeds."""
    user = request.user.writer
    if request.method != 'POST':
        return HttpResponse(json.dumps({'success': False, 'message': 'Invalid request'}))
    try:
        comment = Comment.objects.get(id=int(request.POST['comment_id']))
    except (KeyError, ValueError, Comment.DoesNotExist):
        return HttpResponse(json.dumps({'success': False, 'message': 'Comment not found'}))
    target = comment.content_object
    qset = getattr(target, 'question_set', None)
    if qset is None or not (qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all()):
        return HttpResponse(json.dumps({'success': False, 'message': 'You are not authorized.'}))
    res, created = CommentResolution.objects.get_or_create(
        comment=comment, defaults={'resolved': True, 'resolved_by': user})
    if not created:
        res.resolved = not res.resolved
    res.resolved_by = user
    res.save()
    cache.clear()
    return HttpResponse(json.dumps({'success': True, 'resolved': res.resolved}))


def _question_issue_map(qset):
    """Map 'tossup-<id>'/'bonus-<id>' -> worst repeat-checker severity for every
    flagged question in the set. Cached by the dup-check fingerprint so it only
    recomputes when questions change."""
    key = 'dupissues:{0}:{1}'.format(qset.id, _dup_fingerprint(qset))
    cached = cache.get(key)
    if cached is not None:
        return cached
    rank = {CRITICAL: 3, WARNING: 2, INFO: 1}
    issues = {}

    def add(qtype, qid, sev):
        k = '{0}-{1}'.format(qtype, qid)
        if k not in issues or rank.get(sev, 0) > rank.get(issues[k], 0):
            issues[k] = sev

    for group in find_duplicates(qset):
        for e in group['entries']:
            add(e['type'], e['id'], group['severity'])
    for group in find_topic_repeats(qset):
        for e in group['entries']:
            add(e['type'], e['id'], group['severity'])
    for issue in find_internal_issues(qset):
        add(issue['question_type'], issue['question_id'], issue['severity'])

    cache.set(key, issues, 1800)
    return issues


@login_required
def style_check(request, qset_id):
    """Run the style checker over a set's questions. The style guide is
    selectable via ?guide=; default is Minkowski."""
    from . import style_checker
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)
    if not (qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all()):
        return render(request, 'failure.html',
                      {'message': 'You are not authorized to view this set!',
                       'message_class': 'alert-box alert'})

    guide = request.GET.get('guide', style_checker.DEFAULT_GUIDE)
    if guide not in style_checker.guide_keys():
        guide = style_checker.DEFAULT_GUIDE

    # Editors can turn individual rules on/off for this set.
    if request.method == 'POST' and request.POST.get('action') == 'save_rules':
        if qset.is_owner(user) or user in qset.editor.all():
            configurable = [c for c, _ in style_checker.configurable_rules(guide)]
            disabled = [c for c in configurable if request.POST.get('rule_' + c) != 'on']
            qset.disabled_style_rules = ','.join(disabled)
            qset.save()
            cache.clear()
        return HttpResponseRedirect('/style_check/{0}/?guide={1}'.format(qset.id, guide))

    disabled = qset.disabled_style_rule_set()
    show_dismissed = request.GET.get('show_dismissed') == '1'

    dismissed = set()
    for d in StyleIssueDismissal.objects.filter(question_set=qset):
        dismissed.add((d.question_type, d.question_id, d.code, d.token))
    # Set-wide dismissals hide a suggestion (code, token) on every question.
    rule_dismissed = set(StyleRuleDismissal.objects.filter(question_set=qset)
                         .values_list('code', 'token'))

    results = []
    counts = {'error': 0, 'warning': 0, 'info': 0}
    checked = 0
    flagged = 0
    dismissed_count = 0

    def collect(qtype, q, issues, label, packet, number, edit_url):
        nonlocal flagged, dismissed_count
        active, shown = [], []
        for i in issues:
            i = dict(i, fixable=('fix' in i))
            i.pop('fix', None)  # keep the transform server-side
            if ((qtype, q.id, i['code'], i.get('token', '')) in dismissed
                    or (i['code'], i.get('token', '')) in rule_dismissed):
                dismissed_count += 1
                if show_dismissed:
                    i['dismissed'] = True
                    shown.append(i)
            else:
                counts[i['severity']] = counts.get(i['severity'], 0) + 1
                active.append(i)
        if active:
            flagged += 1
        display = active + shown
        if display:
            results.append({'type': qtype, 'id': q.id, 'edit_url': edit_url, 'label': label,
                            'packet': packet, 'number': number, 'issues': display})

    for tu in qset.tossup_set.select_related('packet').order_by('packet__packet_name', 'question_number'):
        checked += 1
        collect('tossup', tu, style_checker.check_tossup(tu, guide, disabled),
                _grid_answer_preview(tu.tossup_answer),
                tu.packet.packet_name if tu.packet else '', tu.question_number,
                '/edit_tossup/{0}/'.format(tu.id))
    for b in qset.bonus_set.select_related('packet').order_by('packet__packet_name', 'question_number'):
        checked += 1
        collect('bonus', b, style_checker.check_bonus(b, guide, disabled),
                _grid_answer_preview(b.part1_answer, 30),
                b.packet.packet_name if b.packet else '', b.question_number,
                '/edit_bonus/{0}/'.format(b.id))

    return render(request, 'style_check.html',
                  {'qset': qset, 'user': user, 'results': results, 'counts': counts,
                   'checked': checked, 'flagged': flagged, 'guide': guide,
                   'dismissed_count': dismissed_count, 'show_dismissed': show_dismissed,
                   'guides': style_checker.STYLE_GUIDES,
                   'guide_obj': next((g for g in style_checker.STYLE_GUIDES if g['key'] == guide), None),
                   'can_configure': qset.is_owner(user) or user in qset.editor.all(),
                   'is_ai_user': _is_ai_user(request.user),
                   'rule_settings': [{'code': c, 'label': lbl, 'enabled': c not in disabled}
                                     for c, lbl in style_checker.configurable_rules(guide)]})


@login_required
def packet_issues(request, packet_id):
    """JSON map of repeat-checker issue severities for the questions in a packet
    (for the document-view "show issues" toggle)."""
    user = request.user.writer
    packet = Packet.objects.get(id=packet_id)
    qset = packet.question_set
    if not (qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all()):
        return HttpResponse(json.dumps({'issues': {}}))
    full = _question_issue_map(qset)
    keys = set()
    for t in packet.tossup_set.values_list('id', flat=True):
        keys.add('tossup-{0}'.format(t))
    for b in packet.bonus_set.values_list('id', flat=True):
        keys.add('bonus-{0}'.format(b))
    issues = {k: v for k, v in full.items() if k in keys}
    return HttpResponse(json.dumps({'issues': issues}))


@login_required
def packet_style_issues(request, packet_id):
    """JSON map of style-check issues (per question) for a packet, for the
    document-view "show style issues" toggle. Style guide via ?guide=."""
    from . import style_checker
    user = request.user.writer
    packet = Packet.objects.get(id=packet_id)
    qset = packet.question_set
    if not (qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all()):
        return HttpResponse(json.dumps({'issues': {}}))
    guide = request.GET.get('guide', style_checker.DEFAULT_GUIDE)
    if guide not in style_checker.guide_keys():
        guide = style_checker.DEFAULT_GUIDE
    disabled = qset.disabled_style_rule_set()
    issues = {}
    for tu in packet.tossup_set.all():
        found = style_checker.check_tossup(tu, guide, disabled)
        if found:
            issues['tossup-{0}'.format(tu.id)] = found
    for b in packet.bonus_set.all():
        found = style_checker.check_bonus(b, guide, disabled)
        if found:
            issues['bonus-{0}'.format(b.id)] = found
    return HttpResponse(json.dumps({'issues': issues, 'guide': guide}))


def _is_ai_user(user):
    """AI-assisted features are gated to the admin user for now."""
    return user.is_authenticated and user.username == 'admin'


def _clean_for_ai(text):
    """Strip QEMS markup and decode entities so the AI sees readable prose."""
    return re.sub(r'[_~]', '', html.unescape(text or '')).strip()


# Cap how many questions one AI grammar pass covers, to bound latency/cost.
_AI_GRAMMAR_LIMIT = 50


@login_required
def ai_grammar_check(request, qset_id):
    """Admin-only AI grammar/spelling/error pass over a set's questions.
    Returns JSON findings; the Style Check page renders them."""
    from . import ai
    if not _is_ai_user(request.user):
        return HttpResponse(json.dumps({'ok': False, 'message': 'Not authorized.'}), status=403)
    if not ai.ai_enabled():
        return HttpResponse(json.dumps({'ok': False, 'message': 'AI features are not configured.'}))
    user = request.user.writer
    try:
        qset = QuestionSet.objects.get(id=qset_id)
    except QuestionSet.DoesNotExist:
        return HttpResponse(json.dumps({'ok': False, 'message': 'Set not found.'}))

    items, refs = [], {}
    total = 0
    for tu in qset.tossup_set.select_related('packet').order_by('packet__packet_name', 'question_number'):
        total += 1
        if len(items) >= _AI_GRAMMAR_LIMIT:
            continue
        ref = 'tossup-{0}'.format(tu.id)
        text = _clean_for_ai(tu.tossup_text) + '\nANSWER: ' + _clean_for_ai(tu.tossup_answer)
        items.append({'ref': ref, 'text': text})
        refs[ref] = {'edit_url': '/edit_tossup/{0}/'.format(tu.id),
                     'label': _grid_answer_preview(tu.tossup_answer)}
    for b in qset.bonus_set.select_related('packet').order_by('packet__packet_name', 'question_number'):
        total += 1
        if len(items) >= _AI_GRAMMAR_LIMIT:
            continue
        ref = 'bonus-{0}'.format(b.id)
        parts = [_clean_for_ai(b.leadin),
                 _clean_for_ai(b.part1_text), 'ANSWER: ' + _clean_for_ai(b.part1_answer),
                 _clean_for_ai(b.part2_text), 'ANSWER: ' + _clean_for_ai(b.part2_answer),
                 _clean_for_ai(b.part3_text), 'ANSWER: ' + _clean_for_ai(b.part3_answer)]
        items.append({'ref': ref, 'text': '\n'.join(p for p in parts if p.strip())})
        refs[ref] = {'edit_url': '/edit_bonus/{0}/'.format(b.id),
                     'label': _grid_answer_preview(b.part1_answer, 30)}

    findings, error = ai.grammar_check_questions(items)
    if error:
        return HttpResponse(json.dumps({'ok': False, 'message': error}))
    # Attach the edit link + answer label to each finding for rendering.
    for f in findings:
        meta = refs.get(f.get('ref'), {})
        f['edit_url'] = meta.get('edit_url', '')
        f['label'] = meta.get('label', f.get('ref', ''))
    return HttpResponse(json.dumps({'ok': True, 'findings': findings,
                                    'checked': len(items), 'total': total,
                                    'truncated': total > len(items)}))


def _style_question(qtype, qid):
    if qtype == 'tossup':
        return Tossup.objects.filter(id=qid).first()
    if qtype == 'bonus':
        return Bonus.objects.filter(id=qid).first()
    return None


def _can_edit_question(user, qset, question):
    """Edit rights for applying a fix: owners/editors can always edit; a writer
    can edit their own question while it is unlocked."""
    if qset.is_owner(user) or user in qset.editor.all():
        return True
    return (question.author_id == user.id) and not question.locked


@login_required
def apply_style_fix(request):
    """Auto-apply an easy style fix (e.g. insert a missing pronunciation guide)
    to a question. The fix transform is recomputed server-side from the issue's
    (code, token); the client only identifies which issue to fix."""
    from . import style_checker
    if request.method != 'POST':
        return HttpResponse(json.dumps({'ok': False, 'error': 'POST required'}), status=405)
    user = request.user.writer
    qtype = request.POST.get('question_type', '')
    qid = request.POST.get('question_id', '')
    code = request.POST.get('code', '')
    token = request.POST.get('token', '')
    guide = request.POST.get('guide', style_checker.DEFAULT_GUIDE)
    if guide not in style_checker.guide_keys():
        guide = style_checker.DEFAULT_GUIDE

    question = _style_question(qtype, qid)
    if question is None:
        return HttpResponse(json.dumps({'ok': False, 'error': 'No such question'}), status=404)
    qset = question.question_set
    if not (qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all()):
        return HttpResponse(json.dumps({'ok': False, 'error': 'Not authorized'}), status=403)
    if not _can_edit_question(user, qset, question):
        return HttpResponse(json.dumps({'ok': False, 'error': 'This question is locked'}), status=403)

    fix = style_checker.find_fix(question, qtype, code, token, guide)
    if not fix:
        return HttpResponse(json.dumps({'ok': False, 'error': 'Nothing to apply'}), status=400)
    if not style_checker.apply_fix(question, fix):
        return HttpResponse(json.dumps({'ok': False, 'error': 'Could not apply automatically'}), status=400)
    question.save_question(edit_type=QUESTION_EDIT, changer=user)
    return HttpResponse(json.dumps({'ok': True}))


@login_required
def dismiss_style_issue(request):
    """Dismiss (or restore) a style-check issue for a question so it is hidden
    on future runs. Set-wide. POST action=restore to undo."""
    if request.method != 'POST':
        return HttpResponse(json.dumps({'ok': False, 'error': 'POST required'}), status=405)
    user = request.user.writer
    qtype = request.POST.get('question_type', '')
    qid = request.POST.get('question_id', '')
    code = request.POST.get('code', '')
    token = request.POST.get('token', '')
    restore = request.POST.get('action', '') == 'restore'

    question = _style_question(qtype, qid)
    if question is None:
        return HttpResponse(json.dumps({'ok': False, 'error': 'No such question'}), status=404)
    qset = question.question_set
    if not (qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all()):
        return HttpResponse(json.dumps({'ok': False, 'error': 'Not authorized'}), status=403)

    scope_all = request.POST.get('scope', '') == 'all'

    if scope_all:
        # Set-wide: hide every example of this suggestion across the set.
        if restore:
            StyleRuleDismissal.objects.filter(
                question_set=qset, code=code, token=token).delete()
        else:
            StyleRuleDismissal.objects.get_or_create(
                question_set=qset, code=code, token=token,
                defaults={'dismissed_by': user})
    elif restore:
        StyleIssueDismissal.objects.filter(
            question_type=qtype, question_id=question.id, code=code, token=token).delete()
    else:
        StyleIssueDismissal.objects.get_or_create(
            question_type=qtype, question_id=question.id, code=code, token=token,
            defaults={'question_set': qset, 'dismissed_by': user})
    return HttpResponse(json.dumps({'ok': True}))


@login_required
def live_char_count(request):
    """Character count for in-progress edit text, using the set's counting
    rules (pronunciation guides / moderator instructions excluded as configured).
    Body: qset_id and one or more text[] fields (summed, like the model does)."""
    from .utils import get_character_count
    if request.method != 'POST':
        return HttpResponse(json.dumps({'count': 0}))
    ignore = True
    try:
        qset = QuestionSet.objects.get(id=int(request.POST['qset_id']))
        ignore = qset.char_count_ignores_pronunciation_guides
    except (KeyError, ValueError, QuestionSet.DoesNotExist):
        pass
    texts = request.POST.getlist('text[]')
    if not texts:
        texts = [request.POST.get('text', '')]
    total = sum(get_character_count(t, ignore) for t in texts)
    return HttpResponse(json.dumps({'count': total}))


#########################################################################
# Category tags: editor-defined sub-distribution requirements
#########################################################################

def _tag_matches_path(tag_path, question_path):
    """A tag applies to a question when one path is a segment-wise prefix of
    the other: a 'Literature - World' tag covers questions in
    'Literature - World - Drama', and a tag on a deeper path than the
    question's leaf still applies to that leaf."""
    return (question_path == tag_path or
            question_path.startswith(tag_path + ' - ') or
            tag_path.startswith(question_path + ' - '))

def get_applicable_tags(qset, dist_entry):
    """Tags of the set that apply to a question in the given category."""
    if dist_entry is None:
        return []
    path = str(dist_entry)
    return [tag for tag in CategoryTag.objects.filter(question_set=qset).order_by('category_path', 'name')
            if _tag_matches_path(tag.category_path, path)]

def build_tag_checkboxes(qset, question, dist_entry):
    """Context rows for the tag checkboxes on the question edit pages."""
    tags = get_applicable_tags(qset, dist_entry)
    if not tags:
        return []
    if question is None:
        checked_ids = set()
    elif isinstance(question, Tossup):
        checked_ids = set(question.category_tags.values_list('id', flat=True))
    else:
        checked_ids = set(question.category_tags.values_list('id', flat=True))
    return [{'tag': tag, 'checked': tag.id in checked_ids} for tag in tags]

def save_tag_selection(request, qset, question, dist_entry, is_tossup):
    """Apply the 'category_tags' checkbox selection from a question edit POST."""
    selected = set()
    for value in request.POST.getlist('category_tags'):
        if value.isdigit():
            selected.add(int(value))
    for tag in get_applicable_tags(qset, dist_entry):
        relation = tag.tossups if is_tossup else tag.bonuses
        if tag.id in selected:
            relation.add(question)
        else:
            relation.remove(question)

@login_required
def category_tags(request, qset_id):
    user = request.user.writer
    qset = QuestionSet.objects.get(id=qset_id)

    if not qset.is_owner(user) and user not in qset.editor.all() and user not in qset.writer.all():
        return render(request, 'failure.html',
                                 {'message': 'You are not authorized to view this set!',
                                  'message_class': 'alert-box alert'})

    can_edit = qset.is_owner(user) or user in qset.editor.all()
    message = ''
    message_class = ''
    # Remember the category just used so the add form stays on it after saving
    selected_path = ''

    if request.method == 'POST':
        selected_path = request.POST.get('category_path', '')
        if not can_edit:
            message = 'Only editors can change tags!'
            message_class = 'alert-box warning'
        else:
            action = request.POST.get('action', '')
            try:
                if action == 'add':
                    path = request.POST.get('category_path', '').strip()
                    name = request.POST.get('name', '').strip()
                    num_tossups = int(request.POST.get('num_tossups') or 0)
                    num_bonuses = int(request.POST.get('num_bonuses') or 0)
                    if not path or not name:
                        raise ValueError('A category and a tag name are required')
                    tag, created = CategoryTag.objects.get_or_create(
                        question_set=qset, category_path=path, name=name,
                        defaults={'num_tossups': num_tossups, 'num_bonuses': num_bonuses})
                    if not created:
                        tag.num_tossups = num_tossups
                        tag.num_bonuses = num_bonuses
                        tag.save()
                    message = 'Tag "{0}" saved'.format(name)
                    message_class = 'alert-box success'
                elif action == 'delete':
                    CategoryTag.objects.filter(question_set=qset, id=int(request.POST['tag_id'])).delete()
                    message = 'Tag deleted'
                    message_class = 'alert-box success'
            except (ValueError, KeyError) as ex:
                message = str(ex) or 'Invalid request'
                message_class = 'alert-box warning'

    # Category path choices from the set's distribution tree
    path_choices = [row['path'] for row in get_packetization_rows(qset)]

    # Group tags by category path with completion status
    groups = []
    tags_by_path = {}
    for tag in CategoryTag.objects.filter(question_set=qset).order_by('category_path', 'name').prefetch_related('tossups', 'bonuses'):
        tags_by_path.setdefault(tag.category_path, []).append(tag)
    for path in sorted(tags_by_path):
        rows = []
        for tag in tags_by_path[path]:
            tossups = [{'id': t.id,
                        'answer': _grid_answer_preview(t.tossup_answer),
                        'location': '{0} #{1}'.format(t.packet.packet_name, t.question_number) if t.packet else 'Unassigned'}
                       for t in tag.tossups.all().select_related('packet')]
            bonuses = [{'id': b.id,
                        'answer': ' / '.join(filter(None, [
                            _grid_answer_preview(b.part1_answer, 20),
                            _grid_answer_preview(b.part2_answer, 20),
                            _grid_answer_preview(b.part3_answer, 20)])),
                        'location': '{0} #{1}'.format(b.packet.packet_name, b.question_number) if b.packet else 'Unassigned'}
                       for b in tag.bonuses.all().select_related('packet')]
            tu_done = len(tossups)
            bs_done = len(bonuses)
            rows.append({
                'tag': tag,
                'tossups': tossups,
                'bonuses': bonuses,
                'tu_done': tu_done,
                'bs_done': bs_done,
                'tu_complete': tag.num_tossups == 0 or tu_done >= tag.num_tossups,
                'bs_complete': tag.num_bonuses == 0 or bs_done >= tag.num_bonuses,
            })
        groups.append({'path': path, 'rows': rows})

    return render(request, 'category_tags.html',
                             {'qset': qset,
                              'user': user,
                              'groups': groups,
                              'path_choices': path_choices,
                              'can_edit': can_edit,
                              'selected_path': selected_path,
                              'message': message,
                              'message_class': message_class})


def _member_or_403(request, qset):
    """Return the writer if they belong to the set, else None."""
    user = request.user.writer
    if qset.is_owner(user) or user in qset.editor.all() or user in qset.writer.all():
        return user
    return None


@login_required
def recap(request, qset_id):
    """Set-wide "what's changed lately" digest: new questions, recent edits,
    and recent comments across the whole set over a selectable time window."""
    qset = QuestionSet.objects.get(id=qset_id)
    user = _member_or_403(request, qset)
    if user is None:
        return render(request, 'failure.html',
                      {'message': 'You are not authorized to view this set!',
                       'message_class': 'alert-box alert'})

    try:
        days = int(request.GET.get('days', 7))
    except ValueError:
        days = 7
    if days not in (1, 3, 7, 14, 30):
        days = 7
    from datetime import timedelta
    since = timezone.now() - timedelta(days=days)

    def location(q):
        if q.packet_id and q.question_number:
            return '{0} #{1}'.format(q.packet.packet_name, q.question_number)
        return 'Unassigned'

    # --- New questions created in the window ---
    new_questions = []
    new_tu = (Tossup.objects.filter(question_set=qset, created_date__gte=since)
              .select_related('packet', 'author__user', 'category').order_by('-created_date'))
    new_bs = (Bonus.objects.filter(question_set=qset, created_date__gte=since)
              .select_related('packet', 'author__user', 'category').order_by('-created_date'))
    for t in new_tu:
        new_questions.append({
            'date': t.created_date, 'qtype': 'tossup',
            'answer': _grid_answer_preview(t.tossup_answer),
            'author': str(t.author) if t.author else 'unknown',
            'category': str(t.category) if t.category else '',
            'location': location(t), 'edit_url': '/edit_tossup/{0}/'.format(t.id)})
    for b in new_bs:
        new_questions.append({
            'date': b.created_date, 'qtype': 'bonus',
            'answer': _grid_answer_preview(b.part1_answer),
            'author': str(b.author) if b.author else 'unknown',
            'category': str(b.category) if b.category else '',
            'location': location(b), 'edit_url': '/edit_bonus/{0}/'.format(b.id)})
    new_questions.sort(key=lambda x: x['date'], reverse=True)

    # --- Recent edits (history rows that aren't the question's creation) ---
    edit_items = []
    for model, hist_model, edit_url, preview in (
            (Tossup, TossupHistory, '/edit_tossup/', lambda q: _grid_answer_preview(q.tossup_answer)),
            (Bonus, BonusHistory, '/edit_bonus/', lambda q: _grid_answer_preview(q.part1_answer))):
        qs = model.objects.filter(question_set=qset).select_related('packet')
        hist_to_q = {q.question_history_id: q for q in qs if q.question_history_id}
        if not hist_to_q:
            continue
        histories = (hist_model.objects.filter(question_history_id__in=hist_to_q.keys(),
                                               change_date__gte=since)
                     .select_related('changer__user').order_by('-change_date')[:200])
        for h in histories:
            q = hist_to_q.get(h.question_history_id)
            # Skip the very first history row (creation) so this is edits only.
            if q is None or (q.created_date and h.change_date <= q.created_date):
                continue
            edit_items.append({
                'date': h.change_date, 'qtype': 'tossup' if model is Tossup else 'bonus',
                'by': str(h.changer) if h.changer else 'unknown',
                'answer': preview(q), 'location': location(q),
                'edit_url': '{0}{1}/'.format(edit_url, q.id)})
    edit_items.sort(key=lambda x: x['date'], reverse=True)
    edit_items = edit_items[:200]

    # --- Recent comments across the set ---
    tu_ct = ContentType.objects.get_for_model(Tossup)
    bs_ct = ContentType.objects.get_for_model(Bonus)
    tu_ids = [str(i) for i in qset.tossup_set.values_list('id', flat=True)]
    bs_ids = [str(i) for i in qset.bonus_set.values_list('id', flat=True)]
    comments = (Comment.objects.filter(is_removed=False, submit_date__gte=since)
                .filter(Q(content_type=tu_ct, object_pk__in=tu_ids) |
                        Q(content_type=bs_ct, object_pk__in=bs_ids))
                .select_related('user', 'content_type').order_by('-submit_date')[:200])
    comment_items = []
    for c in comments:
        is_tu = c.content_type_id == tu_ct.id
        by = '{0} {1}'.format(c.user.first_name, c.user.last_name).strip() if c.user else ''
        by = by or (c.user_name or 'unknown')
        comment_items.append({
            'date': c.submit_date, 'by': by, 'text': c.comment,
            'qtype': 'tossup' if is_tu else 'bonus',
            'edit_url': '{0}{1}/'.format('/edit_tossup/' if is_tu else '/edit_bonus/', c.object_pk)})

    return render(request, 'recap.html',
                  {'qset': qset, 'user': user, 'days': days,
                   'new_questions': new_questions,
                   'edit_items': edit_items,
                   'comment_items': comment_items,
                   'new_question_count': len(new_questions),
                   'edit_count': len(edit_items),
                   'comment_count': len(comment_items)})


def _tossup_reading(tossup):
    """Turn a tossup into the data the play UI reads clue-by-clue: a list of
    display words, the index of the power boundary (the word containing the
    "(*)" power mark, or -1 if unmarked), and the answer as formatted HTML."""
    plain = strip_markup(tossup.tossup_text or '').replace('\n', ' ').strip()
    raw_words = [w for w in plain.split(' ') if w != '']
    power_index = -1
    words = []
    for i, w in enumerate(raw_words):
        if '(*)' in w and power_index == -1:
            power_index = len(words)
            w = w.replace('(*)', '').strip()
            if w == '':
                continue
        words.append(w)
    answer_html = get_formatted_question_html(tossup.tossup_answer, True, True, False, False)
    return {
        'id': tossup.id,
        'qtype': 'tossup',
        'number': tossup.question_number or 0,
        'words': words,
        'power_index': power_index,
        'answer_html': answer_html,
        'category': str(tossup.category) if tossup.category else '',
    }


def _bonus_reading(bonus):
    """Turn a bonus into the data the play UI reveals part-by-part."""
    is_acf = bonus.get_bonus_type() == ACF_STYLE_BONUS
    leadin_html = get_formatted_question_html(bonus.leadin, False, True, False, False) if is_acf else ''
    parts = []
    fields = [(bonus.part1_text, bonus.part1_answer, bonus.part1_difficulty),
              (bonus.part2_text, bonus.part2_answer, bonus.part2_difficulty),
              (bonus.part3_text, bonus.part3_answer, bonus.part3_difficulty)]
    if not is_acf:
        fields = fields[:1]
    for text, answer, difficulty in fields:
        if text is None or text == '':
            continue
        parts.append({
            'text_html': get_formatted_question_html(text, False, True, False, False),
            'answer_html': get_formatted_question_html(answer, True, True, False, False),
            'difficulty': difficulty or '',
        })
    return {
        'id': bonus.id,
        'qtype': 'bonus',
        'number': bonus.question_number or 0,
        'leadin_html': leadin_html,
        'parts': parts,
        'category': str(bonus.category) if bonus.category else '',
    }


@login_required
def play(request, qset_id):
    """Play the set's questions clue-by-clue (tossups) / part-by-part (bonuses).
    The player chooses what to play: recent questions (optionally filtered by
    category) or questions by category, and whether to play tossups, bonuses, or
    both. Buzzes/results are recorded via AJAX (record_buzz/record_bonus_result)."""
    qset = QuestionSet.objects.get(id=qset_id)
    user = _member_or_403(request, qset)
    if user is None:
        return render(request, 'failure.html',
                      {'message': 'You are not authorized to play this set!',
                       'message_class': 'alert-box alert'})

    mode = request.GET.get('mode', 'recent')
    if mode not in ('recent', 'category'):
        mode = 'recent'
    qtypes = request.GET.get('qtypes', 'both')
    if qtypes not in ('both', 'tossups', 'bonuses'):
        qtypes = 'both'
    if qset.tossups_only:
        qtypes = 'tossups'
    try:
        limit = int(request.GET.get('limit', 30))
    except ValueError:
        limit = 30
    limit = max(1, min(limit, 200))
    # In recent mode, play either the most-recent N questions ('count') or every
    # question created in the last N days ('days').
    recent_by = request.GET.get('recent_by', 'count')
    if recent_by not in ('count', 'days'):
        recent_by = 'count'
    try:
        days = int(request.GET.get('days', 7))
    except ValueError:
        days = 7
    days = max(1, min(days, 365))
    selected_cat_ids = [c for c in request.GET.getlist('cats') if c.isdigit()]

    # Categories actually used by this set's questions, for the filter UI.
    used_cat_ids = set(Tossup.objects.filter(question_set=qset, category__isnull=False)
                       .values_list('category_id', flat=True))
    used_cat_ids |= set(Bonus.objects.filter(question_set=qset, category__isnull=False)
                        .values_list('category_id', flat=True))
    categories = sorted(
        ({'id': de.id, 'name': str(de)}
         for de in DistributionEntry.objects.filter(id__in=used_cat_ids)),
        key=lambda c: c['name'].lower())

    def gather(model):
        qs = model.objects.filter(question_set=qset).select_related(
            'category', 'packet', 'question_type')
        if selected_cat_ids:
            qs = qs.filter(category_id__in=selected_cat_ids)
        if mode == 'recent':
            qs = qs.order_by('-created_date')
            if recent_by == 'days':
                from datetime import timedelta
                since = timezone.now() - timedelta(days=days)
                # Cap to keep the page reasonable; this is the most-recent slice.
                return qs.filter(created_date__gte=since)[:200]
            return qs[:limit]
        return qs.order_by('category__category', 'category__subcategory',
                           'packet__packet_name', 'question_number')[:200]

    tossups, bonuses = [], []
    if qtypes in ('both', 'tossups'):
        for t in gather(Tossup):
            r = _tossup_reading(t)
            r['edit_url'] = '/edit_tossup/{0}/'.format(t.id)
            tossups.append(r)
    if qtypes in ('both', 'bonuses'):
        for b in gather(Bonus):
            r = _bonus_reading(b)
            r['edit_url'] = '/edit_bonus/{0}/'.format(b.id)
            bonuses.append(r)

    return render(request, 'play.html',
                  {'qset': qset, 'user': user,
                   'mode': mode, 'qtypes': qtypes, 'limit': limit,
                   'recent_by': recent_by, 'days': days,
                   'categories': categories,
                   'selected_cat_ids': [int(c) for c in selected_cat_ids],
                   'tossups_only': qset.tossups_only,
                   'questions': {'tossups': tossups, 'bonuses': bonuses}})


def _get_or_create_session(request, qset, session_id):
    """Reuse the play session identified by session_id if it belongs to this
    user and set, otherwise start a new one."""
    user = request.user.writer
    if session_id:
        session = PlaytestSession.objects.filter(id=session_id, question_set=qset,
                                                 player=user).first()
        if session is not None:
            return session
    return PlaytestSession.objects.create(question_set=qset, player=user,
                                          source=PLAYTEST_SOURCE_WEB)


@login_required
def record_buzz(request):
    """Record a tossup buzz from the play UI. Returns the session id so the
    client can group later buzzes from the same sitting."""
    if request.method != 'POST':
        return HttpResponse(json.dumps({'success': False, 'message': 'Invalid request'}))
    user = request.user.writer
    try:
        tossup = Tossup.objects.select_related('question_set').get(id=int(request.POST['tossup_id']))
    except (KeyError, ValueError, Tossup.DoesNotExist):
        return HttpResponse(json.dumps({'success': False, 'message': 'Tossup not found'}))

    qset = tossup.question_set
    if _member_or_403(request, qset) is None:
        return HttpResponse(json.dumps({'success': False, 'message': 'Not authorized'}))

    correct = request.POST.get('correct') == 'true'
    powered = request.POST.get('powered') == 'true'
    neg = request.POST.get('neg') == 'true'
    try:
        buzz_word_index = int(request.POST.get('buzz_word_index') or 0)
        total_words = int(request.POST.get('total_words') or 0)
        char_position = int(request.POST.get('char_position') or 0)
    except ValueError:
        return HttpResponse(json.dumps({'success': False, 'message': 'Bad buzz position'}))

    if correct:
        value = 15 if powered else 10
    elif neg:
        value = -5
    else:
        value = 0

    session = _get_or_create_session(request, qset, request.POST.get('session_id'))
    TossupBuzz.objects.create(
        tossup=tossup, session=session, player=user,
        buzz_word_index=buzz_word_index, total_words=total_words,
        char_position=char_position, correct=correct, powered=powered and correct,
        value=value, answer_given=request.POST.get('answer_given', '')[:1000],
        tossup_history=tossup.latest_history(), source=PLAYTEST_SOURCE_WEB)

    return HttpResponse(json.dumps({'success': True, 'session_id': session.id, 'value': value}))


@login_required
def record_bonus_result(request):
    """Record the result of playing a bonus from the play UI."""
    if request.method != 'POST':
        return HttpResponse(json.dumps({'success': False, 'message': 'Invalid request'}))
    user = request.user.writer
    try:
        bonus = Bonus.objects.select_related('question_set').get(id=int(request.POST['bonus_id']))
    except (KeyError, ValueError, Bonus.DoesNotExist):
        return HttpResponse(json.dumps({'success': False, 'message': 'Bonus not found'}))

    qset = bonus.question_set
    if _member_or_403(request, qset) is None:
        return HttpResponse(json.dumps({'success': False, 'message': 'Not authorized'}))

    p1 = request.POST.get('part1_correct') == 'true'
    p2 = request.POST.get('part2_correct') == 'true'
    p3 = request.POST.get('part3_correct') == 'true'
    total = 10 * sum((p1, p2, p3))

    session = _get_or_create_session(request, qset, request.POST.get('session_id'))
    BonusResult.objects.create(
        bonus=bonus, session=session, player=user,
        part1_correct=p1, part2_correct=p2, part3_correct=p3, total=total,
        bonus_history=bonus.latest_history(), source=PLAYTEST_SOURCE_WEB)

    return HttpResponse(json.dumps({'success': True, 'session_id': session.id, 'total': total}))


def _question_buzz_data(question, qtype):
    """Aggregate playtest results for a single question, shown as a panel on its
    edit page (buzz stats are a property of the question, not a separate list).
    Returns None when nothing has been recorded yet."""
    if question is None:
        return None

    if qtype == 'tossup':
        buzzes = list(question.buzzes.select_related('player__user').order_by('buzz_date'))
        if not buzzes:
            return None
        correct = [b for b in buzzes if b.correct]
        powers = [b for b in correct if b.powered]
        negs = [b for b in buzzes if b.value < 0]
        fracs = [b.buzz_fraction() for b in correct]
        words = _tossup_reading(question)['words']
        rows = []
        for b in buzzes:
            heard = ' '.join(words[:b.buzz_word_index]) if words else ''
            if len(heard) > 140:
                heard = '...' + heard[-140:]
            rows.append({
                'player': b.get_player_name(), 'date': b.buzz_date,
                'correct': b.correct, 'powered': b.powered, 'value': b.value,
                'fraction': '{0:.0f}%'.format(100.0 * b.buzz_fraction()),
                'answer_given': b.answer_given, 'heard': heard, 'source': b.source,
                'history_url': b.history_url()})
        return {
            'qtype': 'tossup', 'plays': len(buzzes), 'correct': len(correct),
            'powers': len(powers), 'negs': len(negs),
            'conversion': '{0:.0f}%'.format(100.0 * len(correct) / len(buzzes)),
            'avg_buzz': '{0:.0f}%'.format(100.0 * sum(fracs) / len(fracs)) if fracs else '—',
            'rows': rows}

    results = list(question.results.select_related('player__user').order_by('answered_date'))
    if not results:
        return None
    n = len(results)
    rows = [{
        'player': r.get_player_name(), 'date': r.answered_date,
        'p1': r.part1_correct, 'p2': r.part2_correct, 'p3': r.part3_correct,
        'total': r.total, 'source': r.source, 'history_url': r.history_url()}
        for r in results]
    return {
        'qtype': 'bonus', 'plays': n,
        'avg_points': '{0:.1f}'.format(sum(r.total for r in results) / n),
        'p1': '{0:.0f}%'.format(100.0 * sum(1 for r in results if r.part1_correct) / n),
        'p2': '{0:.0f}%'.format(100.0 * sum(1 for r in results if r.part2_correct) / n),
        'p3': '{0:.0f}%'.format(100.0 * sum(1 for r in results if r.part3_correct) / n),
        'rows': rows}


@login_required
def api_access(request, qset_id):
    """Owner/co-owner page to view and manage the set's Discord-bot API key."""
    qset = QuestionSet.objects.get(id=qset_id)
    user = request.user.writer
    if not qset.is_owner(user):
        return render(request, 'failure.html',
                      {'message': 'Only the set owner or a co-owner can manage API access.',
                       'message_class': 'alert-box alert'})
    api_key = SetApiKey.objects.filter(question_set=qset).first()
    base_url = request.build_absolute_uri('/').rstrip('/')
    return render(request, 'api_access.html',
                  {'qset': qset, 'user': user, 'api_key': api_key, 'base_url': base_url})


@login_required
def generate_set_api_key(request):
    """Create, rotate, or revoke a set's API key (owner/co-owner only)."""
    if request.method != 'POST':
        return HttpResponseRedirect('/main/')
    user = request.user.writer
    qset = QuestionSet.objects.get(id=int(request.POST['qset_id']))
    if not qset.is_owner(user):
        return render(request, 'failure.html',
                      {'message': 'Only the set owner or a co-owner can manage API access.',
                       'message_class': 'alert-box alert'})
    action = request.POST.get('action', 'generate')
    if action == 'revoke':
        SetApiKey.objects.filter(question_set=qset).delete()
        messages.success(request, 'API key revoked.')
    else:
        api_key, _ = SetApiKey.objects.get_or_create(
            question_set=qset, defaults={'key': SetApiKey.generate_token(), 'created_by': user})
        # Regenerate always issues a fresh token (revoking the old one).
        api_key.key = SetApiKey.generate_token()
        api_key.active = True
        api_key.created_by = user
        api_key.save()
        messages.success(request, 'A new API key has been generated.')
    return HttpResponseRedirect('/api_access/{0}/'.format(qset.id))
