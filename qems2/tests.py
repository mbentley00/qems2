from django.test import TestCase
from django.contrib.auth.models import AnonymousUser, User
from datetime import datetime
from django.utils import timezone

from qems2.qsub.packet_parser import is_answer, is_bpart, is_vhsl_bpart, is_category
from qems2.qsub.packet_parser import parse_packet_data, get_bonus_part_value, remove_category
from qems2.qsub.packet_parser import remove_answer_label
from qems2.qsub.utils import get_character_count, get_formatted_question_html, are_special_characters_balanced
from qems2.qsub.models import *
from qems2.qsub.model_utils import *
from qems2.qsub.packetizer import *
from qems2.qsub.packetizer import _ensure_packets
from decimal import Decimal

class PacketParserTests(TestCase):

    # TODO: Determine if we really need this block of code anymore
    #if django.VERSION[:2] == (1, 7):
    #    # Django 1.7 requires an explicit setup() when running tests in PTVS
    #    @classmethod
    #    def setUpClass(cls):
    #        django.setup()
    #elif django.VERSION[:2] >= (1, 8):
    #    # Django 1.8 requires a different setup. See https://github.com/Microsoft/PTVS-Samples/issues/1
    #    @classmethod
    #    def setUpClass(cls):
    #        super(DjangoTestCase, cls).setUpClass()
    #        django.setup()

    user = None
    writer = None
    dist = None
    qset = None
    pwe = None
    packet = None
    period = None
    ce = None
    cefd = None
    pwce = None
    opce = None
    acf_tossup = None
    acf_bonus = None
    vhsl_bonus = None
    
    acf_reg_dist = None
    acf_tb_dist = None
    acf_qset = None
    
    acf_reg_pwe = None
    acf_tb_pwe = None
    
    acf_packets = []
    acf_reg_periods = []
    acf_tb_periods = []
    
    acf_ces = []
    acf_reg_cefds = []
    acf_tb_cefds = []
    
    #########################################################################
    # Setup and helper methods
    #########################################################################
    
    def setUp(self):
        acfTossup = QuestionType.objects.get_or_create(question_type=ACF_STYLE_TOSSUP)
        acfBonus = QuestionType.objects.get_or_create(question_type=ACF_STYLE_BONUS)
        vhslBonus = QuestionType.objects.get_or_create(question_type=VHSL_BONUS)

    def create_user(self):
        self.user, created = User.objects.get_or_create(username="testuser")
        if (created):
            self.user.email='qems2test@gmail.com'
            self.user.password='top_secret'
            self.user.save()
            
        self.writer = Writer.objects.get(user=self.user.id)

    def create_generic_distribution(self, per_period_totals=20):
        self.dist = Distribution.objects.create(name="Test Distribution", acf_tossup_per_period_count=per_period_totals, acf_bonus_per_period_count=per_period_totals, vhsl_bonus_per_period_count=per_period_totals)
        self.dist.save()
        
    def create_generic_qset(self, num_packets=10):
        self.create_user()
        self.create_generic_distribution()
        self.qset = QuestionSet.objects.create(
            name="new_set",
            date=timezone.now(),
            host="test host",
            owner=self.writer,
            num_packets=num_packets,
            distribution=self.dist)
        self.qset.save()

    def create_generic_period(self, pwe_cur_value=5, pwe_total_value=10, cefd_fraction=5, cefd_min=15, cefd_max=15, period_cur_value=10, pwce_cur_value=10, pwce_total_value=10, opce_cur_value=10):
        self.create_generic_qset()
        
        self.pwe = PeriodWideEntry.objects.create(period_type=ACF_REGULAR_PERIOD, question_set=self.qset, distribution=self.dist)
        self.pwe.acf_tossup_cur=pwe_cur_value
        self.pwe.acf_bonus_cur=pwe_cur_value 
        self.pwe.vhsl_bonus_cur=pwe_cur_value
        self.pwe.acf_tossup_total=pwe_total_value
        self.pwe.acf_bonus_total=pwe_total_value
        self.pwe.vhsl_bonus_total=pwe_total_value
        self.pwe.save()
                    
        self.packet = Packet.objects.create(packet_name="Test Packet", question_set=self.qset, created_by=self.writer)
        self.packet.save()
        
        self.period = Period.objects.create(name="Test Period", packet=self.packet, period_wide_entry=self.pwe, acf_tossup_cur=period_cur_value, acf_bonus_cur=period_cur_value, vhsl_bonus_cur=period_cur_value)
        self.period.save()
        
        self.ce, created = CategoryEntry.objects.get_or_create(category_name="Test Category", category_type=CATEGORY)
        self.ce.save()
        
        self.cefd = CategoryEntryForDistribution.objects.create(
            distribution=self.dist,
            category_entry=self.ce, 
            acf_tossup_fraction=cefd_fraction,
            acf_bonus_fraction=cefd_fraction,
            vhsl_bonus_fraction=cefd_fraction,
            min_total_questions_in_period=cefd_min,
            max_total_questions_in_period=cefd_max)
        self.cefd.save()
        
        self.pwce = PeriodWideCategoryEntry.objects.create(
            period_wide_entry=self.pwe,
            category_entry_for_distribution=self.cefd,
            acf_tossup_cur_across_periods=pwce_cur_value,
            acf_bonus_cur_across_periods=pwce_cur_value,
            vhsl_bonus_cur_across_periods=pwce_cur_value,
            acf_tossup_total_across_periods=pwce_total_value,
            acf_bonus_total_across_periods=pwce_total_value,
            vhsl_bonus_total_across_periods=pwce_total_value)
        self.pwce.save()
        
        self.opce = OnePeriodCategoryEntry.objects.create(
            period=self.period,
            period_wide_category_entry=self.pwce,
            acf_tossup_cur_in_period=opce_cur_value,
            acf_bonus_cur_in_period=opce_cur_value,
            vhsl_bonus_cur_in_period=opce_cur_value)
        self.opce.save()

    # Returns 3 tuple of tuples.  The first layer of tuples is just the list of 
    # each category entry, the second is a CategoryEntry, a CategoryEntryForDistribution, 
    # a PeriodWideCategoryEntry and a OnePeriodCategoryEntry that all relate back to the CategoryEntry
    def create_generic_ce_hierarchy(self):
        
        cat_tuple = self.create_ce_and_dependencies(
            category_type=CATEGORY,
            cefd_fraction=4.2,
            cefd_min=12,
            cefd_max=13,
            pwce_cur_value=0,
            pwce_total_value=13, # TODO: Double check that this is right
            opce_cur_value=0,
            category_name="Literature")

        sub_cat_tuple1 = self.create_ce_and_dependencies(
            category_type=SUB_CATEGORY,
            cefd_fraction=2,
            cefd_min=6,
            cefd_max=6,
            pwce_cur_value=0,
            pwce_total_value=6, # TODO: Double check that this is right            
            opce_cur_value=0,
            category_name="Literature",
            sub_category_name="American")

        sub_cat_tuple2 = self.create_ce_and_dependencies(
            category_type=SUB_CATEGORY,
            cefd_fraction=2.2,
            cefd_min=6,
            cefd_max=7,
            pwce_cur_value=0,
            pwce_total_value=7, # TODO: Double check that this is right            
            opce_cur_value=0,
            category_name="Literature",
            sub_category_name="European")

        sub_sub_cat_tuple1 = self.create_ce_and_dependencies(
            category_type=SUB_SUB_CATEGORY,
            cefd_fraction=2,
            cefd_min=6,
            cefd_max=6,
            pwce_cur_value=0,
            pwce_total_value=6, # TODO: Double check that this is right            
            opce_cur_value=0,
            category_name="Literature",
            sub_category_name="American",
            sub_sub_category_name="Novels")

        sub_sub_cat_tuple2 = self.create_ce_and_dependencies(
            category_type=SUB_SUB_CATEGORY,
            cefd_fraction=1,
            cefd_min=3,
            cefd_max=3,
            pwce_cur_value=0,
            pwce_total_value=3, # TODO: Double check that this is right            
            opce_cur_value=0,
            category_name="Literature",
            sub_category_name="European",
            sub_sub_category_name="Novels")

        sub_sub_cat_tuple3 = self.create_ce_and_dependencies(
            category_type=SUB_SUB_CATEGORY,
            cefd_fraction=1.2,
            cefd_min=3,
            cefd_max=4,
            pwce_cur_value=0,
            pwce_total_value=4, # TODO: Double check that this is right            
            opce_cur_value=0,
            category_name="Literature",
            sub_category_name="European",
            sub_sub_category_name="Poetry")
            
        cats = (cat_tuple, None) # TODO: Change
        sub_cats = (sub_cat_tuple1, sub_cat_tuple2)
        sub_sub_cats = (sub_sub_cat_tuple1, sub_sub_cat_tuple2, sub_sub_cat_tuple3)
        return cats, sub_cats, sub_sub_cats

    # Creates a CategoryEntry, a CategoryEntryForDistribution, a PeriodWideCategoryEntry and a
    # OnePeriodCategoryEntry that all relate back to the CategoryEntry
    def create_ce_and_dependencies(
        self, 
        category_type, 
        cefd_fraction, 
        cefd_min, 
        cefd_max, 
        pwce_cur_value, 
        pwce_total_value, 
        opce_cur_value, 
        category_name, 
        sub_category_name=None, 
        sub_sub_category_name=None,
        dist=None,
        pwe=None,
        period=None):
            
        if (dist is None):
            dist=self.dist
            
        if (pwe is None):
            pwe=self.pwe
            
        if (period is None):
            period=self.period
            
        ce, cefd = self.create_just_ce_and_cefd(
            dist=dist,
            category_type=category_type,
            cefd_fraction=cefd_fraction,
            cefd_min=cefd_min,
            cefd_max=cefd_max,
            category_name=category_name,
            sub_category_name=sub_category_name,
            sub_sub_category_name=sub_sub_category_name)
        
        # TODO: Maybe put these into a separate method since don't always want to create them at same time as above
        pwce = PeriodWideCategoryEntry.objects.create(
            period_wide_entry=pwe,
            category_entry_for_distribution=cefd,
            acf_tossup_cur_across_periods=pwce_cur_value,
            acf_bonus_cur_across_periods=pwce_cur_value,
            vhsl_bonus_cur_across_periods=pwce_cur_value,
            acf_tossup_total_across_periods=pwce_total_value,
            acf_bonus_total_across_periods=pwce_total_value,
            vhsl_bonus_total_across_periods=pwce_total_value)
        pwce.save()
        
        opce = OnePeriodCategoryEntry.objects.create(
            period=period,
            period_wide_category_entry=pwce,
            acf_tossup_cur_in_period=opce_cur_value,
            acf_bonus_cur_in_period=opce_cur_value,
            vhsl_bonus_cur_in_period=opce_cur_value)
        opce.save()
        
        ce_tuple = (ce, cefd, pwce, opce)
        return ce_tuple
        
    # Creats a category entry and a category entry for distribution
    def create_just_ce_and_cefd(
        self,
        dist,
        category_type, 
        cefd_fraction, 
        cefd_min, 
        cefd_max,         
        category_name, 
        sub_category_name=None, 
        sub_sub_category_name=None):
            
        ce = None
        if (category_type==CATEGORY):
            ce, created = CategoryEntry.objects.get_or_create(category_name=category_name, category_type=category_type)            
        elif (category_type==SUB_CATEGORY):
            ce, created = CategoryEntry.objects.get_or_create(category_name=category_name, sub_category_name=sub_category_name, category_type=category_type)                        
        else:
            ce, created = CategoryEntry.objects.get_or_create(category_name=category_name, sub_category_name=sub_category_name, sub_sub_category_name=sub_sub_category_name, category_type=category_type)            
                    
        cefd = CategoryEntryForDistribution.objects.create(
            distribution=dist,
            category_entry=ce, 
            acf_tossup_fraction=cefd_fraction,
            acf_bonus_fraction=cefd_fraction,
            vhsl_bonus_fraction=cefd_fraction,
            min_total_questions_in_period=cefd_min,
            max_total_questions_in_period=cefd_max)
        cefd.save()
        
        return ce, cefd
    
        
    #########################################################################
    # Unit tests that don't depend on the database
    #########################################################################
                            
#    def test_is_answer(self):
#        answers = ["answer:", "Answer:", "ANSWER:", "ANSWER: _underlined_", "ANSWER: no underline", "ANSWER: <u>underline2</u>"]
#        for answer in answers:
#            self.assertTrue(is_answer(answer), msg=answer)
#        non_answers = ["question:", "answer", "ansER"]
#        for non_answer in non_answers:
#            self.assertFalse(is_answer(non_answer), msg=non_answer)
#    def test_remove_answer_label(self):
#        answers = ["ANSWER: <u><b>my answer</b></u>", "answer:      <u><b>my answer</b></u>", "Answer:\t<u><b>my answer</b></u>"]
#        for answer in answers:
#            self.assertEqual(remove_answer_label(answer), '<u><b>my answer</b></u>')
#    def test_are_special_characters_balanced(self):
#        balancedLines = ["", "No special chars", "_Underscores_", "~Italics~", "(Parens)", "_~Several_~ (items) in (one) _question_."]
#        unbalancedLines = ["_", "~", "_test__", "~~test~", "(test", "test)", "((test)", "(", ")", ")test(", "(test))"]
#        for balancedLine in balancedLines:
#            self.assertTrue(are_special_characters_balanced(balancedLine))
#        for unbalancedLine in unbalancedLines:
#            self.assertFalse(are_special_characters_balanced(unbalancedLine))
#    def test_is_bpart(self):
#        bonusParts = ['[10]', '[15]']
#        for bonusPart in bonusParts:
#            self.assertTrue(is_bpart(bonusPart), msg=bonusPart)
#        notBonusParts = ['(10)', '10', '[10', '10]', '(10]', '[10)', '[or foo]', '(not a number)', '[10.5]', '[<i>10</i>]']
#        for notBonus in notBonusParts:
#            self.assertFalse(is_bpart(notBonus), msg=notBonus)
#    def test_is_vhsl_bpart(self):
#        bonusParts = ['[V10]', '[V15]']
#        for bonusPart in bonusParts:
#            self.assertTrue(is_vhsl_bpart(bonusPart), msg=bonusPart)
#        notBonusParts = ['(V10)', '[10]', '(10)', 'V10', '[10', '10]', '(10]', '[V10)', '[or foo]', '(not a number)']
#        for notBonus in notBonusParts:
#            self.assertFalse(is_vhsl_bpart(notBonus), msg=notBonus)
#    def test_is_category(self):
#        categories = ["{History - European}, 'ANSWER: _foo_ {Literature - American}"]
#        for category in categories:
#            self.assertTrue(is_category(category), msg=category)
#        notCategories = ["answer: _foo_", '{History - World', 'History - Other}']
#        for notCategory in notCategories:
#            self.assertFalse(is_category(notCategory), msg=notCategory)
#    def test_remove_category(self):
#        self.assertEqual(remove_category('ANSWER: _foo_ {History - European}'), 'ANSWER: _foo_ ')
#        self.assertEqual(remove_category('ANSWER: _foo_ {History - European'), 'ANSWER: _foo_ {History - European')
#    def test_get_bonus_part_value(self):
#        bonusParts = ['[10]']
#        for bonusPart in bonusParts:
#            self.assertEqual(get_bonus_part_value(bonusPart), '10')
#
#    def test_get_character_count(self):
#        emptyTossup = ""
#        self.assertEqual(get_character_count(emptyTossup), 0)
#
#        noSpecialCharacters = "123456789"
#        self.assertEqual(get_character_count(noSpecialCharacters), 9)
#
#        onlySpecialCharacters = "~~()"
#        self.assertEqual(get_character_count(onlySpecialCharacters), 0)
#
#        mixed = "(~1234~) ~67~"
#        self.assertEqual(get_character_count(mixed), 3)
#
#    def test_get_formatted_question_html(self):
#        emptyLine = ""
#        self.assertEqual(get_formatted_question_html(emptyLine, False, True, False), "")
#        self.assertEqual(get_formatted_question_html(emptyLine, True, True, False), "")
#
#        noSpecialChars = "No special chars"
#        self.assertEqual(get_formatted_question_html(noSpecialChars, False, True, False), noSpecialChars)
#        self.assertEqual(get_formatted_question_html(noSpecialChars, True, True, False), noSpecialChars)
#
#        specialChars = "_Underlines_, ~italics~ and (parens).  And again _Underlines_, ~italics~ and (parens)."
#        self.assertEqual(get_formatted_question_html(specialChars, False, True, False), "_Underlines_, <i>italics</i> and <strong class="pronunciation-guide">(parens)</strong>.  And again _Underlines_, <i>italics</i> and <strong class="pronunciation-guide">(parens)</strong>.")
#        self.assertEqual(get_formatted_question_html(specialChars, True, True, False), "<u><b>Underlines</b></u>, <i>italics</i> and <strong class="pronunciation-guide">(parens)</strong>.  And again <u><b>Underlines</b></u>, <i>italics</i> and <strong class="pronunciation-guide">(parens)</strong>.")
#
#        newLinesNoParens = "(No parens).&lt;br&gt;New line."
#        self.assertEqual(get_formatted_question_html(newLinesNoParens, False, False, False), "(No parens).&lt;br&gt;New line.")
#        self.assertEqual(get_formatted_question_html(newLinesNoParens, False, False, True), "(No parens).<br />New line.")
#
#    def test_does_answerline_have_underlines(self):
#        self.assertFalse(does_answerline_have_underlines("ANSWER: Foo"))
#        self.assertTrue(does_answerline_have_underlines("ANSWER: _Foo_"))
#        self.assertTrue(does_answerline_have_underlines(""))
#
#            
#    #########################################################################
#    # Tests with database dependencies
#    #########################################################################            
#            
#    def test_parse_packet_data(self):
#        self.create_generic_distribution()
#
#        euroHistory = DistributionEntry(category="History", subcategory="European", distribution=self.dist)
#        euroHistory.save()
#
#        americanHistory = DistributionEntry(category="History", subcategory="American", distribution=self.dist)
#        americanHistory.save()
#
#        validTossup = 'This is a valid test tossup.\nANSWER: _My Answer_'
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(validTossup.splitlines())
#        self.assertEqual(len(tossups), 1)
#        self.assertEqual(tossups[0].tossup_text, 'This is a valid test tossup.');
#        self.assertEqual(tossups[0].tossup_answer, '_My Answer_');
#        self.assertEqual(tossups[0].category, None);
#
#        validTossupWithCategory = 'This should be a ~European History~ tossup.\nANSWER: _Charles I_ {History - European}'
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(validTossupWithCategory.splitlines())
#        self.assertEqual(len(tossups), 1)
#        self.assertEqual(tossups[0].tossup_text, 'This should be a ~European History~ tossup.');
#        self.assertEqual(tossups[0].tossup_answer, '_Charles I_ ');
#        self.assertEqual(str(tossups[0].category), 'History - European');
#
#        validBonus = 'This is a valid bonus.  For 10 points each:\n[10] Prompt 1.\nANSWER: _Answer 1_\n[10] Prompt 2.\nANSWER: _Answer 2_\n[10] Prompt 3.\nANSWER: _Answer 3_'
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(validBonus.splitlines())
#        self.assertEqual(len(bonuses), 1)
#        self.assertEqual(str(bonuses[0].question_type), 'ACF-style bonus')
#        self.assertEqual(bonuses[0].leadin, 'This is a valid bonus.  For 10 points each:')
#        self.assertEqual(bonuses[0].part1_text, 'Prompt 1.')
#        self.assertEqual(bonuses[0].part1_answer, '_Answer 1_')
#        self.assertEqual(bonuses[0].part2_text, 'Prompt 2.')
#        self.assertEqual(bonuses[0].part2_answer, '_Answer 2_')
#        self.assertEqual(bonuses[0].part3_text, 'Prompt 3.')
#        self.assertEqual(bonuses[0].part3_answer, '_Answer 3_')
#
#        validBonusWithCategory = validBonus + " {History - American}"
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(validBonusWithCategory.splitlines())
#        self.assertEqual(len(bonuses), 1)
#        self.assertEqual(str(bonuses[0].category), "History - American");
#
#        validVHSLBonus = '[V10] This is a valid VHSL bonus.\nANSWER: _VHSL Answer_'
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(validVHSLBonus.splitlines())
#        self.assertEqual(len(bonuses), 1)
#        self.assertEqual(str(bonuses[0].question_type), 'VHSL bonus')
#        self.assertEqual(bonuses[0].leadin, '')
#        self.assertEqual(bonuses[0].part1_text, 'This is a valid VHSL bonus.')
#        self.assertEqual(bonuses[0].part1_answer, '_VHSL Answer_')
#
#        multipleQuestions = 'This is tossup 1.\nANSWER: _Tossup 1 Answer_\n[V10] This is a VHSL bonus.\nANSWER: _VHSL Answer_\nThis is another tossup.\nANSWER: _Tossup 2 Answer_'
#        multipleQuestions += '\n' + validBonus
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(multipleQuestions.splitlines())
#        self.assertEqual(len(bonuses), 2)
#        self.assertEqual(len(tossups), 2)
#        self.assertEqual(tossups[0].tossup_text, 'This is tossup 1.')
#        self.assertEqual(tossups[0].tossup_answer, '_Tossup 1 Answer_')
#        self.assertEqual(tossups[1].tossup_text, 'This is another tossup.')
#        self.assertEqual(tossups[1].tossup_answer, '_Tossup 2 Answer_')
#        self.assertEqual(bonuses[0].part1_text, 'This is a VHSL bonus.')
#        self.assertEqual(bonuses[0].part1_answer, '_VHSL Answer_')
#        self.assertEqual(str(bonuses[0].question_type), 'VHSL bonus')
#        self.assertEqual(str(bonuses[1].question_type), 'ACF-style bonus')
#        self.assertEqual(bonuses[1].leadin, 'This is a valid bonus.  For 10 points each:')
#        self.assertEqual(bonuses[1].part1_text, 'Prompt 1.')
#        self.assertEqual(bonuses[1].part1_answer, '_Answer 1_')
#        self.assertEqual(bonuses[1].part2_text, 'Prompt 2.')
#        self.assertEqual(bonuses[1].part2_answer, '_Answer 2_')
#        self.assertEqual(bonuses[1].part3_text, 'Prompt 3.')
#        self.assertEqual(bonuses[1].part3_answer, '_Answer 3_')
#
#        multipleQuestionsBlankLines = 'This is tossup 1.\nANSWER: _Tossup 1 Answer_\n\n\nThis is tossup 2.\nANSWER: _Tossup 2 Answer_'
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(multipleQuestionsBlankLines.splitlines())
#        self.assertEqual(len(bonuses), 0)
#        self.assertEqual(len(tossups), 2)
#        self.assertEqual(tossups[0].tossup_text, 'This is tossup 1.')
#        self.assertEqual(tossups[0].tossup_answer, '_Tossup 1 Answer_')
#        self.assertEqual(tossups[1].tossup_text, 'This is tossup 2.')
#        self.assertEqual(tossups[1].tossup_answer, '_Tossup 2 Answer_')
#
#        tossupWithoutAnswer = 'This is not a valid tossup.  It does not have an answer.'
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(tossupWithoutAnswer.splitlines())
#        self.assertEqual(len(tossups), 0)
#        self.assertEqual(len(bonuses), 0)
#
#        tossupWithLineBreaks = 'This is not a valid tossup.\nIt has a line break before its answer.\nANSWER: _foo_'
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(tossupWithLineBreaks.splitlines())
#        self.assertEqual(len(tossups), 1)
#        self.assertEqual(tossups[0].tossup_text, 'It has a line break before its answer.');
#        self.assertEqual(tossups[0].tossup_answer, '_foo_');
#        self.assertEqual(len(bonuses), 0)
#
#        tossupWithoutQuestion = 'ANSWER: This is an answer line without a question'
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(tossupWithoutQuestion.splitlines())
#        self.assertEqual(len(tossups), 0)
#        self.assertEqual(len(bonuses), 0)
#        self.assertEqual(len(tossup_errors), 1)
#
#        tossupWithSingleQuotes = "This is a tossup with 'single quotes' in it.\nANSWER: '_Single Quoted Answer_'"
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(tossupWithSingleQuotes.splitlines())
#        self.assertEqual(len(tossups), 1)
#        self.assertEqual(tossups[0].tossup_text, "This is a tossup with &#39;single quotes&#39; in it.");
#        self.assertEqual(tossups[0].tossup_answer, "&#39;_Single Quoted Answer_&#39;");
#
#        tossupWithDoubleQuotes = 'This is a tossup with "double quotes" in it.\nANSWER: "_Double Quoted Answer_"'
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(tossupWithDoubleQuotes.splitlines())
#        self.assertEqual(len(tossups), 1)
#        self.assertEqual(tossups[0].tossup_text, 'This is a tossup with &quot;double quotes&quot; in it.');
#        self.assertEqual(tossups[0].tossup_answer, '&quot;_Double Quoted Answer_&quot;');
#
#        tossupWithDoubleQuotes = 'This is a tossup with an <i>italic tag</i> in it.\nANSWER: <i>_Italic Answer_</i>'
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(tossupWithDoubleQuotes.splitlines())
#        self.assertEqual(len(tossups), 1)
#        self.assertEqual(tossups[0].tossup_text, 'This is a tossup with an &lt;i&gt;italic tag&lt;/i&gt; in it.');
#        self.assertEqual(tossups[0].tossup_answer, '&lt;i&gt;_Italic Answer_&lt;/i&gt;');
#
#        bonusWithNonIntegerValues = 'This is a bonus with non-integer values.  For 10 points each:\n[A] Prompt 1.\nANSWER: _Answer 1_\n[10.5] Prompt 2.\nANSWER: _Answer 2_\n[10C] Prompt 3.\nANSWER: _Answer 3_'
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(bonusWithNonIntegerValues.splitlines())
#        self.assertEqual(len(bonuses), 0)
#
#        bonusWithHtmlValues = 'This is a bonus with html values.  For 10 points each:\n[<i>10</i>] Prompt 1.\nANSWER: _Answer 1_\n[10] Prompt 2.\nANSWER: _Answer 2_\n[<i>10</i>] Prompt 3.\nANSWER: _Answer 3_'
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(bonusWithHtmlValues.splitlines())
#        self.assertEqual(len(bonuses), 0)
#
#        bonusWithQuotesAndHtml = 'This is a <i>valid</i> "bonus".  For 10 points each:\n[10] <i>Prompt 1</i>.\nANSWER: "_Answer 1_"\n[10] "Prompt 2."\nANSWER: <i>_Answer 2_</i>\n[10] <i>Prompt 3.</i>\nANSWER: <i>_Answer 3_</i>'
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(bonusWithQuotesAndHtml.splitlines())
#        self.assertEqual(len(bonuses), 1)
#        self.assertEqual(str(bonuses[0].question_type), 'ACF-style bonus')
#        self.assertEqual(bonuses[0].leadin, 'This is a &lt;i&gt;valid&lt;/i&gt; &quot;bonus&quot;.  For 10 points each:')
#        self.assertEqual(bonuses[0].part1_text, '&lt;i&gt;Prompt 1&lt;/i&gt;.')
#        self.assertEqual(bonuses[0].part1_answer, '&quot;_Answer 1_&quot;')
#        self.assertEqual(bonuses[0].part2_text, '&quot;Prompt 2.&quot;')
#        self.assertEqual(bonuses[0].part2_answer, '&lt;i&gt;_Answer 2_&lt;/i&gt;')
#        self.assertEqual(bonuses[0].part3_text, '&lt;i&gt;Prompt 3.&lt;/i&gt;')
#        self.assertEqual(bonuses[0].part3_answer, '&lt;i&gt;_Answer 3_&lt;/i&gt;')
#
#        tossupWithUnbalancedSpecialCharsInQuestion = 'This is a tossup question with ~an unclosed tilde.\nANSWER: _foo_'
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(tossupWithUnbalancedSpecialCharsInQuestion.splitlines())
#        self.assertEqual(len(tossup_errors), 1)
#
#        tossupWithUnbalancedSpecialCharsInAnswer = 'This is a tossup question.\nANSWER: _unclosed answer'
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(tossupWithUnbalancedSpecialCharsInAnswer.splitlines())
#        self.assertEqual(len(tossup_errors), 1)
#
#        bonusWithUnbalancedSpecialCharsInLeadin = 'This is a bonus with (unbalanced leadin characters.  For 10 points each:\n[10] <i>Prompt 1</i>.\nANSWER: "_Answer 1_"\n[10] "Prompt 2."\nANSWER: <i>_Answer 2_</i>\n[10] <i>Prompt 3.</i>\nANSWER: <i>_Answer 3_</i>'
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(bonusWithUnbalancedSpecialCharsInLeadin.splitlines())
#        self.assertEqual(len(bonus_errors), 1)
#
#        bonusWithUnbalancedSpecialCharsInPrompts = 'This is a bonus with unbalanced prompt characters.  For 10 points each:\n[10] ~Prompt 1.\nANSWER: "_Answer 1_"\n[10] "Prompt 2."\nANSWER: <i>_Answer 2_</i>\n[10] <i>Prompt 3.</i>\nANSWER: <i>_Answer 3_</i>'
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(bonusWithUnbalancedSpecialCharsInPrompts.splitlines())
#        self.assertEqual(len(bonus_errors), 1)
#
#        bonusWithUnbalancedSpecialCharsInAnswers = 'This is a bonus with unbalanced answer characters.  For 10 points each:\n[10] Prompt 1.\nANSWER: "_Answer 1_"\n[10] "Prompt 2."\nANSWER: _Answer 2\n[10] <i>Prompt 3.</i>\nANSWER: <i>_Answer 3_</i>'
#        tossups, bonuses, tossup_errors, bonus_errors = parse_packet_data(bonusWithUnbalancedSpecialCharsInAnswers.splitlines())
#        self.assertEqual(len(bonus_errors), 1)
#
#    def test_tossup_to_html(self):
#        self.create_generic_distribution()
#
#        acfTossup, created = QuestionType.objects.get_or_create(question_type="ACF-style tossup")
#
#        americanHistory = DistributionEntry(category="History", subcategory="American", distribution=self.dist)
#        americanHistory.save()
#
#        tossup_text = "(Test) ~tossup~."
#        tossup_answer = "_test answer_"
#        tossup_no_category_no_question_type = Tossup(tossup_text=tossup_text, tossup_answer=tossup_answer)
#        expectedOutput = "<p><strong>(Test)</strong> <i>tossup</i>.<br />ANSWER: <u><b>test answer</b></u></p>"
#        self.assertEqual(tossup_no_category_no_question_type.to_html(), expectedOutput)
#        self.assertEqual(tossup_no_category_no_question_type.to_html(include_category=True), expectedOutput)
#        expectedOutputWithCharCount = expectedOutput + "<p><strong>Character Count:</strong> 8</p>"
#        self.assertEqual(tossup_no_category_no_question_type.to_html(include_character_count=True), expectedOutputWithCharCount)
#
#        tossup_with_category = Tossup(tossup_text=tossup_text, tossup_answer=tossup_answer, category=americanHistory)
#        self.assertEqual(tossup_with_category.to_html(), expectedOutput)
#        expectedOutputWithCategory = "<p><strong>(Test)</strong> <i>tossup</i>.<br />ANSWER: <u><b>test answer</b></u></p><p><strong>Category:</strong> History - American</p>"
#        self.assertEqual(tossup_with_category.to_html(include_category=True), expectedOutputWithCategory)
#
#    def test_bonus_to_html(self):
#        self.create_generic_distribution()
#
#        acfBonus, created = QuestionType.objects.get_or_create(question_type="ACF-style bonus")
#
#        vhslBonus, created = QuestionType.objects.get_or_create(question_type="VHSL bonus")
#
#        americanHistory = DistributionEntry(category="History", subcategory="American", distribution=self.dist)
#        americanHistory.save()
#
#        leadin = "Leadin with ~italics~ and (parens) and _underlines_."
#        part1_text = "Part 1 with ~italics~ and (parens) and _underlines_."
#        part1_answer = "_~Answer 1~_ [or foo (bar)]"
#        part2_text = "Part 2."
#        part2_answer = "_answer 2_"
#        part3_text = "Part 3."
#        part3_answer = "_answer 3_"
#        acf_bonus_no_category = Bonus(leadin=leadin, part1_text=part1_text, part1_answer=part1_answer, part2_text=part2_text, part2_answer=part2_answer, part3_text=part3_text, part3_answer=part3_answer)
#        expectedOutput = "<p>Leadin with <i>italics</i> and <strong class="pronunciation-guide">(parens)</strong> and _underlines_.<br />"
#        expectedOutput += "[10] Part 1 with <i>italics</i> and <strong class="pronunciation-guide">(parens)</strong> and _underlines_.<br />"
#        expectedOutput += "ANSWER: <u><b><i>Answer 1</i></b></u> [or foo <strong>(bar)</strong>]<br />"
#        expectedOutput += "[10] Part 2.<br />"
#        expectedOutput += "ANSWER: <u><b>answer 2</b></u><br />"
#        expectedOutput += "[10] Part 3.<br />"
#        expectedOutputWithoutLastLine = expectedOutput
#        expectedOutput += "ANSWER: <u><b>answer 3</b></u></p>"
#        self.assertEqual(acf_bonus_no_category.to_html(), expectedOutput)
#        self.assertEqual(acf_bonus_no_category.to_html(include_category=True), expectedOutput)
#        expectedOutputWithCharCount = expectedOutput + "<p><strong>Character Count:</strong> "
#        expectedOutputWithCharCount += "98</p>"
#        self.assertEqual(acf_bonus_no_category.to_html(include_character_count=True), expectedOutputWithCharCount)
#        acf_bonus_no_category.category = americanHistory
#        self.assertEqual(acf_bonus_no_category.to_html(), expectedOutput)
#        expectedOutputWithCategory = expectedOutputWithoutLastLine + "ANSWER: <u><b>answer 3</b></u></p><p><strong>Category:</strong> History - American</p>"
#        self.assertEqual(acf_bonus_no_category.to_html(include_category=True), expectedOutputWithCategory)
#
#        vhsl_bonus_no_category = Bonus(part1_text=part1_text, part1_answer=part1_answer, question_type=vhslBonus)
#        expectedVhslOutput = "<p>Part 1 with <i>italics</i> and <strong class="pronunciation-guide">(parens)</strong> and _underlines_.<br />"
#        expectedVhslOutput += "ANSWER: <u><b><i>Answer 1</i></b></u> [or foo <strong>(bar)</strong>]</p>"
#        self.assertEqual(vhsl_bonus_no_category.to_html(), expectedVhslOutput)
#        vhsl_bonus_no_category.category = americanHistory
#        self.assertEqual(vhsl_bonus_no_category.to_html(), expectedVhslOutput)
#        expectedVhslOutput = "<p>Part 1 with <i>italics</i> and <strong class="pronunciation-guide">(parens)</strong> and _underlines_.<br />"
#        expectedVhslOutput += "ANSWER: <u><b><i>Answer 1</i></b></u> [or foo <strong>(bar)</strong>]</p><p><strong>Category:</strong> History - American</p>"
#        self.assertEqual(vhsl_bonus_no_category.to_html(include_category=True), expectedVhslOutput)

    def test_character_count_ignores_moderator_instructions(self):
        base = "This is a question about a thing that does stuff."
        base_count = get_character_count(base, True)

        # Leading directive sentences are excluded
        for directive in ("Description acceptable. ",
                          "A description is acceptable. ",
                          "Note to moderator: read the answerline carefully. ",
                          "Note to players: Description acceptable. ",
                          "Two answers required. "):
            self.assertEqual(get_character_count(directive + base, True), base_count, msg=directive)
            self.assertEqual(get_character_count(directive + base, False), len(base), msg=directive)

        # Trailing directives and inline markers are excluded
        self.assertEqual(get_character_count(base + " You have ten seconds.", True), base_count)
        self.assertEqual(get_character_count(base.replace("about", "[emphasize] about"), True), base_count)

        # Mid-sentence content mentions still count
        content = "Critics found the description acceptable in most reviews."
        self.assertEqual(get_character_count(content, True), len(content))

        # Markup-wrapped directives are also excluded
        self.assertEqual(get_character_count("~Description acceptable.~ " + base, True), base_count)

    def test_category_entry_get_requirements_methods(self):
        self.create_generic_period(cefd_fraction=2.2)
                
        self.assertEqual(self.cefd.get_acf_tossup_integer(), 2)
        self.assertEqual(self.cefd.get_acf_tossup_remainder(), 0.20)
        self.assertEqual(self.cefd.get_acf_tossup_upper_bound(), 3)
        self.assertEqual(self.cefd.get_acf_bonus_integer(), 2)
        self.assertEqual(self.cefd.get_acf_bonus_remainder(), 0.20)
        self.assertEqual(self.cefd.get_acf_bonus_upper_bound(), 3)        
        self.assertEqual(self.cefd.get_vhsl_bonus_integer(), 2)
        self.assertEqual(self.cefd.get_vhsl_bonus_remainder(), 0.20)
        self.assertEqual(self.cefd.get_vhsl_bonus_upper_bound(), 3)
        
        self.create_generic_period(cefd_fraction=1)
                
        self.assertEqual(self.cefd.get_acf_tossup_integer(), 1)
        self.assertEqual(self.cefd.get_acf_tossup_remainder(), 0)
        self.assertEqual(self.cefd.get_acf_tossup_upper_bound(), 1)        
        self.assertEqual(self.cefd.get_acf_bonus_integer(), 1)
        self.assertEqual(self.cefd.get_acf_bonus_remainder(), 0)
        self.assertEqual(self.cefd.get_acf_bonus_upper_bound(), 1)                
        self.assertEqual(self.cefd.get_vhsl_bonus_integer(), 1)
        self.assertEqual(self.cefd.get_vhsl_bonus_remainder(), 0)
        self.assertEqual(self.cefd.get_vhsl_bonus_upper_bound(), 1)        
    
    def test_sub_sub_category_to_string(self):
        self.create_generic_distribution()
        
        category_entry = CategoryEntry(category_type=CATEGORY, category_name="Arts")
        category_entry.save()

        sub_category_entry = CategoryEntry(category_type=SUB_CATEGORY, category_name="Arts", sub_category_name="Opera")
        sub_category_entry.save()

        sub_sub_category_entry = CategoryEntry(category_type=SUB_SUB_CATEGORY, category_name="Arts", sub_category_name="Opera", sub_sub_category_name="Baroque")
        sub_sub_category_entry.save()
        
        self.assertEqual(str(category_entry), "Arts")
        self.assertEqual(str(sub_category_entry), "Arts - Opera")                
        self.assertEqual(str(sub_sub_category_entry), "Arts - Opera - Baroque")

    def test_period_wide_entry_reset_current_values(self):        
        self.create_generic_qset()
        
        pwe = PeriodWideEntry.objects.create(period_type=ACF_REGULAR_PERIOD, question_set=self.qset, distribution=self.dist, acf_tossup_cur=5, acf_bonus_cur=5, vhsl_bonus_cur=5, acf_tossup_total=10, acf_bonus_total=10, vhsl_bonus_total=10)
        pwe.save()
        pwe.reset_current_values()
        self.assertEqual(pwe.acf_tossup_cur, 0)
        self.assertEqual(pwe.acf_bonus_cur, 0)
        self.assertEqual(pwe.vhsl_bonus_cur, 0)
        
    def test_period_reset_current_values(self):
        self.create_generic_period()
        self.period.reset_current_values()
        self.assertEqual(self.period.acf_tossup_cur, 0)
        self.assertEqual(self.period.acf_bonus_cur, 0)
        self.assertEqual(self.period.vhsl_bonus_cur, 0)        

    def test_period_wide_category_entry_reset_current_values(self):
        self.create_generic_period()
        self.pwce.reset_current_values()        
        self.assertEqual(self.pwce.acf_tossup_cur_across_periods, 0)
        self.assertEqual(self.pwce.acf_bonus_cur_across_periods, 0)
        self.assertEqual(self.pwce.vhsl_bonus_cur_across_periods, 0)

    def test_period_wide_category_entry_reset_total_values(self):
        self.create_generic_period()
        self.pwce.reset_total_values()
        self.assertEqual(self.pwce.acf_tossup_total_across_periods, 0)
        self.assertEqual(self.pwce.acf_bonus_total_across_periods, 0)
        self.assertEqual(self.pwce.vhsl_bonus_total_across_periods, 0)

    def test_period_wide_category_entry_get_category_type(self):
        self.create_generic_period()
        self.assertEqual(self.pwce.get_category_type(), CATEGORY)        
                
    def test_one_period_category_entry_get_linked_category_entry_for_distribution(self):
        self.create_generic_period()
        linked_entry = self.opce.get_linked_category_entry_for_distribution()
        self.assertEqual(linked_entry, self.cefd)
    
    def test_one_period_category_entry_get_total_questions_all_types(self):
        self.create_generic_period(opce_cur_value=10)
        self.assertEqual(self.opce.get_total_questions_all_types(), 30)
        
    def test_one_period_category_entry_is_over_min_limits(self):
        self.create_generic_period(cefd_fraction=5, opce_cur_value=10)
        self.assertEqual(self.opce.is_over_min_acf_tossup_limit(), True)
        self.assertEqual(self.opce.is_over_min_acf_bonus_limit(), True)
        self.assertEqual(self.opce.is_over_min_vhsl_bonus_limit(), True)
        
        self.create_generic_period(cefd_fraction=4.2, opce_cur_value=4)
        self.assertEqual(self.opce.is_over_min_acf_tossup_limit(), False)
        self.assertEqual(self.opce.is_over_min_acf_bonus_limit(), False)
        self.assertEqual(self.opce.is_over_min_vhsl_bonus_limit(), False)
        
        self.create_generic_period(cefd_fraction=5, opce_cur_value=5)
        self.assertEqual(self.opce.is_over_min_acf_tossup_limit(), False)
        self.assertEqual(self.opce.is_over_min_acf_bonus_limit(), False)
        self.assertEqual(self.opce.is_over_min_vhsl_bonus_limit(), False)        
        
    def test_one_period_category_entry_is_under_max_total_questions_limit(self):
        # It creates (3*10=30) total questions by default, and we've said the max total is 5
        self.create_generic_period(opce_cur_value=10, cefd_max=5)
        self.assertEqual(self.opce.is_under_max_total_questions_limit(), False)
                
        # Now try when the max total is 40        
        self.create_generic_period(opce_cur_value=10, cefd_max=40)
        self.assertEqual(self.opce.is_under_max_total_questions_limit(), True)
        
    def test_one_period_category_entry_is_over_min_total_questions_limit(self):
        # It creates (3*10=30) total questions by default, and we've said the min total is 5
        self.create_generic_period(opce_cur_value=10, cefd_min=5)        
        self.assertEqual(self.opce.is_over_min_total_questions_limit(), True)
                
        # Now try when the min total is 40
        self.create_generic_period(opce_cur_value=10, cefd_min=40)
        self.assertEqual(self.opce.is_over_min_total_questions_limit(), False)
        
    def test_one_period_category_entry_reset_current_values(self):
        self.create_generic_period(opce_cur_value=10)
        self.opce.reset_current_values()
        self.assertEqual(self.opce.acf_tossup_cur_in_period, 0)
        self.assertEqual(self.opce.acf_bonus_cur_in_period, 0)
        self.assertEqual(self.opce.vhsl_bonus_cur_in_period, 0)
    
    #############################################################
    # Packetizer Tests
    #############################################################

    def test_get_parents_from_category_entry(self):
        # Just category entry
        dist = Distribution.objects.create(name="new_distribution", acf_tossup_per_period_count=20, acf_bonus_per_period_count=20, vhsl_bonus_per_period_count=20)
        dist.save()

        dist2 = Distribution.objects.create(name="new_distribution2", acf_tossup_per_period_count=20, acf_bonus_per_period_count=20, vhsl_bonus_per_period_count=20)
        dist2.save()
        
        ce, created = CategoryEntry.objects.get_or_create(category_name="AHistory", category_type=CATEGORY)
        ce.save()
        
        c, sc, ssc = get_parents_from_category_entry(ce)
        self.assertEqual(c, ce)
        self.assertEqual(sc, None)
        self.assertEqual(ssc, None)
        
        # From a subcat with a parent
        
        subcat1, created = CategoryEntry.objects.get_or_create(category_name="AHistory", sub_category_name="European", category_type=SUB_CATEGORY)
        subcat1.save()
        
        c, sc, ssc = get_parents_from_category_entry(subcat1)
        self.assertEqual(c, ce)
        self.assertEqual(sc, subcat1)
        self.assertEqual(ssc, None)        
        
        # From a subcat without a parent for some reason
        
        subcat2, created = CategoryEntry.objects.get_or_create(category_name="ALiterature", sub_category_name="European", category_type=SUB_CATEGORY)
        subcat2.save()
        c, sc, ssc = get_parents_from_category_entry(subcat2)
        self.assertEqual(c, None)
        self.assertEqual(sc, subcat2)
        self.assertEqual(ssc, None)
                
        # From a subsubcat with valid parents
        
        subsubcat1, created = CategoryEntry.objects.get_or_create(category_name="AHistory", sub_category_name="European", sub_sub_category_name="British", category_type=SUB_SUB_CATEGORY)
        subsubcat1.save()
        c, sc, ssc = get_parents_from_category_entry(subsubcat1)                
        self.assertEqual(c, ce)
        self.assertEqual(sc, subcat1)
        self.assertEqual(ssc, subsubcat1)
        
        # From a subsubcat without a parent
        
        subsubcat2, created = CategoryEntry.objects.get_or_create(category_name="AGeography", sub_category_name="World", sub_sub_category_name="French", category_type=SUB_SUB_CATEGORY)
        subsubcat2.save()
        c, sc, ssc = get_parents_from_category_entry(subsubcat2)                        
        self.assertEqual(c, None)
        self.assertEqual(sc, None)
        self.assertEqual(ssc, subsubcat2)

    def test_get_period_entries_from_category_entry(self):
        # Create a category entry and a period
        self.create_generic_period()
        
        ce = CategoryEntry(category_name="History", category_type=CATEGORY)
        ce.save()
                
        pwce, opce = get_period_entries_from_category_entry(self.ce, self.period)
        self.assertEqual(pwce, self.pwce)
        self.assertEqual(opce, self.opce)

    def test_get_parents_from_period_wide_category_entry(self):
        self.create_generic_period()
        
        # (ce, cefd, pwce, opce)
        cats, sub_cats, sub_sub_cats = self.create_generic_ce_hierarchy()
        
        # Try for category
        cat, subcat, subsubcat = get_parents_from_period_wide_category_entry(cats[0][2])
        self.assertEqual(cat, cats[0][2])
        self.assertEqual(subcat, None)
        self.assertEqual(subsubcat, None)
        
        # Try for sub category
        cat, subcat, subsubcat = get_parents_from_period_wide_category_entry(sub_cats[0][2])
        self.assertEqual(cat, cats[0][2])
        self.assertEqual(subcat, sub_cats[0][2])
        self.assertEqual(subsubcat, None)
        
        # Try for sub sub category "Literature - American - Novels" option 
        cat, subcat, subsubcat = get_parents_from_period_wide_category_entry(sub_sub_cats[0][2])
        self.assertEqual(cat, cats[0][2])
        self.assertEqual(subcat, sub_cats[0][2])
        self.assertEqual(subsubcat, sub_sub_cats[0][2])

        # Try for sub sub category "Literature - European - Poetry" option
        cat, subcat, subsubcat = get_parents_from_period_wide_category_entry(sub_sub_cats[2][2])
        self.assertEqual(cat, cats[0][2])
        self.assertEqual(subcat, sub_cats[1][2])
        self.assertEqual(subsubcat, sub_sub_cats[2][2])

    def test_get_parents_from_category_entry_for_distribution(self):
        self.create_generic_period()
        
        # (ce, cefd, pwce, opce)
        cats, sub_cats, sub_sub_cats = self.create_generic_ce_hierarchy()

        # Try for category
        cat, subcat, subsubcat = get_parents_from_category_entry_for_distribution(cats[0][1])
        self.assertEqual(cat, cats[0][1])
        self.assertEqual(subcat, None)
        self.assertEqual(subsubcat, None)
        
        # Try for sub category
        cat, subcat, subsubcat = get_parents_from_category_entry_for_distribution(sub_cats[0][1])
        self.assertEqual(cat, cats[0][1])
        self.assertEqual(subcat, sub_cats[0][1])
        self.assertEqual(subsubcat, None)
        
        # Try for sub sub category "Literature - American - Novels" option 
        cat, subcat, subsubcat = get_parents_from_category_entry_for_distribution(sub_sub_cats[0][1])
        self.assertEqual(cat, cats[0][1])
        self.assertEqual(subcat, sub_cats[0][1])
        self.assertEqual(subsubcat, sub_sub_cats[0][1])

        # Try for sub sub category "Literature - European - Poetry" option
        cat, subcat, subsubcat = get_parents_from_category_entry_for_distribution(sub_sub_cats[2][1])
        self.assertEqual(cat, cats[0][1])
        self.assertEqual(subcat, sub_cats[1][1])
        self.assertEqual(subsubcat, sub_sub_cats[2][1])        

    def test_get_fraction_array(self):
        self.create_generic_period()
                
        fractions = get_fraction_array(self.pwce, 1)
        self.assertEqual(len(fractions), 0)

        fractions = get_fraction_array(self.pwce, 1.101)
        self.assertEqual(len(fractions), 101)        

#        if (c_pwce is not None and c_pwce.acf_tossup_cur_across_periods >= c_pwce.acf_tossup_total_across_periods):
#            return False
#                
#        if (c_opce is not None and c_opce.is_over_min_acf_tossup_limit()):
#            return False
#
#        if (c_opce is not None and c_opce.is_over_min_total_questions_limit()):
#            return False
#        
#        if (sc_pwce is not None and sc_pwce.acf_tossup_cur_across_periods >= sc_pwce.acf_tossup_total_across_periods):
#            return False
#            
#        if (sc_opce is not None and sc_opce.is_over_min_acf_tossup_limit()):
#            return False
#
#        if (sc_opce is not None and sc_opce.is_over_min_total_questions_limit()):
#            return False
#
#        if (ssc_pwce is not None and ssc_pwce.acf_tossup_cur_across_periods >= ssc_pwce.acf_tossup_total_across_periods):
#            return False
#            
#        if (ssc_opce is not None and ssc_opce.is_over_min_acf_tossup_limit()):
#            return False
#
#        if (ssc_opce is not None and ssc_opce.is_over_min_total_questions_limit()):
#            return False
        
        
        pass


class AutoPacketizeTests(TestCase):
    """Tests for packetizer.auto_packetize."""

    def setUp(self):
        self.user = User.objects.create_user(username="packetizeuser", password="top_secret", email="qems2test@gmail.com")
        self.writer = Writer.objects.get(user=self.user.id)

        self.dist = Distribution.objects.create(name="Packetize Test Distribution")
        self.qset = QuestionSet.objects.create(
            name="Packetize Test Set", date=timezone.now(), host="host", address="",
            owner=self.writer, num_packets=4, distribution=self.dist,
            tossups_per_packet=6, bonuses_per_packet=6)

        self.entries = {}
        for category, subcategory in [
                ('History', 'European'), ('History', 'American'),
                ('Literature', 'American'), ('Literature', 'European'),
                ('Science', 'Biology'),
                ('Fine Arts', 'Visual'), ('Fine Arts', 'Music')]:
            entry = DistributionEntry.objects.create(
                distribution=self.dist, category=category, subcategory=subcategory)
            self.entries[(category, subcategory)] = entry

    def _add_tossup(self, entry, text):
        return Tossup.objects.create(
            author=self.writer, question_set=self.qset, tossup_text=text,
            tossup_answer='_answer_', category=entry,
            created_date=datetime.now(), last_changed_date=datetime.now())

    def _add_bonus(self, entry, text):
        return Bonus.objects.create(
            author=self.writer, question_set=self.qset, leadin=text,
            part1_text='p1', part1_answer='_a1_', part2_text='p2', part2_answer='_a2_',
            part3_text='p3', part3_answer='_a3_', category=entry,
            created_date=datetime.now(), last_changed_date=datetime.now())

    def _create_questions(self):
        # Regular load for 4 packets of 6/6:
        # History 2/2 per packet, Literature 1.5/1.5, Science 1/1, Fine Arts 1.5/1.5
        for i in range(4):
            self._add_tossup(self.entries[('History', 'European')], 'HE tu {0}'.format(i))
            self._add_tossup(self.entries[('History', 'American')], 'HA tu {0}'.format(i))
            self._add_bonus(self.entries[('History', 'European')], 'HE bs {0}'.format(i))
            self._add_bonus(self.entries[('History', 'American')], 'HA bs {0}'.format(i))
            self._add_tossup(self.entries[('Science', 'Biology')], 'SB tu {0}'.format(i))
            self._add_bonus(self.entries[('Science', 'Biology')], 'SB bs {0}'.format(i))
        for i in range(3):
            self._add_tossup(self.entries[('Literature', 'American')], 'LA tu {0}'.format(i))
            self._add_tossup(self.entries[('Literature', 'European')], 'LE tu {0}'.format(i))
            self._add_bonus(self.entries[('Literature', 'American')], 'LA bs {0}'.format(i))
            self._add_bonus(self.entries[('Literature', 'European')], 'LE bs {0}'.format(i))
            self._add_tossup(self.entries[('Fine Arts', 'Visual')], 'FV tu {0}'.format(i))
            self._add_tossup(self.entries[('Fine Arts', 'Music')], 'FM tu {0}'.format(i))
            self._add_bonus(self.entries[('Fine Arts', 'Visual')], 'FV bs {0}'.format(i))
            self._add_bonus(self.entries[('Fine Arts', 'Music')], 'FM bs {0}'.format(i))
        # Two extra history tossups that exceed every cap -> tiebreakers
        self._add_tossup(self.entries[('History', 'European')], 'HE extra 1')
        self._add_tossup(self.entries[('History', 'American')], 'HA extra 2')

    def _quotas(self):
        PacketizationEntry.objects.create(
            question_set=self.qset, path='History', depth=0,
            min_tossups=2, max_tossups=2, min_bonuses=2, max_bonuses=2)
        PacketizationEntry.objects.create(
            question_set=self.qset, path='Literature', depth=0,
            min_tossups=Decimal('1.5'), max_tossups=Decimal('1.5'),
            min_bonuses=Decimal('1.5'), max_bonuses=Decimal('1.5'))
        PacketizationEntry.objects.create(
            question_set=self.qset, path='Science', depth=0,
            min_tossups=1, max_tossups=1, min_bonuses=1, max_bonuses=1)
        PacketizationEntry.objects.create(
            question_set=self.qset, path='Fine Arts', depth=0,
            min_tossups=Decimal('1.5'), max_tossups=Decimal('1.5'),
            min_bonuses=Decimal('1.5'), max_bonuses=Decimal('1.5'))
        return build_quota_dict(self.qset)

    def _packetize(self, seed=42):
        report = auto_packetize(self.qset, 4, 6, 6, self._quotas(), created_by=self.writer, seed=seed)
        packets = list(Packet.objects.filter(question_set=self.qset)
                       .exclude(packet_name=EXTRAS_PACKET_NAME).order_by('packet_name'))
        return report, packets

    def _extras_packet(self):
        return Packet.objects.get(question_set=self.qset, packet_name=EXTRAS_PACKET_NAME)

    def test_auto_packetize_basic_structure(self):
        self._create_questions()
        report, packets = self._packetize()

        self.assertEqual(len(packets), 4)
        self.assertTrue(Packet.objects.filter(
            question_set=self.qset, packet_name=EXTRAS_PACKET_NAME).exists())
        for packet in packets:
            regular_tossups = Tossup.objects.filter(packet=packet, question_number__lte=6)
            regular_bonuses = Bonus.objects.filter(packet=packet, question_number__lte=6)
            self.assertEqual(regular_tossups.count(), 6)
            self.assertEqual(regular_bonuses.count(), 6)
            # Sequential numbering from 1
            tu_numbers = sorted(t.question_number for t in Tossup.objects.filter(packet=packet))
            self.assertEqual(tu_numbers[:6], [1, 2, 3, 4, 5, 6])

        # Nothing left unassigned
        self.assertEqual(Tossup.objects.filter(question_set=self.qset, packet=None).count(), 0)
        self.assertEqual(Bonus.objects.filter(question_set=self.qset, packet=None).count(), 0)

    def test_auto_packetize_category_quotas(self):
        self._create_questions()
        report, packets = self._packetize()

        for packet in packets:
            regular_tossups = Tossup.objects.filter(packet=packet, question_number__lte=6)
            regular_bonuses = Bonus.objects.filter(packet=packet, question_number__lte=6)
            tu_by_cat = {}
            bs_by_cat = {}
            for t in regular_tossups:
                tu_by_cat[t.category.category] = tu_by_cat.get(t.category.category, 0) + 1
            for b in regular_bonuses:
                bs_by_cat[b.category.category] = bs_by_cat.get(b.category.category, 0) + 1

            self.assertEqual(tu_by_cat.get('History', 0), 2, msg=str(tu_by_cat))
            self.assertEqual(bs_by_cat.get('History', 0), 2)
            self.assertEqual(tu_by_cat.get('Science', 0), 1)
            self.assertEqual(bs_by_cat.get('Science', 0), 1)
            # Fractional 1.5/1.5: combined total is exactly 3, each type 1 or 2
            for cat in ('Literature', 'Fine Arts'):
                tu = tu_by_cat.get(cat, 0)
                bs = bs_by_cat.get(cat, 0)
                self.assertEqual(tu + bs, 3, msg='{0}: {1}+{2}'.format(cat, tu, bs))
                self.assertIn(tu, (1, 2))
                self.assertIn(bs, (1, 2))

    def test_auto_packetize_subcategory_spread(self):
        self._create_questions()
        report, packets = self._packetize()

        # 4 History-European tossups over 4 packets with 2 History per packet
        # should spread to exactly one per packet
        for packet in packets:
            he = Tossup.objects.filter(packet=packet, question_number__lte=6,
                                       category=self.entries[('History', 'European')])
            self.assertEqual(he.count(), 1)

    def test_auto_packetize_overflow_to_extras(self):
        self._create_questions()
        report, packets = self._packetize()

        # The two questions exceeding every cap land in the Extras packet
        extras = self._extras_packet()
        overflow = Tossup.objects.filter(packet=extras)
        self.assertEqual(overflow.count(), 2)
        self.assertEqual(sorted(t.question_number for t in overflow), [1, 2])
        for tossup in overflow:
            self.assertEqual(tossup.category.category, 'History')
        # Regular packets hold exactly their per-packet limit, nothing extra
        for packet in packets:
            self.assertEqual(Tossup.objects.filter(packet=packet, question_number__gt=6).count(), 0)
        # And the report calls it out
        self.assertTrue(any(EXTRAS_PACKET_NAME in w for w in report['warnings']), report['warnings'])

    def test_auto_packetize_extras_packet_always_exists(self):
        # Even with no overflow, the Extras packet is created (it's the
        # default packet for new questions) and stays out of the rotation
        self._create_questions()
        Tossup.objects.filter(tossup_text__contains='extra').delete()
        report, packets = self._packetize()

        extras = self._extras_packet()
        self.assertEqual(Tossup.objects.filter(packet=extras).count(), 0)
        self.assertEqual(len(packets), 4)
        # Re-running doesn't duplicate it or pull it into the rotation
        report, packets = self._packetize(seed=7)
        self.assertEqual(Packet.objects.filter(
            question_set=self.qset, packet_name=EXTRAS_PACKET_NAME).count(), 1)
        self.assertEqual(len(packets), 4)

    def test_auto_packetize_ordering(self):
        self._create_questions()
        report, packets = self._packetize()

        for packet in packets:
            tossups = list(Tossup.objects.filter(packet=packet, question_number__lte=6).order_by('question_number'))
            cats = [t.category.category for t in tossups]
            adjacencies = sum(1 for i in range(len(cats) - 1) if cats[i] == cats[i + 1])
            self.assertEqual(adjacencies, 0, msg='Packet {0}: {1}'.format(packet.packet_name, cats))
            # History has 2 per packet: no adjacency means at least one
            # question apart.  (Half-balance is not guaranteed: the quarter
            # deal places a 2-question category in consecutive quarters,
            # which can be the same half.)
            history_positions = [i for i, c in enumerate(cats) if c == 'History']
            self.assertEqual(len(history_positions), 2, msg=str(cats))
            self.assertGreaterEqual(history_positions[1] - history_positions[0], 2, msg=str(cats))

    def test_auto_packetize_overwrites_existing_assignments(self):
        self._create_questions()
        report, packets = self._packetize()

        # Cram everything into the first packet with bogus numbers, then re-run
        first = packets[0]
        Tossup.objects.filter(question_set=self.qset).update(packet=first, question_number=999)
        Bonus.objects.filter(question_set=self.qset).update(packet=first, question_number=999)

        report, packets = self._packetize(seed=7)
        for packet in packets:
            self.assertEqual(Tossup.objects.filter(packet=packet, question_number__lte=6).count(), 6)
        self.assertEqual(Tossup.objects.filter(question_number=999).count(), 0)

    def test_auto_packetize_lower_level_max(self):
        self._create_questions()
        # Six more Literature-American tossups (9 total) try to cluster;
        # a subcategory max of 1 per packet must hold for regular slots
        for i in range(3):
            self._add_tossup(self.entries[('Literature', 'American')], 'LA extra {0}'.format(i))
        PacketizationEntry.objects.create(
            question_set=self.qset, path='Literature - American', depth=1,
            max_tossups=1)
        report, packets = self._packetize()

        for packet in packets:
            la = Tossup.objects.filter(packet=packet, question_number__lte=6,
                                       category=self.entries[('Literature', 'American')])
            self.assertLessEqual(la.count(), 1)

    def test_ensure_packets_continues_naming(self):
        Packet.objects.create(question_set=self.qset, packet_name='Round 01', created_by=self.writer)
        Packet.objects.create(question_set=self.qset, packet_name='Round 02', created_by=self.writer)
        packets = _ensure_packets(self.qset, 4, self.writer)
        self.assertEqual([p.packet_name for p in packets],
                         ['Round 01', 'Round 02', 'Round 03', 'Round 04'])


class CrossSetAuthorizationTests(TestCase):
    """A member of one set must not be able to inject into, or manipulate the
    questions of, a set they do not belong to."""

    def setUp(self):
        self.attacker_user = User.objects.create_user('attacker', password='pw', email='atk@test.com')
        self.victim_user = User.objects.create_user('victim', password='pw', email='vic@test.com')
        self.attacker = Writer.objects.get(user=self.attacker_user)
        self.victim = Writer.objects.get(user=self.victim_user)

        self.dist = Distribution.objects.create(name='authz dist')
        self.de = DistributionEntry.objects.create(
            distribution=self.dist, category='History', subcategory='European')

        self.victim_set = self._make_set('Victim Set', self.victim)
        self.attacker_set = self._make_set('Attacker Set', self.attacker)

        self.v_tossup = self._add_tossup(self.victim, self.victim_set, 'secret victim tossup')
        self.v_bonus = self._add_bonus(self.victim, self.victim_set)
        self.atk_packet = Packet.objects.create(
            question_set=self.attacker_set, packet_name='Atk P1', created_by=self.attacker)

        self.client.login(username='attacker', password='pw')

    def _make_set(self, name, owner):
        return QuestionSet.objects.create(
            name=name, date=timezone.now(), host='h', address='', owner=owner,
            num_packets=1, distribution=self.dist, tossups_per_packet=1, bonuses_per_packet=1)

    def _add_tossup(self, author, qset, text):
        return Tossup.objects.create(
            author=author, question_set=qset, tossup_text=text, tossup_answer='_secret_',
            category=self.de, created_date=datetime.now(), last_changed_date=datetime.now(),
            question_number=1)

    def _add_bonus(self, author, qset):
        return Bonus.objects.create(
            author=author, question_set=qset, leadin='secret', part1_text='p1', part1_answer='_a1_',
            part2_text='p2', part2_answer='_a2_', part3_text='p3', part3_answer='_a3_',
            category=self.de, created_date=datetime.now(), last_changed_date=datetime.now(),
            question_number=1)

    def test_complete_upload_rejects_non_member(self):
        before = Tossup.objects.filter(question_set=self.victim_set).count()
        self.client.post('/complete_upload/', {
            'qset-id': self.victim_set.id, 'num-tossups': 1, 'num-bonuses': 0,
            'tossup-text-0': 'INJECTED', 'tossup-answer-0': '_x_',
            'tossup-category-0': 'History - European', 'tossup-type-0': 'ACF-style tossup'})
        self.assertEqual(Tossup.objects.filter(question_set=self.victim_set).count(), before)
        self.assertFalse(Tossup.objects.filter(
            question_set=self.victim_set, tossup_text__contains='INJECTED').exists())

    def test_assign_tossups_ignores_foreign_question(self):
        self.client.post('/assign_tossups_to_packet/', {
            'packet_id': self.atk_packet.id, 'tossup_ids[]': [str(self.v_tossup.id)]})
        self.v_tossup.refresh_from_db()
        self.assertNotEqual(self.v_tossup.packet_id, self.atk_packet.id)

    def test_assign_bonuses_ignores_foreign_question(self):
        self.client.post('/assign_bonuses_to_packet/', {
            'packet_id': self.atk_packet.id, 'bonus_ids[]': [str(self.v_bonus.id)]})
        self.v_bonus.refresh_from_db()
        self.assertNotEqual(self.v_bonus.packet_id, self.atk_packet.id)

    def test_bulk_change_author_ignores_foreign_question(self):
        self.client.post('/bulk_change_set/{0}/'.format(self.attacker_set.id), {
            'confirm': '1', 'change-type': 'author-step2', 'new-author': self.attacker.id,
            'num-tossups': 1, 'num-bonuses': 0, 'tossup-id-0': str(self.v_tossup.id)})
        self.v_tossup.refresh_from_db()
        self.assertEqual(self.v_tossup.author_id, self.victim.id)

    def test_bulk_move_ignores_foreign_question(self):
        self.client.post('/bulk_change_set/{0}/'.format(self.attacker_set.id), {
            'confirm': '1', 'change-type': 'move-step2', 'new-set': self.attacker_set.id,
            'num-tossups': 1, 'num-bonuses': 0, 'tossup-id-0': str(self.v_tossup.id)})
        self.v_tossup.refresh_from_db()
        self.assertEqual(self.v_tossup.question_set_id, self.victim_set.id)

    def test_change_question_order_ignores_foreign_question(self):
        self.client.post('/change_question_order/', {
            'packet_id': self.atk_packet.id, 'num_questions': 1, 'question_type': 'tossup',
            'order_data[0][id]': str(self.v_tossup.id), 'order_data[0][order]': 99})
        self.v_tossup.refresh_from_db()
        self.assertEqual(self.v_tossup.question_number, 1)

    def test_legitimate_assign_still_works(self):
        own = self._add_tossup(self.attacker, self.attacker_set, 'own tossup')
        self.client.post('/assign_tossups_to_packet/', {
            'packet_id': self.atk_packet.id, 'tossup_ids[]': [str(own.id)]})
        own.refresh_from_db()
        self.assertEqual(own.packet_id, self.atk_packet.id)


class PronunciationGuideTests(TestCase):
    """Verified-OL pronunciation-guide suggestions (qsub.pron_dict), used by the
    style checker. These rely on the bundled data/pronunciations.json."""

    def setUp(self):
        from qems2.qsub import pron_dict
        pron_dict.reset_cache()

    def _terms(self, text):
        from qems2.qsub.pron_dict import suggest_guides
        return {term for term, _pron in suggest_guides(text)}

    def test_suggests_known_proper_noun(self):
        self.assertIn('Goethe', self._terms('The poet Goethe wrote this work.'))

    def test_suppressed_when_guide_already_present(self):
        # A parenthetical guide right after the term should suppress the hint.
        self.assertNotIn('Goethe', self._terms('The poet Goethe ("GUR-tuh") wrote this.'))

    def test_suppressed_when_respelling_present_elsewhere(self):
        from qems2.qsub.pron_dict import suggest_guides
        # If the exact respelling already appears, don't re-suggest it.
        text = 'Goethe, pronounced GUR-tuh, wrote this.'
        self.assertNotIn('Goethe', {t for t, _ in suggest_guides(text)})

    def test_no_suggestion_for_common_words(self):
        # Common English words are stoplisted out of the dictionary.
        self.assertEqual(self._terms('They were at the table by the water.'), set())

    def test_each_term_suggested_once(self):
        from qems2.qsub.pron_dict import suggest_guides
        pairs = suggest_guides('Goethe admired Goethe and again Goethe.')
        self.assertEqual(len([t for t, _ in pairs if t == 'Goethe']), 1)


class StyleCheckFixTests(TestCase):
    """Auto-apply and dismiss of style-check issues (pronunciation guides etc.)."""

    def setUp(self):
        from qems2.qsub import pron_dict
        pron_dict.reset_cache()
        self.owner_user = User.objects.create_user('sc_owner', password='pw', email='o@test.com')
        self.owner = Writer.objects.get(user=self.owner_user)
        self.dist = Distribution.objects.create(name='sc dist')
        self.de = DistributionEntry.objects.create(
            distribution=self.dist, category='Literature', subcategory='European')
        self.qset = QuestionSet.objects.create(
            name='SC Set', date=timezone.now(), host='h', address='', owner=self.owner,
            num_packets=1, distribution=self.dist, tossups_per_packet=1, bonuses_per_packet=1)
        self.tu = Tossup.objects.create(
            author=self.owner, question_set=self.qset,
            tossup_text='The poet Goethe wrote this famous work.', tossup_answer='_Faust_',
            category=self.de, created_date=datetime.now(), last_changed_date=datetime.now(),
            question_number=1)
        self.client.login(username='sc_owner', password='pw')

    def test_apply_pronunciation_fix_inserts_guide(self):
        resp = self.client.post('/apply_style_fix/', {
            'question_type': 'tossup', 'question_id': self.tu.id,
            'code': 'pronunciation', 'token': 'Question|Goethe', 'guide': 'minkowski'})
        self.assertEqual(resp.status_code, 200)
        self.tu.refresh_from_db()
        self.assertIn('GUR-tuh', self.tu.tossup_text)

    def test_apply_is_idempotent_after_guide_present(self):
        self.client.post('/apply_style_fix/', {
            'question_type': 'tossup', 'question_id': self.tu.id,
            'code': 'pronunciation', 'token': 'Question|Goethe', 'guide': 'minkowski'})
        # Second apply: the issue no longer exists, so it should report failure.
        resp = self.client.post('/apply_style_fix/', {
            'question_type': 'tossup', 'question_id': self.tu.id,
            'code': 'pronunciation', 'token': 'Question|Goethe', 'guide': 'minkowski'})
        self.assertEqual(resp.status_code, 400)

    def test_dismiss_hides_issue_on_page(self):
        before = self.client.get('/style_check/%d/?guide=minkowski' % self.qset.id)
        self.assertIn(b'GUR-tuh', before.content)
        self.client.post('/dismiss_style_issue/', {
            'question_type': 'tossup', 'question_id': self.tu.id,
            'code': 'pronunciation', 'token': 'Question|Goethe'})
        self.assertTrue(StyleIssueDismissal.objects.filter(
            question_id=self.tu.id, code='pronunciation', token='Question|Goethe').exists())
        after = self.client.get('/style_check/%d/?guide=minkowski' % self.qset.id)
        self.assertNotIn(b'GUR-tuh', after.content)

    def test_restore_undismisses(self):
        self.client.post('/dismiss_style_issue/', {
            'question_type': 'tossup', 'question_id': self.tu.id,
            'code': 'pronunciation', 'token': 'Question|Goethe'})
        self.client.post('/dismiss_style_issue/', {
            'question_type': 'tossup', 'question_id': self.tu.id,
            'code': 'pronunciation', 'token': 'Question|Goethe', 'action': 'restore'})
        self.assertFalse(StyleIssueDismissal.objects.filter(
            question_id=self.tu.id, code='pronunciation').exists())


class YappAnswerCategoryImportTests(TestCase):
    """Categories embedded in YAPP answer lines (e.g. ANSWER: _x_ <Biology>)
    are parsed, mapped onto existing set categories, and generated for new sets."""

    def setUp(self):
        import json as _json
        from django.core.files.uploadedfile import SimpleUploadedFile
        self._json = _json
        self._Upload = SimpleUploadedFile
        QuestionType.objects.get_or_create(question_type=ACF_STYLE_TOSSUP)
        QuestionType.objects.get_or_create(question_type=ACF_STYLE_BONUS)
        QuestionType.objects.get_or_create(question_type=VHSL_BONUS)
        self.user, _ = User.objects.get_or_create(username="yapp_importer")
        self.writer = Writer.objects.get(user=self.user.id)

    def _file(self, name, payload):
        data = self._json.dumps(payload).encode('utf-8')
        return self._Upload(name, data, content_type='application/json')

    def test_metadata_field_author_comma_category(self):
        # The de-facto standard: a "<Author, Category - Subcategory>" line after
        # the answer, which YAPP stores (brackets stripped) in `metadata`.
        from qems2.qsub.packet_set_importer import import_packets_from_files
        payload = {'tossups': [
            {'question': 'A plant process. (*) Sugar.',
             'answer': 'ANSWER: _Photosynthesis_',
             'metadata': 'Jane Smith, Science - Biology'}]}
        summary = import_packets_from_files(
            [self._file('P.json', payload)], 'Meta Set', self.writer)
        tu = Tossup.objects.get(question_set=summary['question_set'])
        self.assertEqual((tu.category.category, tu.category.subcategory),
                         ('Science', 'Biology'))

    def test_metadata_field_category_without_author(self):
        # "<Painting - 1900-2000>" form: no author, no comma.
        from qems2.qsub.packet_set_importer import import_packets_from_files
        payload = {'tossups': [
            {'question': 'A famous work. (*) Oil on canvas.',
             'answer': 'ANSWER: _The Scream_',
             'metadata': 'Painting - 1900-2000'}]}
        summary = import_packets_from_files(
            [self._file('P.json', payload)], 'NoAuthor Set', self.writer)
        tu = Tossup.objects.get(question_set=summary['question_set'])
        self.assertEqual((tu.category.category, tu.category.subcategory),
                         ('Painting', '1900-2000'))

    def test_metadata_lone_token_stays_uncategorized(self):
        # A bare "<Jane Smith>" is an author, not a category.
        from qems2.qsub.packet_set_importer import import_packets_from_files
        payload = {'tossups': [
            {'question': 'Something. (*) here.',
             'answer': 'ANSWER: _Thing_', 'metadata': 'Jane Smith'}]}
        summary = import_packets_from_files(
            [self._file('P.json', payload)], 'Lone Set', self.writer)
        tu = Tossup.objects.get(question_set=summary['question_set'])
        self.assertIsNone(tu.category)

    def test_new_set_generates_categories_from_answer_line(self):
        from qems2.qsub.packet_set_importer import import_packets_from_files
        payload = {
            'tossups': [
                {'question': 'A plant process. (*) It makes sugar.',
                 'answer': 'ANSWER: _Photosynthesis_ &lt;Biology&gt;', 'metadata': 'Jane Doe'},
                {'question': 'A famous physicist. (*) Relativity.',
                 'answer': 'ANSWER: _Einstein_ &lt;Science - Physics&gt;', 'metadata': ''},
            ],
            'bonuses': [
                {'leadin': 'Answer these.', 'parts': ['Part one.', 'Part two.', 'Part three.'],
                 'answers': ['_one_', '_two_', '_three_ &lt;History&gt;'], 'metadata': ''},
            ],
        }
        summary = import_packets_from_files(
            [self._file('Packet 1.json', payload)], 'YAPP Set', self.writer)
        qset = summary['question_set']
        self.assertEqual(summary['tossups'], 2)
        self.assertEqual(summary['bonuses'], 1)

        entries = {(e.category, e.subcategory)
                   for e in DistributionEntry.objects.filter(distribution=qset.distribution)}
        self.assertIn(('Biology', ''), entries)
        self.assertIn(('Science', 'Physics'), entries)
        self.assertIn(('History', ''), entries)

        photo = Tossup.objects.get(tossup_answer__contains='Photosynthesis')
        self.assertIsNotNone(photo.category)
        self.assertEqual(photo.category.category, 'Biology')
        # The trailing tag is stripped from the stored answer.
        self.assertNotIn('Biology', photo.tossup_answer)
        self.assertNotIn('&lt;', photo.tossup_answer)

        bonus = Bonus.objects.get(question_set=qset)
        self.assertEqual(bonus.category.category, 'History')
        self.assertNotIn('History', bonus.part3_answer)

    def test_existing_set_maps_answer_category_to_existing_entry(self):
        from qems2.qsub.packet_set_importer import import_packets_into_set
        dist = Distribution.objects.create(name='Existing Dist')
        qset = QuestionSet.objects.create(
            name='Existing Set', date=timezone.now(), host='', address='',
            owner=self.writer, num_packets=1, distribution=dist)
        existing = DistributionEntry.objects.create(
            distribution=dist, category='Science', subcategory='Biology')
        SetWideDistributionEntry.objects.create(
            question_set=qset, dist_entry=existing, num_tossups=0, num_bonuses=0)

        before = DistributionEntry.objects.filter(distribution=dist).count()
        payload = {'tossups': [
            {'question': 'A plant process. (*) Sugar.',
             'answer': 'ANSWER: _Photosynthesis_ &lt;Biology&gt;', 'metadata': ''}]}
        import_packets_into_set([self._file('Add.json', payload)], qset, self.writer)

        # A bare "Biology" answer tag resolves onto "Science - Biology"; no dup.
        self.assertEqual(DistributionEntry.objects.filter(distribution=dist).count(), before)
        tu = Tossup.objects.get(question_set=qset)
        self.assertEqual(tu.category_id, existing.id)


class BuzzHistoryLinkTests(TestCase):
    """Buzzes/results link to the question version current when they were
    recorded — for web play and for Discord imports — via TossupBuzz.history_url."""

    def setUp(self):
        import json as _json
        self._json = _json
        QuestionType.objects.get_or_create(question_type=ACF_STYLE_TOSSUP)
        QuestionType.objects.get_or_create(question_type=ACF_STYLE_BONUS)
        self.acf_tu = QuestionType.objects.get(question_type=ACF_STYLE_TOSSUP)
        self.acf_bn = QuestionType.objects.get(question_type=ACF_STYLE_BONUS)
        self.owner_user = User.objects.create_user('bz_owner', password='pw', email='b@test.com')
        self.owner = Writer.objects.get(user=self.owner_user)
        self.dist = Distribution.objects.create(name='bz dist')
        self.qset = QuestionSet.objects.create(
            name='BZ Set', date=timezone.now(), host='h', address='', owner=self.owner,
            num_packets=1, distribution=self.dist)
        # Owners are added to the editor set in practice (the importer does this);
        # the history view authorizes by writer/editor membership.
        self.owner.question_set_editor.add(self.qset)
        self.tu = Tossup(
            author=self.owner, question_set=self.qset, question_type=self.acf_tu,
            tossup_text='An early version of the stem. (*) The end.',
            tossup_answer='_Photosynthesis_', question_number=1)
        self.tu.save_question(edit_type=QUESTION_CREATE, changer=self.owner)
        self.bn = Bonus(
            author=self.owner, question_set=self.qset, question_type=self.acf_bn,
            leadin='Name these.', part1_text='P1', part1_answer='_Alpha_',
            part2_text='P2', part2_answer='_Beta_', part3_text='P3', part3_answer='_Gamma_',
            question_number=1)
        self.bn.save_question(edit_type=QUESTION_CREATE, changer=self.owner)

    def test_latest_history_is_most_recent_version(self):
        first = self.tu.latest_history()
        self.tu.tossup_text = 'An edited version of the stem. (*) The end.'
        self.tu.save_question(edit_type=QUESTION_EDIT, changer=self.owner)
        second = self.tu.latest_history()
        self.assertNotEqual(first.id, second.id)
        self.assertGreater(second.id, first.id)

    def test_web_buzz_links_to_current_version(self):
        self.client.login(username='bz_owner', password='pw')
        resp = self.client.post('/record_buzz/', {
            'tossup_id': self.tu.id, 'correct': 'true', 'powered': 'false',
            'neg': 'false', 'buzz_word_index': 3, 'total_words': 8, 'char_position': 20,
            'answer_given': 'photosynthesis'})
        self.assertTrue(self._json.loads(resp.content)['success'])
        buzz = TossupBuzz.objects.get(tossup=self.tu)
        self.assertIsNotNone(buzz.tossup_history_id)
        self.assertEqual(buzz.tossup_history_id, self.tu.latest_history().id)
        self.assertEqual(
            buzz.history_url(),
            '/tossup_history/{0}/?v={1}#version-{1}'.format(self.tu.id, buzz.tossup_history_id))

    def test_buzz_keeps_old_version_after_question_is_edited(self):
        self.client.login(username='bz_owner', password='pw')
        self.client.post('/record_buzz/', {
            'tossup_id': self.tu.id, 'correct': 'true', 'buzz_word_index': 2,
            'total_words': 8, 'char_position': 10})
        buzz = TossupBuzz.objects.get(tossup=self.tu)
        played_version = buzz.tossup_history_id
        # Edit the question after the buzz: the buzz still points at the old text.
        self.tu.tossup_text = 'Totally rewritten. (*) Done.'
        self.tu.save_question(edit_type=QUESTION_EDIT, changer=self.owner)
        buzz.refresh_from_db()
        self.assertEqual(buzz.tossup_history_id, played_version)
        self.assertNotEqual(buzz.tossup_history_id, self.tu.latest_history().id)

    def test_web_bonus_result_links_to_current_version(self):
        self.client.login(username='bz_owner', password='pw')
        resp = self.client.post('/record_bonus_result/', {
            'bonus_id': self.bn.id, 'part1_correct': 'true',
            'part2_correct': 'false', 'part3_correct': 'true'})
        self.assertTrue(self._json.loads(resp.content)['success'])
        result = BonusResult.objects.get(bonus=self.bn)
        self.assertEqual(result.bonus_history_id, self.bn.latest_history().id)
        self.assertEqual(
            result.history_url(),
            '/bonus_history/{0}/?v={1}#version-{1}'.format(self.bn.id, result.bonus_history_id))

    def test_discord_buzz_links_to_current_version(self):
        key = SetApiKey.objects.create(
            question_set=self.qset, key=SetApiKey.generate_token(),
            active=True, created_by=self.owner)
        body = self._json.dumps({'events': [
            {'external_id': 'd1', 'answer': 'Photosynthesis', 'player_name': 'Alice',
             'buzz_word_index': 4, 'total_words': 8, 'correct': True}]})
        resp = self.client.post('/api/v1/buzzes', data=body,
                                content_type='application/json',
                                HTTP_AUTHORIZATION='Bearer ' + key.key)
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(self._json.loads(resp.content)['results'][0]['status'], 'recorded')
        buzz = TossupBuzz.objects.get(external_id='d1')
        self.assertEqual(buzz.source, PLAYTEST_SOURCE_DISCORD)
        self.assertEqual(buzz.tossup_history_id, self.tu.latest_history().id)

    def test_history_view_highlights_linked_version(self):
        self.client.login(username='bz_owner', password='pw')
        hid = self.tu.latest_history().id
        resp = self.client.get('/tossup_history/{0}/?v={1}'.format(self.tu.id, hid))
        self.assertEqual(resp.status_code, 200)
        html = resp.content.decode()
        self.assertIn('id="version-{0}"'.format(hid), html)
        self.assertIn('history-version-highlight', html)


class DiscordCommentDisplayTests(TestCase):
    """Discord (bot) comments are distinguishable from human comments and carry
    a link to the Discord thread."""

    def setUp(self):
        from django.contrib.contenttypes.models import ContentType
        from django.contrib.sites.models import Site
        from django_comments.models import Comment
        self.Comment = Comment
        self.site = Site.objects.get_current()
        QuestionType.objects.get_or_create(question_type=ACF_STYLE_TOSSUP)
        self.owner_user = User.objects.create_user('dc_owner', password='pw', email='d@test.com')
        self.owner = Writer.objects.get(user=self.owner_user)
        self.dist = Distribution.objects.create(name='dc dist')
        self.qset = QuestionSet.objects.create(
            name='DC Set', date=timezone.now(), host='h', address='', owner=self.owner,
            num_packets=1, distribution=self.dist)
        self.tu = Tossup.objects.create(
            author=self.owner, question_set=self.qset,
            tossup_text='Stem. (*) end.', tossup_answer='_Answer_',
            created_date=datetime.now(), last_changed_date=datetime.now(), question_number=1)
        self.tu_ct = ContentType.objects.get_for_model(Tossup)

    def _comment(self, text, user=None, user_name=''):
        return self.Comment.objects.create(
            content_type=self.tu_ct, object_pk=str(self.tu.id), site=self.site,
            user=user, user_name=user_name, comment=text, is_public=True, is_removed=False)

    def test_ref_comment_marked_with_thread_url(self):
        from qems2.qsub.model_utils import mark_discord_comments
        thread = DiscordThread.objects.create(
            question_set=self.qset, tossup=self.tu,
            url='https://discord.com/channels/1/2/3', title='Thread')
        bot = self._comment('From the bot', user=None, user_name=DISCORD_BOT_NAME)
        DiscordCommentRef.objects.create(external_id='c1', comment=bot, question_set=self.qset)
        human = self._comment('From a person', user=self.owner_user, user_name='')

        mark_discord_comments([bot, human])
        self.assertTrue(bot.is_discord)
        self.assertEqual(bot.discord_thread_url, thread.url)
        self.assertFalse(human.is_discord)
        self.assertEqual(human.discord_thread_url, '')

    def test_botname_fallback_without_ref(self):
        from qems2.qsub.model_utils import mark_discord_comments
        # No DiscordCommentRef, but posted under the bot name with no user.
        bot = self._comment('orphaned bot comment', user=None, user_name=DISCORD_BOT_NAME)
        mark_discord_comments([bot])
        self.assertTrue(bot.is_discord)
        # No thread recorded -> empty url, still flagged as Discord.
        self.assertEqual(bot.discord_thread_url, '')

    def test_threaded_comments_tag_marks_discord(self):
        from qems2.qsub.templatetags.filters import get_threaded_comments
        DiscordThread.objects.create(
            question_set=self.qset, tossup=self.tu, url='https://discord.com/x', title='T')
        bot = self._comment('bot', user=None, user_name=DISCORD_BOT_NAME)
        DiscordCommentRef.objects.create(external_id='c2', comment=bot, question_set=self.qset)
        data = get_threaded_comments(self.tu)
        marked = {c.id: c for c in data['top_level']}
        self.assertTrue(marked[bot.id].is_discord)
        self.assertEqual(marked[bot.id].discord_thread_url, 'https://discord.com/x')

    def test_edit_tossup_page_shows_discord_styling(self):
        self.owner.question_set_editor.add(self.qset)
        self.client.login(username='dc_owner', password='pw')
        DiscordThread.objects.create(
            question_set=self.qset, tossup=self.tu, url='https://discord.com/y', title='T')
        bot = self._comment('bot says hi', user=None, user_name=DISCORD_BOT_NAME)
        DiscordCommentRef.objects.create(external_id='c3', comment=bot, question_set=self.qset)
        resp = self.client.get('/edit_tossup/{0}/'.format(self.tu.id))
        html = resp.content.decode()
        self.assertEqual(resp.status_code, 200)
        self.assertIn('comment-discord', html)
        self.assertIn('discord-badge', html)
        self.assertIn('https://discord.com/y', html)
        self.assertIn('toggle-discord-comments', html)  # the filter


class PacketizedWordExportTests(TestCase):
    """The 'Export Packetized Word' output follows the actual packets/order,
    uses Times New Roman 12 + narrow margins, the <Author, Category - Subcategory>
    ~Id~ <Editor: Name> attribution line, keeps each question as one paragraph,
    and attaches open comments as Word comments."""

    def setUp(self):
        import io as _io, zipfile as _zip
        from docx import Document as _Doc
        from django.contrib.contenttypes.models import ContentType
        from django.contrib.sites.models import Site
        from django_comments.models import Comment
        self._io, self._zip, self._Doc = _io, _zip, _Doc
        self.Comment = Comment
        self.site = Site.objects.get_current()
        QuestionType.objects.get_or_create(question_type=ACF_STYLE_TOSSUP)
        QuestionType.objects.get_or_create(question_type=ACF_STYLE_BONUS)
        self.acf_tu = QuestionType.objects.get(question_type=ACF_STYLE_TOSSUP)
        self.acf_bn = QuestionType.objects.get(question_type=ACF_STYLE_BONUS)
        self.ou = User.objects.create_user('pw_owner', password='pw', email='p@test.com')
        self.ou.first_name, self.ou.last_name = 'Pat', 'Writer'
        self.ou.save()
        self.owner = Writer.objects.get(user=self.ou)
        self.dist = Distribution.objects.create(
            name='pw dist', acf_tossup_per_period_count=1, acf_bonus_per_period_count=1)
        self.de = DistributionEntry.objects.create(
            distribution=self.dist, category='Science', subcategory='Biology')
        self.qset = QuestionSet.objects.create(
            name='PW Set', date=timezone.now(), host='h', address='', owner=self.owner,
            num_packets=2, distribution=self.dist)
        self.owner.question_set_editor.add(self.qset)
        # Two packets, natural-sort order check: "Packet 10" must come after "Packet 2".
        self.p2 = Packet.objects.create(question_set=self.qset, packet_name='Packet 2', created_by=self.owner)
        self.p10 = Packet.objects.create(question_set=self.qset, packet_name='Packet 10', created_by=self.owner)
        self.tu_p2 = Tossup.objects.create(
            author=self.owner, editor=self.owner, question_set=self.qset, packet=self.p2,
            question_type=self.acf_tu, category=self.de, edited=True,
            tossup_text='Packet two stem. (*) end.', tossup_answer='_Photosynthesis_',
            created_date=datetime.now(), last_changed_date=datetime.now(), question_number=1)
        self.tu_p10 = Tossup.objects.create(
            author=self.owner, question_set=self.qset, packet=self.p10,
            question_type=self.acf_tu, category=self.de,
            tossup_text='Packet ten stem. (*) end.', tossup_answer='_Mitochondria_',
            created_date=datetime.now(), last_changed_date=datetime.now(), question_number=1)
        self.bn_p2 = Bonus.objects.create(
            author=self.owner, question_set=self.qset, packet=self.p2,
            question_type=self.acf_bn, category=self.de,
            leadin='Lead.', part1_text='P1', part1_answer='_Alpha_',
            part2_text='P2', part2_answer='_Beta_', part3_text='P3', part3_answer='_Gamma_',
            created_date=datetime.now(), last_changed_date=datetime.now(), question_number=1)
        self.tu_ct = ContentType.objects.get_for_model(Tossup)
        self.client.login(username='pw_owner', password='pw')

    def _open_zip(self, output_format):
        resp = self.client.get('/export_question_set/{0}/{1}/'.format(self.qset.id, output_format))
        self.assertEqual(resp.status_code, 200)
        return self._zip.ZipFile(self._io.BytesIO(resp.content))

    def _doc(self, zf, name):
        return self._Doc(self._io.BytesIO(zf.read(name)))

    def test_zip_has_one_docx_per_packet_natural_sorted(self):
        zf = self._open_zip('docx-packetized')
        names = zf.namelist()
        self.assertIn('Packet 2.docx', names)
        self.assertIn('Packet 10.docx', names)
        self.assertIn('Answer Matrix.xlsx', names)

    def test_packet_docx_uses_its_own_questions(self):
        zf = self._open_zip('docx-packetized')
        text2 = '\n'.join(p.text for p in self._doc(zf, 'Packet 2.docx').paragraphs)
        text10 = '\n'.join(p.text for p in self._doc(zf, 'Packet 10.docx').paragraphs)
        self.assertIn('Photosynthesis', text2)
        self.assertNotIn('Mitochondria', text2)
        self.assertIn('Mitochondria', text10)

    def test_style_font_and_margins(self):
        zf = self._open_zip('docx-packetized')
        doc = self._doc(zf, 'Packet 2.docx')
        from docx.shared import Pt, Inches
        self.assertEqual(doc.styles['Normal'].font.name, 'Times New Roman')
        self.assertEqual(doc.styles['Normal'].font.size, Pt(12))
        sec = doc.sections[0]
        self.assertEqual(sec.left_margin, Inches(0.5))
        self.assertEqual(sec.top_margin, Inches(0.5))

    def test_attribution_line_format(self):
        zf = self._open_zip('docx-packetized')
        doc = self._doc(zf, 'Packet 2.docx')
        full = '\n'.join(p.text for p in doc.paragraphs)
        # Edited tossup -> editor included; matches <Author, Cat - Sub> ~id~ <Editor: ...>
        self.assertIn('<{0}, Science - Biology> ~{1}~ <Editor: {0}>'.format(
            self.owner.get_real_name().strip(), self.tu_p2.id), full)

    def test_each_tossup_is_single_paragraph(self):
        zf = self._open_zip('docx-packetized')
        doc = self._doc(zf, 'Packet 2.docx')
        # The tossup stem, ANSWER, and attribution share one paragraph.
        tu_paras = [p for p in doc.paragraphs
                    if 'Packet two stem' in p.text]
        self.assertEqual(len(tu_paras), 1)
        self.assertIn('ANSWER:', tu_paras[0].text)
        self.assertIn('~{0}~'.format(self.tu_p2.id), tu_paras[0].text)
        self.assertTrue(tu_paras[0].paragraph_format.keep_together)

    def test_open_comments_become_word_comments_resolved_excluded(self):
        from qems2.qsub.models import CommentResolution
        open_c = self.Comment.objects.create(
            content_type=self.tu_ct, object_pk=str(self.tu_p2.id), site=self.site,
            user=self.ou, user_name='', comment='please fix the power mark',
            is_public=True, is_removed=False)
        resolved_c = self.Comment.objects.create(
            content_type=self.tu_ct, object_pk=str(self.tu_p2.id), site=self.site,
            user=self.ou, user_name='', comment='this one is handled',
            is_public=True, is_removed=False)
        CommentResolution.objects.create(comment=resolved_c, resolved=True, resolved_by=self.owner)
        zf = self._open_zip('docx-packetized')
        # python-docx exposes comments on the document.
        doc = self._doc(zf, 'Packet 2.docx')
        texts = [c.text for c in doc.comments]
        self.assertTrue(any('please fix the power mark' in t for t in texts))
        self.assertFalse(any('this one is handled' in t for t in texts))

    def test_html_entities_decoded_and_blank_author_omitted(self):
        # A tossup whose stored text has HTML entities (as YAPP import produces)
        # and whose author has no real name.
        blank_user = User.objects.create_user('noname', password='pw', email='n@test.com')
        blank_writer = Writer.objects.get(user=blank_user)
        tu = Tossup.objects.create(
            author=blank_writer, question_set=self.qset, packet=self.p2,
            question_type=self.acf_tu, category=self.de,
            tossup_text='This artist&#x27;s work &amp; legacy. (*) end.',
            tossup_answer='_Church_', created_date=datetime.now(),
            last_changed_date=datetime.now(), question_number=2)
        zf = self._open_zip('docx-packetized')
        doc = self._doc(zf, 'Packet 2.docx')
        full = '\n'.join(p.text for p in doc.paragraphs)
        self.assertIn("This artist's work & legacy.", full)
        self.assertNotIn('&#x27;', full)
        self.assertNotIn('&amp;', full)
        # Blank author -> no "< ," artifact; just the category.
        self.assertIn('<Science - Biology> ~{0}~'.format(tu.id), full)
        self.assertNotIn('< ,', full)


class QuestionSetSettingsSaveTests(TestCase):
    """Editing question-set settings (e.g. max tossup character count) persists.
    Regression: tossups_per_packet/bonuses_per_packet were required by the form
    but never rendered, so every save failed validation silently."""

    def setUp(self):
        self.ou = User.objects.create_user('qs_owner', password='pw', email='q@test.com')
        self.owner = Writer.objects.get(user=self.ou)
        self.dist = Distribution.objects.create(name='qs dist')
        self.qset = QuestionSet.objects.create(
            name='QS Set', date=timezone.now(), host='', address='', owner=self.owner,
            num_packets=5, distribution=self.dist, max_acf_tossup_length=725)
        self.client.login(username='qs_owner', password='pw')

    def test_form_is_valid_with_rendered_fields_only(self):
        from qems2.qsub.forms import QuestionSetForm
        form = QuestionSetForm(data={
            'name': 'QS Set', 'date': '01/01/2026', 'num_packets': '5',
            'distribution': str(self.dist.id), 'max_acf_tossup_length': '198',
            'max_acf_bonus_length': '650',
            'char_count_ignores_pronunciation_guides': 'on', 'tossups_only': ''})
        self.assertTrue(form.is_valid(), form.errors)

    def test_edit_persists_max_tossup_length(self):
        resp = self.client.post('/edit_question_set/{0}/'.format(self.qset.id), {
            'name': 'QS Set', 'date': '01/01/2026', 'num_packets': '5',
            'distribution': str(self.dist.id), 'max_acf_tossup_length': '198',
            'max_acf_bonus_length': '650',
            'char_count_ignores_pronunciation_guides': 'on', 'tossups_only': ''})
        self.assertEqual(resp.status_code, 200)
        self.qset.refresh_from_db()
        self.assertEqual(self.qset.max_acf_tossup_length, 198)
        # The packetization fields kept their defaults (not wiped by the form).
        self.assertEqual(self.qset.tossups_per_packet, 20)


class UnpacketizedAssignmentTests(TestCase):
    """Unpacketized questions: assign into empty slots + new packets, unassign,
    and grid context exposes them."""

    def setUp(self):
        QuestionType.objects.get_or_create(question_type=ACF_STYLE_TOSSUP)
        QuestionType.objects.get_or_create(question_type=ACF_STYLE_BONUS)
        self.acf_tu = QuestionType.objects.get(question_type=ACF_STYLE_TOSSUP)
        self.ou = User.objects.create_user('up_owner', password='pw', email='u@test.com')
        self.owner = Writer.objects.get(user=self.ou)
        self.dist = Distribution.objects.create(name='up dist')
        self.qset = QuestionSet.objects.create(
            name='UP Set', date=timezone.now(), host='', address='', owner=self.owner,
            num_packets=1, distribution=self.dist, tossups_per_packet=2, bonuses_per_packet=2)
        self.p1 = Packet.objects.create(question_set=self.qset, packet_name='Packet 01', created_by=self.owner)
        # Packet 01 has a tossup at slot 1; slot 2 is empty.
        self.placed = self._tu('placed', packet=self.p1, number=1)
        self.client.login(username='up_owner', password='pw')

    def _tu(self, ans, packet=None, number=None):
        return Tossup.objects.create(
            author=self.owner, question_set=self.qset, packet=packet, question_number=number,
            question_type=self.acf_tu, tossup_text='Stem (*) end.', tossup_answer='_%s_' % ans,
            created_date=datetime.now(), last_changed_date=datetime.now())

    def test_assign_fills_empty_slot_then_new_packets(self):
        # 3 unpacketized tossups: 1 fills Packet 01 slot 2, 2 go to a new packet.
        a = self._tu('aaa'); b = self._tu('bbb'); c = self._tu('ccc')
        resp = self.client.post('/assign_unpacketized/', {'qset_id': self.qset.id})
        self.assertTrue(json.loads(resp.content)['success'])
        a.refresh_from_db(); b.refresh_from_db(); c.refresh_from_db()
        # One of them filled Packet 01 slot 2.
        in_p1 = [q for q in (a, b, c) if q.packet_id == self.p1.id]
        self.assertEqual(len(in_p1), 1)
        self.assertEqual(in_p1[0].question_number, 2)
        # A new packet was created for the other two.
        self.assertEqual(Packet.objects.filter(question_set=self.qset).count(), 2)
        self.assertEqual(Tossup.objects.filter(question_set=self.qset, packet=None).count(), 0)
        # The already-placed question is untouched.
        self.placed.refresh_from_db()
        self.assertEqual(self.placed.question_number, 1)

    def test_assign_noop_when_nothing_unpacketized(self):
        resp = self.client.post('/assign_unpacketized/', {'qset_id': self.qset.id})
        data = json.loads(resp.content)
        self.assertTrue(data['success'])
        self.assertEqual(Packet.objects.filter(question_set=self.qset).count(), 1)

    def test_unassign_returns_question_to_pool(self):
        resp = self.client.post('/unassign_packet_question/', {
            'question_type': 'tossup', 'question_id': self.placed.id})
        self.assertTrue(json.loads(resp.content)['success'])
        self.placed.refresh_from_db()
        self.assertIsNone(self.placed.packet_id)
        self.assertIsNone(self.placed.question_number)

    def test_grid_shows_unpacketized(self):
        self._tu('lonely')
        resp = self.client.get('/packet_grid/{0}/'.format(self.qset.id))
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.context['unassigned_tu'], 1)
        self.assertContains(resp, 'unpacketized-panel')
        self.assertContains(resp, 'assign-unpacketized-btn')

    def test_writer_cannot_assign(self):
        wu = User.objects.create_user('up_writer', password='pw', email='w2@test.com')
        writer = Writer.objects.get(user=wu)
        self.qset.writer.add(writer)
        self.client.logout(); self.client.login(username='up_writer', password='pw')
        self._tu('x')
        resp = self.client.post('/assign_unpacketized/', {'qset_id': self.qset.id})
        self.assertFalse(json.loads(resp.content)['success'])

    def test_swap_candidates_includes_unpacketized(self):
        # Source is a packetized tossup; an unpacketized tossup in the same
        # category should show up as a swap candidate flagged unpacketized.
        de = DistributionEntry.objects.create(
            distribution=self.dist, category='Science', subcategory='Biology')
        src = self._tu('source', packet=self.p1, number=2)
        src.category = de; src.save()
        free = self._tu('freebie')
        free.category = de; free.save()
        resp = self.client.get('/swap_candidates/', {
            'question_type': 'tossup', 'question_id': src.id, 'scope': 'top'})
        data = json.loads(resp.content)
        cands = {c['id']: c for c in data['candidates']}
        self.assertIn(free.id, cands)
        self.assertTrue(cands[free.id]['unpacketized'])
        self.assertEqual(cands[free.id]['packet_name'], '(unpacketized)')

    def test_grid_empty_cells_open_fill_dialog(self):
        # Leave a gap (slot 2 empty, slot 3 filled) so a real empty cell renders
        # the "+ place" trigger that opens the swap dialog (not a new-question link).
        self._tu('gap', packet=self.p1, number=3)
        resp = self.client.get('/packet_grid/{0}/'.format(self.qset.id))
        self.assertEqual(resp.status_code, 200)
        html = resp.content.decode()
        self.assertIn('grid-fill', html)
        self.assertNotIn('/add_tossups/', html)

    def test_swap_candidates_fill_mode_lists_unpacketized(self):
        free = self._tu('freebie')  # unpacketized
        self._tu('placed2', packet=self.p1, number=2)
        resp = self.client.get('/swap_candidates/', {
            'question_type': 'tossup', 'qset_id': self.qset.id})
        data = json.loads(resp.content)
        self.assertTrue(data['fill_mode'])
        ids = {c['id'] for c in data['candidates']}
        self.assertIn(free.id, ids)
        self.assertNotIn(self.placed.id, ids)

    def test_doc_view_has_swap_buttons(self):
        self._tu('swapme', packet=self.p1, number=1)
        resp = self.client.get('/view_packet/{0}/'.format(self.p1.id))
        self.assertEqual(resp.status_code, 200)
        html = resp.content.decode()
        self.assertIn('doc-swap-btn', html)
        self.assertIn('doc-swap-dialog', html)

    def test_packet_revision_changes_on_edit_and_move(self):
        # The document-view staleness token changes when a question in the
        # packet is edited or moved out.
        from qems2.qsub.views import _packet_revision
        from datetime import datetime
        tu = self._tu('rev', packet=self.p1, number=2)
        rev1 = _packet_revision(self.p1)
        # Endpoint returns the same token.
        resp = self.client.get('/packet_revision/{0}/'.format(self.p1.id))
        self.assertEqual(json.loads(resp.content)['revision'], rev1)
        # Editing a question's timestamp changes it.
        tu.last_changed_date = datetime(2030, 1, 1, 12, 0, 0)
        tu.save()
        rev2 = _packet_revision(self.p1)
        self.assertNotEqual(rev1, rev2)
        # Moving a question out of the packet changes it again.
        tu.packet = None; tu.question_number = None; tu.save()
        self.assertNotEqual(rev2, _packet_revision(self.p1))


class ImportEntityStorageTests(TestCase):
    """YAPP import stores literal punctuation (apostrophes/ampersands), not
    HTML entities, and the cleanup command repairs old escaped data."""

    def setUp(self):
        import json as _json
        from django.core.files.uploadedfile import SimpleUploadedFile
        self._json = _json
        self._Upload = SimpleUploadedFile
        QuestionType.objects.get_or_create(question_type=ACF_STYLE_TOSSUP)
        QuestionType.objects.get_or_create(question_type=ACF_STYLE_BONUS)
        self.user, _ = User.objects.get_or_create(username='ent_importer')
        self.writer = Writer.objects.get(user=self.user.id)

    def _file(self, name, payload):
        return self._Upload(name, self._json.dumps(payload).encode('utf-8'),
                            content_type='application/json')

    def test_import_stores_literal_apostrophe_and_ampersand(self):
        from qems2.qsub.packet_set_importer import import_packets_from_files
        payload = {'tossups': [
            {'question': "This author&#x27;s novel about Crosby &amp; Nash. (*) end.",
             'answer': 'ANSWER: _Pratt &amp; Whitney_', 'metadata': 'A, Science - Physics'}]}
        summary = import_packets_from_files([self._file('P.json', payload)], 'Ent Set', self.writer)
        tu = Tossup.objects.get(question_set=summary['question_set'])
        self.assertIn("author's novel", tu.tossup_text)
        self.assertIn("Crosby & Nash", tu.tossup_text)
        self.assertNotIn('&#x27;', tu.tossup_text)
        self.assertNotIn('&amp;', tu.tossup_text)
        self.assertIn('Pratt & Whitney', tu.tossup_answer)

    def test_cleanup_command_fixes_existing(self):
        from django.core.management import call_command
        from datetime import datetime
        acf = QuestionType.objects.get(question_type=ACF_STYLE_TOSSUP)
        dist = Distribution.objects.create(name='c dist')
        qset = QuestionSet.objects.create(
            name='C Set', date=timezone.now(), host='', address='', owner=self.writer,
            num_packets=1, distribution=dist)
        tu = Tossup.objects.create(
            author=self.writer, question_set=qset, question_type=acf,
            tossup_text='Bach&#x27;s fugue &amp; chorale &lt;kept&gt;.',
            tossup_answer='_J.S. Bach_', created_date=datetime.now(),
            last_changed_date=datetime.now(), question_number=1)
        call_command('fix_question_entities')
        tu.refresh_from_db()
        self.assertEqual(tu.tossup_text, "Bach's fugue & chorale &lt;kept&gt;.")
        # Angle brackets left escaped so they can't inject a tag.
        self.assertIn('&lt;kept&gt;', tu.tossup_text)


class PacketCommentAndPostTests(TestCase):
    """Packet-level comments via the generic post_comment endpoint, and the
    careful-notes moderator heads-up in the doc view."""

    def setUp(self):
        from django.contrib.contenttypes.models import ContentType
        from django_comments.models import Comment
        from datetime import datetime
        self.Comment = Comment
        self.ContentType = ContentType
        QuestionType.objects.get_or_create(question_type=ACF_STYLE_TOSSUP)
        self.acf = QuestionType.objects.get(question_type=ACF_STYLE_TOSSUP)
        self.ou = User.objects.create_user('pc_owner', password='pw', email='pc@test.com')
        self.owner = Writer.objects.get(user=self.ou)
        self.dist = Distribution.objects.create(name='pc dist')
        self.qset = QuestionSet.objects.create(
            name='PC Set', date=timezone.now(), host='', address='', owner=self.owner,
            num_packets=1, distribution=self.dist)
        self.owner.question_set_editor.add(self.qset)
        self.packet = Packet.objects.create(question_set=self.qset, packet_name='Packet 01', created_by=self.owner)
        self.tu = Tossup.objects.create(
            author=self.owner, question_set=self.qset, packet=self.packet, question_type=self.acf,
            tossup_text='Stem (*) end.', tossup_answer='_Tricky Answer_', read_carefully=True,
            created_date=datetime.now(), last_changed_date=datetime.now(), question_number=1)
        self.client.login(username='pc_owner', password='pw')

    def test_post_packet_comment(self):
        resp = self.client.post('/post_comment/', {
            'target_type': 'packet', 'target_id': self.packet.id,
            'qset_id': self.qset.id, 'comment_text': 'Packet-wide note'})
        self.assertTrue(json.loads(resp.content)['success'])
        ct = self.ContentType.objects.get_for_model(Packet)
        self.assertTrue(self.Comment.objects.filter(
            content_type=ct, object_pk=str(self.packet.id), comment='Packet-wide note').exists())

    def test_post_tossup_comment_no_security_form(self):
        # No django_comments timestamp/security_hash needed (the stale-tab fix).
        resp = self.client.post('/post_comment/', {
            'target_type': 'tossup', 'target_id': self.tu.id,
            'qset_id': self.qset.id, 'comment_text': 'looks good'})
        self.assertTrue(json.loads(resp.content)['success'])
        ct = self.ContentType.objects.get_for_model(Tossup)
        self.assertTrue(self.Comment.objects.filter(
            content_type=ct, object_pk=str(self.tu.id)).exists())

    def test_post_comment_rejects_non_member(self):
        other = User.objects.create_user('pc_outsider', password='pw', email='o3@test.com')
        Writer.objects.get(user=other)
        self.client.logout(); self.client.login(username='pc_outsider', password='pw')
        resp = self.client.post('/post_comment/', {
            'target_type': 'packet', 'target_id': self.packet.id,
            'qset_id': self.qset.id, 'comment_text': 'nope'})
        self.assertFalse(json.loads(resp.content)['success'])

    def test_doc_view_shows_packet_comments_and_careful_notes(self):
        self.client.post('/post_comment/', {
            'target_type': 'packet', 'target_id': self.packet.id,
            'qset_id': self.qset.id, 'comment_text': 'Packet-wide note'})
        resp = self.client.get('/view_packet/{0}/'.format(self.packet.id))
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode()
        self.assertIn('packet-comments-panel', body)
        self.assertIn('Packet-wide note', body)
        # Careful-notes heads-up for the read_carefully tossup.
        self.assertIn('read these answer lines carefully', body)
        self.assertIn('Tricky Answer', body)

    def test_packet_status_shows_subcategories(self):
        de = DistributionEntry.objects.create(
            distribution=self.dist, category='Painting', subcategory='1900-2000')
        SetWideDistributionEntry.objects.create(
            question_set=self.qset, dist_entry=de, num_tossups=4, num_bonuses=0)
        resp = self.client.get('/edit_packet/{0}/'.format(self.packet.id))
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode()
        self.assertIn('Painting', body)        # top-level row
        self.assertIn('1900-2000', body)       # subcategory detail row
        self.assertIn('status-subcat', body)


class TossupsOnlyCompletionTests(TestCase):
    """packet_completion omits bonuses for tossups-only sets."""

    def setUp(self):
        QuestionType.objects.get_or_create(question_type=ACF_STYLE_TOSSUP)
        self.acf = QuestionType.objects.get(question_type=ACF_STYLE_TOSSUP)
        self.ou = User.objects.create_user('to_owner', password='pw', email='to@test.com')
        self.owner = Writer.objects.get(user=self.ou)
        self.dist = Distribution.objects.create(name='to dist')

    def _make_set(self, tossups_only):
        from datetime import datetime
        qset = QuestionSet.objects.create(
            name='TO Set', date=timezone.now(), host='', address='', owner=self.owner,
            num_packets=1, distribution=self.dist, tossups_only=tossups_only,
            tossups_per_packet=2, bonuses_per_packet=2)
        packet = Packet.objects.create(question_set=qset, packet_name='P1', created_by=self.owner)
        for i in (1, 2):
            Tossup.objects.create(author=self.owner, question_set=qset, packet=packet,
                question_type=self.acf, tossup_text='S (*) e.', tossup_answer='_A_',
                created_date=datetime.now(), last_changed_date=datetime.now(), question_number=i)
        return packet

    def test_tossups_only_completion_has_no_bonus(self):
        from qems2.qsub.templatetags.filters import packet_completion
        s = packet_completion(self._make_set(True))
        self.assertIn('TU', s)
        self.assertNotIn(' B)', s)
        self.assertIn('100%', s)

    def test_normal_completion_shows_bonus(self):
        from qems2.qsub.templatetags.filters import packet_completion
        s = packet_completion(self._make_set(False))
        self.assertIn('TU', s)
        self.assertIn(' B)', s)


class PacketOrderingTests(TestCase):
    """Natural packet ordering, extras-last, and user-set custom order."""

    def setUp(self):
        self.ou = User.objects.create_user('po_owner', password='pw', email='po@test.com')
        self.owner = Writer.objects.get(user=self.ou)
        self.dist = Distribution.objects.create(name='po dist')
        self.qset = QuestionSet.objects.create(
            name='PO Set', date=timezone.now(), host='', address='', owner=self.owner,
            num_packets=1, distribution=self.dist)
        self.owner.question_set_editor.add(self.qset)
        names = ['Round 10', 'Round 2', 'Round 1', 'Extras', 'Round 9']
        self.packets = {n: Packet.objects.create(question_set=self.qset, packet_name=n, created_by=self.owner)
                        for n in names}
        self.client.login(username='po_owner', password='pw')

    def test_natural_sort_extras_last(self):
        from qems2.qsub.model_utils import sorted_packets
        order = [p.packet_name for p in sorted_packets(self.qset)]
        self.assertEqual(order, ['Round 1', 'Round 2', 'Round 9', 'Round 10', 'Extras'])

    def test_custom_order_overrides(self):
        ids = [self.packets['Round 2'].id, self.packets['Extras'].id, self.packets['Round 1'].id,
               self.packets['Round 9'].id, self.packets['Round 10'].id]
        resp = self.client.post('/set_packet_order/', {
            'qset_id': self.qset.id, 'packet_ids[]': [str(i) for i in ids]})
        self.assertTrue(json.loads(resp.content)['success'])
        from qems2.qsub.model_utils import sorted_packets
        order = [p.packet_name for p in sorted_packets(self.qset)]
        self.assertEqual(order, ['Round 2', 'Extras', 'Round 1', 'Round 9', 'Round 10'])

    def test_grid_has_reorder_and_unassign_controls(self):
        resp = self.client.get('/packet_grid/{0}/'.format(self.qset.id))
        html = resp.content.decode()
        self.assertIn('reorder-packets-btn', html)
        self.assertIn('reorder-packets-dialog', html)


class GridUnassignTests(TestCase):
    """The per-cell unassign control removes a question from its packet."""

    def setUp(self):
        from datetime import datetime
        QuestionType.objects.get_or_create(question_type=ACF_STYLE_TOSSUP)
        self.acf = QuestionType.objects.get(question_type=ACF_STYLE_TOSSUP)
        self.ou = User.objects.create_user('gu_owner', password='pw', email='gu@test.com')
        self.owner = Writer.objects.get(user=self.ou)
        self.dist = Distribution.objects.create(name='gu dist')
        self.qset = QuestionSet.objects.create(
            name='GU Set', date=timezone.now(), host='', address='', owner=self.owner,
            num_packets=1, distribution=self.dist)
        self.owner.question_set_editor.add(self.qset)
        self.packet = Packet.objects.create(question_set=self.qset, packet_name='Round 1', created_by=self.owner)
        self.tu = Tossup.objects.create(author=self.owner, question_set=self.qset, packet=self.packet,
            question_type=self.acf, tossup_text='S (*) e.', tossup_answer='_A_',
            created_date=datetime.now(), last_changed_date=datetime.now(), question_number=1)
        self.client.login(username='gu_owner', password='pw')

    def test_unassign_button_rendered_and_works(self):
        resp = self.client.get('/packet_grid/{0}/'.format(self.qset.id))
        self.assertIn('unassign-btn', resp.content.decode())
        r = self.client.post('/unassign_packet_question/', {
            'question_type': 'tossup', 'question_id': self.tu.id})
        self.assertTrue(json.loads(r.content)['success'])
        self.tu.refresh_from_db()
        self.assertIsNone(self.tu.packet_id)


class AccountAgeAndDistributionPermissionTests(TestCase):
    """New accounts can't create sets/distributions for 2 days; distributions
    are only visible/editable for sets the user belongs to."""

    def setUp(self):
        from datetime import timedelta
        QuestionType.objects.get_or_create(question_type=ACF_STYLE_TOSSUP)
        QuestionType.objects.get_or_create(question_type=ACF_STYLE_BONUS)
        QuestionType.objects.get_or_create(question_type=VHSL_BONUS)
        # Old-enough member who owns a set.
        self.old_user = User.objects.create_user('old_u', password='pw', email='o@test.com')
        self.old_user.date_joined = timezone.now() - timedelta(days=5)
        self.old_user.save()
        self.old = Writer.objects.get(user=self.old_user)
        self.my_dist = Distribution.objects.create(name='Mine')
        self.qset = QuestionSet.objects.create(
            name='Mine Set', date=timezone.now(), host='', address='', owner=self.old,
            num_packets=1, distribution=self.my_dist)
        self.old.question_set_editor.add(self.qset)
        # Someone else's distribution the old user has no part in.
        self.other_dist = Distribution.objects.create(name='Theirs')
        other_user = User.objects.create_user('other_u', password='pw', email='x@test.com')
        other = Writer.objects.get(user=other_user)
        QuestionSet.objects.create(
            name='Other Set', date=timezone.now(), host='', address='', owner=other,
            num_packets=1, distribution=self.other_dist)
        # Brand-new account.
        self.new_user = User.objects.create_user('new_u', password='pw', email='n@test.com')
        self.new_user.date_joined = timezone.now()
        self.new_user.save()
        Writer.objects.get(user=self.new_user)

    def test_new_account_cannot_create_question_set(self):
        self.client.login(username='new_u', password='pw')
        resp = self.client.get('/create_question_set/')
        self.assertContains(resp, '2 days after sign-up')

    def test_old_account_can_reach_create_question_set(self):
        self.client.login(username='old_u', password='pw')
        resp = self.client.get('/create_question_set/')
        self.assertNotContains(resp, '2 days after sign-up')

    def test_new_account_cannot_create_distribution(self):
        self.client.login(username='new_u', password='pw')
        resp = self.client.get('/edit_distribution/')  # no id = new dist
        self.assertContains(resp, '2 days after sign-up')

    def test_distributions_list_only_shows_own(self):
        self.client.login(username='old_u', password='pw')
        resp = self.client.get('/distributions/')
        self.assertEqual(resp.status_code, 200)
        dist_ids = {d.id for d in resp.context['dists']}
        self.assertIn(self.my_dist.id, dist_ids)
        self.assertNotIn(self.other_dist.id, dist_ids)

    def test_cannot_edit_other_distribution(self):
        self.client.login(username='old_u', password='pw')
        resp = self.client.get('/edit_distribution/{0}/'.format(self.other_dist.id))
        self.assertContains(resp, 'only view or edit distributions for your own sets')

    def test_can_edit_own_distribution(self):
        self.client.login(username='old_u', password='pw')
        resp = self.client.get('/edit_distribution/{0}/'.format(self.my_dist.id))
        self.assertNotContains(resp, 'only view or edit distributions for your own sets')


class PublicSetAndJoinRequestTests(TestCase):
    """Public sets are listed for non-members, who can request to join (email)."""

    def setUp(self):
        from django.core import mail
        self.mail = mail
        self.ou = User.objects.create_user('ps_owner', password='pw', email='owner@test.com')
        self.owner = Writer.objects.get(user=self.ou)
        self.dist = Distribution.objects.create(name='ps dist')
        self.public_set = QuestionSet.objects.create(
            name='Public Set', date=timezone.now(), host='', address='', owner=self.owner,
            num_packets=1, distribution=self.dist, public=True)
        self.private_set = QuestionSet.objects.create(
            name='Private Set', date=timezone.now(), host='', address='', owner=self.owner,
            num_packets=1, distribution=Distribution.objects.create(name='pd'), public=False)
        self.outsider_u = User.objects.create_user('ps_out', password='pw', email='out@test.com')
        self.outsider_u.first_name = 'Out'; self.outsider_u.save()
        Writer.objects.get(user=self.outsider_u)

    def test_public_set_listed_for_non_member(self):
        self.client.login(username='ps_out', password='pw')
        resp = self.client.get('/')
        self.assertEqual(resp.status_code, 200)
        ids = {qs.id for qs in resp.context['public_sets']}
        self.assertIn(self.public_set.id, ids)
        self.assertNotIn(self.private_set.id, ids)

    def test_public_set_not_listed_for_owner(self):
        self.client.login(username='ps_owner', password='pw')
        resp = self.client.get('/')
        ids = {qs.id for qs in resp.context['public_sets']}
        self.assertNotIn(self.public_set.id, ids)  # already a member

    def test_request_to_join_emails_owner(self):
        self.client.login(username='ps_out', password='pw')
        self.mail.outbox = []
        resp = self.client.post('/request_to_join/', {
            'qset_id': self.public_set.id, 'message': 'I would love to help write science.'})
        self.assertTrue(json.loads(resp.content)['success'])
        self.assertEqual(len(self.mail.outbox), 1)
        msg = self.mail.outbox[0]
        self.assertIn('owner@test.com', msg.to)
        self.assertIn('Public Set', msg.subject)
        self.assertIn('love to help write science', msg.body)
        self.assertEqual(msg.reply_to, ['out@test.com'])

    def test_request_to_join_rejected_for_private(self):
        self.client.login(username='ps_out', password='pw')
        resp = self.client.post('/request_to_join/', {'qset_id': self.private_set.id})
        self.assertFalse(json.loads(resp.content)['success'])

    def test_member_cannot_request_to_join(self):
        self.client.login(username='ps_owner', password='pw')
        resp = self.client.post('/request_to_join/', {'qset_id': self.public_set.id})
        self.assertFalse(json.loads(resp.content)['success'])
